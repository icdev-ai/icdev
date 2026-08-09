#!/usr/bin/env python3
# CUI // SP-CTI
"""
ICDEV™ Web Dashboard - Flask Application
========================================
Provides a web interface for monitoring projects, agents, compliance,
and system health within the ICDEV™ framework.

Usage:
    python tools/dashboard/app.py [--port 5050] [--debug]
"""

import argparse
import importlib
import json
import os  # noqa: F811 — needed directly (not just as _os)
import sys
import threading
import time
import uuid
from datetime import datetime, timedelta, timezone
from pathlib import Path

# ---------------------------------------------------------------------------
# Path setup  (so `tools.dashboard.config` is importable when run directly)
# ---------------------------------------------------------------------------
BASE_DIR = Path(__file__).resolve().parent.parent.parent
if str(BASE_DIR) not in sys.path:
    sys.path.insert(0, str(BASE_DIR))

from tools.logging.icdev_logger import get_logger  # noqa: E402
from tools.db.storage import get_connection, sql_placeholder, table_exists  # noqa: E402

from flask import (
    Flask,
    render_template,
    jsonify,
    request as flask_request,
    g,
    session as flask_session,
    redirect,
    url_for,
    send_from_directory,
    make_response,
    Response,
)  # noqa: E402

from tools.dashboard.config import (  # noqa: E402
    DB_PATH,
    CUI_BANNER_TOP,
    CUI_BANNER_BOTTOM,
    CUI_DESIGNATION,
    CUI_BANNER_ENABLED,
    DEFAULT_CLASSIFICATION,
    BYOK_ENABLED,
    PORT,
    HOST,
    DEBUG,
)
from tools.dashboard.brand import brand_context_processor  # noqa: E402
from tools.dashboard.lazy_canvas import install_lazy_canvas_loader  # noqa: E402
from tools.dashboard.auth import register_dashboard_auth, validate_api_key, log_auth_event, require_role  # noqa: E402
from tools.dashboard.websocket import init_socketio, get_socketio  # noqa: E402
from tools.dashboard.findings_aggregator import (  # noqa: E402
    aggregate_findings as _aggregate_findings,
    close_canvas_connections as _close_canvas_connections,
)
# P1.1: Centralized API blueprint registration (replaces 50+ individual imports)
from tools.dashboard.api import register_api_blueprints  # noqa: E402

logger = get_logger("icdev.dashboard.app")
try:
    from tools.usage_analytics.event_collector import track_request as _track_request
except ImportError:
    _track_request = None


def _session_actor(default: str = "dashboard") -> str:
    """Resolved actor identity for audit fields (nav-sec-06).

    Always sourced from the authenticated session user (``g.current_user``),
    never from a spoofable request-body ``actor`` field. Falls back to
    ``default`` only when no user is on the request context (which the
    ``@require_role`` gate on the calling route already prevents).
    """
    user = getattr(g, "current_user", None)
    if isinstance(user, dict):
        return str(user.get("id") or user.get("email") or default)
    if user is not None:
        try:
            return str(user["id"])
        except Exception:
            return default
    return default


def _current_user_id() -> "str | None":
    """Return the authenticated user's id, or None for anonymous requests.

    nav-misc-01: used to persist per-user preferences (e.g. the last-seen
    release version) instead of relying on a per-browser cookie.
    """
    user = getattr(g, "current_user", None)
    if isinstance(user, dict):
        uid = user.get("id")
        return str(uid) if uid else None
    return None


# nav-sec-06: editorial (state-changing) Pulse blog actions — approve / reject /
# publish / unpublish / delete — are restricted to an approver/editor-class role.
# A lowest-privilege ``developer`` must not be able to approve or publish
# externally visible content. Reads (GET) stay open to any authenticated user.
# Every role here also appears in ``VALID_DASHBOARD_ROLES`` in
# tools/dashboard/auth.py.

# nav-intel-06: the Ask-ICDEV Q&A endpoints (/api/components-map/ask and
# /api/ask-icdev/sessions/<id>/message) route every narrate=true request through
# the LLMRouter — a real cost surface. Both endpoints are already
# authenticated-only (the global before_request 401s anon), but any
# authenticated user could otherwise trigger unbounded LLM calls. Role-gating
# narration to admin/pm would break the chat UX for the developer/isso/co roles
# that legitimately use it, so we rate-limit the narration per user instead
# (defense-in-depth against authenticated abuse) and degrade to raw evidence
# when the budget is exhausted — the non-narration retrieval branches stay live.
_NARRATION_RATE_LOCK = threading.Lock()
_NARRATION_CALLS: dict[str, list[float]] = {}


def _narration_budget_ok(key: str, *, max_per_min: int | None = None) -> bool:
    """Per-key sliding-window budget for the LLM narration cost surface.

    Records a call and returns True while under the per-minute cap; returns
    False (→ caller falls back to raw evidence) once the cap is exhausted.
    Cap defaults to ICDEV_ASK_NARRATION_MAX_PER_MIN (20).
    """
    try:
        cap = max_per_min if max_per_min is not None else int(
            os.environ.get("ICDEV_ASK_NARRATION_MAX_PER_MIN", "20")
        )
    except (TypeError, ValueError):
        cap = 20
    now = time.monotonic()
    with _NARRATION_RATE_LOCK:
        recent = [t for t in _NARRATION_CALLS.get(key, ()) if now - t < 60.0]
        if len(recent) >= cap:
            _NARRATION_CALLS[key] = recent
            return False
        recent.append(now)
        _NARRATION_CALLS[key] = recent
        return True


# Air-gap mode: hide cloud-dependent pages (Pulse, ClawHub, Genesis, GovCon, etc.)
_AIRGAP_MODE = os.environ.get("ICDEV_AIRGAP", "").lower() in ("true", "1", "yes")
# Demo mode: read-only enforcement (POST/PUT/DELETE to /api/* blocked except onboarding + IQE)
_DEMO_MODE = os.environ.get("ICDEV_DEMO_MODE", "").lower() in ("true", "1", "yes")
# Pages disabled in air-gap mode (routes → friendly message instead of 404)
_AIRGAP_DISABLED_ROUTES = frozenset(
    {
        "/pulse",
        "/clawhub",
        "/research",
        "/autoresearch",
        "/genesis",
        "/govcon",
        "/proposals",
        "/cpmp",
        "/proposal-genesis",
        "/leads",
        "/studio/marketplace",
        "/fathomdesk",
    }
)
# Legacy canvas feature flags (derived from the component registry — see _CANVAS_FLAGS below)
_HAS_STRATEGOS = os.environ.get("ICDEV_STRATEGOS_ENABLED", "true").lower() in ("true", "1", "yes")
_HAS_NETWORK = os.environ.get("ICDEV_NETWORK_ENABLED", "false").lower() == "true"
_HAS_PIPELINE = os.environ.get("ICDEV_PIPELINE_ENABLED", "false").lower() == "true"
_HAS_SECURITY_CANVAS = os.environ.get("ICDEV_SECURITY_ENABLED", "false").lower() in ("true", "1", "yes")
_HAS_INFRA_CANVAS = os.environ.get("ICDEV_INFRA_ENABLED", "false").lower() in ("true", "1", "yes")
_HAS_DATA_CANVAS = os.environ.get("ICDEV_DATA_CANVAS_ENABLED", "false").lower() in ("true", "1", "yes")
_HAS_BOUNDARY_CANVAS = os.environ.get("ICDEV_BOUNDARY_ENABLED", "false").lower() in ("true", "1", "yes")
_HAS_OBSERVABILITY_CANVAS = os.environ.get("ICDEV_OBSERVABILITY_ENABLED", "false").lower() in ("true", "1", "yes")
# Canvas Knowledge Graph: feature-flagged
_CANVAS_KG_ENABLED = os.environ.get("ICDEV_CANVAS_KG_ENABLED", "false").lower() in ("true", "1", "yes")
_HAS_CANVAS_KG = False
if _CANVAS_KG_ENABLED:
    try:
        from tools.canvas.blueprint import create_canvas_kg_blueprint  # noqa: E402

        _HAS_CANVAS_KG = True
    except ImportError:
        _HAS_CANVAS_KG = False

# D-CHILD-6: GovProposal/CPMP/GovCon conditionally loaded.
# Opt-in: default is OFF. Operators set ICDEV_GOVCON_ENABLED=true to enable.
# Air-gap installs (ICDEV_AIRGAP=true) force this off regardless so the
# GovCon Python modules are never imported, not just route-blocked.
_GOVCON_ENABLED = (
    os.environ.get("ICDEV_GOVCON_ENABLED", "false").lower() in ("true", "1", "yes")
    and not _AIRGAP_MODE
)
# Feature flags for page-route registration (blueprints registered via register_api_blueprints)
_HAS_GOVCON = False
if _GOVCON_ENABLED:
    import importlib.util as _ilu  # noqa: E402
    _HAS_GOVCON = all(
        _ilu.find_spec(m) is not None
        for m in (
            "tools.dashboard.api.proposals",
            "tools.dashboard.api.govcon",
            "tools.dashboard.api.cpmp",
        )
    )
    _HAS_PROPOSAL_GENESIS = _ilu.find_spec("tools.dashboard.api.proposal_genesis") is not None
else:
    _HAS_PROPOSAL_GENESIS = False
# Feature flags for finetune and chat API (always available — routes exist unconditionally)
_HAS_FINETUNE_API = True
_HAS_CHAT_API = True
from tools.dashboard.ux_helpers import register_ux_filters  # noqa: E402

# ── Component Registry (single source of truth) ─────────────────────────────
# Replaces legacy _CANVAS_DEFS / _APP_DEFS hardcoded lists with data from
# args/component_registry.yaml. See tools/config/component_registry.py.
from tools.config.component_registry import get_registry  # noqa: E402

_REGISTRY = get_registry()

# cvx-nav-01: build the IQE canvas dispatch map and the client-side path→canvas
# map ONCE at import (the registry is load-once — no hot-reload) instead of
# rebuilding per request in iqe_dispatch(). Both are cached at module scope.
_IQE_CANVAS_MAP: dict[str, tuple[str, list[str]]] = _REGISTRY.get_iqe_mapping()
# Supplement with canvases registered in app.py (not in component_registry.yaml)
# and the ai-brief-only alias cpmp_deliverables (served by the cpmp adapter,
# whose collections include cpmp.deliverables).
_IQE_CANVAS_MAP.setdefault("updates", ("tools.iqe.adapters.updates", ["updates.releases"]))
_IQE_CANVAS_MAP.setdefault("logs", ("tools.iqe.adapters.logs", ["logs.entries", "logs.errors"]))
if "cpmp" in _IQE_CANVAS_MAP:
    _IQE_CANVAS_MAP.setdefault("cpmp_deliverables", _IQE_CANVAS_MAP["cpmp"])

# Ordered [regex_source, canvas_key] list, JSON-serialized once for injection
# into base.html as window.__ICDEV_PATH_CANVAS__ (consumed by the IQE mini-bar
# and the AI-brief banner — single source of truth).
_IQE_PATH_CANVAS_JSON: str = json.dumps(_REGISTRY.get_iqe_path_canvas())


def _get_active_tier_safe() -> str:
    """Return the active license tier; falls back to 'community' if unavailable."""
    try:
        from tools.billing.tier import get_active_tier
        return get_active_tier()
    except Exception:
        return "community"


# URL prefixes for all canvases; used by base.html to highlight the Canvases menu.
_CANVAS_URL_PREFIXES = tuple(
    comp.url_prefix
    for comp in _REGISTRY.iter_canvases()
    if comp.url_prefix
)

# ── Design Canvases (conditional registration) ────────────────────────────
# cvx-net-02: eager canvas blueprint import is the dominant one-time startup
# cost (each factory imports its module and, for several canvases, seeds a DB).
# When ICDEV_LAZY_CANVASES is set, the heavy module import is DEFERRED until the
# first request to each canvas (see install_lazy_canvas_loader). Default OFF so
# existing behavior — and url_map-at-import expectations in the test suite —
# stay byte-for-byte unchanged unless an operator opts in.
_LAZY_CANVASES_ENABLED = os.environ.get(
    "ICDEV_LAZY_CANVASES", "false"
).strip().strip('"').strip("'").lower() in ("true", "1", "yes", "on")

_CANVAS_FLAGS: dict[str, bool] = {}
_CANVAS_BLUEPRINTS: dict[str, object] = {}

_CANVAS_URL_PREFIX_MAP = _REGISTRY.get_url_prefixes()
for _comp in _REGISTRY.iter_enabled(kind="canvas"):
    # Only canvases that declare a non-empty url_prefix can be matched by request
    # path, so only those are deferred. Empty-prefix canvases self-prefix inside
    # their blueprint (unknowable without importing) and are cheap — they stay
    # eager. The heavy canvas (network / ndc, '/network') declares a prefix.
    _canvas_pfx = (_CANVAS_URL_PREFIX_MAP.get(_comp.key) or "").strip()
    if _LAZY_CANVASES_ENABLED and _canvas_pfx:
        # Deferred: nav/context flag on, but do NOT import the module now.
        # install_lazy_canvas_loader() registers a first-hit loader in create_app.
        _CANVAS_FLAGS[_comp.key] = True
        continue
    try:
        _bp = _comp.get_blueprint()
        if _bp:
            _CANVAS_BLUEPRINTS[_comp.key] = _bp
            _CANVAS_FLAGS[_comp.key] = True
    except Exception as _exc:
        get_logger("icdev.dashboard").warning(
            "Canvas %s import failed (%s): %s", _comp.key.upper(), _comp.module, _exc
        )

# ── Application Modules (conditional registration) ─────────────────────────
_APP_FLAGS: dict[str, bool] = {}
_APP_BLUEPRINTS: dict[str, object] = {}

for _comp in _REGISTRY.iter_enabled(kind="child_app"):
    try:
        _bp = _comp.get_blueprint()
        if _bp:
            _APP_BLUEPRINTS[_comp.key] = _bp
            _APP_FLAGS[_comp.key] = True
    except Exception as _exc:
        get_logger("icdev.dashboard").warning(
            "App module %s import failed (%s): %s", _comp.key, _comp.module, _exc
        )

# ── Core Extensions (registry-driven, e.g. admin_console) ──────────────────
_CORE_EXT_BLUEPRINTS: dict[str, object] = {}

for _comp in _REGISTRY.iter_enabled(kind="core_extension"):
    if not getattr(_comp, "blueprint_attr", None):
        continue
    try:
        _bp = _comp.get_blueprint()
        if _bp:
            _CORE_EXT_BLUEPRINTS[_comp.key] = _bp
    except Exception as _exc:
        get_logger("icdev.dashboard").warning(
            "Core extension %s import failed (%s): %s", _comp.key, _comp.module, _exc
        )

# ---------------------------------------------------------------------------
# GovCon/CPMP/Proposals page registration (D-CHILD-6: isolated)
# ---------------------------------------------------------------------------


def _register_govcon_pages(app: "Flask", _get_db):
    """Register GovProposal/CPMP/GovCon SSR page routes on the Flask app.

    Called only when _HAS_GOVCON is True.  Extracted from create_app() so that
    child apps (and parent apps with ICDEV_GOVCON_ENABLED=false) never register
    these routes.
    """

    @app.route("/cpmp")
    @require_role("admin", "pm", "developer", "isso", "co", "contract_mgr")
    def cpmp_portfolio_page():
        """CPMP Portfolio — contract performance overview, health scoring."""
        try:
            from tools.govcon.portfolio_manager import get_portfolio_summary

            portfolio_data = get_portfolio_summary()
            pf = portfolio_data.get("portfolio", {})
            contracts = pf.get("contracts", [])
            upcoming = pf.get("upcoming_deliverables", [])
            portfolio = {
                "total_contracts": pf.get("total_contracts", 0),
                "active_contracts": pf.get("active_contracts", 0),
                "total_value": pf.get("total_value", 0),
                "burn_rate": pf.get("burn_rate_pct", 0),
                "overdue_deliverables": pf.get("overdue_deliverables", 0),
                "at_risk": pf.get("at_risk_contracts", 0),
                "health_distribution": pf.get("health_distribution", {"green": 0, "yellow": 0, "red": 0}),
            }
            try:
                from tools.govcon.risk_manager import get_portfolio_risk_summary

                risk_summary = get_portfolio_risk_summary().get("summary", {})
            except Exception:
                risk_summary = {}

            # Health matrix: per-contract dimension breakdown (fast DB-only calls)
            health_matrix = []
            try:
                from tools.govcon.portfolio_manager import compute_contract_health
                from tools.govcon.option_period_tracker import get_portfolio_countdown

                def _dim_tier(score):
                    if score is None:
                        return "unknown"
                    try:
                        s = float(score)
                    except (TypeError, ValueError):
                        return "unknown"
                    return "green" if s >= 0.75 else "yellow" if s >= 0.50 else "red"

                for c in contracts:
                    cid = c.get("id") or c.get("contract_id")
                    if not cid:
                        continue
                    try:
                        h = compute_contract_health(cid)
                        dims = h.get("dimension_scores") or h.get("dimensions") or {}
                        raw_score = h.get("health_score")
                        # Normalize: score may be 0-1 or 0-100
                        if raw_score is not None and raw_score <= 1.0:
                            display_score = round(raw_score * 100)
                        else:
                            display_score = round(raw_score) if raw_score is not None else None
                        health_matrix.append({
                            "contract_id": cid,
                            "contract_number": c.get("contract_number", "—"),
                            "title": c.get("title", ""),
                            "agency": c.get("agency", ""),
                            "overall": h.get("health", "unknown"),
                            "health_score": display_score,
                            "evm": _dim_tier(dims.get("evm")),
                            "deliverables": _dim_tier(dims.get("deliverables")),
                            "cpars": _dim_tier(dims.get("cpars")),
                            "funding": _dim_tier(dims.get("funding")),
                            "negative_events": _dim_tier(dims.get("negative_events")),
                        })
                    except Exception:
                        health_matrix.append({
                            "contract_id": cid,
                            "contract_number": c.get("contract_number", "—"),
                            "title": c.get("title", ""),
                            "agency": c.get("agency", ""),
                            "overall": "unknown",
                            "health_score": None,
                            "evm": "unknown", "deliverables": "unknown",
                            "cpars": "unknown", "funding": "unknown", "negative_events": "unknown",
                        })

                option_countdown = get_portfolio_countdown().get("options", [])
            except Exception:
                option_countdown = []

            return render_template(
                "cpmp/portfolio.html",
                portfolio=portfolio,
                contracts=contracts,
                upcoming_deliverables=upcoming,
                risk_summary=risk_summary,
                health_matrix=health_matrix,
                option_countdown=option_countdown,
            )
        except Exception as e:
            import traceback

            traceback.print_exc()
            return render_template(
                "cpmp/portfolio.html",
                portfolio={
                    "total_contracts": 0,
                    "active_contracts": 0,
                    "total_value": 0,
                    "burn_rate": 0,
                    "overdue_deliverables": 0,
                    "health_distribution": {"green": 0, "yellow": 0, "red": 0},
                },
                contracts=[],
                upcoming_deliverables=[],
                risk_summary={},
                health_matrix=[],
                option_countdown=[],
                error=str(e),
            )

    @app.route("/cpmp/<contract_id>")
    @require_role("admin", "pm", "contract_mgr", "co", "cor", "isso")
    def cpmp_detail_page(contract_id):
        """CPMP Contract Detail — 7-tab view."""
        try:
            from tools.govcon.contract_manager import get_contract, list_clins, list_wbs, list_deliverables

            contract_result = get_contract(contract_id)
            if contract_result.get("status") == "error":
                return render_template("404.html", message="Contract not found"), 404
            contract = contract_result.get("contract", contract_result)
            clins = list_clins(contract_id).get("clins", [])
            wbs_elements = list_wbs(contract_id).get("wbs_elements", [])
            deliverables = list_deliverables(contract_id).get("deliverables", [])
            try:
                from tools.govcon.subcontractor_tracker import list_subcontractors

                subcontractors = list_subcontractors(contract_id).get("subcontractors", [])
            except Exception:
                subcontractors = []
            try:
                from tools.govcon.evm_engine import aggregate_contract_evm

                evm = aggregate_contract_evm(contract_id)
                if "indicators" in evm and isinstance(evm["indicators"], dict):
                    evm.update(evm["indicators"])
            except Exception:
                evm = {}
            try:
                from tools.govcon.cpars_predictor import predict_cpars, list_assessments

                cpars_prediction = predict_cpars(contract_id)
                if "dimension_scores" in cpars_prediction:
                    cpars_prediction["dimensions"] = {
                        k: round(v * 5, 2) for k, v in cpars_prediction["dimension_scores"].items()
                    }
                cpars_assessments = list_assessments(contract_id).get("assessments", [])
            except Exception:
                cpars_prediction = {}
                cpars_assessments = []
            try:
                from tools.govcon.milestone_manager import list_milestones, list_deps

                milestones = list_milestones(contract_id).get("milestones", [])
                milestone_deps = list_deps(contract_id).get("deps", [])
            except Exception:
                milestones = []
                milestone_deps = []
            try:
                from tools.govcon.risk_manager import list_risks, get_risk_matrix

                risks = list_risks(contract_id).get("risks", [])
                risk_matrix = get_risk_matrix(contract_id)
            except Exception:
                risks = []
                risk_matrix = {}
            try:
                from tools.govcon.contract_periods_manager import list_periods, get_obligation_summary

                periods = list_periods(contract_id).get("periods", [])
                obligation_summary = get_obligation_summary(contract_id)
            except Exception:
                periods = []
                obligation_summary = {}
            return render_template(
                "cpmp/detail.html",
                contract=contract,
                clins=clins,
                wbs_elements=wbs_elements,
                deliverables=deliverables,
                subcontractors=subcontractors,
                evm=evm,
                cpars_prediction=cpars_prediction,
                cpars_assessments=cpars_assessments,
                milestones=milestones,
                milestone_deps=milestone_deps,
                risks=risks,
                risk_matrix=risk_matrix,
                periods=periods,
                obligation_summary=obligation_summary,
            )
        except Exception as e:
            import traceback

            traceback.print_exc()
            return render_template(
                "cpmp/detail.html",
                contract={},
                clins=[],
                wbs_elements=[],
                deliverables=[],
                subcontractors=[],
                evm={},
                cpars_prediction={},
                cpars_assessments=[],
                milestones=[],
                milestone_deps=[],
                risks=[],
                risk_matrix={},
                periods=[],
                obligation_summary={},
                error=str(e),
            ), 200

    @app.route("/cpmp/<contract_id>/deliverables/<deliverable_id>")
    @require_role("admin", "pm", "contract_mgr", "co", "cor", "isso")
    def cpmp_deliverable_detail_page(contract_id, deliverable_id):
        """CPMP Deliverable Detail — status pipeline, CDRL generation."""
        try:
            from tools.govcon.contract_manager import get_contract, get_deliverable

            contract_result = get_contract(contract_id)
            contract = contract_result.get("contract", contract_result) if contract_result.get("status") == "ok" else {}
            deliv_result = get_deliverable(deliverable_id)
            if deliv_result.get("status") == "error":
                return render_template("404.html", message="Deliverable not found"), 404
            deliverable = deliv_result.get("deliverable", deliv_result)
            generations = deliverable.get("generations", []) if isinstance(deliverable, dict) else []
            status_history = deliverable.get("status_history", []) if isinstance(deliverable, dict) else []
            return render_template(
                "cpmp/deliverable_detail.html",
                contract=contract,
                deliverable=deliverable,
                generations=generations,
                status_history=status_history,
            )
        except Exception as e:
            import traceback

            traceback.print_exc()
            return render_template(
                "cpmp/deliverable_detail.html",
                contract={},
                deliverable={},
                generations=[],
                status_history=[],
                error=str(e),
            ), 200

    @app.route("/cpmp/cor")
    @require_role("admin", "pm", "isso", "co", "cor", "contract_mgr")
    def cpmp_cor_portal_page():
        """COR Portal — read-only government view of assigned contracts."""
        user = getattr(g, "current_user", None)
        cor_email = user.get("email", "") if user else ""
        conn = _get_db()
        try:
            if cor_email:
                rows = conn.execute(
                    "SELECT * FROM cpmp_contracts WHERE cor_email = %s ORDER BY created_at DESC",
                    (cor_email,),
                ).fetchall()
                contracts = [dict(r) for r in rows]
            else:
                contracts = []
            return render_template("cpmp/cor_portal.html", contracts=contracts, cor_email=cor_email)
        except Exception:
            import traceback

            traceback.print_exc()
            return render_template("cpmp/cor_portal.html", contracts=[], cor_email=cor_email)
        finally:
            conn.close()

    @app.route("/cpmp/cor/<contract_id>")
    @require_role("admin", "pm", "isso", "co", "cor", "contract_mgr")
    def cpmp_cor_detail_page(contract_id):
        """COR Contract Detail — read-only, no internal cost data."""
        user = getattr(g, "current_user", None)
        cor_email = user.get("email", "") if user else ""
        conn = _get_db()
        try:
            from tools.govcon.contract_manager import get_contract, list_deliverables

            contract_result = get_contract(contract_id)
            if contract_result.get("status") == "error":
                return render_template("404.html", message="Contract not found"), 404
            contract = contract_result.get("contract", contract_result)
            deliverables = list_deliverables(contract_id).get("deliverables", [])
            try:
                from tools.govcon.evm_engine import aggregate_contract_evm

                evm = aggregate_contract_evm(contract_id)
                if "indicators" in evm and isinstance(evm["indicators"], dict):
                    evm.update(evm["indicators"])
                if "total_bac" in evm:
                    evm.setdefault("bac", evm["total_bac"])
                if "total_pv" in evm:
                    evm.setdefault("pv", evm["total_pv"])
                if "total_ev" in evm:
                    evm.setdefault("ev", evm["total_ev"])
                if "percent_complete" in evm:
                    evm.setdefault(
                        "percent_complete_schedule",
                        evm["percent_complete"] / 100 if evm["percent_complete"] > 1 else evm["percent_complete"],
                    )
            except Exception:
                evm = {}
            try:
                from tools.govcon.cpars_predictor import list_assessments

                cpars_assessments = list_assessments(contract_id).get("assessments", [])
            except Exception:
                cpars_assessments = []
            try:
                conn.execute(
                    "INSERT INTO cpmp_cor_access_log (user_id, contract_id, action) "
                    "VALUES (%s, %s, %s)",
                    (
                        cor_email,
                        contract_id,
                        "view_contract",
                    ),
                )
                conn.commit()
            except Exception as exc:  # noqa: BLE001 - best-effort persistence; logged, never raised
                logger.warning(
                    "cpmp_cor_detail_page: best-effort INSERT into cpmp_cor_access_log failed (non-blocking): %s",
                    exc,
                )
            return render_template(
                "cpmp/cor_detail.html",
                contract=contract,
                deliverables=deliverables,
                evm=evm,
                cpars_assessments=cpars_assessments,
                cor_email=cor_email,
            )
        except Exception as e:
            return render_template("404.html", message=f"Error: {e}"), 500
        finally:
            conn.close()

    @app.route("/cpmp/deliverables")
    @require_role("admin", "pm", "developer", "isso", "co", "contract_mgr")
    def cpmp_deliverable_center_page():
        """Deliverable Command Center — all deliverables across all active contracts."""
        return render_template("cpmp/deliverable_center.html")

    @app.route("/api/ai-brief/<canvas>")
    def api_ai_brief(canvas):
        """Return rendered AI brief banner HTML for a given canvas key."""
        from flask import jsonify
        try:
            from tools.dashboard.components.ai_brief_banner import render_ai_brief
            html_content = render_ai_brief(canvas, {})
        except Exception as exc:
            html_content = (
                f'<aside class="ai-brief-banner card border-0 shadow-sm mb-3">'
                f'<div class="card-body py-2 px-3 text-muted small">'
                f'AI brief unavailable: {exc}'
                f'</div></aside>'
            )
        return jsonify({"html": html_content, "canvas": canvas})

    @app.route("/cpmp/reports")
    @require_role("admin", "pm", "capture_mgr", "bd")
    def cpmp_reports_page():
        """CPMP Reports — exportable contract performance reports and analytics."""
        try:
            from tools.govcon.portfolio_manager import get_portfolio_summary

            portfolio_data = get_portfolio_summary()
            pf = portfolio_data.get("portfolio", {})
            contracts = pf.get("contracts", [])
            upcoming = pf.get("upcoming_deliverables", [])
            portfolio = {
                "total_contracts": pf.get("total_contracts", 0),
                "active_contracts": pf.get("active_contracts", 0),
                "total_value": pf.get("total_value", 0),
                "burn_rate": pf.get("burn_rate_pct", 0),
                "overdue_deliverables": pf.get("overdue_deliverables", 0),
                "at_risk": pf.get("at_risk_contracts", 0),
                "health_distribution": pf.get("health_distribution", {"green": 0, "yellow": 0, "red": 0}),
            }
            try:
                from tools.govcon.risk_manager import get_portfolio_risk_summary

                risk_summary = get_portfolio_risk_summary().get("summary", {})
            except Exception:
                risk_summary = {}
            return render_template(
                "cpmp/reports.html",
                portfolio=portfolio,
                contracts=contracts,
                upcoming_deliverables=upcoming,
                risk_summary=risk_summary,
            )
        except Exception as e:
            import traceback

            traceback.print_exc()
            return render_template(
                "cpmp/reports.html",
                portfolio={
                    "total_contracts": 0,
                    "active_contracts": 0,
                    "total_value": 0,
                    "burn_rate": 0,
                    "overdue_deliverables": 0,
                    "health_distribution": {"green": 0, "yellow": 0, "red": 0},
                },
                contracts=[],
                upcoming_deliverables=[],
                risk_summary={},
                error=str(e),
            )

    @app.route("/proposals")
    @require_role("admin", "bd", "capture_mgr", "pm", "reviewer")
    def proposals_list_page():
        """Proposal Opportunities — GovCon proposal writing lifecycle tracker."""
        conn = _get_db()
        try:
            rows = conn.execute("SELECT * FROM proposal_opportunities ORDER BY due_date ASC").fetchall()
            opportunities = [dict(r) for r in rows]
            from datetime import date

            # Attach section count per opportunity so we can bubble up proposals with content
            try:
                sec_counts = conn.execute(
                    "SELECT opportunity_id, COUNT(*) as cnt FROM proposal_sections GROUP BY opportunity_id"
                ).fetchall()
                sec_map = {r["opportunity_id"]: r["cnt"] for r in sec_counts}
            except Exception:
                sec_map = {}
            for opp in opportunities:
                opp["section_count"] = sec_map.get(opp["id"], 0)

            # Sort: proposals with content first (by section_count desc), then by due_date asc
            opportunities.sort(key=lambda o: (-o["section_count"], o.get("due_date") or "9999-99-99"))

            today = date.today()
            nearest_deadline = None
            for opp in opportunities:
                if opp.get("due_date") and opp["status"] not in ("submitted", "won", "lost", "cancelled", "no_bid"):
                    try:
                        dd = date.fromisoformat(opp["due_date"])
                        days_left = (dd - today).days
                        opp["days_left"] = days_left
                        if nearest_deadline is None or days_left < nearest_deadline:
                            nearest_deadline = days_left
                    except (ValueError, TypeError):
                        opp["days_left"] = None
                else:
                    opp["days_left"] = None

            # Attach most-recent pWin assessment per opportunity
            try:
                from tools.govcon.bayesian_bid_scorer import pipeline_value_rollup
                rollup = pipeline_value_rollup()
                pwin_map = {item["opportunity_id"]: item for item in rollup.get("opportunities", [])}
                for opp in opportunities:
                    item = pwin_map.get(opp["id"])
                    if item:
                        opp["computed_pwin_pct"] = item["pwin_pct"]
                        opp["weighted_value"] = item["weighted_value"]
                        opp["has_pwin_model"] = item["has_pwin_model"]
                        opp["pwin_factors"] = item.get("factor_breakdown")
                    else:
                        opp["computed_pwin_pct"] = None
                        opp["weighted_value"] = None
                        opp["has_pwin_model"] = False
                        opp["pwin_factors"] = None
            except Exception:
                rollup = {"total_weighted_pipeline_value": 0, "total_potential_value": 0, "scored_count": 0, "unscored_count": 0}
                for opp in opportunities:
                    opp["computed_pwin_pct"] = None
                    opp["weighted_value"] = None
                    opp["has_pwin_model"] = False
                    opp["pwin_factors"] = None

            return render_template(
                "proposals/list.html",
                opportunities=opportunities,
                nearest_deadline=nearest_deadline,
                pipeline_rollup=rollup,
            )
        finally:
            conn.close()

    @app.route("/proposals/<opp_id>")
    @require_role("admin", "bd", "capture_mgr", "pm", "reviewer")
    def proposals_detail_page(opp_id):
        """Proposal Opportunity Detail — 6-tab view with sections, compliance, reviews."""
        conn = _get_db()
        try:
            opp = conn.execute("SELECT * FROM proposal_opportunities WHERE id = %s", (opp_id,)).fetchone()
            if not opp:
                return render_template("404.html", message="Opportunity not found"), 404
            opp = dict(opp)
            sections = [
                dict(r)
                for r in conn.execute(
                    """SELECT s.*, v.volume_number, v.title as volume_title,
                          latest_draft.status as draft_status,
                          latest_draft.id as draft_id
                   FROM proposal_sections s
                   LEFT JOIN proposal_volumes v ON s.volume_id = v.id
                   LEFT JOIN (
                       SELECT section_id, id, status,
                              ROW_NUMBER() OVER (PARTITION BY section_id ORDER BY created_at DESC) as rn
                       FROM proposal_section_drafts
                       WHERE opportunity_id = %s
                   ) latest_draft ON s.id = latest_draft.section_id AND latest_draft.rn = 1
                   WHERE s.opportunity_id = %s
                   ORDER BY v.volume_number, s.section_number""",
                    (opp_id, opp_id),
                ).fetchall()
            ]
            from datetime import date

            today = date.today()
            for s in sections:
                s["overdue"] = False
                if s.get("due_date") and s["status"] not in ("final", "submitted"):
                    try:
                        s["overdue"] = date.fromisoformat(s["due_date"]) < today
                    except (ValueError, TypeError):
                        pass
            volumes = [
                dict(r)
                for r in conn.execute(
                    "SELECT * FROM proposal_volumes WHERE opportunity_id = %s ORDER BY volume_number", (opp_id,)
                ).fetchall()
            ]
            compliance_items = [
                dict(r)
                for r in conn.execute(
                    "SELECT * FROM proposal_compliance_matrix WHERE opportunity_id = %s", (opp_id,)
                ).fetchall()
            ]
            reviews = [
                dict(r)
                for r in conn.execute(
                    "SELECT * FROM proposal_reviews WHERE opportunity_id = %s ORDER BY scheduled_date", (opp_id,)
                ).fetchall()
            ]
            findings = [
                dict(r)
                for r in conn.execute(
                    """SELECT f.*, r.review_type FROM proposal_review_findings f
                   JOIN proposal_reviews r ON f.review_id = r.id
                   WHERE r.opportunity_id = %s""",
                    (opp_id,),
                ).fetchall()
            ]
            total_sections = len(sections)
            completed_sections = len([s for s in sections if s["status"] in ("final", "submitted")])
            total_compliance = len(compliance_items)
            compliant_count = len([c for c in compliance_items if c.get("compliance_status") == "compliant"])
            coverage_pct = (compliant_count / total_compliance * 100) if total_compliance > 0 else 0
            open_findings = len([f for f in findings if f.get("status") in ("open", "in_progress")])
            critical_findings = len(
                [f for f in findings if f.get("severity") == "critical" and f.get("status") in ("open", "in_progress")]
            )
            section_status_dist = {}
            for s in sections:
                st = s.get("status", "not_started")
                section_status_dist[st] = section_status_dist.get(st, 0) + 1
            finding_severity_dist = {}
            for f in findings:
                if f.get("status") in ("open", "in_progress"):
                    sev = f.get("severity", "minor")
                    finding_severity_dist[sev] = finding_severity_dist.get(sev, 0) + 1
            cm_compliant = len([c for c in compliance_items if c.get("compliance_status") == "compliant"])
            cm_partial = len([c for c in compliance_items if c.get("compliance_status") == "partial"])
            cm_non_compliant = len([c for c in compliance_items if c.get("compliance_status") == "non_compliant"])
            cm_not_addressed = len([c for c in compliance_items if c.get("compliance_status") == "not_addressed"])
            cm_not_applicable = len([c for c in compliance_items if c.get("compliance_status") == "not_applicable"])
            cm_gap_pct = round(cm_not_addressed / total_compliance * 100) if total_compliance > 0 else 0
            compliance_stats = {
                "total": total_compliance,
                "compliant": cm_compliant,
                "partial": cm_partial,
                "non_compliant": cm_non_compliant,
                "not_addressed": cm_not_addressed,
                "not_applicable": cm_not_applicable,
                "gap_pct": cm_gap_pct,
            }
            findings_by_review = {}
            for f in findings:
                rid = f.get("review_id")
                if rid:
                    findings_by_review.setdefault(rid, []).append(f)
            # Load reviewer assignments per review (prop-rev-09)
            assignments_by_review = {}
            try:
                all_asgns = conn.execute(
                    """SELECT a.* FROM proposal_reviewer_assignments a
                       JOIN proposal_reviews r ON a.review_id = r.id
                       WHERE r.opportunity_id = %s ORDER BY a.created_at DESC""",
                    (opp_id,),
                ).fetchall()
                for a in all_asgns:
                    assignments_by_review.setdefault(a["review_id"], []).append(dict(a))
            except Exception:
                conn.rollback()
            reviews_data = []
            for rev in reviews:
                rd = dict(rev)
                rd["findings"] = findings_by_review.get(rev["id"], [])
                rd["assignments"] = assignments_by_review.get(rev["id"], [])
                reviews_data.append(rd)
            days_left = None
            if opp.get("due_date"):
                try:
                    days_left = (date.fromisoformat(opp["due_date"]) - today).days
                except (ValueError, TypeError):
                    pass
            stats = {
                "sections_total": total_sections,
                "sections_complete": completed_sections,
                "compliance_coverage_pct": round(coverage_pct),
                "open_findings": open_findings,
                "critical_findings": critical_findings,
                "section_status_distribution": section_status_dist,
                "finding_severity_distribution": finding_severity_dist,
            }
            questions = [
                dict(r)
                for r in conn.execute(
                    "SELECT * FROM proposal_questions WHERE opportunity_id = %s ORDER BY question_number ASC",
                    (opp_id,),
                ).fetchall()
            ]
            question_stats = {
                "total": len(questions),
                "high_priority": len([q for q in questions if q.get("priority") == "high"]),
                "draft": len([q for q in questions if q.get("status") == "draft"]),
                "approved": len([q for q in questions if q.get("status") == "approved"]),
                "submitted": len([q for q in questions if q.get("status") == "submitted"]),
                "answered": len([q for q in questions if q.get("status") == "answered"]),
            }
            questions_days_left = None
            if opp.get("questions_due_date"):
                try:
                    questions_days_left = (date.fromisoformat(opp["questions_due_date"]) - today).days
                except (ValueError, TypeError):
                    pass
            amendments = [
                dict(r)
                for r in conn.execute(
                    "SELECT * FROM proposal_amendments WHERE opportunity_id = %s ORDER BY version_number ASC",
                    (opp_id,),
                ).fetchall()
            ]
            responses = {}
            for q in questions:
                if q.get("status") == "answered":
                    resp = conn.execute(
                        "SELECT * FROM proposal_question_responses WHERE question_id = %s ORDER BY created_at DESC LIMIT 1",
                        (q["id"],),
                    ).fetchone()
                    if resp:
                        responses[q["id"]] = dict(resp)
            return render_template(
                "proposals/detail.html",
                opp=opp,
                sections=sections,
                volumes=volumes,
                compliance_items=compliance_items,
                reviews=reviews_data,
                findings=findings,
                stats=stats,
                compliance_stats=compliance_stats,
                reviews_data=reviews_data,
                days_left=days_left,
                questions=questions,
                question_stats=question_stats,
                questions_days_left=questions_days_left,
                amendments=amendments,
                responses=responses,
            )
        finally:
            conn.close()

    @app.route("/proposals/<opp_id>/sections/<sec_id>")
    @require_role("admin", "bd", "capture_mgr", "pm", "reviewer")
    def proposals_section_detail_page(opp_id, sec_id):
        """Proposal Section Detail — status pipeline, notes, compliance, findings, history."""
        conn = _get_db()
        try:
            section = conn.execute(
                """SELECT s.*, v.volume_number, v.title as volume_title
                   FROM proposal_sections s
                   LEFT JOIN proposal_volumes v ON s.volume_id = v.id
                   WHERE s.id = %s AND s.opportunity_id = %s""",
                (sec_id, opp_id),
            ).fetchone()
            if not section:
                return render_template("404.html", message="Section not found"), 404
            section = dict(section)
            opp = conn.execute("SELECT title FROM proposal_opportunities WHERE id = %s", (opp_id,)).fetchone()
            opp_title = opp["title"] if opp else "Unknown"
            from tools.dashboard.api.proposals import SECTION_TRANSITIONS

            section["valid_transitions"] = SECTION_TRANSITIONS.get(section["status"], [])
            from datetime import date

            section["overdue"] = False
            if section.get("due_date") and section["status"] not in ("final", "submitted"):
                try:
                    section["overdue"] = date.fromisoformat(section["due_date"]) < date.today()
                except (ValueError, TypeError):
                    pass
            # Latest draft for this section
            draft_row = conn.execute(
                """SELECT * FROM proposal_section_drafts
                   WHERE section_id = %s
                   ORDER BY created_at DESC LIMIT 1""", (sec_id,)
            ).fetchone()
            draft = None
            if draft_row:
                draft = dict(draft_row)
                import json as _json

                meta_str = draft.get("metadata", "{}")
                try:
                    draft["metadata_parsed"] = _json.loads(meta_str) if isinstance(meta_str, str) else meta_str
                except Exception:
                    draft["metadata_parsed"] = {}
                notes_str = draft.get("review_notes", "{}")
                try:
                    draft["review_notes_parsed"] = _json.loads(notes_str) if isinstance(notes_str, str) else notes_str
                except Exception:
                    draft["review_notes_parsed"] = {}
                # WriteGuard may be in review_notes (new pipeline) or metadata (legacy)
                if draft["review_notes_parsed"] and draft["review_notes_parsed"].get("writeguard"):
                    draft["writeguard"] = draft["review_notes_parsed"]["writeguard"]
                elif draft["metadata_parsed"] and draft["metadata_parsed"].get("writeguard"):
                    draft["writeguard"] = draft["metadata_parsed"]["writeguard"]
                else:
                    draft["writeguard"] = None

            section["compliance_items"] = [
                dict(r)
                for r in conn.execute(
                    "SELECT * FROM proposal_compliance_matrix WHERE proposal_section_id = %s", (sec_id,)
                ).fetchall()
            ]
            section["findings"] = [
                dict(r)
                for r in conn.execute(
                    """SELECT f.*, r.review_type FROM proposal_review_findings f
                   JOIN proposal_reviews r ON f.review_id = r.id
                   WHERE f.section_id = %s""",
                    (sec_id,),
                ).fetchall()
            ]
            deps = conn.execute(
                """SELECT d.*, s.title as depends_on_title, s.status as depends_on_status
                   FROM proposal_section_dependencies d
                   JOIN proposal_sections s ON d.depends_on_section_id = s.id
                   WHERE d.section_id = %s""",
                (sec_id,),
            ).fetchall()
            dep_list = []
            for d in deps:
                d = dict(d)
                from tools.dashboard.api.proposals import SECTION_STATUS_ORDER

                req_idx = (
                    SECTION_STATUS_ORDER.index(d["required_status"])
                    if d["required_status"] in SECTION_STATUS_ORDER
                    else 0
                )
                cur_idx = (
                    SECTION_STATUS_ORDER.index(d["depends_on_status"])
                    if d["depends_on_status"] in SECTION_STATUS_ORDER
                    else 0
                )
                d["met"] = cur_idx >= req_idx
                dep_list.append(d)
            section["dependencies"] = dep_list
            section["history"] = [
                dict(r)
                for r in conn.execute(
                    "SELECT * FROM proposal_status_history WHERE entity_id = %s ORDER BY created_at DESC", (sec_id,)
                ).fetchall()
            ]
            return render_template("proposals/section_detail.html", section=section, opp_title=opp_title, draft=draft)
        finally:
            conn.close()

    @app.route("/proposals/reviews-dashboard")
    @require_role("admin", "pm", "reviewer")
    def proposals_reviews_dashboard():
        """Executive cross-proposal review dashboard (prop-rev-07)."""
        from datetime import date

        conn = _get_db()
        try:
            raw_opps = conn.execute(
                "SELECT * FROM proposal_opportunities WHERE status NOT IN ('won','lost','no_bid','cancelled') ORDER BY due_date ASC"
            ).fetchall()
            opps = []
            unresolved_critical = 0
            passed = 0
            pass_with_findings = 0
            failed = 0
            today = date.today()
            for row in raw_opps:
                opp = dict(row)
                reviews = conn.execute(
                    "SELECT * FROM proposal_reviews WHERE opportunity_id = %s", (opp["id"],)
                ).fetchall()
                rev_map = {r["review_type"]: dict(r) for r in reviews}
                opp["review_map"] = rev_map
                # Count outcomes
                for r in rev_map.values():
                    if r.get("overall_rating") == "pass":
                        passed += 1
                    elif r.get("overall_rating") == "pass_with_findings":
                        pass_with_findings += 1
                    elif r.get("overall_rating") == "fail":
                        failed += 1
                # Critical open findings
                crit = conn.execute(
                    """SELECT COUNT(*) as cnt FROM proposal_review_findings f
                       JOIN proposal_reviews r ON f.review_id = r.id
                       WHERE r.opportunity_id = %s AND f.severity = 'critical' AND f.status IN ('open','in_progress')""",
                    (opp["id"],),
                ).fetchone()
                opp["critical_open"] = crit["cnt"] if crit else 0
                unresolved_critical += opp["critical_open"]
                try:
                    opp["days_left"] = (date.fromisoformat(opp["due_date"]) - today).days
                except Exception:
                    opp["days_left"] = None
                opps.append(opp)
            # Count pending reviewer assignments (prop-rev-09)
            pending_assignments = 0
            try:
                row = conn.execute(
                    "SELECT COUNT(*) as cnt FROM proposal_reviewer_assignments WHERE status = 'pending'"
                ).fetchone()
                pending_assignments = row["cnt"] if row else 0
            except Exception:
                pass
            return render_template(
                "proposals/reviews_dashboard.html",
                opportunities=opps,
                unresolved_critical=unresolved_critical,
                passed=passed,
                pass_with_findings=pass_with_findings,
                failed=failed,
                pending_assignments=pending_assignments,
            )
        finally:
            conn.close()

    @app.route("/proposals/<opp_id>/language")
    @require_role("admin", "bd", "capture_mgr", "pm", "reviewer")
    def proposals_language_page(opp_id):
        """Proposal Language Settings — glossary, wall of truth, taxonomy, style templates."""
        conn = _get_db()
        try:
            opp = conn.execute("SELECT * FROM proposal_opportunities WHERE id = %s", (opp_id,)).fetchone()
            if not opp:
                return render_template("404.html", message="Opportunity not found"), 404
            opp = dict(opp)
            glossary = [
                dict(r) for r in conn.execute(
                    "SELECT * FROM wg_glossary WHERE scope = 'project' AND scope_id = %s AND is_active = 1 ORDER BY term_type, term",
                    (opp_id,),
                ).fetchall()
            ]
            taxonomy = [
                dict(r) for r in conn.execute(
                    "SELECT * FROM proposal_taxonomy WHERE opportunity_id = %s AND is_active = 1 ORDER BY label",
                    (opp_id,),
                ).fetchall()
            ]
            style_guides = [
                dict(r) for r in conn.execute(
                    "SELECT id, guide_name, version, created_at FROM wg_style_guides WHERE scope = 'project' AND scope_id = %s AND is_active = 1 ORDER BY guide_name",
                    (opp_id,),
                ).fetchall()
            ]
            return render_template(
                "proposals/language.html",
                opp=opp,
                opp_id=opp_id,
                glossary=glossary,
                taxonomy=taxonomy,
                style_guides=style_guides,
            )
        finally:
            conn.close()

    @app.route("/proposals/<opp_id>/ptw")
    @require_role("admin", "capture_mgr", "pm")
    def proposals_ptw_page(opp_id):
        """Black-hat / PTW workspace — competitor intelligence + price-to-win (prop-cap-13)."""
        conn = _get_db()
        try:
            opp = conn.execute("SELECT * FROM proposal_opportunities WHERE id = %s", (opp_id,)).fetchone()
            if not opp:
                return render_template("404.html", message="Opportunity not found"), 404
            opp = dict(opp)
            # Load saved black-hat assessments
            try:
                conn.execute("""
                    CREATE TABLE IF NOT EXISTS proposal_blackhat_assessments (
                        id TEXT PRIMARY KEY, opportunity_id TEXT NOT NULL,
                        competitor_name TEXT NOT NULL, approach_hypothesis TEXT,
                        price_estimate_low REAL, price_estimate_high REAL,
                        strengths TEXT, weaknesses TEXT, win_strategy TEXT,
                        differentiators TEXT, risk_factors TEXT,
                        ptw_posture TEXT DEFAULT 'competitive',
                        leaderboard_rank INTEGER, award_count INTEGER,
                        total_award_value REAL, naics_diversity INTEGER,
                        agency_diversity INTEGER, classification TEXT DEFAULT 'CUI',
                        created_at TEXT NOT NULL, updated_at TEXT, created_by TEXT
                    )
                """)
                conn.commit()
                bh_rows = conn.execute(
                    "SELECT * FROM proposal_blackhat_assessments WHERE opportunity_id = %s ORDER BY created_at DESC",
                    (opp_id,),
                ).fetchall()
                assessments = [dict(r) for r in bh_rows]
            except Exception:
                assessments = []
        finally:
            conn.close()
        # PTW analysis (rate_benchmarker)
        ptw = {}
        try:
            from tools.govcon.rate_benchmarker import ptw_analysis
            from tools.security.column_security import current_role, mask_ptw_payload
            # prop-sec-01: reviewer/co are denied the stored ptw_estimate_* columns,
            # so they must not receive the same figures recomputed here either.
            ptw = mask_ptw_payload(ptw_analysis(opp_id), current_role())
        except Exception:
            ptw = {}
        # Bayesian bid score with default mid-scores
        bid_score = None
        try:
            from tools.govcon.bayesian_bid_scorer import score_opportunity, DIMENSIONS
            dims = {d: 0.5 for d in DIMENSIONS}
            bid_score = score_opportunity(opp_id, dims)
        except Exception:
            bid_score = None
        return render_template(
            "proposals/ptw.html",
            opp=opp,
            assessments=assessments,
            ptw=ptw,
            bid_score=bid_score,
        )

    @app.route("/proposals/<opp_id>/compliance/gaps")
    @require_role("admin", "bd", "capture_mgr", "pm", "reviewer")
    def proposals_compliance_gaps(opp_id):
        """Compliance gap drill-down — all not_addressed requirements (prop-cmp-10)."""
        conn = _get_db()
        try:
            opp = conn.execute("SELECT * FROM proposal_opportunities WHERE id = %s", (opp_id,)).fetchone()
            if not opp:
                return render_template("404.html", message="Opportunity not found"), 404
            opp = dict(opp)
            orphaned = conn.execute(
                """SELECT * FROM proposal_compliance_matrix
                   WHERE opportunity_id = %s AND (proposal_section_id IS NULL OR proposal_section_id = '')
                     AND compliance_status != 'not_applicable'
                   ORDER BY sort_order""",
                (opp_id,),
            ).fetchall()
            orphaned = [dict(r) for r in orphaned]
            compliance_total = conn.execute(
                "SELECT COUNT(*) as cnt FROM proposal_compliance_matrix WHERE opportunity_id = %s",
                (opp_id,),
            ).fetchone()["cnt"]
            sections = conn.execute(
                """SELECT s.id, s.section_number, s.title FROM proposal_sections s
                   WHERE s.opportunity_id = %s ORDER BY s.section_number""",
                (opp_id,),
            ).fetchall()
            sections = [dict(s) for s in sections]
            gap_pct = round(len(orphaned) / compliance_total * 100, 1) if compliance_total > 0 else 0
            return render_template(
                "proposals/compliance_gaps.html",
                opp=opp,
                orphaned=orphaned,
                compliance_total=compliance_total,
                sections=sections,
                gap_pct=gap_pct,
            )
        finally:
            conn.close()

    # GovCon RBAC (prop-fix-09, roles per prop-fix-08): GovCon is the BD/capture
    # pre-award intelligence domain. View pages are read-only dashboards open to
    # the capture roles + management; deny -> 403 + audit via require_role().
    @app.route("/govcon")
    @require_role("admin", "bd", "capture_mgr", "pm")
    def govcon_pipeline_page():
        """GovCon Intelligence — pipeline status, recent opportunities, domain distribution."""
        conn = _get_db()
        try:
            from tools.govcon.govcon_engine import get_status
            from tools.govcon.bayesian_bid_scorer import pipeline_value_rollup
            from tools.govcon.crm_heat import get_engagement_heat_by_agency
            from tools.govcon.sam_scanner import list_forecast_notices

            stats = get_status()
            try:
                opps = conn.execute("SELECT * FROM sam_gov_opportunities ORDER BY posted_date DESC LIMIT 25").fetchall()
                opportunities = [dict(r) for r in opps]
            except Exception:
                opportunities = []
            linked_opp_ids = set()
            try:
                linked = conn.execute(
                    "SELECT sam_gov_opportunity_id FROM proposal_opportunities WHERE sam_gov_opportunity_id IS NOT NULL"
                ).fetchall()
                linked_opp_ids = {r["sam_gov_opportunity_id"] for r in linked}
            except Exception:
                pass

            # pWin-weighted pipeline roll-up (optional — failures don't block the proposals table)
            pipeline_rollup = {
                "total_weighted_pipeline_value": 0, "total_potential_value": 0,
                "scored_count": 0, "unscored_count": 0, "opportunities": [],
            }
            try:
                pipeline_rollup = pipeline_value_rollup()
            except Exception:
                pass

            # Active proposals — always runs independently of rollup
            active_proposals = []
            try:
                rows = conn.execute(
                    "SELECT id, solicitation_number, title, agency, due_date, estimated_value_low, estimated_value_high, "
                    "win_probability, status, capture_manager, proposal_manager, capture_phase "
                    "FROM proposal_opportunities WHERE status NOT IN ('won','lost','no_bid','cancelled') "
                    "ORDER BY due_date ASC"
                ).fetchall()
                pwin_map = {item["opportunity_id"]: item for item in pipeline_rollup.get("opportunities", [])}
                for row in rows:
                    opp = dict(row)
                    item = pwin_map.get(opp["id"])
                    if item:
                        opp["computed_pwin_pct"] = item["pwin_pct"]
                        opp["weighted_value"] = item["weighted_value"]
                        opp["has_pwin_model"] = item["has_pwin_model"]
                        opp["pwin_factors"] = item.get("factor_breakdown")
                    else:
                        opp["computed_pwin_pct"] = opp.get("win_probability")
                        opp["weighted_value"] = None
                        opp["has_pwin_model"] = False
                        opp["pwin_factors"] = None
                    try:
                        from datetime import date
                        dd = date.fromisoformat(opp["due_date"])
                        opp["days_left"] = (dd - date.today()).days
                    except Exception:
                        opp["days_left"] = None
                    active_proposals.append(opp)
            except Exception:
                pass

            # BD view: CRM engagement heat per agency (prop-cap-14)
            try:
                heat_by_agency = get_engagement_heat_by_agency(
                    [opp.get("agency") for opp in active_proposals]
                )
                for opp in active_proposals:
                    opp["engagement_heat"] = heat_by_agency.get(opp.get("agency"))
            except Exception:
                for opp in active_proposals:
                    opp["engagement_heat"] = None

            # BD view: SAM.gov forecast/presolicitation notice feed (prop-cap-14)
            try:
                forecast_notices = list_forecast_notices(limit=10)
            except Exception:
                forecast_notices = {"notices": [], "count": 0}

            return render_template(
                "govcon/pipeline.html", stats=stats, opportunities=opportunities, linked_opp_ids=linked_opp_ids,
                pipeline_rollup=pipeline_rollup, active_proposals=active_proposals,
                forecast_notices=forecast_notices, degraded=False,
            )
        except Exception as exc:
            # nav-misc-02: a total pipeline failure previously rendered an all-zeros
            # page indistinguishable from an empty-but-healthy pipeline. Log it and
            # flag the page as degraded so operators know the numbers are stale.
            get_logger("icdev.dashboard").warning("govcon_pipeline_page: pipeline load failed: %s", exc)
            stats = {
                "total_opportunities": 0,
                "total_requirements": 0,
                "total_patterns": 0,
                "total_capability_maps": 0,
                "total_drafts": 0,
                "total_awards": 0,
                "knowledge_blocks": 0,
                "linked_proposals": 0,
                "domain_distribution": {},
                "last_pipeline_run": None,
            }
            return render_template(
                "govcon/pipeline.html", stats=stats, opportunities=[], linked_opp_ids=set(),
                forecast_notices={"notices": [], "count": 0}, degraded=True,
            )
        finally:
            conn.close()

    @app.route("/govcon/requirements")
    @require_role("admin", "bd", "capture_mgr", "pm")
    def govcon_requirements_page():
        """GovCon Requirements — pattern frequency, domain heatmap, statement types."""
        conn = _get_db()
        try:
            total_requirements = 0
            try:
                r = conn.execute("SELECT COUNT(*) as cnt FROM rfp_shall_statements").fetchone()
                total_requirements = r["cnt"] if r else 0
            except Exception:
                pass
            total_patterns = 0
            try:
                r = conn.execute("SELECT COUNT(*) as cnt FROM rfp_requirement_patterns").fetchone()
                total_patterns = r["cnt"] if r else 0
            except Exception:
                pass
            domain_stats = {}
            try:
                rows = conn.execute(
                    "SELECT domain_category, COUNT(*) as cnt FROM rfp_shall_statements GROUP BY domain_category ORDER BY cnt DESC"
                ).fetchall()
                domain_stats = {r["domain_category"]: {"count": r["cnt"]} for r in rows}
            except Exception:
                pass
            domain_count = len(domain_stats)
            patterns = []
            min_frequency = 3
            try:
                rows = conn.execute(
                    "SELECT * FROM rfp_requirement_patterns WHERE frequency >= %s ORDER BY frequency DESC LIMIT 30",
                    (min_frequency,),
                ).fetchall()
                patterns = [dict(r) for r in rows]
            except Exception:
                pass
            top_frequency = patterns[0]["frequency"] if patterns else 0
            type_stats = {}
            try:
                rows = conn.execute(
                    "SELECT statement_type, COUNT(*) as cnt FROM rfp_shall_statements GROUP BY statement_type ORDER BY cnt DESC"
                ).fetchall()
                type_stats = {r["statement_type"]: r["cnt"] for r in rows}
            except Exception:
                pass
            return render_template(
                "govcon/requirements.html",
                total_requirements=total_requirements,
                total_patterns=total_patterns,
                domain_stats=domain_stats,
                domain_count=domain_count,
                patterns=patterns,
                top_frequency=top_frequency,
                type_stats=type_stats,
                min_frequency=min_frequency,
            )
        finally:
            conn.close()

    @app.route("/govcon/capabilities")
    @require_role("admin", "bd", "capture_mgr", "pm")
    def govcon_capabilities_page():
        """GovCon Capabilities — coverage by domain, gap list, enhancement recommendations."""
        conn = _get_db()
        try:
            coverage = {"L": 0, "M": 0, "N": 0, "rate": 0}
            try:
                rows = conn.execute(
                    """SELECT
                        SUM(CASE WHEN m.coverage_score >= 0.80 THEN 1 ELSE 0 END) as "L",
                        SUM(CASE WHEN m.coverage_score >= 0.40 AND m.coverage_score < 0.80 THEN 1 ELSE 0 END) as "M",
                        SUM(CASE WHEN m.coverage_score < 0.40 OR m.coverage_score IS NULL THEN 1 ELSE 0 END) as "N",
                        COUNT(*) as total
                    FROM rfp_shall_statements s
                    LEFT JOIN icdev_capability_map m ON s.id = m.pattern_id"""
                ).fetchone()
                if rows:
                    r = dict(rows)
                    total = r.get("total", 0) or 0
                    if total > 0:
                        coverage["L"] = r.get("L", r.get("l", 0)) or 0
                        coverage["M"] = r.get("M", r.get("m", 0)) or 0
                        coverage["N"] = r.get("N", r.get("n", 0)) or 0
                        coverage["rate"] = round(coverage["L"] / total * 100)
            except Exception:
                pass
            domain_coverage = []
            try:
                rows = conn.execute(
                    """SELECT s.domain_category as domain, COUNT(*) as total,
                        SUM(CASE WHEN m.coverage_score >= 0.80 THEN 1 ELSE 0 END) as "L",
                        SUM(CASE WHEN m.coverage_score >= 0.40 AND m.coverage_score < 0.80 THEN 1 ELSE 0 END) as "M",
                        SUM(CASE WHEN m.coverage_score < 0.40 OR m.coverage_score IS NULL THEN 1 ELSE 0 END) as "N"
                    FROM rfp_shall_statements s
                    LEFT JOIN icdev_capability_map m ON s.id = m.pattern_id
                    GROUP BY s.domain_category ORDER BY total DESC"""
                ).fetchall()
                domain_coverage = [
                    {
                        "domain": dict(r).get("domain", ""),
                        "total": dict(r).get("total", 0) or 0,
                        "L": dict(r).get("L", dict(r).get("l", 0)) or 0,
                        "M": dict(r).get("M", dict(r).get("m", 0)) or 0,
                        "N": dict(r).get("N", dict(r).get("n", 0)) or 0,
                    }
                    for r in rows
                ]
            except Exception:
                pass
            gaps = []
            total_gaps = 0
            try:
                rows = conn.execute(
                    """SELECT p.pattern_name as requirement, p.domain_category as domain,
                        p.frequency, COALESCE(m.coverage_score, 0) as coverage,
                        p.frequency * (1 - COALESCE(m.coverage_score, 0)) as priority
                    FROM rfp_requirement_patterns p
                    LEFT JOIN icdev_capability_map m ON p.id = m.pattern_id
                    WHERE COALESCE(m.coverage_score, 0) < 0.40
                    ORDER BY priority DESC LIMIT 20"""
                ).fetchall()
                gaps = [dict(r) for r in rows]
                total_gaps = len(gaps)
            except Exception:
                pass
            recommendations = []
            try:
                from tools.govcon.gap_analyzer import generate_recommendations

                rec_result = generate_recommendations()
                recommendations = rec_result.get("recommendations", [])[:15]
            except Exception:
                pass
            # Cross-RFI capability-gap demand signals (tools/govcon/rfi_demand.py):
            # unmet RFI requirements aggregated across opportunities, highest demand
            # first. Best-effort — absent table / disabled feature yields an empty list.
            rfi_demand_signals = []
            try:
                from tools.govcon.rfi_demand import list_demand_signals

                rfi_demand_signals = list_demand_signals(limit=20)
            except Exception:
                pass
            return render_template(
                "govcon/capabilities.html",
                coverage=coverage,
                domain_coverage=domain_coverage,
                gaps=gaps,
                total_gaps=total_gaps,
                recommendations=recommendations,
                rfi_demand_signals=rfi_demand_signals,
            )
        finally:
            conn.close()


# ---------------------------------------------------------------------------
# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


# require_installed / _module_not_installed moved to tools/dashboard/route_utils.py
# (nav-misc-03). Re-imported here to preserve the public symbol.
from tools.dashboard.route_utils import (  # noqa: F401
    require_installed,
    _module_not_installed,
)




# App factory
# ---------------------------------------------------------------------------


# OPT-12 — in-process result caches for the two canvas-heavy chart endpoints.
# Canvas state changes infrequently so a 45-second TTL gives a good balance
# between freshness and eliminating the ~50 ms cold-open-and-parse cost that
# hits every poll cycle.
_CANVAS_COMPLIANCE_CACHE: dict = {}   # {"ts": float, "data": list}
_CANVAS_TREND_CACHE: dict = {}        # {"ts": float, "data": list}
_CANVAS_CACHE_TTL = 45.0              # seconds

_PATTERN_LABELS: dict[str, str] = {
    "double_top": "▼ Double Top",
    "double_bottom": "▲ Double Bottom",
    "triple_top": "▼ Triple Top",
    "triple_bottom": "▲ Triple Bottom",
    "rising_wedge": "⋀ Rising Wedge",
    "falling_wedge": "⋁ Falling Wedge",
}


def _enrich_chart_patterns(patterns: list[dict]) -> list[dict]:
    """Attach label, breakout_bar, confidence, price_low/high to each pattern."""
    enriched = []
    for p in patterns:
        ep = dict(p)
        ep["label"] = _PATTERN_LABELS.get(p["type"], p["type"])
        ep["breakout_bar"] = p["end_bar"]

        tol = p.get("tolerance_pct", 3.0) or 3.0

        if p["type"] in ("double_top", "double_bottom"):
            avg = p.get("avg_price") or 0.0
            if p["type"] == "double_top":
                s1, s2 = p["high_1"], p["high_2"]
                ep["price_high"] = round(avg, 4)
                ep["price_low"] = round(p["neckline"]["price"], 4)
            else:
                s1, s2 = p["low_1"], p["low_2"]
                ep["price_low"] = round(avg, 4)
                ep["price_high"] = round(p["neckline"]["price"], 4)
            max_dev_pct = (max(abs(s1["price"] - avg), abs(s2["price"] - avg)) / avg * 100) if avg else 0
            ep["confidence"] = round(max(0.10, min(0.99, 1.0 - max_dev_pct / tol)), 2)

        elif p["type"] in ("triple_top", "triple_bottom"):
            avg = p.get("avg_price") or 0.0
            prices = [p["swing_1"]["price"], p["swing_2"]["price"], p["swing_3"]["price"]]
            if p["type"] == "triple_top":
                ep["price_high"] = round(avg, 4)
                ep["price_low"] = round(min(prices), 4)
            else:
                ep["price_low"] = round(avg, 4)
                ep["price_high"] = round(max(prices), 4)
            max_dev_pct = (max(abs(pr - avg) for pr in prices) / avg * 100) if avg else 0
            ep["confidence"] = round(max(0.10, min(0.99, 1.0 - max_dev_pct / tol)), 2)

        elif p["type"] in ("rising_wedge", "falling_wedge"):
            sb, eb = p["start_bar"], p["end_bar"]
            sh = p.get("slope_high", 0)
            ih = p.get("intercept_high", 0)
            sl = p.get("slope_low", 0)
            il = p.get("intercept_low", 0)
            res_prices = [sh * sb + ih, sh * eb + ih]
            sup_prices = [sl * sb + il, sl * eb + il]
            ep["price_high"] = round(max(res_prices), 4)
            ep["price_low"] = round(min(sup_prices), 4)
            ep["confidence"] = 0.65

        else:
            ep["confidence"] = 0.50
            ep.setdefault("price_low", 0.0)
            ep.setdefault("price_high", 0.0)

        enriched.append(ep)
    return enriched


def _derive_chart_provenance(bars: list[dict]) -> dict:
    """Derive top-level data-provenance flags for a chart response.

    ``market_data.fetch_bars`` tags every bar with a per-bar ``source``
    ("alpaca" / "alpaca_crypto" / "sample") plus ``as_of``. The chart frontend
    ignores those per-bar markers, so synthetic ("sample") bars can render as
    real market data in a financial UI. This helper collapses the per-bar
    markers into top-level ``data_source`` / ``simulated`` / ``as_of`` fields
    the UI can surface as a prominent "SIMULATED DATA" banner (nav-plat-01).

    Args:
        bars: OHLCV bar dicts, each optionally carrying ``source`` / ``as_of``.

    Returns:
        Dict with keys::

            {
              "data_source": str,   # "alpaca"|"sample"|...|"mixed"|"unknown"
              "simulated": bool,    # True iff any bar is synthetic ("sample")
              "as_of": str | None,  # most-recent per-bar as_of, if present
            }
    """
    if not bars:
        return {"data_source": "unknown", "simulated": False, "as_of": None}
    sources = {str(b.get("source") or "unknown") for b in bars if isinstance(b, dict)}
    simulated = "sample" in sources
    if len(sources) == 1:
        data_source = next(iter(sources))
    elif sources:
        data_source = "mixed"
    else:
        data_source = "unknown"
    as_of = None
    for b in reversed(bars):
        if isinstance(b, dict) and b.get("as_of"):
            as_of = b["as_of"]
            break
    return {"data_source": data_source, "simulated": simulated, "as_of": as_of}


def _get_chat_models() -> tuple[list[dict], str]:
    """Read available chat models from args/llm_config.yaml.

    Returns (models_list, default_model_key) where models_list is
    [{value, label, provider}] and default_model_key is the first entry
    in the chat_response routing chain.
    """
    import yaml

    config_path = BASE_DIR / "args" / "llm_config.yaml"
    try:
        with open(config_path, encoding="utf-8") as _f:
            cfg = yaml.safe_load(_f)
    except Exception:
        return [{"value": "default", "label": "Default", "provider": ""}], "default"

    _PROVIDER_LABELS = {
        "ollama": "Local (Ollama)",
        "anthropic": "Anthropic",
        "openai": "OpenAI",
        "gemini": "Google Gemini",
        "bedrock": "AWS Bedrock",
        "azure_openai": "Azure OpenAI",
        "ibm_watsonx": "IBM watsonx",
        "mistral": "Mistral",
        "mistral_vllm": "Mistral (vLLM)",
        "vllm": "vLLM",
    }

    # Collect embedding model IDs so we can skip them
    embed_ids: set[str] = set()
    try:
        for _v in cfg.get("embedding", {}).get("models", {}).values():
            if isinstance(_v, dict):
                embed_ids.add(_v.get("model_id", ""))
    except Exception:
        pass

    result: list[dict] = []
    for key, mcfg in (cfg.get("models") or {}).items():
        if not isinstance(mcfg, dict):
            continue
        if key.startswith("agent_"):
            continue
        provider = mcfg.get("provider", "")
        model_id = mcfg.get("model_id", key)
        if model_id in embed_ids:
            continue
        provider_label = _PROVIDER_LABELS.get(provider, provider.replace("_", " ").title())
        result.append({
            "value": key,
            "label": f"{key}  [{provider_label}]",
            "provider": provider,
        })

    # Determine default from chat_response routing chain
    default_model = result[0]["value"] if result else "default"
    try:
        chain = cfg.get("routing", {}).get("chat_response", {}).get("chain", [])
        if chain:
            default_model = chain[0]
    except Exception:
        pass

    return result, default_model


def _aggregate_chat_sources(conn, tenant_id: str, context_id: str) -> list[dict]:
    """Aggregate chat-upload RAG chunks by source_id.

    Parses JSON metadata in Python so no SQLite-only JSON SQL function appears
    in the SQL — making the query run on PostgreSQL without modification.
    Mirrors the logic of the original GROUP-BY json_extract query.
    """
    import json as _json
    from tools.db.storage import sql_placeholder as _sqlph
    ph = _sqlph(conn)

    params: list = ["chat_upload"]
    sql = (
        "SELECT source_id, metadata, created_at "
        "FROM rag_chunks "
        f"WHERE source_type = {ph}"
    )
    if tenant_id:
        sql += f" AND tenant_id = {ph}"
        params.append(tenant_id)
    sql += " ORDER BY created_at DESC"

    rows = conn.execute(sql, params).fetchall()

    agg: dict = {}
    for row in rows:
        r = dict(row)
        sid = r["source_id"]
        try:
            meta = _json.loads(r.get("metadata") or "{}")
        except (ValueError, TypeError):
            meta = {}
        row_ctx = meta.get("context_id", "")
        if context_id and row_ctx != context_id:
            continue
        if sid not in agg:
            agg[sid] = {
                "source_id": sid,
                "filename": meta.get("filename", ""),
                "context_id": row_ctx,
                "chunk_count": 0,
                "indexed_at": r.get("created_at"),
            }
        agg[sid]["chunk_count"] += 1
        if not agg[sid]["filename"] and meta.get("filename"):
            agg[sid]["filename"] = meta.get("filename", "")
        cur_ts = r.get("created_at")
        if cur_ts and (not agg[sid]["indexed_at"] or cur_ts > agg[sid]["indexed_at"]):
            agg[sid]["indexed_at"] = cur_ts

    result = sorted(agg.values(), key=lambda x: x.get("indexed_at") or "", reverse=True)
    return result[:50]


def create_app(testing: bool = False) -> Flask:
    app = Flask(
        __name__,
        template_folder=str(Path(__file__).resolve().parent / "templates"),
        static_folder=str(Path(__file__).resolve().parent / "static"),
    )
    if testing:
        app.config["TESTING"] = True

    # Auto-reload templates on change (no server restart needed)
    app.config["TEMPLATES_AUTO_RELOAD"] = True
    app.jinja_env.auto_reload = True
    app.config["EXCALIDRAW_HOST"] = os.environ.get("EXCALIDRAW_HOST", "")

    # cnr-plat-02: global upload size cap. Flask rejects any request body larger
    # than MAX_CONTENT_LENGTH with 413 before it is buffered into memory. Tunable
    # via ICDEV_MAX_UPLOAD_MB (default 50 MB). Tighter per-route caps (docgen,
    # workflow_canvas) already validate content_length and stay authoritative for
    # their routes; this is the platform-wide backstop.
    try:
        _max_upload_mb = float(os.environ.get("ICDEV_MAX_UPLOAD_MB", "50"))
    except (TypeError, ValueError):
        _max_upload_mb = 50.0
    if _max_upload_mb > 0:
        app.config["MAX_CONTENT_LENGTH"] = int(_max_upload_mb * 1024 * 1024)

    # Release cached canvas DB connections after each request (OPT-06).
    @app.teardown_appcontext
    def _teardown_canvas_connections(exc):  # noqa: ANN001
        _close_canvas_connections()

    # Register UX filters (glossary, timestamps, error recovery, quick paths)
    register_ux_filters(app)

    # Register CLI bridge toggle middleware (silences 404 in cli_bridge_indicator)
    try:
        from tools.dashboard.api.cli_bridge_api import register_cli_bridge
        register_cli_bridge(app)
    except Exception as _exc:
        app.logger.warning("CLI bridge middleware skipped: %s", _exc)

    # Register dashboard auth middleware (D169-D172)
    register_dashboard_auth(app)

    # cnr-plat-01: CSRF protection for cookie-authenticated mutating JSON APIs.
    # Registered right after auth so its before_request runs after the auth hook
    # (g.current_user / cortex_binding are resolved first). Enforced by default
    # (ICDEV_CSRF_ENFORCE); token-auth / API-key / ICDEV_AUTH_BYPASS paths exempt.
    try:
        from tools.security.csrf import register_csrf
        register_csrf(app)
    except Exception as _csrf_exc:
        app.logger.warning("CSRF protection not registered: %s", _csrf_exc)

    # Register field-level security middleware (CUI enforcement on JSON responses)
    try:
        from tools.security.middleware import init_security
        init_security(app, classification="CUI")
    except ImportError:
        pass

    # Initialize WebSocket (D170 — optional, graceful fallback)
    init_socketio(app)

    # Register geospatial dashboard SocketIO handlers (task-a866147c27-d4)
    try:
        from src.routes.dashboard import register_socketio_handlers as _register_geo_ws
        _register_geo_ws()
    except Exception as _exc:
        app.logger.warning("Geospatial SocketIO handlers skipped: %s", _exc)

    # Correlation ID middleware (D149)
    try:
        from tools.resilience.correlation import register_correlation_middleware

        register_correlation_middleware(app)
    except ImportError:
        pass

    # Budget monitor + throttling controller — background daemon wired into the
    # generative intelligence and predictive analysis pipeline (services package).
    try:
        from services import start_budget_services
        app.extensions["throttle_controller"] = start_budget_services()
    except Exception as _exc:
        app.logger.warning("Budget services skipped: %s", _exc)

    # Health / readiness / liveness probes (ECR-OBS-03)
    try:
        from tools.observability.health_blueprint import health_bp as _health_bp
        app.register_blueprint(_health_bp)
    except Exception as _exc:
        app.logger.warning("Health blueprint skipped: %s", _exc)

    # Distributed tracing activation (obx-trc-01, D290). Gated by
    # ICDEV_TRACING_ENABLED (default on) inside the helper. Wrapped so tracing
    # never blocks app startup. Span store routes to the primary backend via
    # tools.db.storage.
    try:
        from tools.observability import enable_tracing_if_enabled
        enable_tracing_if_enabled()
    except Exception as _exc:
        app.logger.warning("Tracing activation skipped: %s", _exc)

    @app.route("/api/_introspect/routes", methods=["GET"])
    def _introspect_routes():
        """Internal: real GET-able, parameter-free page routes for the health prober.

        Sourced from the live Flask url_map so the prober tests the actual mounted
        paths (with blueprint url_prefixes applied) instead of decorator-relative
        paths. Excludes /static, /api/* (endpoints, not pages — many are POST/param),
        and parameterized rules (can't HEAD without a real id).
        """
        routes = set()
        for rule in app.url_map.iter_rules():
            if rule.arguments:
                continue  # parameterized — skip
            if "GET" not in (rule.methods or set()):
                continue  # not GET-probeable
            path = rule.rule
            if path.startswith("/static") or path.startswith("/api/"):
                continue  # static assets / API endpoints, not user-facing pages
            routes.add(path)
        return jsonify({"routes": sorted(routes), "count": len(routes)})

    @app.route("/favicon.ico")
    def favicon():
        return make_response("", 204)

    @app.route("/api/live-check", methods=["GET"])
    def api_live_check():
        """Scheduler heartbeat + in_progress task count for the Live Activity panel."""
        import time as _t
        from flask import jsonify as _j
        from tools.db.storage import get_connection as _gc

        hb_path = BASE_DIR / ".tmp" / "kanban_scheduler.heartbeat"
        log_path = BASE_DIR / ".tmp" / "kanban_scheduler.log"
        sched_secs = None
        for _p in (hb_path, log_path):
            if _p.exists():
                sched_secs = int(_t.time() - _p.stat().st_mtime)
                break

        if sched_secs is None:
            staleness = "unknown"
        elif sched_secs > 600:
            staleness = "stale"
        elif sched_secs > 180:
            staleness = "warning"
        else:
            staleness = "active"

        conn = _gc()
        pending_count = 0
        # kax-obs-02: the scheduler's liveness above is a log-file mtime, which
        # is exactly what goes missing when logging breaks. The PR watcher's is
        # a DB row written by each completed poll instead.
        watcher = {}
        try:
            rows = conn.execute(
                "SELECT id, title, priority, task_type, failure_count, status, "
                "last_failure_at, updated_at, dispatch_source, executor_type "
                "FROM kanban_tasks WHERE status = 'in_progress' ORDER BY updated_at DESC"
            ).fetchall()
            tasks = [dict(r) for r in rows]
            # kv-viz: add attempt_count, current_attempt_started_at, last_reaped_reason
            from tools.dashboard.api.kanban import _annotate_in_progress_tasks
            _annotate_in_progress_tasks(conn, tasks)
            # Count queued (scheduled + backlog) tasks so the Projects in Flight
            # section stays visible even when no tasks are actively in_progress.
            prow = conn.execute(
                "SELECT COUNT(*) AS n FROM kanban_tasks "
                "WHERE status IN ('scheduled', 'backlog')"
            ).fetchone()
            pending_count = int((prow or {}).get("n") or 0)
        except Exception:
            tasks = []
        finally:
            try:
                from tools.kanban.metrics import watcher_heartbeat as _wh
                watcher = _wh(conn=conn)
            except Exception:
                watcher = {}
            conn.close()

        # Enrich each task with per-task liveness from its agent log file.
        # .tmp/kanban/{task_id}.log is written by the Claude CLI subprocess
        # while the agent is running. A stale or missing log while the task
        # is still in_progress means the subprocess is dead or hung.
        #
        # task_log_age_secs thresholds (independent of scheduler heartbeat):
        #   null          → log file not found yet (task just dispatched)
        #   0 – 300       → active   (agent writing output in last 5 min)
        #   300 – 900     → quiet    (agent silent 5-15 min; may be thinking)
        #   > 900         → suspect  (likely zombie — log older than task timeout)
        kanban_dir = BASE_DIR / ".tmp" / "kanban"
        now_ts = _t.time()
        for task in tasks:
            log_file = kanban_dir / f"{task['id']}.log"
            if log_file.exists():
                task["task_log_age_secs"] = int(now_ts - log_file.stat().st_mtime)
                task["task_log_size"] = log_file.stat().st_size
            else:
                task["task_log_age_secs"] = None
                task["task_log_size"] = None

        return _j({
            "scheduler_last_seen_secs": sched_secs,
            "staleness": staleness,
            "tasks": tasks,
            "pending_tasks_count": pending_count,
            "pr_watcher": watcher,
        })

    @app.route("/api/kanban/scheduler/status")
    def kanban_scheduler_status():
        """GET — whether the kanban scheduler is paused (manual flag or auto)."""
        from tools.kanban.scheduler_control import status as _sched_status  # noqa: PLC0415
        return jsonify(_sched_status())

    @app.route("/api/kanban/scheduler/pause", methods=["POST"])
    @require_role("admin", "pm")
    def kanban_scheduler_pause():
        """POST — manually pause the kanban scheduler cycle."""
        from tools.kanban.scheduler_control import pause as _sched_pause  # noqa: PLC0415
        body = flask_request.get_json(silent=True) or {}
        return jsonify(_sched_pause(actor=_session_actor(),
                                    reason=body.get("reason", "manual (dashboard)")))

    @app.route("/api/kanban/scheduler/resume", methods=["POST"])
    @require_role("admin", "pm")
    def kanban_scheduler_resume():
        """POST — resume the kanban scheduler cycle."""
        from tools.kanban.scheduler_control import resume as _sched_resume  # noqa: PLC0415
        return jsonify(_sched_resume(actor=_session_actor()))

    @app.route("/api/kanban/build-mode", methods=["GET", "POST"])
    @require_role("admin", "pm")
    def kanban_build_mode():
        """Manual Build: promote and track as normal, but do not auto-dispatch.

        NOT the same as Pause Scheduler. Pausing freezes the whole cycle — nothing
        promotes, nothing is tracked, and a build you then do by hand is invisible to the
        board. Manual Build stops only the dispatch: backlog still promotes to scheduled,
        project cards still show progress, and a CLI session does the building. So when a
        manual build is interrupted, the board still says exactly where it got to.
        """
        from tools.kanban.build_mode import set_manual, status as _bm_status  # noqa: PLC0415

        if flask_request.method == "GET":
            return jsonify(_bm_status())

        body = flask_request.get_json(silent=True) or {}
        manual = bool(body.get("manual"))
        return jsonify(set_manual(
            manual,
            actor=_session_actor(),
            reason=body.get("reason", ""),
        ))

    @app.route("/api/kanban/build-model", methods=["GET", "POST"])
    @require_role("admin", "pm")
    def kanban_build_model():
        """Which model the runner builds with. Live — no scheduler restart.

        POST {"model": null} clears back to config-driven routing.

        Selecting a model the Claude CLI cannot serve (Kimi, Ollama, GPT...) removes
        claude_cli from the executor chain, so the choice actually takes effect rather
        than being silently ignored while Claude keeps building.
        """
        from tools.kanban.model_override import (  # noqa: PLC0415
            available as _models_available,
            set_model as _set_model,
            status as _model_status,
        )

        if flask_request.method == "GET":
            return jsonify({**_model_status(), "available": _models_available()})

        body = flask_request.get_json(silent=True) or {}
        try:
            return jsonify(_set_model(body.get("model"), actor=_session_actor()))
        except ValueError as exc:
            return jsonify({"error": str(exc)}), 400

    @app.route("/api/kanban/recent-events")
    def api_kanban_recent_events():
        """GET /api/kanban/recent-events — Last 24h kanban task events from notifications.

        Returns done/failed/in_progress/unverified events with validation gate detail.
        Used by the Projects in Flight panel on the home dashboard.
        """
        from tools.db.storage import get_connection as _gc
        from datetime import timedelta as _td
        from flask import jsonify as _j2

        limit = min(int(flask_request.args.get("limit", 30)), 100)
        since_hours = int(flask_request.args.get("hours", 24))

        try:
            cutoff = (datetime.now(timezone.utc) - _td(hours=since_hours)).isoformat()
            with _gc() as conn:
                rows = conn.execute(
                    "SELECT id, title, message, severity, created_at "
                    "FROM notifications "
                    "WHERE source = 'genesis.kanban' "
                    "  AND created_at > %s "
                    "ORDER BY created_at DESC LIMIT %s",
                    (cutoff, limit),
                ).fetchall()
            events = []
            for r in rows:
                d = dict(r)
                title = d.get("title") or ""
                # Parse event type from title: "Task {event}: {task_title}"
                event_type = "unknown"
                task_title = title
                if title.upper().startswith("UNVERIFIED"):
                    event_type = "unverified"
                    task_title = title
                elif title.startswith("Task "):
                    rest = title[5:]
                    colon = rest.find(":")
                    if colon != -1:
                        event_type = rest[:colon].strip().lower()
                        task_title = rest[colon + 1:].strip()
                events.append({
                    "id": d["id"],
                    "title": task_title,
                    "event_type": event_type,
                    "message": (d.get("message") or "")[:600],
                    "severity": d.get("severity") or "info",
                    "created_at": d["created_at"],
                })
            return _j2({"events": events})
        except Exception as exc:
            return _j2({"events": [], "error": str(exc)})

    _NOTIFY_SETTINGS_PATH = BASE_DIR / "args" / "kanban_notify.json"
    _NOTIFY_CHANNELS = [
        {"id": "telegram",   "label": "Telegram",   "env_keys": ["TELEGRAM_BOT_TOKEN", "TELEGRAM_CHAT_ID"]},
        {"id": "slack",      "label": "Slack",      "env_keys": ["SLACK_BOT_TOKEN", "SLACK_WEBHOOK_URL"]},
        {"id": "teams",      "label": "MS Teams",   "env_keys": ["TEAMS_WEBHOOK_URL"]},
        {"id": "email",      "label": "Email",      "env_keys": ["SMTP_HOST", "EMAIL_HOST"]},
        {"id": "webhook",    "label": "Webhook",    "env_keys": ["WEBHOOK_URL"]},
        {"id": "mattermost", "label": "Mattermost", "env_keys": ["MATTERMOST_WEBHOOK_URL", "MATTERMOST_URL"]},
    ]

    @app.route("/api/kanban/notify-channel", methods=["GET", "PUT"])
    def api_kanban_notify_channel():
        """GET/PUT /api/kanban/notify-channel — Read or set the active notification channel.

        GET returns {available: [{id, label, configured}], current}.
        PUT body {channel: "slack"} persists to args/kanban_notify.json.
        """
        import os as _os
        from flask import jsonify as _jnc

        def _load_settings():
            if _NOTIFY_SETTINGS_PATH.exists():
                try:
                    import json as _j
                    return _j.loads(_NOTIFY_SETTINGS_PATH.read_text(encoding="utf-8"))
                except Exception:
                    pass
            return {"current": "telegram"}

        def _save_settings(settings):
            _NOTIFY_SETTINGS_PATH.parent.mkdir(parents=True, exist_ok=True)
            import json as _j
            _NOTIFY_SETTINGS_PATH.write_text(
                _j.dumps(settings, indent=2), encoding="utf-8", newline=""
            )

        if flask_request.method == "PUT":
            body = flask_request.get_json(silent=True) or {}
            channel = body.get("channel", "").strip()
            valid_ids = {ch["id"] for ch in _NOTIFY_CHANNELS}
            if channel not in valid_ids:
                return _jnc({"error": f"Unknown channel: {channel}"}), 400
            settings = _load_settings()
            settings["current"] = channel
            _save_settings(settings)
            return _jnc({"ok": True, "current": channel})

        # GET
        settings = _load_settings()
        current = settings.get("current", "telegram")
        available = []
        for ch in _NOTIFY_CHANNELS:
            configured = any(_os.environ.get(k) for k in ch["env_keys"])
            available.append({"id": ch["id"], "label": ch["label"], "configured": configured})
        return _jnc({"available": available, "current": current})

    # Role-based view configuration
    ROLE_VIEWS = {
        "pm": {
            "label": "Program Manager",
            "show_tabs": ["overview", "compliance", "deployments", "audit"],
            "hide_columns": ["stig_id", "finding_id"],
        },
        "developer": {
            "label": "Developer / Architect",
            "show_tabs": ["overview", "security", "deployments", "audit"],
            "hide_columns": [],
        },
        "isso": {
            "label": "ISSO / Security Officer",
            "show_tabs": ["overview", "compliance", "security", "audit"],
            "hide_columns": [],
        },
        "co": {
            "label": "Contracting Officer",
            "show_tabs": ["overview", "compliance", "deployments"],
            "hide_columns": ["stig_id", "finding_id", "source"],
        },
        "analyst": {
            "label": "Analyst",
            "show_tabs": ["overview", "compliance", "security", "audit"],
            "hide_columns": [],
        },
        "solutions_architect": {
            "label": "Solutions Architect",
            "show_tabs": ["overview", "security", "deployments", "audit"],
            "hide_columns": [],
        },
        "sales_engineer": {
            "label": "Sales Engineer",
            "show_tabs": ["overview", "compliance", "deployments"],
            "hide_columns": ["stig_id", "finding_id"],
        },
        "innovator": {
            "label": "Innovator",
            "show_tabs": ["overview", "security", "deployments", "audit"],
            "hide_columns": [],
        },
        "biz_dev": {
            "label": "Business Development",
            "show_tabs": ["overview", "compliance", "deployments"],
            "hide_columns": ["stig_id", "finding_id", "source"],
        },
        "cor": {
            "label": "Contracting Officer Representative",
            "show_tabs": ["overview", "compliance"],
            "hide_columns": ["stig_id", "finding_id", "source", "internal_cost_details"],
        },
    }

    # Make CUI config, role, and user info available in all templates
    @app.context_processor
    def inject_cui():
        role = flask_request.args.get("role", "")
        role_config = ROLE_VIEWS.get(role, None)
        current_user = getattr(g, "current_user", None)
        # Route-to-module map for assistant widget auto-scoping (D-CA-4)
        try:
            from tools.dashboard.assistant_config import ROUTE_MODULE_MAP

            _route_map = ROUTE_MODULE_MAP
        except ImportError:
            _route_map = {}

        theme_pref = flask_request.cookies.get("icdev_theme", "dark")

        # ECR-CL-02 / nav-misc-01: unseen-release badge. Logged-in users get a
        # per-user last-seen version persisted in user_preferences (synced across
        # devices, survives cookie clears); anonymous users fall back to the
        # icdev_seen_version cookie.
        _unseen_release = False
        try:
            from tools.dashboard.brand import get_brand as _get_brand_ctx
            _bv = _get_brand_ctx().get("version", "")
            if _bv:
                _uid = _current_user_id()
                if _uid:
                    from tools.auth.onboarding import get_last_seen_version
                    _seen_ver = get_last_seen_version(_uid) or ""
                else:
                    _seen_ver = flask_request.cookies.get("icdev_seen_version", "")
                _unseen_release = _seen_ver != _bv
        except Exception:
            pass

        return {
            "cui_banner_top": CUI_BANNER_TOP,
            "cui_banner_bottom": CUI_BANNER_BOTTOM,
            "cui_banner_enabled": CUI_BANNER_ENABLED,
            "cui_designation": CUI_DESIGNATION,
            "auth_bypass": os.environ.get("ICDEV_AUTH_BYPASS", "").lower() in ("1", "true", "yes"),
            "current_role": role,
            "theme_pref": theme_pref,
            "role_config": role_config,
            "ROLE_VIEWS": ROLE_VIEWS,
            "current_user": current_user,
            "byok_enabled": BYOK_ENABLED,
            "strategos_enabled": _HAS_STRATEGOS,
            "govcon_enabled": _HAS_GOVCON and not _AIRGAP_MODE,
            "network_enabled": _HAS_NETWORK,
            "pipeline_enabled": _HAS_PIPELINE,
            "security_canvas_enabled": _HAS_SECURITY_CANVAS,
            "infra_canvas_enabled": _HAS_INFRA_CANVAS,
            "data_canvas_enabled": _HAS_DATA_CANVAS,
            "boundary_canvas_enabled": _HAS_BOUNDARY_CANVAS,
            "observability_canvas_enabled": _HAS_OBSERVABILITY_CANVAS,
            "canvas_kg_enabled": _HAS_CANVAS_KG,
            "qdc_enabled": _CANVAS_FLAGS.get("qdc", False),
            "integrity_enabled": _CANVAS_FLAGS.get("integrity", False),
            "foundry_enabled": _CANVAS_FLAGS.get("foundry", False),
            "migration_canvas_enabled": _CANVAS_FLAGS.get("mdc", False),
            "aadc_enabled": _CANVAS_FLAGS.get("aadc", False),
            "aimc_enabled": _CANVAS_FLAGS.get("aimc", False),
            "ohc_enabled": _CANVAS_FLAGS.get("ohc", False),
            "nocc_enabled": _CANVAS_FLAGS.get("nocc", False),
            "pmc_enabled": _CANVAS_FLAGS.get("pmc", False),
            "ccc_enabled": _CANVAS_FLAGS.get("ccc", False),
            "dsoc_enabled": _CANVAS_FLAGS.get("dsoc", False),
            "aiify_enabled": _CANVAS_FLAGS.get("aiify", False),
            "dic_enabled": _CANVAS_FLAGS.get("dic", False),
            "demo_runner_enabled": _CANVAS_FLAGS.get("demo_runner", False),
            "govlift_enabled": _CANVAS_FLAGS.get("govlift", False),
            "mission_canvas_enabled": _CANVAS_FLAGS.get("mission_canvas", False),
            "canvas_flags": _CANVAS_FLAGS,
            "hitl_enabled": _APP_FLAGS.get("hitl_workflow", False),
            "academy_enabled": _APP_FLAGS.get("forge_academy", False),
            "gameday_enabled": _APP_FLAGS.get("gameday", False),
            "airgap_mode": _AIRGAP_MODE,
            "route_module_map": _route_map,
            "nav_tree": _REGISTRY.get_nav_context(),
            "component_registry": _REGISTRY,
            # cvx-nav-01: single registry-derived path→canvas map (JSON) injected
            # once into base.html as window.__ICDEV_PATH_CANVAS__.
            "iqe_path_canvas_json": _IQE_PATH_CANVAS_JSON,
            "canvas_menu_active": any(
                flask_request.path.startswith(prefix)
                for prefix in _CANVAS_URL_PREFIXES
            ),
            # License tier context — used by nav lock indicators and tier_gate.html
            "active_tier": _get_active_tier_safe(),
            "tier_order": {"community": 0, "professional": 1, "enterprise": 2},
            "unseen_release": _unseen_release,
        }

    # ---- Brand + banner context processor (DSW-1) ----
    app.context_processor(brand_context_processor)

    # ---- Air-gap route guard: friendly message for disabled pages ----
    if _AIRGAP_MODE:

        @app.before_request
        def _airgap_route_guard():
            path = flask_request.path.rstrip("/") or "/"
            # Check exact match or prefix match for nested routes
            for disabled in _AIRGAP_DISABLED_ROUTES:
                if path == disabled or path.startswith(disabled + "/"):
                    if flask_request.is_json or path.startswith("/api/"):
                        return jsonify(
                            {
                                "error": "unavailable",
                                "message": "This feature is not available in air-gap mode.",
                            }
                        ), 503
                    return render_template(
                        "airgap_unavailable.html",
                        feature_name=disabled.strip("/").replace("-", " ").title(),
                    ), 200

    # ---- ECR-DEMO-03: Demo mode read-only guard ----
    if _DEMO_MODE:

        @app.before_request
        def _demo_mode_guard():
            if flask_request.method not in ("POST", "PUT", "DELETE", "PATCH"):
                return None
            path = flask_request.path
            if not path.startswith("/api/"):
                return None
            # Allow onboarding wizard and IQE queries through
            if path.startswith("/api/onboarding/") or path == "/api/iqe-query":
                return None
            from tools.dashboard.brand import get_brand
            upgrade_url = get_brand().get("support_url", "")
            return jsonify({"error": "Demo mode: read-only", "upgrade_url": upgrade_url}), 403

    # ---- ECR-BILL-02: API call metering (fire-and-forget after_request) ----
    @app.after_request
    def _meter_api_call(response):
        try:
            if flask_request.path.startswith("/api/"):
                from flask import g as _g
                _tenant = getattr(_g, "tenant_id", None) or "system"
                from tools.billing.metering import record_usage as _rec
                _rec(_tenant, "api_call")
        except Exception:
            pass
        return response

    # ---- Auto-register A2A agents from card files ----
    try:
        from tools.a2a.agent_registry import register_all_from_cards

        registered = register_all_from_cards()
        if registered:
            app.logger.info("Auto-registered %d agents from card files", len(registered))
    except Exception as exc:
        app.logger.debug("Agent auto-registration skipped: %s", exc)

    # ---- Studio Artifact Download — standalone route outside blueprint namespace ----
    @app.route("/api/artifacts/<path:filepath>", methods=["GET"])
    def studio_artifact_download(filepath: str):
        import mimetypes
        from flask import send_file as _sf, jsonify as _jf, request as _req
        _artifacts_root = BASE_DIR / "data" / "studio_artifacts"
        safe = (BASE_DIR / Path(filepath.replace("\\", "/"))).resolve()
        if not str(safe).startswith(str(_artifacts_root.resolve())):
            return _jf({"error": "Access denied"}), 403
        if not safe.exists():
            return _jf({"error": "Artifact not found", "path": str(safe)}), 404
        mime, _ = mimetypes.guess_type(str(safe))
        as_attachment = _req.args.get("download", "0") == "1"
        return _sf(safe, mimetype=mime or "text/plain", as_attachment=as_attachment,
                   download_name=safe.name)

    # ---- Studio Run History Delete — standalone routes (blueprint legacy drops DELETE) ----
    @app.route("/api/studio/workflows/runs/<run_id>", methods=["DELETE"])
    def studio_delete_run(run_id: str):
        from tools.studio.workflow_runner import delete_run as _delete_run
        from flask import jsonify as _jf
        deleted = _delete_run(run_id)
        if not deleted:
            return _jf({"error": "Run not found"}), 404
        return _jf({"status": "deleted", "run_id": run_id})

    @app.route("/api/studio/workflows/runs", methods=["DELETE"])
    def studio_delete_all_runs():
        from tools.studio.workflow_runner import delete_all_runs as _delete_all_runs
        from flask import jsonify as _jf, request as _req
        workflow_id = _req.args.get("workflow_id")
        count = _delete_all_runs(workflow_id=workflow_id or None)
        return _jf({"status": "deleted", "count": count})

    # ---- Register API blueprints (P1.1: centralized via register_api_blueprints) ----
    # All 55+ blueprints are mounted under /api/v1/* with /api/* legacy aliases.
    # See tools/dashboard/api/__init__.py for the full registration sequence.
    register_api_blueprints(app)

    # ---- Studio DB init (kanban/ci-fix-26594490171) ----
    try:
        from tools.studio.init_db import init_studio_tables
        init_studio_tables()
    except Exception as _exc:
        app.logger.warning("Studio DB init skipped: %s", _exc)

    # ---- Studio run reconciliation (dwo-dur-02) ----
    # Re-attach runs parked on an approval gate, expire the ones past their
    # window, and fail steps whose subprocess died with the previous process.
    # Called here rather than at import time so importing workflow_runner has
    # no database side effects.
    try:
        from tools.studio.workflow_runner import reconcile_runs_on_boot
        _rec = reconcile_runs_on_boot()
        if _rec.get("resumed") or _rec.get("expired"):
            app.logger.info(
                "Studio runs reconciled: resumed=%d expired=%d",
                len(_rec.get("resumed", [])), len(_rec.get("expired", [])),
            )
    except Exception as _exc:
        app.logger.warning("Studio run reconciliation skipped: %s", _exc)

    # ---- Kanban DB init (ci-fix-26601155261) ----
    try:
        from tools.kanban.init_db import init_kanban_tables as _init_kanban
        _init_kanban()
    except Exception as _exc:
        app.logger.warning("Kanban DB init skipped: %s", _exc)

    # ---- E2E demo session seed (ME conflict lifecycle tests) ----
    try:
        from tools.db.storage import get_connection as _gc
        _seed_conn = _gc()
        try:
            _seed_conn.execute(
                "INSERT OR IGNORE INTO intake_sessions "
                "(id, customer_name, customer_org, session_status, classification, context_summary) "
                "VALUES (%s, %s, %s, %s, %s, %s)",
                ("sess-9cc6891cb548", "E2E Test User", "ICDEV CI", "active", "CUI", "{}"),
            )
            _seed_conn.commit()
        finally:
            _seed_conn.close()
    except Exception as _exc:
        app.logger.debug("E2E session seed skipped: %s", _exc)
    # ---- Geospatial Dashboard (task-a866147c27-d4) ----
    try:
        from src.routes.dashboard import bp as _geo_bp
        app.register_blueprint(_geo_bp)
        app.logger.info("Geospatial dashboard registered at /geospatial")
    except Exception as _exc:
        app.logger.warning("Geospatial dashboard blueprint skipped: %s", _exc)

    # ---- SRE Dashboard Page ----
    @app.route("/sre")
    def sre_dashboard_page():
        return render_template("sre/dashboard.html")

    # ---- NDC SOPs Dashboard Page ----
    @app.route("/ndc/sops")
    def ndc_sops_page():
        return render_template("ndc_sops.html")

    # ---- Canvas Knowledge Graph Blueprint ----
    if _HAS_CANVAS_KG:
        try:
            ckg_bp = create_canvas_kg_blueprint()
            if ckg_bp:
                app.register_blueprint(ckg_bp)
                app.logger.info("Canvas KG registered at /canvas-kg")
        except Exception as exc:
            app.logger.warning("Canvas KG failed to register: %s", exc)

        @app.route("/canvas-kg")
        def canvas_kg_page():
            return render_template("canvas_kg.html")

    # ---- Unified Canvas Compliance Dashboard (Canvas Posture) ----
    @app.route("/canvas-compliance")
    def canvas_compliance_page():
        """Unified compliance posture across the design canvases (Canvas Posture).

        Distinct from Canvas Health (/health/canvases), which reports
        file-existence QA rather than runtime compliance posture.
        """
        # cnr-cc-01(b): explicit, lightweight login gate (defense-in-depth).
        # This standalone @app.route is not a registry canvas blueprint, so it
        # never picks up guard_component_access; assert an authenticated user
        # rather than relying solely on the global before_request hook.
        if not getattr(g, "current_user", None):
            if flask_request.is_json or flask_request.path.startswith("/api/"):
                return jsonify({"error": "Authentication required"}), 401
            return redirect(url_for("login_page"))
        try:
            from tools.canvas_compliance.compliance import get_all_cards
            cards = get_all_cards()
        except Exception as _exc:
            app.logger.warning("canvas_compliance: get_all_cards failed: %s", _exc)
            cards = []
        return render_template("canvas_compliance.html", cards=cards)

    # ---- Design Canvases (registry-driven) ----
    try:
        from tools.security.canvas_access import guard_component_access
    except Exception as _guard_exc:
        app.logger.warning("Canvas access guard unavailable: %s", _guard_exc)
        guard_component_access = None  # type: ignore[assignment]

    # Eager registration for whatever is already imported in _CANVAS_BLUEPRINTS.
    # In eager mode that is every enabled canvas; in lazy mode it is only the
    # empty-prefix (self-prefixing) canvases that cannot be matched by path.
    _CANVAS_ROUTES = _REGISTRY.get_url_prefixes()
    for _ck, _cbp in _CANVAS_BLUEPRINTS.items():
        try:
            prefix = _CANVAS_ROUTES.get(_ck, f"/{_ck}")

            # Attach RBAC + canvas-access guard (backward-compatible unless
            # ICDEV_ENFORCE_CANVAS_ACCESS is set) BEFORE registering the
            # blueprint on the app. Flask forbids adding before_request hooks to a
            # blueprint that has already been registered once (process-wide, not
            # per-app), and the hook persists on the blueprint across app instances
            # anyway — so attach it ONCE. Without this guard, a second import of
            # this module / a second create_app() re-ran the attach and raised on
            # every canvas, logged as a spurious "registration failed" warning even
            # though routes still worked (the hook was already on the blueprint).
            if guard_component_access and not getattr(_cbp, "_icdev_guard_attached", False):
                _comp_meta = _REGISTRY.get(_ck)
                if _comp_meta:
                    _cbp.before_request(
                        guard_component_access(_ck, _comp_meta.min_il)
                    )
                    _cbp._icdev_guard_attached = True

            # Register on THIS app only if not already present on it (idempotent
            # across re-entrant setup); a blueprint may be registered on several apps.
            if _cbp.name not in app.blueprints:
                if not _cbp.url_prefix:
                    app.register_blueprint(_cbp, url_prefix=prefix)
                else:
                    app.register_blueprint(_cbp)
                app.logger.info("Canvas %s registered at %s/", _ck.upper(), prefix)
        except Exception as exc:
            app.logger.warning("Canvas %s registration failed: %s", _ck.upper(), exc)

    if _LAZY_CANVASES_ENABLED:
        # cvx-net-02: defer heavy canvas blueprint module imports until the first
        # request to each prefixed canvas. Nav/context flags were populated at
        # module import; only the module import + route registration is deferred.
        _lazy_comps = [
            _c
            for _c in _REGISTRY.iter_enabled(kind="canvas")
            if (_CANVAS_ROUTES.get(_c.key) or "").strip()
        ]
        install_lazy_canvas_loader(
            app, _lazy_comps, _CANVAS_ROUTES, guard_component_access
        )
        app.logger.info(
            "Lazy canvas loading active — %d prefixed canvases deferred to first-hit import",
            len(_lazy_comps),
        )

    # ---- Simulation Chat Blueprint ----
    try:
        from tools.simulation.blueprint import create_simulation_blueprint
        _sim_bp = create_simulation_blueprint()
        if _sim_bp:
            app.register_blueprint(_sim_bp)
            app.logger.info("Simulation chat blueprint registered")
    except Exception as _exc:
        app.logger.warning("Simulation blueprint failed to register: %s", _exc)

    # ---- Infra IaC Generator Blueprint ----
    try:
        from tools.infra.blueprint import bp as _infra_iac_bp
        app.register_blueprint(_infra_iac_bp, url_prefix="/infra")
        app.logger.info("Infra IaC blueprint registered at /infra")
    except Exception as _exc:
        app.logger.warning("Infra IaC blueprint failed to register: %s", _exc)

    # ---- Migration Intelligence Engine Blueprint ----
    try:
        from tools.migration_intelligence.blueprint import create_migration_intel_blueprint
        _mi_bp = create_migration_intel_blueprint()
        app.register_blueprint(_mi_bp)
        app.logger.info("Migration Intelligence blueprint registered at /migration-intel")
    except Exception as _exc:
        app.logger.warning("Migration Intelligence blueprint failed to register: %s", _exc)

    # ---- Supply Chain Intelligence Blueprint ----
    try:
        from tools.supply_chain.blueprint import create_supply_chain_blueprint
        _sc_bp = create_supply_chain_blueprint()
        app.register_blueprint(_sc_bp)
        app.logger.info("Supply Chain blueprint registered at /supply_chain")
    except Exception as _exc:
        app.logger.warning("Supply Chain blueprint failed to register: %s", _exc)

    # ---- SIPA Software Integrity & Provenance Assessor Blueprint ----
    # Registered via the registry-driven canvas loop above (key "integrity",
    # empty url_prefix so the blueprint's explicit /integrity + /api/integrity
    # paths are not double-prefixed). No manual registration here.

    # ---- Strategos Blueprint ----
    if _HAS_STRATEGOS:
        try:
            from apps.strategos.blueprint import (
                create_strategos_blueprint,
                create_strategos_api_blueprint,
            )
            _sg_bp = create_strategos_blueprint()
            if _sg_bp:
                app.register_blueprint(_sg_bp, url_prefix="/strategos")
                app.logger.info("Strategos blueprint registered at /strategos")
            _sg_api_bp = create_strategos_api_blueprint()
            if _sg_api_bp:
                app.register_blueprint(_sg_api_bp, url_prefix="/api/strategos")
                app.logger.info("Strategos API blueprint registered at /api/strategos")

            # Canonical DAT aliases (issue #18) — the Diplomatic Activity Tracker
            # is implemented under Strategos; expose the issue's documented paths
            # (/dat, /api/dat/dti) without duplicating the engine/page logic.
            def _dat_page_alias():
                from flask import redirect, request as _rq
                qs = _rq.query_string.decode()
                return redirect("/strategos/dat" + (f"?{qs}" if qs else ""))

            def _dat_dti_api():
                from flask import jsonify, make_response, request as _rq
                from tools.strategos.dat import (
                    compute_dti, ensure_tables, get_dti_history, get_latest_dti,
                )
                ensure_tables()
                theater = _rq.args.get("theater", "global")
                score = get_latest_dti(theater) or compute_dti(theater)
                try:
                    limit = min(int(_rq.args.get("limit", 48)), 200)
                except (TypeError, ValueError):
                    limit = 48
                history = get_dti_history(theater, limit=limit)
                resp = make_response(jsonify({
                    "theater": theater,
                    "dti": score,
                    "history": history,
                    "total": len(history),
                }))
                resp.headers["X-Classification"] = "CUI"
                return resp

            app.add_url_rule("/dat", "dat_page_alias", _dat_page_alias)
            app.add_url_rule("/mcip", "mcip_page_alias", _dat_page_alias)
            app.add_url_rule("/api/dat/dti", "dat_dti_api", _dat_dti_api)
            app.logger.info("DAT aliases registered at /dat, /mcip, and /api/dat/dti")
        except Exception as _exc:
            app.logger.warning("Strategos blueprint failed to register: %s", _exc)
    else:
        app.logger.info("Strategos disabled (ICDEV_STRATEGOS_ENABLED=false)")

    # ---- GeoSIGINT Blueprint ----
    try:
        import importlib.util as _ilu
        import sys as _sys
        _geo_bp_path = BASE_DIR / "apps" / "geosigint" / "blueprint.py"
        if not _geo_bp_path.exists():
            raise FileNotFoundError(f"GeoSIGINT blueprint not found at {_geo_bp_path}")
        # Ensure apps and apps.geosigint are in sys.modules for intra-package imports.
        # __path__ must be set so that sub-module imports (e.g. from apps.geosigint.a2ad_mapper)
        # resolve correctly at request time.
        for _pkg, _pkg_path, _pkg_dir in [
            ("apps", BASE_DIR / "apps" / "__init__.py", BASE_DIR / "apps"),
            ("apps.geosigint", BASE_DIR / "apps" / "geosigint" / "__init__.py", BASE_DIR / "apps" / "geosigint"),
        ]:
            if _pkg not in _sys.modules:
                _spec = _ilu.spec_from_file_location(_pkg, str(_pkg_path))
                _mod = _ilu.module_from_spec(_spec)
                _mod.__path__ = [str(_pkg_dir)]
                _mod.__package__ = _pkg
                _sys.modules[_pkg] = _mod
                _spec.loader.exec_module(_mod)
        _spec = _ilu.spec_from_file_location("apps.geosigint.blueprint", str(_geo_bp_path))
        _geo_mod = _ilu.module_from_spec(_spec)
        _sys.modules["apps.geosigint.blueprint"] = _geo_mod
        _spec.loader.exec_module(_geo_mod)
        _geo_bp = _geo_mod.create_geosigint_blueprint()
        _geo_api_bp = _geo_mod.create_geosigint_api_blueprint()
        app.register_blueprint(_geo_bp)
        app.register_blueprint(_geo_api_bp)
        app.logger.info("GeoSIGINT blueprints registered at /geosigint and /api/geosigint")
    except Exception as _exc:
        import traceback as _tb
        app.logger.warning("GeoSIGINT blueprint failed to register: %s\n%s", _exc, _tb.format_exc())

    # ---- TA Patterns Blueprint ----
    try:
        from tools.trading.ta.blueprint import create_ta_blueprint
        _ta_bp = create_ta_blueprint()
        if _ta_bp:
            app.register_blueprint(_ta_bp)
            app.logger.info("TA Patterns blueprint registered at /api/ta/patterns")
    except Exception as _exc:
        app.logger.warning("TA Patterns blueprint failed to register: %s", _exc)

    # ---- App Module Blueprints (registry-driven child_app components) ----
    for _ak, _abp in _APP_BLUEPRINTS.items():
        try:
            app.register_blueprint(_abp)
            app.logger.info("App module %s registered", _ak)
        except Exception as _exc:
            app.logger.warning("App module %s registration failed: %s", _ak, _exc)

    # ---- Core Extension Blueprints (registry-driven, e.g. admin_console) ----
    for _ek, _ebp in _CORE_EXT_BLUEPRINTS.items():
        try:
            app.register_blueprint(_ebp)
            app.logger.info("Core extension %s registered", _ek)
        except Exception as _exc:
            app.logger.warning("Core extension %s registration failed: %s", _ek, _exc)

    # ---- HITL Workflow API Blueprint (always registered when importable; gated per-route) ----
    try:
        from tools.workflow_hitl.blueprint import create_wf_blueprint
        _wf_bp = create_wf_blueprint()
        app.register_blueprint(_wf_bp, url_prefix="/api/wf")
        app.logger.info("HITL Workflow API blueprint registered at /api/wf")
    except Exception as _exc:
        app.logger.warning("HITL Workflow API blueprint failed to register: %s", _exc)

    # ---- Autonomous Coder Blueprint ----
    try:
        try:
            from icdev.apps.autonomous_coder.blueprint import ac_bp as _ac_bp
        except ImportError:
            from apps.autonomous_coder.blueprint import ac_bp as _ac_bp
        app.register_blueprint(_ac_bp)
        app.logger.info("Autonomous Coder blueprint registered at /autonomous-coder")
    except Exception as _exc:
        app.logger.warning("Autonomous Coder blueprint failed to register: %s", _exc)

    # ---- System Graph (Unified Sigma.js graph) ----
    try:
        from tools.system_graph.blueprint import bp as _sysgraph_bp
        app.register_blueprint(_sysgraph_bp)
        app.logger.info("System Graph blueprint registered at /system-graph")
    except Exception as _exc:
        app.logger.warning("System Graph blueprint failed to register: %s", _exc)

    # ---- ZTA LAC Simulator Blueprint (irad-lac-06 / irad-lac-07) ----
    try:
        from tools.zta.blueprint import create_zta_blueprint
        _zta_bp = create_zta_blueprint()
        app.register_blueprint(_zta_bp, url_prefix="/zta")
        app.logger.info("ZTA LAC Simulator blueprint registered at /zta")
    except Exception as _exc:
        app.logger.warning("ZTA LAC Simulator blueprint failed to register: %s", _exc)

    # ---- GovLift Cloud Migration Tool ----
    try:
        from tools.govlift.blueprint import create_govlift_blueprint as _gv_factory
        from tools.govlift.db.init_db import init_govlift_db as _gv_init
        _gv_init()
        _gv_bp = _gv_factory()
        app.register_blueprint(_gv_bp)
        _CANVAS_FLAGS["govlift"] = True
        app.logger.info("GovLift blueprint registered at /govlift")
    except Exception as _exc:
        app.logger.warning("GovLift blueprint failed to register: %s", _exc)

    # ---- AI Observatory (cross-canvas AI decision traceability) ----
    try:
        from tools.ai_observatory.blueprint import bp as _ao_bp
        app.register_blueprint(_ao_bp)
        app.logger.info("AI Observatory blueprint registered at /ai-observatory")
    except Exception as _exc:
        app.logger.warning("AI Observatory blueprint failed to register: %s", _exc)

    try:
        from tools.ontology.blueprint import bp as _ont_bp
        app.register_blueprint(_ont_bp)
        app.logger.info("Ontology Explorer blueprint registered at /ontology")
    except Exception as _exc:
        app.logger.warning("Ontology Explorer blueprint failed to register: %s", _exc)

    # ---- Cache Savings Blueprint ----
    try:
        from tools.cache_savings.blueprint import bp as _cache_savings_bp
        app.register_blueprint(_cache_savings_bp)
        app.logger.info("Cache Savings blueprint registered at /cache-savings")
    except Exception as _exc:
        app.logger.warning("Cache Savings blueprint failed to register: %s", _exc)

    # ---- JISE Portal Blueprint ----
    try:
        from tools.intelligence.jise_portal import jise_bp as _jise_bp
        app.register_blueprint(_jise_bp)
        app.logger.info("JISE Portal blueprint registered at /api/v1/jise")
    except Exception as _exc:
        app.logger.warning("JISE Portal blueprint failed to register: %s", _exc)

    # ---- SaaS Tenant Admin Portal Blueprint ----
    try:
        from tools.saas.portal.app import portal_bp as _portal_bp
        app.register_blueprint(_portal_bp)
        app.logger.info("SaaS Portal blueprint registered at /portal")
    except Exception as _exc:
        app.logger.warning("SaaS Portal blueprint failed to register: %s", _exc)

    # ---- Enterprise SSO Blueprint (SAML 2.0 + OIDC) ----
    try:
        from tools.auth.blueprint import bp as _auth_bp
        app.register_blueprint(_auth_bp)
        app.logger.info("Enterprise SSO blueprint registered at /auth/saml")
    except Exception as _exc:
        app.logger.warning("Enterprise SSO blueprint failed to register: %s", _exc)

    # ---- Onboarding API Blueprint ----
    try:
        from tools.auth.blueprint import onboarding_bp as _onboarding_bp
        app.register_blueprint(_onboarding_bp)
        app.logger.info("Onboarding API blueprint registered at /api/onboarding")
    except Exception as _exc:
        app.logger.warning("Onboarding API blueprint failed to register: %s", _exc)

    # ---- Stripe Webhook Blueprint (no auth, signature-verified) ----
    try:
        from tools.admin.blueprint import create_stripe_webhook_blueprint as _stripe_wh_factory
        _stripe_wh_bp = _stripe_wh_factory()
        app.register_blueprint(_stripe_wh_bp)
        app.logger.info("Stripe webhook blueprint registered at /webhooks/stripe")
    except Exception as _exc:
        app.logger.warning("Stripe webhook blueprint failed to register: %s", _exc)

    # ---- Centralized Logs viewer ----
    # Registered by the component-registry loop above, from the `logs` entry in
    # args/component_registry.yaml. The hand-rolled registration that used to sit
    # here always lost the race to it and logged "The name 'logs' is already
    # registered for a different blueprint" on every boot — and, had it ever won,
    # it would have bypassed the registry's enablement and IL gating.

    # ---- Platform Updates (CHANGELOG.md viewer) ----
    try:
        from tools.dashboard.changelog import parse_changelog as _parse_changelog
        import os as _os

        @app.route("/updates")
        def updates_page():
            # nav-misc-01: record the current version as seen for logged-in users
            # (server-side, so the badge clears in a fresh cookieless session and
            # syncs across devices). Anonymous users rely on the cookie set by the
            # PATCH endpoint fired from this page's JS.
            try:
                _uid = _current_user_id()
                if _uid:
                    from tools.dashboard.brand import get_brand as _gb_upd
                    from tools.auth.onboarding import set_last_seen_version
                    _uv = _gb_upd().get("version", "")
                    if _uv:
                        set_last_seen_version(_uid, _uv)
            except Exception as _exc:
                app.logger.warning("seen-version record on /updates failed: %s", _exc)
            _changelog_path = _os.path.join(_os.path.dirname(__file__), "..", "..", "CHANGELOG.md")
            try:
                releases = _parse_changelog(_changelog_path)
            except Exception as _exc:
                app.logger.warning("Failed to parse CHANGELOG.md: %s", _exc)
                releases = []
                return render_template("updates/page.html", releases=releases, error=str(_exc))
            return render_template("updates/page.html", releases=releases, error=None)

        app.logger.info("Updates route registered at /updates")
    except Exception as _exc:
        app.logger.warning("Updates route failed to register: %s", _exc)

    # ---- ECR-CL-02: Mark current version as seen ----
    @app.route("/api/user/prefs/seen-version", methods=["PATCH"])
    def api_user_prefs_seen_version():
        try:
            from tools.dashboard.brand import get_brand as _get_brand_sv
            _brand_ver = _get_brand_sv().get("version", "")
        except Exception as _exc:
            app.logger.warning("seen-version PATCH failed: %s", _exc)
            return jsonify({"error": str(_exc)}), 500
        # nav-misc-01: persist per-user when logged in; keep the cookie for
        # anonymous visitors (and as a harmless redundant signal for users).
        _uid = _current_user_id()
        if _uid and _brand_ver:
            try:
                from tools.auth.onboarding import set_last_seen_version
                set_last_seen_version(_uid, _brand_ver)
            except Exception as _exc:
                app.logger.warning("seen-version persist failed: %s", _exc)
        resp = jsonify({"ok": True, "seen_version": _brand_ver})
        resp.set_cookie("icdev_seen_version", _brand_ver, max_age=31536000, samesite="Lax")
        return resp

    # ---- Convenience JSON routes that match the spec ----

    @app.route("/api/alerts", methods=["GET"])
    def api_alerts_shortcut():
        """Shortcut: GET /api/alerts -> delegates to metrics alerts."""
        conn = _get_db()
        try:
            rows = conn.execute("SELECT * FROM alerts ORDER BY created_at DESC LIMIT 50").fetchall()
            return jsonify({"alerts": [dict(r) for r in rows], "total": len(rows)})
        finally:
            conn.close()

    @app.route("/api/notifications", methods=["GET"])
    def api_notifications():
        """Return current notification-worthy items (firing alerts, overdue POAMs)."""
        conn = _get_db()
        try:
            notifications = []
            firing = conn.execute("SELECT COUNT(*) as cnt FROM alerts WHERE status = 'firing'").fetchone()["cnt"]
            if firing > 0:
                notifications.append(
                    {
                        "type": "error",
                        "message": f"{firing} alert{'s' if firing > 1 else ''} currently firing",
                        "link": "/monitoring",
                    }
                )
            open_poam = conn.execute("SELECT COUNT(*) as cnt FROM poam_items WHERE status = 'open'").fetchone()["cnt"]
            if open_poam > 5:
                notifications.append(
                    {
                        "type": "warning",
                        "message": f"{open_poam} open POA&M items need attention",
                        "link": "/projects",
                    }
                )
            inactive = conn.execute("SELECT COUNT(*) as cnt FROM agents WHERE status != 'active'").fetchone()["cnt"]
            if inactive > 0:
                notifications.append(
                    {
                        "type": "info",
                        "message": f"{inactive} agent{'s' if inactive > 1 else ''} inactive",
                        "link": "/agents",
                    }
                )
            return jsonify({"notifications": notifications})
        finally:
            conn.close()

    @app.route("/api/dashboard/autonomous-feed", methods=["GET"])
    def api_dashboard_autonomous_feed():
        """GET /api/dashboard/autonomous-feed — Recent autonomous agent activity."""
        limit = min(max(int(flask_request.args.get("limit", 20)), 1), 100)
        conn = _get_db()
        try:
            feed = []

            def _tbl_exists(c, t):
                try:
                    c.execute(f"SELECT 1 FROM {t} LIMIT 1")  # nosec B608
                    return True
                except Exception:
                    return False

            tables_to_check = [
                ("canvas_remediation_proposals", "id", "created_at", "remediation", "canvas_type"),
                ("oracle_convergence_events", "id", "created_at", "convergence", "convergence_type"),
                ("oracle_predictions", "id", "created_at", "prediction", "lens_id"),
            ]
            for table, id_col, ts_col, event_type, label_col in tables_to_check:
                if not _tbl_exists(conn, table):
                    continue
                try:
                    rows = conn.execute(
                        f"SELECT {id_col}, {ts_col}, {label_col} FROM {table} ORDER BY {ts_col} DESC LIMIT %s",  # nosec B608
                        (limit,),
                    ).fetchall()
                    for r in rows:
                        feed.append({"type": event_type, "id": r[0], "ts": r[1], "label": r[2]})
                except Exception:
                    pass
            feed.sort(key=lambda x: x.get("ts") or "", reverse=True)
            return jsonify({"feed": feed[:limit], "total": len(feed)})
        except Exception as exc:
            return jsonify({"error": str(exc)}), 500
        finally:
            conn.close()

    @app.route("/api/charts/overview", methods=["GET"])
    def api_charts_overview():
        """Aggregate chart data for the home dashboard."""

        conn = _get_db()
        try:
            # ----------------------------------------------------------------
            # 1. Task Board Status (donut) — replaces empty projects table
            # ----------------------------------------------------------------
            task_statuses = conn.execute("SELECT status, COUNT(*) as cnt FROM kanban_tasks GROUP BY status").fetchall()

            # ----------------------------------------------------------------
            # 2. Activity Trend: last 7 days (line chart) — from audit_trail
            # ----------------------------------------------------------------
            _is_pg = getattr(conn, "_backend", "sqlite") == "postgresql"
            if _is_pg:
                activity_trend = conn.execute(
                    "SELECT DATE(created_at) as day, COUNT(*) as cnt "
                    "FROM audit_trail WHERE created_at >= NOW() - INTERVAL '7 days' "
                    "GROUP BY DATE(created_at) ORDER BY day"
                ).fetchall()
            else:
                activity_trend = conn.execute(
                    "SELECT DATE(created_at) as day, COUNT(*) as cnt "
                    "FROM audit_trail WHERE created_at >= DATE('now', '-7 days') "
                    "GROUP BY DATE(created_at) ORDER BY day"
                ).fetchall()

            # ----------------------------------------------------------------
            # 3. Compliance Posture — aggregate across canvas assessment DBs
            # OPT-12: cache canvas DB scans for 45 s (canvas state is slow-moving).
            # ----------------------------------------------------------------
            _now = time.monotonic()
            _cached_entry = _CANVAS_COMPLIANCE_CACHE.get("entry")
            if _cached_entry and (_now - _cached_entry["ts"]) < _CANVAS_CACHE_TTL:
                canvas_compliance = _cached_entry["canvas_compliance"]
                overall_score = _cached_entry["overall_score"]
            else:
                # cnr-cc-02: canvas posture aggregation lives in ONE module
                # (tools/canvas_compliance/posture.py), consumed by this index
                # route and by canvas_aggregator.get_canvas_compliance_summary().
                from tools.canvas_compliance.posture import compute_canvas_posture
                canvas_compliance, overall_score = compute_canvas_posture(conn)
                _CANVAS_COMPLIANCE_CACHE["entry"] = {
                    "ts": _now,
                    "canvas_compliance": canvas_compliance,
                    "overall_score": overall_score,
                }

            # ----------------------------------------------------------------
            # 4. Oracle compliance-risk predictions (forward-looking posture)
            # Query oracle_predictions for lens_id='oracle-compliance-risk' with
            # outcome='pending', grouped into CAT1/CAT2/CAT3 by severity.
            # ----------------------------------------------------------------
            oracle_cat1 = oracle_cat2 = oracle_cat3 = 0
            try:
                oracle_rows = conn.execute(
                    "SELECT severity, COUNT(*) as cnt FROM oracle_predictions "
                    "WHERE lens_id = 'oracle-compliance-risk' AND outcome = 'pending' "
                    "GROUP BY severity"
                ).fetchall()
                for _orow in oracle_rows:
                    _sev = ((_orow["severity"] or "")).lower()
                    _cnt = int(_orow["cnt"] or 0)
                    if _sev in ("critical", "high"):
                        oracle_cat1 += _cnt
                    elif _sev == "medium":
                        oracle_cat2 += _cnt
                    else:
                        oracle_cat3 += _cnt
            except Exception:
                pass
            oracle_total = oracle_cat1 + oracle_cat2 + oracle_cat3
            canvases_with_oracle = canvas_compliance + [
                {
                    "name": "Oracle",
                    "score": oracle_total,
                    "open_findings": oracle_total,
                    "closed_findings": 0,
                    "cat1": oracle_cat1,
                    "cat2": oracle_cat2,
                    "cat3": oracle_cat3,
                    "is_oracle": True,
                }
            ]

            # ----------------------------------------------------------------
            # 5. Agent health (gauge: % active) — unchanged
            # ----------------------------------------------------------------
            total_agents = conn.execute("SELECT COUNT(*) as cnt FROM agents").fetchone()["cnt"]
            active_agents = conn.execute("SELECT COUNT(*) as cnt FROM agents WHERE status = 'active'").fetchone()["cnt"]

            # ----------------------------------------------------------------
            # 5b. POA&M / STIG open-closed counts for the home bar chart
            # live.js renders #chart-compliance from d.compliance.poam / d.compliance.stig
            # ----------------------------------------------------------------
            poam_open = 0
            poam_closed = 0
            try:
                poam_row = conn.execute(
                    "SELECT "
                    "SUM(CASE WHEN status IN ('open', 'in_progress') THEN 1 ELSE 0 END) as open_cnt, "
                    "SUM(CASE WHEN status IN ('completed', 'accepted_risk') THEN 1 ELSE 0 END) as closed_cnt "
                    "FROM poam_items"
                ).fetchone()
                poam_open = int(poam_row["open_cnt"] or 0)
                poam_closed = int(poam_row["closed_cnt"] or 0)
            except Exception:
                pass

            stig_open = 0
            stig_closed = 0
            try:
                stig_row = conn.execute(
                    "SELECT "
                    "SUM(CASE WHEN status = 'Open' THEN 1 ELSE 0 END) as open_cnt, "
                    "SUM(CASE WHEN status IN ('NotAFinding', 'Not_Applicable') THEN 1 ELSE 0 END) as closed_cnt "
                    "FROM stig_findings"
                ).fetchone()
                stig_open = int(stig_row["open_cnt"] or 0)
                stig_closed = int(stig_row["closed_cnt"] or 0)
            except Exception:
                pass

            return jsonify(
                {
                    "task_statuses": [dict(r) for r in task_statuses],
                    "activity_trend": [dict(r) for r in activity_trend],
                    "compliance": {
                        "canvases": canvases_with_oracle,
                        "overall_score": overall_score,
                        "poam": {"open": poam_open, "closed": poam_closed},
                        "stig": {"open": stig_open, "closed": stig_closed},
                    },
                    "agent_health": {
                        "total": total_agents,
                        "active": active_agents,
                        "ratio": active_agents / total_agents if total_agents > 0 else 1.0,
                    },
                }
            )
        finally:
            conn.close()

    @app.route("/api/charts/compliance-trend", methods=["GET"])
    def api_charts_compliance_trend():
        """Return 30-day score history per canvas for sparkline overlays."""

        # OPT-12: cache canvas DB scans for 45 s (canvas state is slow-moving).
        _now = time.monotonic()
        _cached_entry = _CANVAS_TREND_CACHE.get("entry")
        if _cached_entry and (_now - _cached_entry["ts"]) < _CANVAS_CACHE_TTL:
            return jsonify({"canvases": _cached_entry["data"]})

        _TREND_CANVASES = [
            ("Security", "tools.security_canvas.db.init_db", "sc_assessments", "risk_score", "ran_at", "inverted"),
            ("Network", None, None, None, None, "skip"),
            ("Pipeline", None, None, None, None, "skip"),
            ("Infra", "tools.infra_canvas.db.init_db", "idc_assessments", "score", "created_at", "direct"),
            ("Data", "tools.data_canvas.db.init_db", "dd_assessments", "score", "created_at", "direct"),
            ("Boundary", "tools.boundary_canvas.db.init_db", "bd_assessments", "score", "created_at", "direct"),
            ("Observability", "tools.observability_canvas.db.init_db", "od_assessments", "score", "created_at", "direct"),
            ("Agentic AI", "tools.agentic_ai_canvas.db.init_db", "aadc_assessments", "score", "created_at", "direct"),
            ("AI/ML", "tools.aiml_canvas.db.init_db", "aiml_assessments", "score", "created_at", "direct"),
            ("QDC", "tools.qdc_canvas.db.init_db", "qdc_assessments", "score", "created_at", "direct"),
            ("Migration", None, None, None, None, "skip"),
            ("GovLift", None, None, None, None, "skip"),
        ]

        def _trend_canvas_conn(module_name):
            if not module_name:
                return None
            try:
                mod = importlib.import_module(module_name)
                cconn = mod.get_connection()
                try:
                    cconn.set_security_context(None)  # rls-bypass: canvas tables lack tenant_id/classification; use module-level canvas connection with RLS disabled
                except Exception:
                    pass
                return cconn
            except Exception:
                return None

        _cutoff = (datetime.now(timezone.utc) - timedelta(days=30)).isoformat()

        results = []
        for canvas_name, module_name, table, score_col, ts_col, mode in _TREND_CANVASES:
            if mode == "skip" or not module_name:
                results.append({"name": canvas_name, "scores": [], "direction": "flat", "delta": 0.0})
                continue
            cconn = _trend_canvas_conn(module_name)
            if not cconn:
                results.append({"name": canvas_name, "scores": [], "direction": "flat", "delta": 0.0})
                continue
            try:
                ph = sql_placeholder(cconn)
                rows = cconn.execute(
                    f"SELECT {score_col} as raw_score, DATE({ts_col}) as day "  # nosec B608 -- whitelist from _TREND_CANVASES
                    f"FROM {table} "  # nosec B608
                    f"WHERE {ts_col} >= {ph} "  # nosec B608
                    f"ORDER BY {ts_col} DESC LIMIT 30",  # nosec B608
                    (_cutoff,),
                ).fetchall()
                scores = []
                for r in rows:
                    raw = float(r["raw_score"] or 0)
                    s = round(max(0.0, 100.0 - raw), 1) if mode == "inverted" else round(raw, 1)
                    scores.append({"score": s, "date": r["day"]})
                direction = "flat"
                delta = 0.0
                if len(scores) >= 2:
                    latest = scores[0]["score"]
                    oldest = scores[-1]["score"]
                    delta = round(latest - oldest, 1)
                    if delta >= 2.0:
                        direction = "up"
                    elif delta <= -2.0:
                        direction = "down"
                results.append({"name": canvas_name, "scores": scores, "direction": direction, "delta": delta})
            except Exception:
                results.append({"name": canvas_name, "scores": [], "direction": "flat", "delta": 0.0})
            finally:
                try:
                    cconn.close()
                except Exception:
                    pass

        # ZIG trend — daily average of pillar scores from zig_maturity_scores
        try:
            zconn = _trend_canvas_conn("tools.security_canvas.db.init_db")
            if zconn:
                try:
                    ph = sql_placeholder(zconn)
                    rows = zconn.execute(
                        "SELECT AVG(score) * 100 as raw_score, DATE(assessment_run_at) as day "
                        "FROM zig_maturity_scores "
                        "WHERE assessment_run_at >= %s "
                        "GROUP BY DATE(assessment_run_at) ORDER BY day DESC LIMIT 30",
                        (_cutoff,),
                    ).fetchall()
                    scores = [{"score": round(float(r["raw_score"] or 0), 1), "date": r["day"]} for r in rows]
                    direction, delta = "flat", 0.0
                    if len(scores) >= 2:
                        delta = round(scores[0]["score"] - scores[-1]["score"], 1)
                        direction = "up" if delta >= 2.0 else "down" if delta <= -2.0 else "flat"
                    results.append({"name": "Zero Trust", "scores": scores, "direction": direction, "delta": delta})
                finally:
                    zconn.close()
            else:
                results.append({"name": "Zero Trust", "scores": [], "direction": "flat", "delta": 0.0})
        except Exception:
            results.append({"name": "Zero Trust", "scores": [], "direction": "flat", "delta": 0.0})

        # AI-ify trend — per-scan average composite_score grouped by day
        try:
            from tools.aiify.db.init_db import get_connection as _aiify_trend_cn
            aconn = _aiify_trend_cn()
            try:
                aconn.set_security_context(None)  # rls-bypass: aiify canvas tables lack tenant_id/classification; use canvas connection with RLS disabled
            except Exception:
                pass
            try:
                ph = sql_placeholder(aconn)
                rows = aconn.execute(
                    "SELECT AVG(s.composite_score) * 100 as raw_score, DATE(sc.created_at) as day "
                    "FROM aiify_scores s "
                    "JOIN aiify_opportunities o ON o.opportunity_id = s.opportunity_id "
                    "JOIN aiify_scans sc ON sc.scan_id = o.scan_id "
                    f"WHERE sc.created_at >= {ph} "
                    "GROUP BY DATE(sc.created_at) ORDER BY day DESC LIMIT 30",
                    (_cutoff,),
                ).fetchall()
                scores = [{"score": round(float(r["raw_score"] or 0), 1), "date": r["day"]} for r in rows]
                direction, delta = "flat", 0.0
                if len(scores) >= 2:
                    delta = round(scores[0]["score"] - scores[-1]["score"], 1)
                    direction = "up" if delta >= 2.0 else "down" if delta <= -2.0 else "flat"
                results.append({"name": "AI-ify", "scores": scores, "direction": direction, "delta": delta})
            finally:
                aconn.close()
        except Exception:
            results.append({"name": "AI-ify", "scores": [], "direction": "flat", "delta": 0.0})

        _CANVAS_TREND_CACHE["entry"] = {"ts": _now, "data": results}
        return jsonify({"canvases": results})

    @app.route("/api/charts/project/<project_id>", methods=["GET"])
    def api_charts_project(project_id):
        """Chart data for a specific project detail page."""
        conn = _get_db()
        try:
            # STIG by severity (donut)
            stig_sev = conn.execute(
                "SELECT severity, status, COUNT(*) as cnt "
                "FROM stig_findings WHERE project_id = %s "
                "GROUP BY severity, status",
                (project_id,),
            ).fetchall()

            # POAM by severity (bar)
            poam_sev = conn.execute(
                "SELECT severity, status, COUNT(*) as cnt "
                "FROM poam_items WHERE project_id = %s "
                "GROUP BY severity, status",
                (project_id,),
            ).fetchall()

            # Deployment history (line — status over time)
            deploys = conn.execute(
                "SELECT DATE(created_at) as day, status, COUNT(*) as cnt "
                "FROM deployments WHERE project_id = %s "
                "GROUP BY DATE(created_at), status ORDER BY day",
                (project_id,),
            ).fetchall()

            # Alert trend for project
            alerts = conn.execute(
                "SELECT DATE(created_at) as day, severity, COUNT(*) as cnt "
                "FROM alerts WHERE project_id = %s "
                "GROUP BY DATE(created_at), severity ORDER BY day",
                (project_id,),
            ).fetchall()

            return jsonify(
                {
                    "stig_by_severity": [dict(r) for r in stig_sev],
                    "poam_by_severity": [dict(r) for r in poam_sev],
                    "deployment_history": [dict(r) for r in deploys],
                    "alert_trend": [dict(r) for r in alerts],
                }
            )
        finally:
            conn.close()

    # ---- HTML page routes ----

    @app.route("/")
    def index():
        """Dashboard home page with Kanban board."""
        conn = _get_db()
        try:
            # All projects for Kanban board
            projects = conn.execute(
                "SELECT id, name, type, status, classification FROM projects ORDER BY updated_at DESC, created_at DESC"
            ).fetchall()
            projects = [dict(r) for r in projects]

            # Agent counts (stat bar)
            total_agents = conn.execute("SELECT COUNT(*) as cnt FROM agents").fetchone()["cnt"]
            active_agents = conn.execute("SELECT COUNT(*) as cnt FROM agents WHERE status = 'active'").fetchone()["cnt"]
            inactive_agents = total_agents - active_agents

            # Recent audit entries (for existing audit trail section)
            recent_audit = conn.execute("SELECT * FROM audit_trail ORDER BY created_at DESC LIMIT 10").fetchall()

            # --- Recent Activity & Findings: audit_trail + canvas CAT1 findings ---
            # Canvas findings now flow through _aggregate_findings() helper
            # below, so no local sqlite3 import is needed here anymore.
            _audit_rows = conn.execute(
                "SELECT event_type, actor, action, project_id, created_at "
                "FROM audit_trail ORDER BY created_at DESC LIMIT 10"
            ).fetchall()
            _activity = []
            for _e in _audit_rows:
                _e = dict(_e)
                _activity.append(
                    {
                        "event_type": _e.get("event_type") or "AUDIT",
                        "source": "System",
                        "details": _e.get("action") or "",
                        "severity": "info",
                        "created_at": _e.get("created_at") or "",
                    }
                )

            # Canvas findings (POA&M) — single source of truth via aggregator.
            # Replaces the prior inline triple-counting loops; the /poam page
            # uses the same helper so the index counter and the list match.
            try:
                _all_findings = _aggregate_findings(
                    get_db_conn=_get_db, include_remediated=False
                )
            except Exception:
                _all_findings = []

            _excluded = {"declined", "accepted_risk", "remediated"}
            cat1_count = 0
            for _f in _all_findings:
                if _f.get("severity") == "CAT1" and _f.get("decision") not in _excluded:
                    cat1_count += 1
                    _activity.append(
                        {
                            "event_type": "Canvas Finding",
                            "source": _f.get("canvas_label") or "",
                            "details": (
                                f"{_f.get('title', '')}: {_f.get('affected_entity', '')}"
                                if _f.get("affected_entity")
                                else _f.get("title", "")
                            ),
                            "severity": "CAT1",
                            "created_at": _f.get("discovered_at") or "",
                        }
                    )

            # Sort merged activity by created_at DESC, limit 10
            _activity.sort(key=lambda x: str(x.get("created_at") or ""), reverse=True)
            recent_activity = _activity[:10]

            # Firing alert count = open CAT1 canvas findings (same exclusion as open_poam)
            firing_alerts = cat1_count

            # Open POAM count = open canvas findings (excludes declined/accepted_risk/remediated)
            open_poam = sum(1 for _f in _all_findings if _f.get("decision") not in _excluded)

            # Board-health alerts (kax-stall-01) — rendered as a banner directly
            # above the Task Board. Board throughput going to zero for four days
            # was previously visible nowhere a human actually looks.
            try:
                board_alerts = [
                    dict(r)
                    for r in conn.execute(
                        "SELECT id, severity, title, description, created_at FROM alerts "
                        "WHERE status = 'firing' AND source LIKE %s "
                        "ORDER BY created_at DESC LIMIT 5",
                        ("board_throughput:%",),
                    ).fetchall()
                ]
            except Exception:
                board_alerts = []

            # Group projects by status for Kanban columns
            kanban_columns = {
                "planning": [],
                "active": [],
                "completed": [],
                "inactive": [],
            }
            for p in projects:
                status = p.get("status", "inactive")
                if status in kanban_columns:
                    kanban_columns[status].append(p)
                else:
                    kanban_columns["inactive"].append(p)

            return render_template(
                "index.html",
                projects=projects,
                kanban_columns=kanban_columns,
                total_projects=len(projects),
                total_agents=total_agents,
                active_agents=active_agents,
                inactive_agents=inactive_agents,
                recent_audit=[dict(r) for r in recent_audit],
                recent_activity=recent_activity,
                firing_alerts=firing_alerts,
                board_alerts=board_alerts,
                open_poam=open_poam,
            )
        finally:
            conn.close()

    @app.route("/kanban")
    def kanban_page():
        """Task Board — Kanban view for scheduled and planned work."""
        return render_template("kanban.html")

    @app.route("/poam")
    def poam_page():
        """POA&M — Canvas findings approval workflow.

        Aggregates findings from all 7 canvas DBs (security, infra, observability,
        boundary, data, network, pipeline) and lets a reviewer approve, decline,
        accept risk, or mark remediated. Approval state lives in finding_approvals.
        """
        return render_template("poam/list.html")

    @app.route("/projects")
    def projects_list():
        """Project listing page."""
        conn = _get_db()
        try:
            projects = conn.execute(
                "SELECT id, name, type, status, classification, created_at FROM projects ORDER BY created_at DESC"
            ).fetchall()
            return render_template("projects/list.html", projects=[dict(r) for r in projects])
        finally:
            conn.close()

    @app.route("/projects/<project_id>")
    def project_detail(project_id):
        """Project detail page with tabs."""
        conn = _get_db()
        try:
            # Project info
            project = conn.execute("SELECT * FROM projects WHERE id = %s", (project_id,)).fetchone()
            if not project:
                return render_template("404.html", message="Project not found"), 404
            project = dict(project)

            # SSP documents
            ssps = conn.execute(
                "SELECT * FROM ssp_documents WHERE project_id = %s ORDER BY created_at DESC",
                (project_id,),
            ).fetchall()

            # POAM items
            poams = conn.execute(
                "SELECT * FROM poam_items WHERE project_id = %s ORDER BY severity, created_at DESC",
                (project_id,),
            ).fetchall()

            # STIG findings
            stigs = conn.execute(
                "SELECT * FROM stig_findings WHERE project_id = %s ORDER BY severity, created_at DESC",
                (project_id,),
            ).fetchall()

            # SBOM records
            sboms = conn.execute(
                "SELECT * FROM sbom_records WHERE project_id = %s ORDER BY generated_at DESC",
                (project_id,),
            ).fetchall()

            # Deployments
            deployments = conn.execute(
                "SELECT * FROM deployments WHERE project_id = %s ORDER BY created_at DESC",
                (project_id,),
            ).fetchall()

            # Audit trail
            audit_entries = conn.execute(
                "SELECT * FROM audit_trail WHERE project_id = %s ORDER BY created_at DESC LIMIT 50",
                (project_id,),
            ).fetchall()

            # Alerts
            alerts = conn.execute(
                "SELECT * FROM alerts WHERE project_id = %s ORDER BY created_at DESC LIMIT 20",
                (project_id,),
            ).fetchall()

            # Summaries
            poam_open = sum(1 for p in poams if dict(p)["status"] == "open")
            stig_open = sum(1 for s in stigs if dict(s)["status"] == "Open")

            stig_by_severity = {}
            for s in stigs:
                sd = dict(s)
                sev = sd.get("severity", "unknown")
                if sev not in stig_by_severity:
                    stig_by_severity[sev] = {"open": 0, "closed": 0}
                if sd["status"] == "Open":
                    stig_by_severity[sev]["open"] += 1
                else:
                    stig_by_severity[sev]["closed"] += 1

            return render_template(
                "projects/detail.html",
                project=project,
                ssps=[dict(r) for r in ssps],
                poams=[dict(r) for r in poams],
                poam_open=poam_open,
                stigs=[dict(r) for r in stigs],
                stig_open=stig_open,
                stig_by_severity=stig_by_severity,
                sboms=[dict(r) for r in sboms],
                deployments=[dict(r) for r in deployments],
                audit_entries=[dict(r) for r in audit_entries],
                alerts=[dict(r) for r in alerts],
            )
        finally:
            conn.close()

    @app.route("/agents")
    def agents_list():
        """Agent status page."""
        agents: list = []
        conn = None
        try:
            conn = _get_db()
            rows = conn.execute("SELECT * FROM agents ORDER BY name").fetchall()
            agent_ids = [r["id"] for r in rows]
            task_counts: dict = {}
            if agent_ids:
                placeholders = ",".join("?" * len(agent_ids))
                tc_rows = conn.execute(
                    f"SELECT target_agent_id, COUNT(*) as cnt FROM a2a_tasks "  # nosec B608 — placeholders contains only "?" bind params, no user input
                    f"WHERE target_agent_id IN ({placeholders}) "
                    f"AND status IN ('submitted', 'working') "
                    f"GROUP BY target_agent_id",
                    agent_ids,
                ).fetchall()
                task_counts = {r["target_agent_id"]: r["cnt"] for r in tc_rows}
            for r in rows:
                agent = dict(r)
                agent["active_task_count"] = task_counts.get(agent["id"], 0)
                agents.append(agent)
        except Exception:
            agents = []
        finally:
            if conn is not None:
                conn.close()

        active = sum(1 for a in agents if a.get("status") == "active")
        inactive = len(agents) - active

        return render_template(
            "agents/list.html",
            agents=agents,
            active_count=active,
            inactive_count=inactive,
        )

    @app.route("/api/core/iqe-query", methods=["POST"])
    def core_iqe_query():
        """Natural-language IQE query against agents and projects collections."""
        import logging as _log
        import tools.iqe.adapters.core_agents  # noqa: F401 — registers agents.* + projects.* collections
        from tools.iqe.nl_to_iqe import nl_to_iqe
        from tools.iqe.parser import Parser
        from tools.iqe.executor import execute_query

        data = flask_request.get_json(silent=True) or {}
        question = (data.get("question") or "").strip()
        if not question:
            return jsonify({"error": "question is required"}), 400

        collections = ["agents.registry", "projects.list"]
        iqe_str = ""
        try:
            result = nl_to_iqe(question, collections)
            iqe_str = result.get("iqe", "")
            explanation = result.get("explanation", "")
            ast = Parser().parse(iqe_str)
            conn = _get_db()
            try:
                rows = execute_query(ast, conn)
            finally:
                conn.close()
            return jsonify({"ok": True, "iqe": iqe_str, "explanation": explanation,
                            "results": rows, "row_count": len(rows)})
        except Exception as exc:
            _log.getLogger(__name__).warning("core IQE error: %s", exc)
            return jsonify({"error": str(exc), "iqe": iqe_str}), 500

    @app.route("/api/iqe/dispatch", methods=["POST"])
    def iqe_dispatch():
        """Canvas-aware IQE dispatcher — routes question to correct adapter by canvas name."""
        import logging as _dlog
        from tools.iqe.nl_to_iqe import nl_to_iqe
        from tools.iqe.parser import parse as _iqe_parse, IQESyntaxError as _IQESyntaxError
        from tools.iqe.executor import execute_query

        # cvx-nav-01: use the module-level cache built once at import (registry
        # is load-once) instead of rebuilding the map on every request.
        _CANVAS_MAP = _IQE_CANVAS_MAP

        data = flask_request.get_json(silent=True) or {}
        question = (data.get("question") or "").strip()
        canvas = (data.get("canvas") or "").strip().lower()
        if not question:
            return jsonify({"error": "question is required"}), 400

        if canvas not in _CANVAS_MAP:
            return jsonify({"error": f"unknown canvas '{canvas}'. Valid: {sorted(_CANVAS_MAP)}"}), 400

        adapter_module, collections = _CANVAS_MAP[canvas]
        iqe_str = ""
        try:
            import importlib
            importlib.import_module(adapter_module)
            result = nl_to_iqe(question, collections)
            iqe_str = result.get("iqe", "")
            explanation = result.get("explanation", "")
            try:
                ast = _iqe_parse(iqe_str)
                rows = execute_query(ast, conn=None)
            except _IQESyntaxError:
                rows = []
            return jsonify({"ok": True, "canvas": canvas, "iqe": iqe_str,
                            "explanation": explanation, "results": rows, "row_count": len(rows)})
        except Exception as exc:
            _dlog.getLogger(__name__).warning("IQE dispatch error [%s]: %s", canvas, exc)
            return jsonify({"error": str(exc), "canvas": canvas, "iqe": iqe_str}), 500

    @app.route("/monitoring")
    def monitoring_overview():
        """Monitoring overview page."""
        conn = _get_db()
        try:
            # All firing alerts (uncapped — operator must see every one)
            firing_alerts = conn.execute(
                "SELECT * FROM alerts WHERE status = 'firing' "
                "ORDER BY "
                "CASE severity "
                "  WHEN 'critical' THEN 0 "
                "  WHEN 'high' THEN 1 "
                "  WHEN 'medium' THEN 2 "
                "  WHEN 'low' THEN 3 "
                "  ELSE 4 END, "
                "created_at DESC"
            ).fetchall()

            # Recent alerts across all statuses (history view, capped)
            alerts = conn.execute("SELECT * FROM alerts ORDER BY created_at DESC LIMIT 50").fetchall()

            # Self-healing events
            healing_events = conn.execute(
                "SELECT she.*, kp.description as pattern_description "
                "FROM self_healing_events she "
                "LEFT JOIN knowledge_patterns kp ON she.pattern_id = kp.id "
                "ORDER BY she.created_at DESC LIMIT 20"
            ).fetchall()

            # Health stats
            firing = len(firing_alerts)
            resolved = conn.execute("SELECT COUNT(*) as cnt FROM alerts WHERE status = 'resolved'").fetchone()["cnt"]
            unresolved_failures = conn.execute("SELECT COUNT(*) as cnt FROM failure_log WHERE resolved = 0").fetchone()[
                "cnt"
            ]

            health = "healthy"
            if firing > 0 or unresolved_failures > 5:
                health = "degraded"
            if firing > 5:
                health = "critical"

            return render_template(
                "monitoring/overview.html",
                firing_alerts=[dict(r) for r in firing_alerts],
                alerts=[dict(r) for r in alerts],
                healing_events=[dict(r) for r in healing_events],
                firing_count=firing,
                resolved_count=resolved,
                unresolved_failures=unresolved_failures,
                health_status=health,
            )
        except Exception as exc:  # noqa: BLE001
            import logging as _mon_log
            _mon_log.getLogger(__name__).error("monitoring_overview DB error: %s", exc)
            return render_template(
                "monitoring/overview.html",
                firing_alerts=[],
                alerts=[],
                healing_events=[],
                firing_count=0,
                resolved_count=0,
                unresolved_failures=0,
                health_status="unknown",
            )
        finally:
            conn.close()

    # ---- Events & NLQ page routes ----

    @app.route("/events")
    def events_page():
        """Real-time event timeline page (SSE-powered)."""
        conn = _get_db()
        try:
            recent_events = conn.execute("SELECT * FROM hook_events ORDER BY created_at DESC LIMIT 50").fetchall()
            return render_template(
                "events/timeline.html",
                recent_events=[dict(r) for r in recent_events],
                degraded=False,
            )
        except Exception as exc:
            # nav-misc-02: an empty timeline on read failure looks like "no events".
            get_logger("icdev.dashboard").warning("events_page: hook_events read failed: %s", exc)
            return render_template("events/timeline.html", recent_events=[], degraded=True)
        finally:
            conn.close()

    @app.route("/oracle")
    def oracle_page():
        """Oracle — predictive intelligence dashboard."""
        return render_template("oracle.html")

    @app.route("/activity")
    def activity_page():
        """Activity feed — merged audit + hook events with real-time updates."""
        return render_template("activity.html")

    @app.route("/usage")
    def usage_page():
        """Usage tracking + cost dashboard."""
        return render_template("usage.html")

    @app.route("/il5")
    def il5_page():
        """IL5 data ingestion — parsed records and SLA compliance."""
        return render_template("il5/page.html")

    @app.route("/wizard")
    def wizard_page():
        """Getting Started wizard — guides new users to the right workflow."""
        import yaml as _yaml
        uc_meta = {}
        try:
            uc_path = BASE_DIR / "args" / "use_cases.yaml"
            if uc_path.exists():
                with open(uc_path, "r", encoding="utf-8") as _fh:
                    data = _yaml.safe_load(_fh) or {}
                for uc in data.get("use_cases", []):
                    uc_meta[uc["id"]] = {
                        "label": uc.get("label", ""),
                        "description": uc.get("description", "").strip(),
                        "canvas_wiring": uc.get("canvas_wiring", []),
                        "req_count": len(uc.get("template_requirements", [])),
                        "fast_track": uc.get("fast_track", False),
                        "badge": uc.get("badge", ""),
                        "category": uc.get("category", ""),
                    }
        except Exception as _e:
            app.logger.warning("wizard_page: failed to load use_cases.yaml: %s", _e)
        return render_template("wizard.html", use_case_meta=uc_meta)

    @app.route("/chat")
    def chat_new():
        """Start a new requirements chat — wizard params set context."""
        goal = flask_request.args.get("goal", "")
        role = flask_request.args.get("role", "")
        classification = flask_request.args.get("classification", "")
        frameworks = flask_request.args.get("frameworks", "")
        custom_role_name = flask_request.args.get("custom_role_name", "")
        custom_role_desc = flask_request.args.get("custom_role_desc", "")
        use_case_id = flask_request.args.get("use_case", "")
        skip_fast_track = flask_request.args.get("skip_fast_track", "") == "1"
        from_wizard = flask_request.args.get("from_wizard", "") == "1"
        # ?canvas= deep link — pre-selects canvas mode (forwarded from /simulate/chat redirect)
        canvas = flask_request.args.get("canvas", "") or flask_request.args.get("canvas_type", "")
        _allowed = {"cam", "ndc", "sdc", "eda", "ddc", "pdc", "bdc", "odc", "idc"}
        wizard_canvas = canvas if canvas in _allowed else ""
        llm_models, llm_default_model = _get_chat_models()
        return render_template(
            "chat.html",
            session_id=None,
            messages=[],
            wizard_goal=goal,
            wizard_role=role,
            wizard_classification=classification,
            wizard_frameworks=frameworks,
            wizard_custom_role_name=custom_role_name,
            wizard_custom_role_desc=custom_role_desc,
            wizard_canvas=wizard_canvas,
            wizard_use_case=use_case_id,
            wizard_skip_fast_track=skip_fast_track,
            from_wizard=from_wizard,
            llm_models=llm_models,
            llm_default_model=llm_default_model,
        )

    @app.route("/chat/<session_id>")
    def chat_session(session_id):
        """Resume an existing requirements chat session."""
        conn = _get_db()
        try:
            session = None
            session_err = None
            try:
                session = conn.execute("SELECT * FROM intake_sessions WHERE id = %s", (session_id,)).fetchone()
            except Exception as exc:
                session_err = str(exc)
                app.logger.warning("chat_session: intake_sessions lookup failed for %s: %s", session_id, session_err)
            if not session:
                # Fallback: try chat_contexts directly (session may exist as a standalone context)
                try:
                    ctx = conn.execute("SELECT * FROM chat_contexts WHERE id = %s", (session_id,)).fetchone()
                    if ctx:
                        # Fabricate a minimal session dict so the template can render
                        ctx_d = dict(ctx)
                        session = {
                            "id": session_id,
                            "session_status": ctx_d.get("status", "active"),
                            "created_at": ctx_d.get("created_at"),
                            "updated_at": ctx_d.get("updated_at"),
                            "context_summary": "{}",
                        }
                        app.logger.info("chat_session: fell back to chat_contexts for %s", session_id)
                except Exception as exc:
                    app.logger.warning("chat_session: chat_contexts fallback failed for %s: %s", session_id, exc)
            if not session:
                app.logger.error("chat_session: returning 404 for %s (intake_sessions error: %s)", session_id, session_err or "no row")
                return render_template("404.html", message="Session not found"), 404
            messages = []
            try:
                messages = conn.execute(
                    "SELECT turn_number, role, content, content_type, created_at "
                    "FROM intake_conversation WHERE session_id = %s ORDER BY turn_number",
                    (session_id,),
                ).fetchall()
            except Exception as exc:
                app.logger.warning("chat_session: intake_conversation query failed for %s: %s", session_id, exc)
            # Look up the linked chat context so JS can auto-select it
            auto_context_id = None
            try:
                ctx_row = conn.execute(
                    "SELECT id FROM chat_contexts WHERE intake_session_id = %s LIMIT 1",
                    (session_id,),
                ).fetchone()
                if ctx_row:
                    auto_context_id = ctx_row["id"]
            except Exception:
                pass
            # Extract context for sidebar display
            import json as _json

            session_dict = dict(session)
            ctx = {}
            try:
                ctx = _json.loads(session_dict.get("context_summary") or "{}")
            except (ValueError, TypeError):
                pass
            llm_models, llm_default_model = _get_chat_models()
            return render_template(
                "chat.html",
                session_id=session_id,
                auto_context_id=auto_context_id,
                session=session_dict,
                messages=[dict(m) for m in messages],
                wizard_goal=None,
                wizard_role=None,
                wizard_classification=None,
                wizard_frameworks=",".join(ctx.get("selected_frameworks", [])),
                wizard_custom_role_name="",
                wizard_custom_role_desc="",
                session_context=ctx,
                llm_models=llm_models,
                llm_default_model=llm_default_model,
            )
        finally:
            conn.close()

    # ── Requirements human-readable view ─────────────────────────────────
    @app.route("/intake/requirements/<session_id>")
    def intake_requirements_view(session_id):
        """Professional requirements document view for an intake session."""
        import json as _json
        from datetime import datetime as _dt, timezone as _tz
        from flask import abort as _abort
        conn = _get_db()
        try:
            session = conn.execute(
                "SELECT * FROM intake_sessions WHERE id = %s", (session_id,)
            ).fetchone()
            if not session:
                app.logger.warning("intake_requirements_view: session %s not found", session_id)
                _abort(404, description=f"Session '{session_id}' not found.")
            session = dict(session)

            reqs = conn.execute(
                "SELECT * FROM intake_requirements WHERE session_id = %s ORDER BY priority, created_at",
                (session_id,),
            ).fetchall()
            reqs = [dict(r) for r in reqs]

            readiness_row = conn.execute(
                "SELECT readiness_score, readiness_breakdown FROM intake_sessions WHERE id = %s",
                (session_id,),
            ).fetchone()
        finally:
            conn.close()

        # Build readiness dict
        readiness = None
        if readiness_row and readiness_row["readiness_score"] is not None:
            breakdown = {}
            try:
                breakdown = _json.loads(readiness_row["readiness_breakdown"] or "{}")
            except Exception:
                pass
            readiness = {
                "overall": float(readiness_row["readiness_score"]),
                "dimensions": {
                    "Completeness": breakdown.get("completeness", 0),
                    "Clarity": breakdown.get("clarity", 0),
                    "Feasibility": breakdown.get("feasibility", 0),
                    "Compliance": breakdown.get("compliance", 0),
                    "Testability": breakdown.get("testability", 0),
                },
            }

        # Group requirements by type
        _type_order = ["functional", "performance", "interface", "data", "compliance", "security", "non_functional"]
        _type_labels = {
            "functional": "Functional", "performance": "Performance",
            "interface": "Interface", "data": "Data", "compliance": "Compliance",
            "security": "Security", "non_functional": "Non-Functional",
        }
        groups = {}
        for req in reqs:
            t = (req.get("requirement_type") or "functional").lower().replace(" ", "_")
            groups.setdefault(t, []).append(req)

        grouped = []
        for t in _type_order:
            if t in groups:
                grouped.append((t, _type_labels.get(t, t.replace("_", " ").title()), groups[t]))
        for t, items in groups.items():
            if t not in _type_order:
                grouped.append((t, t.replace("_", " ").title(), items))

        # Frameworks
        frameworks = []
        try:
            ctx = _json.loads(session.get("context_summary") or "{}")
            raw_fw = ctx.get("selected_frameworks", [])
            frameworks = [fw.replace("_", " ").replace("-", " ").title() for fw in raw_fw]
        except Exception:
            pass

        req_types = list(groups.keys())
        cui_banner = getattr(app, "_cui_banner", None) or "CUI // SP-CTI"
        generated_at = _dt.now(_tz.utc).strftime("%Y-%m-%d %H:%M UTC")

        try:
            return render_template(
                "intake_requirements.html",
                session=session,
                requirements=reqs,
                grouped_requirements=grouped,
                req_types=req_types,
                readiness=readiness,
                frameworks=frameworks,
                classification_banner=cui_banner,
                generated_at=generated_at,
            )
        except Exception as _exc:
            app.logger.error("intake_requirements_view render error: %s", _exc)
            return jsonify({"error": str(_exc)}), 500

    # ── PRD rendered HTML view ────────────────────────────────────────────
    @app.route("/intake/prd/<session_id>/view")
    def intake_prd_view(session_id):
        """Render PRD as a styled HTML page instead of downloading raw markdown."""
        from datetime import datetime as _dt, timezone as _tz
        try:
            import markdown as _md_lib
            _HAS_MARKDOWN_LIB = True
        except ImportError:
            _HAS_MARKDOWN_LIB = False

        # Fetch customer name from session
        conn = _get_db()
        try:
            _sess_row = conn.execute(
                "SELECT customer_name FROM intake_sessions WHERE id = %s", (session_id,)
            ).fetchone()
            session_customer = (_sess_row["customer_name"] if _sess_row else "") or ""
        except Exception:
            session_customer = ""
        finally:
            conn.close()

        try:
            from tools.requirements.prd_generator import generate_prd as _gen_prd
            result = _gen_prd(session_id)
        except Exception as exc:
            result = {"status": "error", "error": str(exc)}

        cui_banner = getattr(app, "_cui_banner", None) or "CUI // SP-CTI"
        generated_at = _dt.now(_tz.utc).strftime("%Y-%m-%d %H:%M UTC")

        if result.get("status") != "ok":
            return render_template(
                "intake_prd_view.html",
                session_id=session_id,
                session_customer=session_customer,
                error=result.get("error", "PRD could not be generated."),
                prd_html="",
                prd_markdown_raw="",
                classification_banner=cui_banner,
                generated_at=generated_at,
            )

        raw_md = result.get("prd_markdown", "")

        if _HAS_MARKDOWN_LIB:
            prd_html = _md_lib.markdown(
                raw_md,
                extensions=["tables", "fenced_code", "nl2br", "toc"],
            )
            # nav-sec-07: python-markdown passes raw HTML through, and PRD content is
            # LLM/user-derived → stored/reflected XSS. Sanitize the rendered HTML with
            # the repo's canonical air-gap-safe sanitizer (no new dependency).
            try:
                from tools.docgen.workflow import _sanitize_html as _sanitize_prd_html
                prd_html = _sanitize_prd_html(prd_html)
            except Exception:
                import html as _html_lib
                prd_html = "<pre style='white-space:pre-wrap'>" + _html_lib.escape(raw_md) + "</pre>"
        else:
            import html as _html_lib
            prd_html = "<pre style='white-space:pre-wrap'>" + _html_lib.escape(raw_md) + "</pre>"

        return render_template(
            "intake_prd_view.html",
            session_id=session_id,
            session_customer=session_customer,
            error=None,
            prd_html=prd_html,
            # nav-sec-07: pass raw markdown; the template serializes it with Jinja
            # |tojson so a </script> sequence in the PRD cannot break out of the block.
            prd_markdown_raw=raw_md,
            classification_banner=cui_banner,
            generated_at=generated_at,
        )

    @app.route("/skillhub")
    def skillhub_page():
        """SkillHub — redirect to Marketplace skill catalog."""
        from flask import redirect
        return redirect("/studio/marketplace")

    @app.route("/quick-paths")
    def quick_paths_page():
        """Quick Path workflow templates — pre-built shortcuts for common tasks."""
        return render_template("quick_paths.html")

    @app.route("/batch")
    def batch_page():
        """Batch operations — run multi-tool workflows from the dashboard."""
        return render_template("batch.html")

    @app.route("/connector-forge")
    def connector_forge_page():
        """Connector Forge — generate API connectors from OpenAPI specs."""
        return render_template("connector_forge.html")

    @app.route("/api/connector-forge/list")
    def api_connector_forge_list():
        """List all generated/registered connectors."""
        try:
            from tools.databridge.registry import list_registered

            registered = list_registered()
            connectors = [{"name": k, "type": "registered", "status": "active"} for k in registered]
            return jsonify({"connectors": connectors, "total": len(connectors)})
        except Exception as exc:
            get_logger("icdev.dashboard").warning(
                "api_connector_forge_list: failed to list connectors: %s", exc
            )
            return jsonify(
                {
                    "connectors": [],
                    "total": 0,
                    "error": True,
                    "detail": "Connector registry is temporarily unavailable.",
                }
            ), 503

    @app.route("/diagrams")
    def diagrams_page():
        """Interactive Mermaid diagrams — catalog, viewer, and editor."""
        return render_template("diagrams.html")

    @app.route("/cicd")
    def cicd_page():
        """CI/CD pipeline status, conversations, and connector health."""
        return render_template("cicd.html")

    @app.route("/gateway")
    def gateway_page():
        """Remote Command Gateway admin — bindings, command log, channel status."""
        import yaml as _yaml

        # Load gateway config
        gateway_config_path = BASE_DIR / "args" / "remote_gateway_config.yaml"
        gw_config = {}
        if gateway_config_path.exists():
            with open(gateway_config_path) as f:
                gw_config = _yaml.safe_load(f) or {}

        env_mode = gw_config.get("environment", {}).get("mode", "connected")
        channels = gw_config.get("channels", {})

        # Determine active channels
        active_channels = []
        for name, ch in channels.items():
            enabled = ch.get("enabled", False)
            req_internet = ch.get("requires_internet", False)
            available = enabled and not (env_mode == "air_gapped" and req_internet)
            active_channels.append(
                {
                    "name": name,
                    "enabled": enabled,
                    "available": available,
                    "max_il": ch.get("max_il", "IL4"),
                    "description": ch.get("description", ""),
                }
            )

        # Load bindings and recent commands
        conn = _get_db()
        try:
            bindings = conn.execute("SELECT * FROM remote_user_bindings ORDER BY created_at DESC LIMIT 50").fetchall()
            bindings = [dict(r) for r in bindings]
        except Exception:
            bindings = []

        try:
            commands = conn.execute("SELECT * FROM remote_command_log ORDER BY created_at DESC LIMIT 50").fetchall()
            commands = [dict(r) for r in commands]
        except Exception:
            commands = []

        conn.close()

        return render_template(
            "gateway.html",
            environment_mode=env_mode,
            channels=active_channels,
            bindings=bindings,
            commands=commands,
            command_allowlist=gw_config.get("command_allowlist", []),
        )

    @app.route("/query")
    def query_page():
        """Natural language compliance query page."""
        conn = _get_db()
        try:
            recent_queries = conn.execute("SELECT * FROM nlq_queries ORDER BY created_at DESC LIMIT 20").fetchall()
            return render_template(
                "query/nlq.html",
                recent_queries=[dict(r) for r in recent_queries],
                degraded=False,
            )
        except Exception as exc:
            # nav-misc-02: an empty query history on read failure looks like "no queries".
            get_logger("icdev.dashboard").warning("query_page: nlq_queries read failed: %s", exc)
            return render_template("query/nlq.html", recent_queries=[], degraded=True)
        finally:
            conn.close()

    # ---- Tour configuration ----

    @app.route("/api/tour/steps", methods=["GET"])
    def api_tour_steps():
        """Return tour step definitions for the onboarding walkthrough.

        Steps are served from DB (tour_config table) first, falling back
        to built-in defaults. tour.js fetches this endpoint on init and
        falls back to built-in defaults if the fetch fails (air-gap safe).
        """
        # Try DB first
        conn = _get_db()
        try:
            rows = conn.execute("SELECT selector, title, description FROM tour_config ORDER BY sort_order").fetchall()
            if rows:
                db_steps = [{"selector": r["selector"], "title": r["title"], "desc": r["description"]} for r in rows]
                return jsonify(
                    {"steps": db_steps, "version": 2, "source": "db", "classification": DEFAULT_CLASSIFICATION}
                )
        except Exception:
            pass  # Table may not exist yet — fall through to defaults
        finally:
            conn.close()

        # Built-in defaults
        steps = [
            {
                "selector": ".navbar",
                "title": "Navigation Bar",
                "desc": (
                    "Navigate between pages: Home, Projects, Agents, "
                    "Monitoring, Quick Paths, Batch Operations, and "
                    "the Getting Started wizard."
                ),
            },
            {
                "selector": ".kanban-board",
                "title": "Project Kanban Board",
                "desc": (
                    "Projects organized by workflow stage: Planning, Active, "
                    "Completed, and Inactive. Click any card to view details."
                ),
            },
            {
                "selector": ".chart-grid",
                "title": "Visual Dashboards",
                "desc": (
                    "Visual dashboards: compliance posture, alert trends, project status, and agent health charts."
                ),
            },
            {
                "selector": ".table-container",
                "title": "Data Tables",
                "desc": ("Detailed data tables with search, sort, filter, and CSV export capabilities."),
            },
            {
                "selector": "#role-select",
                "title": "Role Selector",
                "desc": (
                    "Switch views: Program Manager, Developer, ISSO, or "
                    "Contracting Officer to see role-relevant information."
                ),
            },
            {
                "selector": "a[href*='quick-paths'], a[href*='/quick-paths']",
                "title": "Quick Paths",
                "desc": (
                    "Pre-built workflow shortcuts for common tasks like "
                    "ATO generation, project creation, and security scanning."
                ),
            },
            {
                "selector": "a[href*='/batch']",
                "title": "Batch Operations",
                "desc": (
                    "Run multi-step batch operations: Full ATO Package, "
                    "Security Scan Suite, Multi-Framework Check, or "
                    "Build & Validate from a single click."
                ),
            },
            {
                "selector": "a[href*='/events']",
                "title": "Live Events",
                "desc": (
                    "Real-time event timeline showing hook events, "
                    "agent activity, and system notifications with "
                    "severity filtering."
                ),
            },
        ]
        return jsonify(
            {
                "steps": steps,
                "version": 2,
                "classification": DEFAULT_CLASSIFICATION,
            }
        )

    # ---- Profile routes (D172, D175-D178) ----

    @app.route("/settings")
    def settings_redirect():
        return redirect(url_for("profile_page"))

    @app.route("/profile")
    def profile_page():
        """User profile page with BYOK key management."""
        return render_template("profile.html")

    @app.route("/profile/api/theme", methods=["GET", "POST"])
    def profile_theme():
        """Get or set the user's theme preference (stored as a cookie)."""
        if flask_request.method == "POST":
            data = flask_request.get_json(silent=True) or {}
            theme = data.get("theme", "dark")
            if theme not in ("dark", "light"):
                theme = "dark"
            resp = make_response(jsonify({"theme": theme}))
            resp.set_cookie("icdev_theme", theme, max_age=60 * 60 * 24 * 365, samesite="Lax")
            return resp
        theme = flask_request.cookies.get("icdev_theme", "dark")
        return jsonify({"theme": theme})

    @app.route("/profile/api/keys")
    def profile_api_keys():
        """List current user's dashboard API keys."""
        from tools.dashboard.auth import list_api_keys_for_user

        user = getattr(g, "current_user", None)
        if not user:
            return jsonify({"keys": []})
        keys = list_api_keys_for_user(user["id"])
        return jsonify({"keys": keys})

    @app.route("/profile/api/llm-keys", methods=["GET"])
    def profile_llm_keys():
        """List current user's BYOK LLM keys."""
        from tools.dashboard.byok import list_llm_keys

        user = getattr(g, "current_user", None)
        if not user:
            return jsonify({"keys": []})
        keys = list_llm_keys(user["id"])
        return jsonify({"keys": keys})

    @app.route("/profile/api/llm-keys", methods=["POST"])
    def profile_add_llm_key():
        """Store a new BYOK LLM key for the current user."""
        from tools.dashboard.byok import store_llm_key

        user = getattr(g, "current_user", None)
        if not user:
            return jsonify({"error": "Not authenticated"}), 401
        data = flask_request.get_json(force=True)
        provider = data.get("provider", "").strip()
        api_key = data.get("api_key", "").strip()
        label = data.get("label", "").strip()
        if not provider or not api_key:
            return jsonify({"error": "provider and api_key required"}), 400
        result = store_llm_key(user["id"], provider, api_key, key_label=label)
        return jsonify(result), 201

    @app.route("/profile/api/llm-keys/<key_id>/revoke", methods=["POST"])
    def profile_revoke_llm_key(key_id):
        """Revoke a BYOK LLM key (ownership-scoped)."""
        from tools.dashboard.byok import revoke_llm_key

        user = getattr(g, "current_user", None)
        if not user:
            return jsonify({"error": "Not authenticated"}), 401
        success = revoke_llm_key(key_id, user_id=user["id"])
        if not success:
            return jsonify({"error": "Key not found"}), 404
        return jsonify({"status": "revoked"})

    # ---- WriteGuard route (Phase 1 — Content Quality Dashboard) ----

    @app.route("/writeguard")
    def writeguard_page():
        """WriteGuard — content quality & AI detection dashboard.

        Query params:
          opp_id  — pre-load all section drafts for this proposal opportunity
          sec_id  — pre-load a specific section draft
        """
        from flask import request as _req
        opp_id = _req.args.get("opp_id")
        sec_id = _req.args.get("sec_id")
        preload_sections = []
        preload_title = ""
        if opp_id or sec_id:
            try:
                conn = _get_db()
                if sec_id:
                    row = conn.execute(
                        "SELECT d.draft_content, s.title as section_title, o.title as opp_title "
                        "FROM proposal_section_drafts d "
                        "JOIN proposal_sections s ON d.section_id = s.id "
                        "JOIN proposal_opportunities o ON s.opportunity_id = o.id "
                        "WHERE d.id = %s ORDER BY d.created_at DESC LIMIT 1",
                        (sec_id,),
                    ).fetchone()
                    if row:
                        preload_sections = [{"title": row["section_title"], "content": row["draft_content"]}]
                        preload_title = row["opp_title"]
                elif opp_id:
                    opp = conn.execute(
                        "SELECT title FROM proposal_opportunities WHERE id = %s", (opp_id,)
                    ).fetchone()
                    if opp:
                        preload_title = opp["title"]
                    drafts = conn.execute(
                        "SELECT d.draft_content, s.title as section_title, s.section_number "
                        "FROM proposal_section_drafts d "
                        "JOIN proposal_sections s ON d.section_id = s.id "
                        "WHERE s.opportunity_id = %s "
                        "AND d.created_at = ("
                        "  SELECT MAX(d2.created_at) FROM proposal_section_drafts d2 "
                        "  WHERE d2.section_id = d.section_id"
                        ") "
                        "ORDER BY s.section_number ASC",
                        (opp_id,),
                    ).fetchall()
                    preload_sections = [
                        {"title": f"{r['section_number']}: {r['section_title']}", "content": r["draft_content"]}
                        for r in drafts
                    ]
                conn.close()
            except Exception:
                pass
        return render_template(
            "writeguard.html",
            preload_sections=preload_sections,
            preload_title=preload_title,
            opp_id=opp_id or "",
        )

    # ---- Phase roadmap route ----

    @app.route("/phases")
    def phases_page():
        """Phase roadmap — all ICDEV™ phases with status, categories, and progress."""
        from tools.dashboard.phase_loader import (
            load_phases,
            load_categories,
            load_statuses,
            get_phase_summary,
        )

        phases = load_phases()
        categories = load_categories()
        statuses = load_statuses()
        summary = get_phase_summary(phases)

        # Optional category filter from query param
        cat_filter = flask_request.args.get("category", "")
        if cat_filter:
            phases = [p for p in phases if p.get("category") == cat_filter]

        return render_template(
            "phases.html",
            phases=phases,
            categories=categories,
            statuses=statuses,
            summary=summary,
            category_filter=cat_filter,
        )

    # ---- Dev profile routes (Phase 34, D183-D188) ----

    @app.route("/dev-profiles")
    def dev_profiles_page():
        """Dev profile management — list, create, view profiles."""
        return render_template("dev_profiles.html")

    # ---- Child application routes (Phase 19 + Evolutionary Intelligence) ----

    @app.route("/children")
    def children_page():
        """Child application registry — health, genome, capabilities, heartbeats."""
        conn = _get_db()
        degraded = False
        try:
            # Fetch all registered child applications
            try:
                children_rows = conn.execute("SELECT * FROM child_app_registry ORDER BY created_at DESC").fetchall()
                children_rows = [dict(r) for r in children_rows]
            except Exception as exc:
                # nav-misc-02: registry read failed — an empty roster here would
                # falsely imply "no child apps registered".
                children_rows = []
                degraded = True
                get_logger("icdev.dashboard").warning("children_page: child_app_registry read failed: %s", exc)

            # Fetch latest heartbeat per child from telemetry
            heartbeat_map = {}
            try:
                heartbeats = conn.execute(
                    "SELECT child_id, MAX(reported_at) as last_heartbeat FROM child_telemetry GROUP BY child_id"
                ).fetchall()
                for hb in heartbeats:
                    hb_dict = dict(hb)
                    heartbeat_map[hb_dict["child_id"]] = hb_dict["last_heartbeat"]
            except Exception:
                pass

            # Fetch capability count per child
            capability_map = {}
            try:
                caps = conn.execute(
                    "SELECT child_id, COUNT(*) as cnt FROM child_capabilities GROUP BY child_id"
                ).fetchall()
                for c in caps:
                    c_dict = dict(c)
                    capability_map[c_dict["child_id"]] = c_dict["cnt"]
            except Exception:
                pass

            # Enrich children with heartbeat and capability data
            children = []
            for child in children_rows:
                child["last_heartbeat"] = heartbeat_map.get(child.get("id"), child.get("last_heartbeat"))
                child["capability_count"] = capability_map.get(child.get("id"), child.get("capability_count", 0))
                child["pending_upgrades"] = child.get("pending_upgrades", 0)
                child["genome_version"] = child.get("genome_version", None)
                child["health_status"] = child.get("health_status", "unhealthy")
                children.append(child)

            # Compute summary counts
            healthy_count = sum(1 for c in children if c["health_status"] == "healthy")
            degraded_count = sum(1 for c in children if c["health_status"] == "degraded")
            unhealthy_count = sum(1 for c in children if c["health_status"] not in ("healthy", "degraded"))

            return render_template(
                "children.html",
                children=children,
                total_count=len(children),
                healthy_count=healthy_count,
                degraded_count=degraded_count,
                unhealthy_count=unhealthy_count,
                degraded=degraded,
            )
        finally:
            conn.close()

    @app.route("/dev-profiles/api/list")
    def dev_profiles_api_list():
        """List all dev profiles (JSON)."""
        conn = _get_db()
        try:
            rows = conn.execute(
                """SELECT id, scope, scope_id, version, is_active, inherits_from,
                          created_by, created_at, change_summary
                   FROM dev_profiles WHERE is_active = 1
                   ORDER BY created_at DESC LIMIT 50"""
            ).fetchall()
            return jsonify({"profiles": [dict(r) for r in rows]})
        except Exception as e:
            return jsonify({"profiles": [], "error": str(e)})
        finally:
            conn.close()

    @app.route("/dev-profiles/api/resolve/<scope>/<scope_id>")
    def dev_profiles_api_resolve(scope, scope_id):
        """Resolve 5-layer cascade for a scope (JSON)."""
        try:
            from tools.builder.dev_profile_manager import resolve_profile

            result = resolve_profile(scope, scope_id)
            return jsonify(result)
        except (ImportError, Exception) as e:
            return jsonify({"error": str(e)})

    @app.route("/dev-profiles/api/templates")
    def dev_profiles_api_templates():
        """List available starter templates (JSON)."""
        templates = []
        _degraded = False
        _detail = None
        templates_dir = Path(__file__).resolve().parent.parent.parent / "context" / "profiles"
        if templates_dir.exists():
            try:
                import yaml

                for f in sorted(templates_dir.glob("*.yaml")):
                    with open(f, "r", encoding="utf-8") as fh:
                        data = yaml.safe_load(fh)
                        templates.append(
                            {
                                "name": data.get("name", f.stem),
                                "file": f.name,
                                "description": data.get("description", ""),
                                "impact_levels": data.get("impact_levels", []),
                            }
                        )
            except Exception as exc:
                # nav-misc-02: a malformed template file must not silently hide the
                # whole starter-template list.
                _degraded = True
                _detail = str(exc)
                get_logger("icdev.dashboard").warning("dev_profiles_api_templates: template load failed: %s", exc)
        _resp = {"templates": templates}
        if _degraded:
            _resp["error"] = True
            _resp["detail"] = _detail
        return jsonify(_resp)

    @app.route("/dev-profiles/api/create", methods=["POST"])
    def dev_profiles_api_create():
        """Create a dev profile from template or data (JSON)."""
        try:
            from tools.builder.dev_profile_manager import create_profile

            data = flask_request.get_json(silent=True) or {}
            result = create_profile(
                scope=data.get("scope", "project"),
                scope_id=data.get("scope_id", ""),
                template_name=data.get("template"),
                created_by=data.get("created_by", "dashboard"),
            )
            return jsonify(result), 201 if "error" not in result else 400
        except (ImportError, Exception) as e:
            return jsonify({"error": str(e)}), 500

    @app.route("/dev-profiles/api/export/cursor/<scope>/<scope_id>")
    def dev_profiles_api_export_cursor(scope, scope_id):
        """Export a resolved dev profile as Cursor AI .cursorrules or .mdc."""
        try:
            from tools.builder.cursor_profile_generator import generate

            fmt = flask_request.args.get("format", "cursorrules")
            result = generate(scope, scope_id, fmt=fmt)
            if "error" in result:
                return jsonify(result), 404 if "not found" in result["error"].lower() else 400

            content = result["content"]
            filename = result["filename"]

            if flask_request.args.get("download"):
                response = Response(
                    content,
                    mimetype="text/plain",
                    headers={"Content-Disposition": f'attachment; filename="{filename}"'},
                )
                return response
            return jsonify({
                "status": "generated",
                "format": fmt,
                "filename": filename,
                "scope": scope,
                "scope_id": scope_id,
                "dimensions": result.get("dimensions", []),
                "content": content,
            })
        except (ImportError, Exception) as e:
            return jsonify({"error": str(e)}), 500

    @app.route("/dev-profiles/api/import/cursor", methods=["POST"])
    def dev_profiles_api_import_cursor():
        """Scan .cursor/rules/*.mdc files and seed a dev profile from them."""
        try:
            from tools.builder.cursor_profile_importer import seed_profile

            data = flask_request.get_json(silent=True) or {}
            scope = data.get("scope", "platform")
            scope_id = data.get("scope_id", "default")
            directory = data.get("directory")
            result = seed_profile(
                scope=scope,
                scope_id=scope_id,
                directory=directory,
                created_by=data.get("created_by", "dashboard"),
            )
            return jsonify(result), 201 if "error" not in result else 400
        except (ImportError, Exception) as e:
            return jsonify({"error": str(e)}), 500

    # ---- Auth routes (D169-D172) ----

    @app.route("/login", methods=["GET", "POST"])
    def login_page():
        """Login page — accepts API key via form or header."""
        # Auto-login when .env key is configured
        env_key = os.environ.get("ICDEV_DASHBOARD_API_KEY", "")
        if env_key:
            try:
                user = validate_api_key(env_key)
                if not user:
                    from tools.dashboard.auth import bootstrap_env_user

                    user = bootstrap_env_user(env_key)
                if user:
                    flask_session["user_id"] = user["id"]
                    return redirect(url_for("index"))
                else:
                    app.logger.warning("ICDEV_DASHBOARD_API_KEY set but validation failed")
            except Exception as exc:
                app.logger.error(f"Auto-login failed: {exc}")
        if flask_request.method == "POST":
            raw_key = flask_request.form.get("api_key", "").strip()
            user = validate_api_key(raw_key)
            # Fallback: accept ICDEV_DASHBOARD_API_KEY from .env
            if not user:
                env_key = os.environ.get("ICDEV_DASHBOARD_API_KEY", "")
                if env_key and raw_key == env_key:
                    from tools.dashboard.auth import bootstrap_env_user

                    user = bootstrap_env_user(env_key)
            if user:
                flask_session["user_id"] = user["id"]
                # cnr-plat-01: issue a fresh CSRF token at login (rotate on each
                # successful authentication to avoid token fixation).
                try:
                    from tools.security.csrf import issue_csrf_token
                    issue_csrf_token()
                except Exception:
                    pass
                log_auth_event(
                    user["id"],
                    "login_success",
                    ip_address=flask_request.remote_addr,
                    user_agent=flask_request.headers.get("User-Agent", "")[:256],
                    details="via_login_form",
                )
                return redirect(url_for("index"))
            else:
                log_auth_event(
                    None,
                    "login_failed",
                    ip_address=flask_request.remote_addr,
                    user_agent=flask_request.headers.get("User-Agent", "")[:256],
                    details="via_login_form",
                )
                return render_template("login.html", error="Invalid API key. Please try again.")
        return render_template("login.html", error=None)

    @app.route("/logout")
    def logout():
        """Clear session and redirect to login."""
        user_id = flask_session.get("user_id")
        if user_id:
            log_auth_event(
                user_id,
                "logout",
                ip_address=flask_request.remote_addr,
            )
        flask_session.clear()
        return redirect(url_for("login_page"))

    # ---- Error handlers ----

    # ---- Cross-Language Translation routes (Phase 43) ----

    @app.route("/translations")
    def translations_page():
        """Translation jobs — list, status, validation scores."""
        conn = _get_db()
        try:
            try:
                jobs = conn.execute(
                    """SELECT id, project_id, source_language, target_language,
                              status, total_units, translated_units, mocked_units,
                              failed_units, gate_result, llm_model, llm_tokens_input,
                              llm_tokens_output, elapsed_seconds, created_at
                       FROM translation_jobs ORDER BY created_at DESC LIMIT 100"""
                ).fetchall()
                jobs = [dict(r) for r in jobs]
            except Exception:
                jobs = []

            # Summary stats
            total = len(jobs)
            completed = sum(1 for j in jobs if j.get("status") == "completed")
            in_progress = sum(
                1
                for j in jobs
                if j.get("status") in ("pending", "extracting", "translating", "assembling", "validating")
            )
            failed = sum(1 for j in jobs if j.get("status") in ("failed", "partial"))

            # Average API surface score from validations
            avg_api_score = None
            try:
                row = conn.execute(
                    """SELECT AVG(score) as avg_score FROM translation_validations
                       WHERE check_type = 'api_surface' AND passed = 1"""
                ).fetchone()
                if row and row["avg_score"]:
                    avg_api_score = round(row["avg_score"] * 100, 1)
            except Exception:
                pass

            return render_template(
                "translations.html",
                jobs=jobs,
                total=total,
                completed=completed,
                in_progress=in_progress,
                failed=failed,
                avg_api_score=avg_api_score,
            )
        finally:
            conn.close()

    @app.route("/translations/<job_id>")
    def translation_detail_page(job_id):
        """Translation job detail — units, validations, dependencies."""
        conn = _get_db()
        try:
            # Fetch job
            try:
                job = conn.execute("SELECT * FROM translation_jobs WHERE id = %s", (job_id,)).fetchone()
                job = dict(job) if job else None
            except Exception:
                job = None

            if not job:
                return render_template("404.html", message="Translation job not found"), 404

            # Fetch units
            try:
                units = conn.execute(
                    """SELECT unit_name, unit_kind, source_file, status,
                              source_complexity, target_complexity,
                              repair_count, candidate_selected, created_at
                       FROM translation_units WHERE job_id = %s
                       ORDER BY created_at""",
                    (job_id,),
                ).fetchall()
                units = [dict(u) for u in units]
            except Exception:
                units = []

            # Fetch validations
            try:
                validations = conn.execute(
                    """SELECT check_type, passed, score, findings, created_at
                       FROM translation_validations WHERE job_id = %s
                       ORDER BY created_at""",
                    (job_id,),
                ).fetchall()
                validations = [dict(v) for v in validations]
            except Exception:
                validations = []

            # Fetch dependency mappings
            try:
                deps = conn.execute(
                    """SELECT source_import, target_import, mapping_source,
                              confidence, domain
                       FROM translation_dependency_mappings WHERE job_id = %s
                       ORDER BY domain, source_import""",
                    (job_id,),
                ).fetchall()
                deps = [dict(d) for d in deps]
            except Exception:
                deps = []

            return render_template(
                "translation_detail.html",
                job=job,
                units=units,
                validations=validations,
                deps=deps,
            )
        finally:
            conn.close()

    @app.route("/api/charts/translations")
    def api_charts_translations():
        """Chart data for translations page."""
        conn = _get_db()
        degraded = False
        detail = None
        try:
            # Status distribution
            status_dist = {}
            try:
                rows = conn.execute("SELECT status, COUNT(*) as cnt FROM translation_jobs GROUP BY status").fetchall()
                for r in rows:
                    r_dict = dict(r)
                    status_dist[r_dict["status"]] = r_dict["cnt"]
            except Exception as exc:
                # nav-misc-02: chart empties silently otherwise — flag the outage.
                degraded = True
                detail = str(exc)
                get_logger("icdev.dashboard").warning("api_charts_translations: status query failed: %s", exc)

            # Language pair frequency
            lang_pairs = {}
            try:
                rows = conn.execute(
                    """SELECT source_language || ' → ' || target_language as pair,
                              COUNT(*) as cnt
                       FROM translation_jobs GROUP BY pair ORDER BY cnt DESC LIMIT 10"""
                ).fetchall()
                for r in rows:
                    r_dict = dict(r)
                    lang_pairs[r_dict["pair"]] = r_dict["cnt"]
            except Exception as exc:
                degraded = True
                detail = str(exc)
                get_logger("icdev.dashboard").warning("api_charts_translations: lang-pair query failed: %s", exc)

            payload = {
                "status_distribution": status_dist,
                "language_pair_frequency": lang_pairs,
            }
            if degraded:
                payload["error"] = True
                payload["detail"] = detail
            return jsonify(payload)
        finally:
            conn.close()

    # ---- Phase 46: Observability pages ----

    @app.route("/traces")
    def traces_page():
        """Trace explorer — distributed tracing across MCP, A2A, LLM."""
        return render_template("traces.html")

    @app.route("/provenance")
    def provenance_page():
        """Provenance graph — W3C PROV-AGENT artifact lineage."""
        return render_template("provenance.html")

    @app.route("/xai")
    def xai_page():
        """XAI dashboard — explainability, SHAP attribution, compliance."""
        return render_template("xai.html")

    # ---- OSCAL & Production Audit pages ----

    @app.route("/oscal")
    def oscal_page():
        """OSCAL ecosystem — validation, catalog, format conversion (D302-D306)."""
        return render_template("oscal.html")

    @app.route("/prod-audit")
    def prod_audit_page():
        """Production readiness audit — 30 checks, 6 categories (D291-D300)."""
        return render_template("prod_audit.html")

    @app.route("/ai-transparency")
    def ai_transparency_page():
        """AI Transparency — OMB M-25-21, M-26-04, NIST AI 600-1, GAO-21-519SP (Phase 48, D307-D315)."""
        return render_template("ai_transparency.html")


    @app.route("/ai-accountability")
    def ai_accountability_page():
        """AI Accountability — oversight, appeals, CAIO, incidents, ethics (Phase 49, D316-D321)."""
        return render_template("ai_accountability.html")

    @app.route("/code-quality")
    def code_quality_page():
        """Code Quality Intelligence — AST analysis, smells, maintainability, runtime feedback (Phase 52, D331-D337)."""
        return render_template("code_quality.html")

    @app.route("/fedramp-20x")
    def fedramp_20x_page():
        """FedRAMP 20x KSI Dashboard — KSI evidence generation, maturity levels, authorization package (Phase 53, D338)."""
        return render_template("fedramp_20x.html")

    @app.route("/evidence")
    def evidence_page():
        """Evidence Collection — universal evidence auto-collection across all frameworks (Phase 56, D347)."""
        from tools.compliance.evidence_collector import FRAMEWORK_EVIDENCE_MAP, _get_connection, table_exists

        stats = {"total_frameworks": len(FRAMEWORK_EVIDENCE_MAP), "required_frameworks": 0, "frameworks": []}
        degraded = False
        try:
            conn = _get_connection()
            for fw_id, fw_config in FRAMEWORK_EVIDENCE_MAP.items():
                if fw_config["required"]:
                    stats["required_frameworks"] += 1
                total = 0
                for table_name in fw_config["tables"]:
                    if table_exists(conn, table_name):
                        row = conn.execute(f"SELECT COUNT(*) FROM {table_name}").fetchone()  # nosec B608 -- table/column names are internal constants, not user input
                        total += row[0]
                stats["frameworks"].append(
                    {
                        "id": fw_id,
                        "description": fw_config["description"],
                        "required": fw_config["required"],
                        "total_records": total,
                    }
                )
            conn.close()
        except Exception as exc:
            # nav-misc-05: surface an evidence DB outage instead of a silent empty page.
            degraded = True
            get_logger("icdev.dashboard").warning("evidence_page: DB read failed: %s", exc)
        return render_template("evidence.html", stats=stats, degraded=degraded)

    @app.route("/lineage")
    def lineage_page():
        """Artifact Lineage — unified DAG visualization of digital thread, provenance, audit trail, SBOM (Phase 56, D348)."""
        return render_template("lineage.html")

    # ---- Database helper ----
    def _get_db():
        import os
        from flask import has_request_context
        if os.environ.get("ICDEV_STORAGE_BACKEND", "").lower() == "postgresql":
            conn = get_connection()
        else:
            conn = get_connection(db_path=str(DB_PATH))
        try:
            if not has_request_context():
                conn.set_security_context(None)  # rls-bypass: CLI / background tasks run without a user session; no tenant context available.
            # In a request context: _attach_flask_security_context() already wired
            # g.security_context (set by auth middleware) into the connection.
        except Exception:
            pass
        return conn

    # ---- CPMP / Proposals / GovCon Pages (D-CHILD-6: guarded) ----
    if _HAS_GOVCON:
        _register_govcon_pages(app, _get_db)

    # ---- Phase 61: Orchestration Dashboard ----

    @app.route("/orchestration")
    def orchestration_dashboard():
        """Real-time multi-agent orchestration dashboard — agent grid, DAG, mailbox (Phase 61)."""
        return render_template("orchestration/dashboard.html")

    # ---- NDC Lab Backend Health ----

    @app.route("/network/labs")
    def network_labs_page():
        """NDC — Lab backend health (GNS3, Containerlab, EVE-NG)."""
        return render_template("network_labs.html")

    # ---- Digital Program Twin — Simulation Dashboard ----

    @app.route("/simulation")
    def simulation_page():
        """Digital Program Twin — 6-dimension what-if simulation, Monte Carlo, COA analysis."""
        stats = {"total_scenarios": 0, "running": 0, "completed": 0, "monte_carlo_runs": 0, "coas_generated": 0}
        scenarios = []
        degraded = False
        try:
            conn = _get_db()
            stats["total_scenarios"] = conn.execute(
                "SELECT COUNT(*) FROM simulation_scenarios WHERE status != 'archived'"
            ).fetchone()[0]
            stats["running"] = conn.execute(
                "SELECT COUNT(*) FROM simulation_scenarios WHERE status = 'running'"
            ).fetchone()[0]
            stats["completed"] = conn.execute(
                "SELECT COUNT(*) FROM simulation_scenarios WHERE status = 'completed'"
            ).fetchone()[0]
            stats["monte_carlo_runs"] = conn.execute("SELECT COUNT(*) FROM monte_carlo_runs").fetchone()[0]
            stats["coas_generated"] = conn.execute("SELECT COUNT(*) FROM coa_definitions").fetchone()[0]
            scenarios = [
                dict(r)
                for r in conn.execute(
                    "SELECT id, project_id, scenario_name, scenario_type, status, created_at, completed_at "
                    "FROM simulation_scenarios WHERE status != 'archived' ORDER BY created_at DESC LIMIT 100"
                ).fetchall()
            ]
            conn.close()
        except Exception as exc:
            # nav-misc-02: surface a simulation DB outage instead of an empty board.
            degraded = True
            get_logger("icdev.dashboard").warning("simulation_page: DB read failed: %s", exc)
        return render_template("simulation.html", stats=stats, scenarios=scenarios, degraded=degraded)

    # ------------------------------------------------------------------
    # Phase 73: Cloud Migration Security Pages (5 new dashboard pages)
    # ------------------------------------------------------------------

    @app.route("/security-scan")
    def security_scan_page():
        """Security Scan Results — multi-layer scanning dashboard (SAST, dependency, secret, container)."""
        return render_template("security_scan.html")

    @app.route("/migration")
    def migration_page():
        """Application Migration Tracker — 7R strategy assessment with compliance impact scoring."""
        return render_template("migration.html")

    @app.route("/sbd")
    def sbd_page():
        """CISA Secure by Design Assessment — 8-pillar assessment with automated gating."""
        return render_template("sbd.html")

    @app.route("/pr-intel")
    def pr_intel_page():
        """PR Intelligence — compliance drift detection at the pull request level."""
        return render_template("pr_intel.html")

    @app.route("/iac")
    def iac_page():
        """IaC Gallery — STIG-hardened Infrastructure as Code for multi-cloud IL2-IL6."""
        return render_template("iac.html")

    @app.route("/cato")
    def cato_page():
        """Continuous ATO — real-time ATO health score and evidence freshness."""
        return render_template("cato.html")

    @app.route("/control-inheritance")
    def control_inheritance_page():
        """Control Inheritance Visualizer — CSP vs customer responsibility mapping."""
        return render_template("control_inheritance.html")

    @app.route("/safety")
    def safety_monitor_page():
        """Safety Monitor — circuit breaker states, failure counts, and system resilience."""
        return render_template("safety_monitor/page.html")

    @app.route("/mosa")
    def mosa_page():
        """MOSA Compliance — 10 U.S.C. §4401 modular open systems approach assessment."""
        return render_template("mosa.html")

    @app.route("/api/mosa/summary")
    def api_mosa_summary():
        """MOSA summary — module coupling, cohesion, circular dependency data."""
        try:
            from tools.compliance.mosa_assessor import get_latest_assessment

            data = get_latest_assessment()
            return jsonify(data)
        except Exception:
            return jsonify(
                {
                    "modules": [],
                    "summary": {
                        "total_modules": 0,
                        "avg_coupling": 0,
                        "avg_cohesion": 0,
                        "circular_deps": 0,
                    },
                }
            )

    @app.route("/migration-cost")
    def migration_cost_page():
        """Migration Cost Estimator — 7R ROI calculator with compliance cost."""
        return render_template("migration_cost.html")

    @app.route("/compliance")
    def compliance_page():
        """Compliance Hub — unified posture across all compliance modules."""
        return render_template("compliance.html")

    @app.route("/api/compliance/posture")
    def api_compliance_posture():
        """Return aggregate compliance posture for the hub page."""
        conn = _get_db()
        result = {"controls_implemented": 0, "open_poams": 0, "cat1_findings": 0, "ato_status": "--", "frameworks": []}
        try:
            # 1. Controls implemented — prefer ssp_controls, fall back to project_controls
            try:
                row = conn.execute(
                    "SELECT COUNT(*) as cnt FROM ssp_controls WHERE implementation_status = 'implemented'"
                ).fetchone()
                result["controls_implemented"] = row["cnt"]
            except Exception:
                conn.rollback()
                try:
                    row = conn.execute(
                        "SELECT COUNT(*) as cnt FROM project_controls WHERE implementation_status = 'implemented'"
                    ).fetchone()
                    result["controls_implemented"] = row["cnt"]
                except Exception:
                    conn.rollback()
            # 2. Open POAMs
            try:
                row = conn.execute(
                    "SELECT COUNT(*) as cnt FROM poam_items WHERE status NOT IN ('completed', 'closed')"
                ).fetchone()
                result["open_poams"] = row["cnt"]
            except Exception:
                conn.rollback()
            # 3. CAT I STIG findings
            try:
                row = conn.execute(
                    "SELECT COUNT(*) as cnt FROM stig_findings WHERE severity = 'CAT1' AND status = 'Open'"
                ).fetchone()
                result["cat1_findings"] = row["cnt"]
            except Exception:
                conn.rollback()
            # 4. ATO status
            try:
                row = conn.execute(
                    "SELECT authorization_status FROM ato_packages ORDER BY created_at DESC LIMIT 1"
                ).fetchone()
                result["ato_status"] = row["authorization_status"] if row else "Not Started"
            except Exception:
                conn.rollback()
                result["ato_status"] = "Not Started"
            # 5. Framework summaries — prefer ssp_controls, fall back to project_controls
            frameworks = [
                ("NIST 800-53", "ssp_controls", "implementation_status", "implemented"),
                ("FedRAMP", "ssp_controls", "implementation_status", "implemented"),
            ]
            for name, table, col, val in frameworks:
                try:
                    total = conn.execute(f"SELECT COUNT(*) as cnt FROM {table}").fetchone()["cnt"]  # nosec B608 -- table/column names are internal constants, not user input
                    impl = conn.execute(
                        f"SELECT COUNT(*) as cnt FROM {table} WHERE {col} = %s",  # nosec B608 -- table/column from internal frameworks list, not user input
                        (val,),
                    ).fetchone()["cnt"]
                    result["frameworks"].append({"name": name, "total": total, "implemented": impl, "status": "Active"})
                except Exception:
                    conn.rollback()
                    try:
                        total = conn.execute("SELECT COUNT(*) as cnt FROM project_controls").fetchone()["cnt"]
                        impl = conn.execute(
                            "SELECT COUNT(*) as cnt FROM project_controls WHERE implementation_status = %s",
                            (val,),
                        ).fetchone()["cnt"]
                        result["frameworks"].append({"name": name, "total": total, "implemented": impl, "status": "Active"})
                    except Exception:
                        conn.rollback()
        finally:
            conn.close()
        return jsonify(result)

    @app.route("/api/compliance/unified-posture")
    def api_compliance_unified_posture():
        """Unified compliance posture from PDC + NDC + SDC with NIST 800-53 heatmap."""

        NIST_FAMILIES = [
            ("AC", "Access Control"),
            ("AU", "Audit & Accountability"),
            ("AT", "Awareness & Training"),
            ("CA", "Assessment & Authorization"),
            ("CM", "Configuration Mgmt"),
            ("CP", "Contingency Planning"),
            ("IA", "ID & Authentication"),
            ("IR", "Incident Response"),
            ("MA", "Maintenance"),
            ("MP", "Media Protection"),
            ("PE", "Physical & Environmental"),
            ("PL", "Planning"),
            ("PM", "Program Management"),
            ("PS", "Personnel Security"),
            ("PT", "PII Processing"),
            ("RA", "Risk Assessment"),
            ("SA", "System & Services Acq"),
            ("SC", "System & Comms Protection"),
            ("SI", "System & Info Integrity"),
            ("SR", "Supply Chain Risk Mgmt"),
        ]

        result = {
            "sdc": {
                "available": False,
                "design_count": 0,
                "risk_score": None,
                "posture_grade": "--",
                "open_threats": 0,
                "controls_implemented": 0,
                "nist_coverage_pct": 0,
                "nist_families": {},
            },
            "ndc": {
                "available": False,
                "topology_count": 0,
                "cat1_open": 0,
                "cat2_open": 0,
                "cat3_open": 0,
                "total_findings": 0,
                "pass_rate": 0,
            },
            "pdc": {
                "available": False,
                "pipeline_count": 0,
                "slsa_level": "--",
                "ssdf_pct": 0,
                "owasp_pct": 0,
                "total_findings": 0,
            },
            "aiify": {
                "available": False,
                "grade": "--",
                "overall_score": None,
                "posture": "unrated",
                "scan_count": 0,
                "opportunity_count": 0,
                "weakest_dimension": None,
            },
            "aadc": {
                "available": False,
                "grade": "--",
                "design_count": 0,
                "scored_designs": 0,
                "avg_score": None,
                "nist_rmf_avg": None,
                "owasp_avg": None,
                "omb_compliant": 0,
                "safety_impacting": 0,
                "rights_impacting": 0,
            },
            "heatmap": [],
        }

        sdc_family_pcts: dict = {}
        main_family_pcts: dict = {}

        # --- SDC: security_canvas.db ---
        sdc_db = BASE_DIR / "data" / "security_canvas.db"
        if sdc_db.exists():
            try:
                with get_connection(str(sdc_db)) as sc:
                    result["sdc"]["design_count"] = sc.execute("SELECT COUNT(*) FROM security_designs").fetchone()[0]
                    result["sdc"]["available"] = result["sdc"]["design_count"] > 0
                    row = sc.execute(
                        "SELECT risk_score, posture_grade FROM sc_assessments ORDER BY ran_at DESC LIMIT 1"
                    ).fetchone()
                    if row:
                        result["sdc"]["risk_score"] = row["risk_score"]
                        result["sdc"]["posture_grade"] = row["posture_grade"]
                    result["sdc"]["open_threats"] = sc.execute(
                        "SELECT COUNT(*) FROM sc_threats WHERE status != 'mitigated'"
                    ).fetchone()[0]
                    for family, _ in NIST_FAMILIES:
                        total = sc.execute(
                            "SELECT COUNT(*) FROM sc_controls WHERE control_family = %s", (family,)
                        ).fetchone()[0]
                        impl = sc.execute(
                            "SELECT COUNT(*) FROM sc_controls WHERE control_family = %s"
                            " AND implementation_status IN ('implemented','tested')",
                            (family,),
                        ).fetchone()[0]
                        if total > 0:
                            sdc_family_pcts[family] = round(impl / total * 100)
                            result["sdc"]["controls_implemented"] += impl
                    if sdc_family_pcts:
                        result["sdc"]["nist_coverage_pct"] = round(sum(sdc_family_pcts.values()) / len(sdc_family_pcts))
                        result["sdc"]["nist_families"] = sdc_family_pcts
            except Exception:
                pass

        # --- NDC: network_canvas.db ---
        ndc_db = BASE_DIR / "data" / "network_canvas.db"
        if ndc_db.exists():
            try:
                with get_connection(str(ndc_db)) as nc:
                    result["ndc"]["topology_count"] = nc.execute("SELECT COUNT(*) FROM topologies").fetchone()[0]
                    result["ndc"]["available"] = result["ndc"]["topology_count"] > 0
                    result["ndc"]["cat1_open"] = nc.execute(
                        "SELECT COUNT(*) FROM nc_compliance_findings WHERE severity = 'CAT1' AND status = 'open'"
                    ).fetchone()[0]
                    result["ndc"]["cat2_open"] = nc.execute(
                        "SELECT COUNT(*) FROM nc_compliance_findings WHERE severity = 'CAT2' AND status = 'open'"
                    ).fetchone()[0]
                    result["ndc"]["cat3_open"] = nc.execute(
                        "SELECT COUNT(*) FROM nc_compliance_findings WHERE severity = 'CAT3' AND status = 'open'"
                    ).fetchone()[0]
                    total_f = nc.execute("SELECT COUNT(*) FROM nc_compliance_findings").fetchone()[0]
                    remediated_f = nc.execute(
                        "SELECT COUNT(*) FROM nc_compliance_findings WHERE status = 'remediated'"
                    ).fetchone()[0]
                    result["ndc"]["total_findings"] = total_f
                    result["ndc"]["pass_rate"] = round(remediated_f / total_f * 100) if total_f > 0 else 0
            except Exception:
                pass

        # --- PDC: pipeline_canvas.db ---
        pdc_db = BASE_DIR / "data" / "pipeline_canvas.db"
        if pdc_db.exists():
            try:
                with get_connection(str(pdc_db)) as pc:
                    result["pdc"]["pipeline_count"] = pc.execute("SELECT COUNT(*) FROM pipelines").fetchone()[0]
                    result["pdc"]["available"] = result["pdc"]["pipeline_count"] > 0
                    slsa_row = pc.execute(
                        "SELECT slsa_level FROM pc_snippets WHERE slsa_level IS NOT NULL"
                        " GROUP BY slsa_level ORDER BY COUNT(*) DESC LIMIT 1"
                    ).fetchone()
                    if slsa_row:
                        result["pdc"]["slsa_level"] = slsa_row["slsa_level"]
                    chk = pc.execute(
                        "SELECT findings_json FROM pc_compliance_checks ORDER BY ran_at DESC LIMIT 1"
                    ).fetchone()
                    if chk and chk["findings_json"]:
                        try:
                            findings = json.loads(chk["findings_json"])
                            result["pdc"]["total_findings"] = len(findings) if isinstance(findings, list) else 0
                        except Exception:
                            pass
                    # OWASP coverage — derive from node types in all pipelines
                    try:
                        from tools.pipeline.constants import compute_owasp_coverage

                        all_node_types: list = []
                        for g_row in pc.execute("SELECT graph_json FROM pipelines").fetchall():
                            try:
                                g = json.loads(g_row["graph_json"] or "{}")
                                for n in g.get("nodes", []):
                                    t = n.get("type") or n.get("data", {}).get("type", "")
                                    if t:
                                        all_node_types.append(t)
                            except Exception:
                                pass
                        if all_node_types:
                            owasp = compute_owasp_coverage(all_node_types)
                            result["pdc"]["owasp_pct"] = int(owasp.get("coverage_pct", 0))
                    except Exception:
                        pass
                    # SSDF coverage — remediation rate of SSDF framework findings
                    try:
                        if table_exists(pc, "pc_compliance_findings"):
                            ssdf_total = pc.execute(
                                "SELECT COUNT(*) FROM pc_compliance_findings WHERE framework LIKE 'SSDF%'"
                            ).fetchone()[0]
                            ssdf_rem = pc.execute(
                                "SELECT COUNT(*) FROM pc_compliance_findings"
                                " WHERE framework LIKE 'SSDF%' AND status = 'remediated'"
                            ).fetchone()[0]
                            if ssdf_total > 0:
                                result["pdc"]["ssdf_pct"] = round(ssdf_rem / ssdf_total * 100)
                            else:
                                # No findings means passing — treat as 100%
                                result["pdc"]["ssdf_pct"] = 100
                    except Exception:
                        pass
            except Exception:
                pass

        # --- AI-ify: aiify_canvas — AI-governance posture (live compute) ---
        try:
            from tools.aiify.posture import compute_posture as _aiify_posture
            from tools.aiify.db.init_db import get_connection as _aiify_conn

            ac = _aiify_conn()
            try:
                ap = _aiify_posture(ac)
            finally:
                ac.close()
            counts = ap.get("counts", {})
            applicable = [d for d in ap.get("dimensions", []) if d.get("score") is not None]
            weakest = min(applicable, key=lambda d: d["score"], default=None)
            result["aiify"] = {
                "available": counts.get("total_scans", 0) > 0,
                "grade": ap.get("grade", "--"),
                "overall_score": ap.get("overall_score"),
                "posture": ap.get("posture", "unrated"),
                "scan_count": counts.get("total_scans", 0),
                "opportunity_count": counts.get("total_opportunities", 0),
                "weakest_dimension": (
                    {"label": weakest["label"], "score": weakest["score"]} if weakest else None
                ),
            }
        except Exception:
            pass

        # --- AADC: agentic_ai_canvas — portfolio assessment posture ---
        try:
            from tools.agentic_ai_canvas.db.init_db import get_connection as _aadc_conn

            gc = _aadc_conn()
            try:
                design_count = gc.execute("SELECT COUNT(*) FROM aadc_designs").fetchone()[0]
                result["aadc"]["design_count"] = design_count
                result["aadc"]["available"] = design_count > 0
                # Latest assessment per design
                rows = gc.execute(
                    "SELECT a.score, a.nist_rmf_score, a.owasp_score, a.omb_compliant, "
                    "a.safety_impacting, a.rights_impacting "
                    "FROM aadc_assessments a "
                    "JOIN (SELECT design_id, MAX(created_at) AS mx FROM aadc_assessments "
                    "      GROUP BY design_id) m "
                    "  ON a.design_id = m.design_id AND a.created_at = m.mx"
                ).fetchall()
            finally:
                gc.close()
            if rows:
                def _col(r, key, idx):
                    try:
                        return r[key]
                    except (KeyError, IndexError, TypeError):
                        return r[idx]
                scores = [_col(r, "score", 0) or 0 for r in rows]
                nists = [_col(r, "nist_rmf_score", 1) or 0 for r in rows]
                owasps = [_col(r, "owasp_score", 2) or 0 for r in rows]
                avg = round(sum(scores) / len(scores), 1)
                result["aadc"]["scored_designs"] = len(rows)
                result["aadc"]["avg_score"] = avg
                result["aadc"]["nist_rmf_avg"] = round(sum(nists) / len(nists), 1)
                result["aadc"]["owasp_avg"] = round(sum(owasps) / len(owasps), 1)
                result["aadc"]["omb_compliant"] = sum(1 for r in rows if _col(r, "omb_compliant", 3))
                result["aadc"]["safety_impacting"] = sum(1 for r in rows if _col(r, "safety_impacting", 4))
                result["aadc"]["rights_impacting"] = sum(1 for r in rows if _col(r, "rights_impacting", 5))
                result["aadc"]["grade"] = (
                    "A" if avg >= 90 else "B" if avg >= 80 else "C" if avg >= 70
                    else "D" if avg >= 60 else "F"
                )
        except Exception:
            pass

        # --- NIST 800-53 heatmap from icdev.db project_controls ---
        try:
            with get_connection(db_path=str(DB_PATH)) as mc:
                for family, _ in NIST_FAMILIES:
                    total = mc.execute(
                        "SELECT COUNT(*) as cnt FROM project_controls WHERE control_id LIKE %s",
                        (f"{family}-%",),
                    ).fetchone()["cnt"]
                    impl = mc.execute(
                        "SELECT COUNT(*) as cnt FROM project_controls"
                        " WHERE control_id LIKE %s AND implementation_status = 'implemented'",
                        (f"{family}-%",),
                    ).fetchone()["cnt"]
                    if total > 0:
                        main_family_pcts[family] = round(impl / total * 100)
        except Exception:
            pass

        for family, name in NIST_FAMILIES:
            sdc_pct = sdc_family_pcts.get(family)
            main_pct = main_family_pcts.get(family)
            vals = [v for v in [sdc_pct, main_pct] if v is not None]
            avg_pct = round(sum(vals) / len(vals)) if vals else 0
            result["heatmap"].append(
                {
                    "family": family,
                    "name": name,
                    "sdc_pct": sdc_pct,
                    "main_pct": main_pct,
                    "avg_pct": avg_pct,
                }
            )

        return jsonify(result)

    @app.route("/api/compliance/evidence-chain")
    def api_compliance_evidence_chain():
        """Evidence chain summary — PDC/NDC/SDC audit trail mapped to NIST 800-53."""
        since_hours = float(flask_request.args.get("since_hours", 168))  # 7 days default
        project_id = flask_request.args.get("project_id") or None
        try:
            from tools.compliance.evidence_chain import build_evidence_chain

            chain = build_evidence_chain(
                project_id=project_id,
                since_hours=since_hours,
            )
            # Return lightweight summary (drop full event list for dashboard perf)
            summary = {
                "chain_id": chain["chain_id"],
                "built_at": chain["built_at"],
                "since_hours": since_hours,
                "total_events": chain["timeline"]["total_events"],
                "first_event": chain["timeline"]["first_event"],
                "last_event": chain["timeline"]["last_event"],
                "sources": chain["sources"],
                "coverage": chain["coverage"],
                "evidence_types": chain["evidence_types"],
                "gate": chain["gate"],
                "recent_events": chain["timeline"]["events"][-10:],
            }
            return jsonify(summary)
        except Exception as exc:
            return jsonify({"error": str(exc), "total_events": 0, "sources": {}}), 200

    @app.route("/compliance-debt")
    def compliance_debt_page():
        """Compliance Debt Burndown — POAM, control, and STIG debt tracking."""
        return render_template("compliance_debt.html")

    @app.route("/stig-manager")
    def stig_manager_page():
        """STIG Benchmark Manager — import, track, and assess DISA STIG findings."""
        return render_template("stig_manager.html")

    @app.route("/ato-package")
    def ato_package_page():
        """ATO Package Builder — wizard to assemble SSP/SAR/POAM/SBOM package."""
        return render_template("ato_package.html")

    @app.route("/ato-compliance")
    def ato_compliance_page():
        """ATO Compliance Dashboard — control tracking, RMF stages, artifact readiness, crosswalk."""
        return render_template("ato_compliance.html")

    @app.route("/analytics")
    def analytics_page():
        """Compliance Funnel Analytics — redirects to usage dashboard (template pending)."""
        return redirect("/usage")

    @app.route("/api/simulation/scenarios", methods=["POST"])
    def api_simulation_create():
        """Create a new simulation scenario."""
        data = flask_request.get_json(silent=True) or {}
        try:
            from tools.simulation.simulation_engine import create_scenario

            result = create_scenario(
                project_id=data.get("project_id", ""),
                scenario_name=data.get("scenario_name", ""),
                scenario_type=data.get("scenario_type", "what_if"),
                modifications=data.get("modifications", {}),
            )
            return jsonify(result)
        except Exception as exc:
            return jsonify({"error": str(exc)}), 400

    @app.route("/api/simulation/scenarios/<scenario_id>/run", methods=["POST"])
    def api_simulation_run(scenario_id):
        """Run simulation across all 6 dimensions."""
        try:
            from tools.simulation.simulation_engine import run_simulation

            result = run_simulation(scenario_id)
            return jsonify(result)
        except Exception as exc:
            return jsonify({"error": str(exc)}), 400

    @app.route("/api/simulation/scenarios/<scenario_id>/summary")
    def api_simulation_summary(scenario_id):
        """Get scenario summary with results and MC runs."""
        try:
            from tools.simulation.scenario_manager import get_scenario_summary

            result = get_scenario_summary(scenario_id)
            return jsonify(result)
        except Exception as exc:
            return jsonify({"error": str(exc)}), 400

    @app.route("/api/simulation/scenarios/<scenario_id>/monte-carlo", methods=["POST"])
    def api_simulation_monte_carlo(scenario_id):
        """Run Monte Carlo estimation for a dimension."""
        data = flask_request.get_json(silent=True) or {}
        try:
            from tools.simulation.monte_carlo import run_monte_carlo

            result = run_monte_carlo(
                scenario_id=scenario_id,
                dimension=data.get("dimension", "schedule"),
                iterations=data.get("iterations", 10000),
            )
            return jsonify(result)
        except Exception as exc:
            return jsonify({"error": str(exc)}), 400

    @app.route("/api/simulation/scenarios/<scenario_id>/fork", methods=["POST"])
    def api_simulation_fork(scenario_id):
        """Fork a scenario with optional modifications."""
        data = flask_request.get_json(silent=True) or {}
        try:
            from tools.simulation.scenario_manager import fork_scenario

            result = fork_scenario(
                scenario_id=scenario_id,
                new_name=data.get("new_name", "Forked scenario"),
                additional_modifications=data.get("modifications"),
            )
            return jsonify(result)
        except Exception as exc:
            return jsonify({"error": str(exc)}), 400

    @app.route("/api/simulation/scenarios/<scenario_id>/coas")
    def api_simulation_coas(scenario_id):
        """Get COAs linked to a scenario."""
        try:
            conn = _get_db()
            coas = [
                dict(r)
                for r in conn.execute(
                    "SELECT * FROM coa_definitions WHERE simulation_scenario_id = %s ORDER BY coa_type",
                    (scenario_id,),
                ).fetchall()
            ]
            conn.close()
            return jsonify({"coas": coas})
        except Exception as exc:
            return jsonify({"coas": [], "error": str(exc)})

    @app.route("/api/simulation/coas/<coa_id>/select", methods=["POST"])
    def api_simulation_coa_select(coa_id):
        """Select a COA."""
        try:
            conn = _get_db()
            conn.execute(
                "UPDATE coa_definitions SET status = 'selected', selected_at = datetime('now') WHERE id = %s", (coa_id,)
            )
            conn.commit()
            conn.close()
            return jsonify({"status": "selected", "coa_id": coa_id})
        except Exception as exc:
            return jsonify({"error": str(exc)}), 400

    @app.route("/api/simulation/coas/<coa_id>/reject", methods=["POST"])
    def api_simulation_coa_reject(coa_id):
        """Reject a COA."""
        try:
            conn = _get_db()
            conn.execute("UPDATE coa_definitions SET status = 'rejected' WHERE id = %s", (coa_id,))
            conn.commit()
            conn.close()
            return jsonify({"status": "rejected", "coa_id": coa_id})
        except Exception as exc:
            return jsonify({"error": str(exc)}), 400

    @app.route("/api/simulation/nlq", methods=["POST"])
    def api_simulation_nlq():
        """Parse a natural language query into simulation modifications."""
        data = flask_request.get_json(silent=True) or {}
        query = data.get("query", "")
        if not query:
            return jsonify({"error": "query is required"}), 400
        try:
            from tools.simulation.query_parser import parse_query

            result = parse_query(query)
            return jsonify(result)
        except Exception as exc:
            return jsonify({"error": str(exc)}), 400

    @app.route("/api/simulation/nlq/run", methods=["POST"])
    def api_simulation_nlq_run():
        """Parse NLQ, create scenario, and run simulation in one call."""
        data = flask_request.get_json(silent=True) or {}
        query = data.get("query", "")
        project_id = data.get("project_id", "")
        if not query or not project_id:
            return jsonify({"error": "query and project_id are required"}), 400
        try:
            from tools.simulation.query_parser import parse_query
            from tools.simulation.simulation_engine import create_scenario, run_simulation

            parsed = parse_query(query)
            scenario = create_scenario(
                project_id=project_id,
                scenario_name=parsed["scenario_name"],
                scenario_type=parsed["scenario_type"],
                modifications=parsed["modifications"],
            )
            sim_result = None
            try:
                sim_result = run_simulation(scenario["scenario_id"])
            except Exception:
                pass  # simulation may fail if no SysML data; return parsed+scenario anyway
            return jsonify(
                {
                    "parsed": parsed,
                    "scenario": scenario,
                    "simulation": sim_result,
                }
            )
        except Exception as exc:
            return jsonify({"error": str(exc)}), 400

    @app.route("/api/simulation/cascade", methods=["POST"])
    def api_simulation_cascade():
        """Run cascade analysis through the simulation KG."""
        data = flask_request.get_json(silent=True) or {}
        project_id = data.get("project_id", "")
        trigger = data.get("trigger", "")
        node_id = data.get("node_id")
        max_depth = data.get("max_depth", 7)
        max_width = data.get("max_width", 10)
        if not project_id:
            return jsonify({"error": "project_id is required"}), 400
        try:
            from tools.simulation.cascade_bridge import run_cascade

            start_ids = [node_id] if node_id else None
            result = run_cascade(
                project_id=project_id,
                start_node_ids=start_ids,
                trigger_text=trigger or None,
                max_depth=max_depth,
                max_width=max_width,
            )
            return jsonify(result)
        except Exception as exc:
            return jsonify({"error": str(exc)}), 400

    @app.route("/api/simulation/risk/composite", methods=["POST"])
    def api_simulation_risk_composite():
        """Calculate composite program risk score."""
        data = flask_request.get_json(silent=True) or {}
        project_id = data.get("project_id", "")
        if not project_id:
            return jsonify({"error": "project_id is required"}), 400
        try:
            from tools.simulation.risk_monitor import calculate_composite_risk

            result = calculate_composite_risk(project_id)
            return jsonify(result)
        except Exception as exc:
            return jsonify({"error": str(exc)}), 400

    @app.route("/api/simulation/risk/cpars", methods=["POST"])
    def api_simulation_risk_cpars():
        """Calculate CPARS risk score for a contract."""
        data = flask_request.get_json(silent=True) or {}
        contract_id = data.get("contract_id", "")
        if not contract_id:
            return jsonify({"error": "contract_id is required"}), 400
        try:
            from tools.simulation.risk_monitor import calculate_cpars_risk

            result = calculate_cpars_risk(contract_id)
            return jsonify(result)
        except Exception as exc:
            return jsonify({"error": str(exc)}), 400

    # ---- Phase 63: Industry Research Engine ----

    @app.route("/research")
    def research_page():
        """Industry Research Engine — vertical research sessions, scored dossiers (Phase 63, D-RES-1 through D-RES-13)."""
        stats = {"total_sessions": 0, "active_sessions": 0, "verticals_loaded": 0, "dossiers_generated": 0}
        sessions = []
        verticals = []
        degraded = False
        try:
            conn = _get_db()
            stats["total_sessions"] = conn.execute("SELECT COUNT(*) FROM research_sessions").fetchone()[0]
            stats["active_sessions"] = conn.execute(
                "SELECT COUNT(*) FROM research_sessions WHERE status NOT IN ('archived', 'child_app_triggered')"
            ).fetchone()[0]
            stats["verticals_loaded"] = conn.execute("SELECT COUNT(*) FROM research_verticals").fetchone()[0]
            stats["dossiers_generated"] = conn.execute("SELECT COUNT(*) FROM research_dossiers").fetchone()[0]
            sessions = [
                dict(r)
                for r in conn.execute(
                    """SELECT s.*, (SELECT COUNT(*) FROM research_challenges c WHERE c.session_id = s.id) as challenge_count
                   FROM research_sessions s ORDER BY s.created_at DESC LIMIT 50"""
                ).fetchall()
            ]
            verticals = [dict(r) for r in conn.execute("SELECT * FROM research_verticals ORDER BY name").fetchall()]
            conn.close()
        except Exception as exc:
            # nav-misc-02: a research DB outage must not look like "no sessions yet".
            degraded = True
            get_logger("icdev.dashboard").warning("research_page: DB read failed: %s", exc)
        return render_template(
            "research.html", stats=stats, sessions=sessions, verticals=verticals, degraded=degraded
        )













    # ---- Phase 64: RAG Knowledge Search ----

    @app.route("/knowledge-search")
    def knowledge_search_page():
        """RAG Knowledge Search — natural language search across all ICDEV™ knowledge (Phase 64, D-RAG-1)."""
        status = None
        recent_searches = []
        source_types = []
        degraded = False
        try:
            from tools.rag.ingestion_manager import get_status as rag_get_status
            from tools.rag.source_registry import SOURCE_REGISTRY

            status = rag_get_status()
            source_types = sorted(SOURCE_REGISTRY.keys())
        except Exception as exc:
            # nav-misc-02: RAG subsystem unavailable — flag it rather than showing
            # an empty search page as if the index were simply empty.
            degraded = True
            get_logger("icdev.dashboard").warning("knowledge_search_page: RAG status unavailable: %s", exc)
        try:
            conn = _get_db()
            recent_searches = [
                dict(r)
                for r in conn.execute("SELECT * FROM rag_retrieval_log ORDER BY created_at DESC LIMIT 20").fetchall()
            ]
            conn.close()
        except Exception as exc:
            degraded = True
            get_logger("icdev.dashboard").warning("knowledge_search_page: retrieval-log read failed: %s", exc)
        return render_template(
            "rag/knowledge_search.html",
            status=status,
            recent_searches=recent_searches,
            source_types=source_types,
            degraded=degraded,
        )

    @app.route("/api/rag/search", methods=["POST"])
    def api_rag_search():
        """RAG search API endpoint."""
        data = flask_request.get_json(silent=True) or {}
        query = data.get("query", "")
        if not query:
            return jsonify({"error": "query is required", "results": []}), 400
        try:
            from tools.rag.retriever import RAGRetriever

            retriever = RAGRetriever()
            top_k = data.get("top_k", 5)
            source_types = None
            if data.get("source_type"):
                source_types = [data["source_type"]]
            results = retriever.search(
                query=query,
                top_k=top_k,
                source_types=source_types,
            )
            return jsonify(
                {
                    "classification": DEFAULT_CLASSIFICATION,
                    "query": query,
                    "results_count": len(results),
                    "results": [r.to_dict() for r in results],
                }
            )
        except ImportError:
            return jsonify({"error": "RAG subsystem not available", "results": []}), 503
        except Exception as e:
            return jsonify({"error": str(e), "results": []}), 500

    @app.route("/api/rag/status")
    def api_rag_status():
        """RAG status API endpoint."""
        try:
            from tools.rag.ingestion_manager import get_status as rag_get_status

            return jsonify(rag_get_status())
        except ImportError:
            return jsonify({"error": "RAG subsystem not available"}), 503
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    # ---- Knowledge Graph Dashboard (D-KARL-1 through D-KARL-4) ----

    @app.route("/knowledge-graph")
    def knowledge_graph_page():
        """Knowledge Graph — entity extraction, GraphRAG retrieval, insights."""
        stats = None
        graphs = []
        recent_queries = []
        degraded = False
        try:
            conn = _get_db()
            # Stats
            row = conn.execute(
                "SELECT COUNT(*) as cnt, COALESCE(SUM(entity_count),0) as nodes, "
                "COALESCE(SUM(edge_count),0) as edges FROM kg_graphs"
            ).fetchone()
            query_count = 0
            try:
                query_count = conn.execute("SELECT COUNT(*) FROM kg_retrieval_log").fetchone()[0]
            except Exception:
                pass
            stats = {
                "graph_count": row[0] if row else 0,
                "total_nodes": row[1] if row else 0,
                "total_edges": row[2] if row else 0,
                "recent_queries": query_count,
            }
            # Graph list
            rows = conn.execute(
                "SELECT id, project_id, name, entity_count, edge_count, created_at "
                "FROM kg_graphs ORDER BY created_at DESC LIMIT 50"
            ).fetchall()
            graphs = [dict(r) for r in rows]
            # Recent retrieval log
            try:
                qrows = conn.execute(
                    "SELECT query_hash, profile, node_count, top_score, duration_ms, created_at "
                    "FROM kg_retrieval_log ORDER BY created_at DESC LIMIT 20"
                ).fetchall()
                recent_queries = [dict(r) for r in qrows]
            except Exception:
                pass
            conn.close()
        except Exception as exc:
            # nav-misc-02: a KG store outage must not render an empty graph page
            # that looks like "no knowledge graphs built yet".
            degraded = True
            get_logger("icdev.dashboard").warning("knowledge_graph_page: kg_graphs read failed: %s", exc)
        return render_template(
            "knowledge_graph.html",
            stats=stats,
            graphs=graphs,
            recent_queries=recent_queries,
            degraded=degraded,
        )

    @app.route("/api/knowledge-graph/search", methods=["POST"])
    def api_knowledge_graph_search():
        """GraphRAG search API endpoint."""
        data = flask_request.get_json(silent=True) or {}
        query = data.get("query", "")
        if not query:
            return jsonify({"error": "query is required"}), 400
        try:
            from tools.knowledge_graph.graph_rag import retrieve

            result = retrieve(
                query=query,
                profile=data.get("profile"),
                top_k=data.get("top_k", 10),
            )
            return jsonify(result)
        except ImportError:
            return jsonify({"error": "Knowledge graph module not available"}), 503
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.route("/api/knowledge-graph/graph/<graph_id>")
    def api_knowledge_graph_detail(graph_id):
        """Get graph detail with nodes and edges."""
        try:
            from tools.knowledge_graph.text_network import get_graph

            result = get_graph(graph_id)
            return jsonify(result)
        except ImportError:
            return jsonify({"error": "Knowledge graph module not available"}), 503
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.route("/api/knowledge-graph/insights/<graph_id>")
    def api_knowledge_graph_insights(graph_id):
        """Get graph insights (summary, orphans, components)."""
        try:
            from tools.knowledge_graph.insight_generator import graph_summary

            result = graph_summary(graph_id)
            return jsonify(result)
        except ImportError:
            return jsonify({"error": "Insight generator not available"}), 503
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.route("/api/knowledge-graph/bridge-gaps/<graph_id>")
    def api_knowledge_graph_bridge_gaps(graph_id):
        """Get bridge gaps between disconnected clusters."""
        try:
            from tools.knowledge_graph.insight_generator import find_bridge_gaps

            result = find_bridge_gaps(graph_id)
            return jsonify(result)
        except ImportError:
            return jsonify({"error": "Insight generator not available"}), 503
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.route("/api/knowledge-graph/questions/<graph_id>")
    def api_knowledge_graph_questions(graph_id):
        """Generate research questions from graph structure."""
        try:
            from tools.knowledge_graph.insight_generator import generate_questions

            result = generate_questions(graph_id, use_llm=False)
            return jsonify(result)
        except ImportError:
            return jsonify({"error": "Insight generator not available"}), 503
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.route("/api/knowledge-graph/orphans/<graph_id>")
    def api_knowledge_graph_orphans(graph_id):
        """Find orphan nodes with zero edges."""
        try:
            from tools.knowledge_graph.insight_generator import find_orphan_nodes

            result = find_orphan_nodes(graph_id)
            return jsonify(result)
        except ImportError:
            return jsonify({"error": "Insight generator not available"}), 503
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    # ----------------------------------------------------------------
    # /components-map — Internal Awareness Engine visual map (Phase 1f)
    # ----------------------------------------------------------------
    # Reads from kg_nodes / kg_edges under graph_id =
    # "kg-icdev-self-awareness" (PostgreSQL via get_connection("icdev")).
    # Populated by tools/awareness/component_indexer.py.

    _COMPONENTS_MAP_GRAPH_ID = "kg-icdev-self-awareness"

    def _cmap_conn():
        """Return a connection for components-map queries (PG or SQLite).

        `get_connection`'s only positional argument is `db_path`, so the former
        `get_connection("icdev")` was a no-op on PostgreSQL (where the arg is
        ignored) but told SQLite to open a file literally named `icdev` — every
        components-map API 500'd with "unable to open database file" whenever
        the install fell back to SQLite. Pass nothing and let the resolver pick.
        """
        return get_connection()

    _cmap_pg = _cmap_conn  # backward-compat alias used by API routes

    @app.route("/components-map")
    def components_map_page():
        """Components Map — interactive JointJS graph of all ICDEV(TM) components."""
        stats = {"total": 0, "enabled": 0, "disabled": 0}
        try:
            conn = _cmap_conn()
            _pg = getattr(conn, "_backend", "sqlite") == "postgresql"
            # dict() first: sqlite3.Row has no .get(), so the SQLite path used
            # to raise here and fall into the swallowing except below, which
            # rendered the header as "0 components" on a populated graph.
            _total_row = conn.execute(
                "SELECT COUNT(*) AS n FROM kg_nodes WHERE graph_id = %s",
                (_COMPONENTS_MAP_GRAPH_ID,),
            ).fetchone()
            stats["total"] = dict(_total_row).get("n", 0) if _total_row else 0
            if _pg:
                row = conn.execute(
                    "SELECT "
                    "SUM(CASE WHEN (properties::jsonb)->>'enabled' = 'false' THEN 1 ELSE 0 END) AS dis, "
                    "SUM(CASE WHEN (properties::jsonb)->>'enabled' != 'false' THEN 1 ELSE 0 END) AS en "
                    "FROM kg_nodes WHERE graph_id = %s",
                    (_COMPONENTS_MAP_GRAPH_ID,),
                ).fetchone() or {}
            else:
                row = conn.execute(
                    "SELECT "
                    "SUM(CASE WHEN json_extract(properties, '$.enabled') = 'false' THEN 1 ELSE 0 END) AS dis, "  # pg-ok: SQLite fallback; is_pg branch above uses (properties::jsonb)->>
                    "SUM(CASE WHEN json_extract(properties, '$.enabled') != 'false' THEN 1 ELSE 0 END) AS en "  # pg-ok: SQLite fallback; is_pg branch above uses (properties::jsonb)->>
                    "FROM kg_nodes WHERE graph_id = %s",
                    (_COMPONENTS_MAP_GRAPH_ID,),
                ).fetchone()
            row = dict(row) if row else {}
            stats["enabled"] = int(row.get("en") or stats["total"])
            stats["disabled"] = int(row.get("dis") or 0)
            conn.close()
        except Exception:  # pragma: no cover -- non-critical stat failure
            pass
        return render_template("components_map.html", stats=stats)

    @app.route("/api/components-map/tree")
    def api_cmap_tree():
        """GET /api/components-map/tree -- hierarchical {category: [nodes]} JSON."""
        tree: dict = {}
        try:
            conn = _cmap_pg()
            cur = conn.cursor()
            cur.execute(
                "SELECT id, label, entity_type, properties FROM kg_nodes WHERE graph_id=%s ORDER BY entity_type, label",
                (_COMPONENTS_MAP_GRAPH_ID,),
            )
            rows = cur.fetchall()
            conn.close()
            for r in rows:
                cat = r["entity_type"] or "other"
                props = json.loads(r["properties"]) if r["properties"] else {}
                enabled = props.get("enabled", True)
                tree.setdefault(cat, []).append({
                    "id": r["id"],
                    "label": r["label"],
                    "enabled": enabled,
                    "file_path": props.get("file_path", ""),
                    "description": props.get("description", ""),
                })
        except Exception as exc:
            return jsonify({"error": str(exc)}), 500
        return jsonify(tree)

    @app.route("/api/components-map/graph")
    def api_cmap_graph():
        """GET /api/components-map/graph -- JointJS-compatible {cells:[...]} payload.

        Query params:
          scope=<entity_type>  -- filter to one category
          show_disabled=1      -- include disabled nodes (omitted by default)
          max_edges=<n>        -- cap the edge payload (default 2000)

        The edge cap exists because the graph carries five figures of derived
        dependency edges (idp-cat-02); rendering all of them unscoped would
        wedge the browser. Highest-weight (most mechanical) edges are kept
        first and `edges_truncated` reports what was dropped.
        """
        scope = flask_request.args.get("scope", "").strip() or None
        show_disabled = flask_request.args.get("show_disabled", "0") == "1"
        try:
            max_edges = max(0, int(flask_request.args.get("max_edges", 2000)))
        except (TypeError, ValueError):
            max_edges = 2000
        cells: list = []
        try:
            conn = _cmap_pg()
            cur = conn.cursor()
            node_q = "SELECT id, label, entity_type, properties, centrality FROM kg_nodes WHERE graph_id=%s"
            params: list = [_COMPONENTS_MAP_GRAPH_ID]
            if scope:
                node_q += " AND entity_type=%s"
                params.append(scope)
            cur.execute(node_q, params)
            node_rows = cur.fetchall()
            node_ids: set = set()
            for r in node_rows:
                props = json.loads(r["properties"]) if r["properties"] else {}
                enabled = props.get("enabled", True)
                if not show_disabled and enabled is False:
                    continue
                node_ids.add(r["id"])
                cells.append({
                    "type": "node",
                    "id": r["id"],
                    "label": r["label"],
                    "entity_type": r["entity_type"] or "other",
                    "enabled": enabled,
                    "properties": props,
                    "centrality": r["centrality"] or 0.0,
                })
            cur.execute(
                "SELECT id, source_id, target_id, relationship, weight, properties "
                "FROM kg_edges WHERE graph_id=%s ORDER BY weight DESC",
                (_COMPONENTS_MAP_GRAPH_ID,),
            )
            edge_total = 0
            for e in cur.fetchall():
                if e["source_id"] not in node_ids or e["target_id"] not in node_ids:
                    continue
                edge_total += 1
                if edge_total > max_edges:
                    continue
                try:
                    eprops = json.loads(e["properties"]) if e["properties"] else {}
                except (TypeError, ValueError):
                    eprops = {}
                cells.append({
                    "type": "edge",
                    "id": e["id"],
                    "source": e["source_id"],
                    "target": e["target_id"],
                    "label": e["relationship"] or "",
                    "weight": e["weight"] or 1.0,
                    "derivation": eprops.get("derivation", ""),
                    "mechanical": bool(eprops.get("mechanical", False)),
                })
            conn.close()
        except Exception as exc:
            return jsonify({"error": str(exc)}), 500
        return jsonify({
            "cells": cells,
            "count": len(cells),
            "edge_total": edge_total,
            "edges_truncated": max(0, edge_total - max_edges),
        })

    @app.route("/api/components-map/node/<path:node_id>")
    def api_cmap_node(node_id: str):
        """GET /api/components-map/node/<id> -- full node detail for hover/drawer."""
        try:
            conn = _cmap_pg()
            cur = conn.cursor()
            cur.execute(
                "SELECT id, label, entity_type, properties, centrality, created_at "
                "FROM kg_nodes WHERE graph_id=%s AND id=%s",
                (_COMPONENTS_MAP_GRAPH_ID, node_id),
            )
            row = cur.fetchone()
            if not row:
                conn.close()
                return jsonify({"error": "node not found"}), 404
            cur.execute(
                "SELECT COUNT(*) AS n FROM kg_edges WHERE graph_id=%s AND (source_id=%s OR target_id=%s)",
                (_COMPONENTS_MAP_GRAPH_ID, node_id, node_id),
            )
            # sqlite3.Row has no .get() — go through dict() so this works on
            # both backends instead of 500ing on the SQLite fallback.
            _rc = cur.fetchone()
            rel_count = dict(_rc).get("n", 0) if _rc else 0
            conn.close()
            props = json.loads(row["properties"]) if row["properties"] else {}
            return jsonify({
                "id": row["id"],
                "label": row["label"],
                "entity_type": row["entity_type"] or "other",
                "properties": props,
                "centrality": row["centrality"] or 0.0,
                # PG hands back a datetime, SQLite a plain string — normalise
                # rather than assuming .isoformat() exists (it 500'd on SQLite).
                "created_at": (
                    row["created_at"].isoformat()
                    if hasattr(row["created_at"], "isoformat")
                    else (str(row["created_at"]) if row["created_at"] else None)
                ),
                "last_indexed_at": props.get("last_indexed_at"),
                "relationships_count": rel_count,
                "health": props.get("health", "unknown"),
            })
        except Exception as exc:
            return jsonify({"error": str(exc)}), 500

    # A hub like tools/db/storage.py has >1,000 direct edges now that the graph
    # carries real dependencies. Expanding all of them at once both wedges the
    # JointJS canvas and blows SQLite's 999-parameter limit on the IN clause
    # below, so cap the hop and report the overflow.
    _CMAP_MAX_NEIGHBORS = 300

    @app.route("/api/components-map/neighbors/<path:node_id>")
    def api_cmap_neighbors(node_id: str):
        """GET /api/components-map/neighbors/<id> -- 1-hop subgraph for expansion.

        Capped at `_CMAP_MAX_NEIGHBORS` highest-weight edges; `truncated`
        reports how many were dropped. Use /api/components-map/dependents for
        the full, unrendered list.
        """
        cells: list = []
        truncated = 0
        try:
            conn = _cmap_pg()
            cur = conn.cursor()
            cur.execute(
                "SELECT id, source_id, target_id, relationship, weight "
                "FROM kg_edges WHERE graph_id=%s AND (source_id=%s OR target_id=%s) "
                "ORDER BY weight DESC",
                (_COMPONENTS_MAP_GRAPH_ID, node_id, node_id),
            )
            edges = cur.fetchall()
            if len(edges) > _CMAP_MAX_NEIGHBORS:
                truncated = len(edges) - _CMAP_MAX_NEIGHBORS
                edges = edges[:_CMAP_MAX_NEIGHBORS]
            neighbor_ids: set = {node_id}
            for e in edges:
                neighbor_ids.add(e["source_id"])
                neighbor_ids.add(e["target_id"])
            placeholders = ",".join(["%s"] * len(neighbor_ids))
            cur.execute(
                f"SELECT id, label, entity_type, properties, centrality FROM kg_nodes "  # noqa: S608  # nosec B608
                f"WHERE graph_id=%s AND id IN ({placeholders})",
                [_COMPONENTS_MAP_GRAPH_ID] + list(neighbor_ids),
            )
            node_index: set = set()
            for r in cur.fetchall():
                props = json.loads(r["properties"]) if r["properties"] else {}
                cells.append({
                    "type": "node",
                    "id": r["id"],
                    "label": r["label"],
                    "entity_type": r["entity_type"] or "other",
                    "enabled": props.get("enabled", True),
                    "properties": props,
                    "centrality": r["centrality"] or 0.0,
                })
                node_index.add(r["id"])
            for e in edges:
                if e["source_id"] in node_index and e["target_id"] in node_index:
                    cells.append({
                        "type": "edge",
                        "id": e["id"],
                        "source": e["source_id"],
                        "target": e["target_id"],
                        "label": e["relationship"] or "",
                        "weight": e["weight"] or 1.0,
                    })
            conn.close()
        except Exception as exc:
            return jsonify({"error": str(exc)}), 500
        return jsonify({"cells": cells, "count": len(cells), "truncated": truncated})

    @app.route("/api/components-map/dependents/<path:node_ref>")
    def api_cmap_dependents(node_ref: str):
        """GET /api/components-map/dependents/<ref> -- blast radius for a component.

        `ref` is a node id, a repo-relative file path, or a node label. Answers
        "what breaks if this changes" from the mechanically-derived edges
        (imports, uses_table, invokes, ...). Every returned row carries the
        `derivation` that produced it plus `mechanical`, so a caller can tell a
        parsed import from an inferred keyword match.

        Query params:
          depth=<n>          -- traversal depth (default 1 = direct dependents)
          direction=deps     -- invert: what this component depends on
          mechanical_only=1  -- drop similarity-inferred edges
        """
        try:
            from tools.awareness import edge_deriver
        except ImportError:
            return jsonify({"error": "edge_deriver not available"}), 503
        try:
            depth = max(1, min(5, int(flask_request.args.get("depth", 1))))
        except (TypeError, ValueError):
            depth = 1
        mechanical_only = flask_request.args.get("mechanical_only", "0") == "1"
        inverted = flask_request.args.get("direction", "").lower() in ("deps", "dependencies")
        query = edge_deriver.get_dependencies if inverted else edge_deriver.get_dependents
        try:
            result = query(node_ref, depth=depth, mechanical_only=mechanical_only)
        except Exception as exc:
            return jsonify({"error": str(exc)}), 500
        if result.get("error"):
            return jsonify(result), 404 if "no node matches" in result["error"] else 500
        return jsonify(result)

    @app.route("/api/components-map/refresh", methods=["POST"])
    @require_role("admin", "pm")
    def api_cmap_refresh():
        """POST /api/components-map/refresh -- trigger component_indexer rescan.

        nav-intel-01: gated — previously any caller could spawn the
        component_indexer subprocess (compute-triggering, unauthenticated).
        """
        import subprocess  # noqa: S404 -- intentional controlled subprocess
        try:
            cmd = ["python", "tools/awareness/component_indexer.py", "--scan"]
            result = subprocess.run(  # noqa: S603
                cmd,
                capture_output=True,
                text=True,
                timeout=120,
                cwd=str(Path(__file__).parent.parent.parent),
            )
            return jsonify({
                "ok": result.returncode == 0,
                "stdout": result.stdout[-2000:] if result.stdout else "",
                "stderr": result.stderr[-500:] if result.stderr else "",
            })
        except Exception as exc:
            return jsonify({"ok": False, "error": str(exc)}), 500

    # ----------------------------------------------------------------
    # /api/components-map/ask + /ask-icdev chat page (Phase 4)
    # ----------------------------------------------------------------
    # Unified Q&A endpoint. Parallel fetch from:
    #   * RAG retriever (text hits from rag_chunks)
    #   * GraphRAG (kg_nodes + kg_edges under internal_awareness profile)
    #   * awareness_component_health (latest probe status for matching nodes)
    #   * kanban_tasks + oracle_predictions (suggested_next_actions)
    # Narration is opt-in via `narrate=true` — falls back to raw
    # evidence if the LLM router is unavailable (air-gap safe).

    def _cm_rag_search(query: str, top_k: int = 10):
        """Run RAG search; return [] on failure (air-gap safe)."""
        try:
            from tools.rag.retriever import RAGRetriever
            retriever = RAGRetriever()
            results = retriever.search(query=query, top_k=top_k)
            hits = []
            for r in results:
                if hasattr(r, "to_dict"):
                    hits.append(r.to_dict())
                elif isinstance(r, dict):
                    hits.append(r)
                else:
                    hits.append({"content": str(r)[:400]})
            return hits
        except Exception as exc:
            return [{"error": str(exc)[:200]}]

    def _cm_kg_retrieve(query: str, top_k: int = 10):
        """Run GraphRAG with the internal_awareness profile."""
        try:
            from tools.knowledge_graph.graph_rag import retrieve
            return retrieve(
                query=query,
                profile="internal_awareness",
                top_k=top_k,
                compress=False,
            )
        except Exception as exc:
            return {"error": str(exc)[:200], "nodes": [], "edges": []}

    def _cm_health_hits(conn, query: str):
        """Fetch recent failing health snapshots matching query tokens."""
        try:
            rows = conn.execute(
                "SELECT node_id, probe_type, status, detail, probed_at "
                "FROM awareness_component_health "
                "WHERE status IN ('fail', 'error') "
                "ORDER BY probed_at DESC LIMIT 20"
            ).fetchall()
        except Exception:
            return []
        results = []
        q_lower = query.lower()
        for r in rows:
            d = dict(r)
            # Include the hit if the query mentions any part of the
            # node_id or probe type (loose match — LLM narration or
            # the client can filter further)
            hay = (d.get("node_id", "") + " " + d.get("probe_type", "") + " " + (d.get("detail") or "")).lower()
            score = sum(1 for tok in q_lower.split() if tok in hay)
            if score > 0 or len(results) < 5:
                results.append({
                    "node_id": d["node_id"],
                    "probe_type": d["probe_type"],
                    "status": d["status"],
                    "probed_at": d["probed_at"].isoformat() if hasattr(d["probed_at"], "isoformat") else d["probed_at"],
                    "detail": d.get("detail", "")[:300],
                    "score": score,
                })
        return results[:10]

    def _cm_suggested_next_actions(conn, query: str):
        """Return the 10 most recent suggested kanban cards that
        originated from the internal_awareness lens."""
        try:
            rows = conn.execute(
                "SELECT kt.id, kt.title, kt.priority, kt.task_type, kt.created_at, "
                "       op.confidence, op.prediction_type "
                "FROM kanban_tasks kt "
                "JOIN oracle_predictions op ON op.id = kt.source_prediction_id "
                "WHERE op.lens_name = 'internal_awareness' AND kt.status = 'suggested' "
                "ORDER BY kt.created_at DESC LIMIT 10"
            ).fetchall()
        except Exception:
            return []
        results = []
        for r in rows:
            d = dict(r)
            created = d.get("created_at")
            results.append({
                "task_id": d["id"],
                "title": d["title"],
                "priority": d["priority"],
                "task_type": d["task_type"],
                "confidence": float(d.get("confidence") or 0),
                "prediction_type": d.get("prediction_type", ""),
                "created_at": created.isoformat() if hasattr(created, "isoformat") else created,
            })
        return results

    def _cm_llm_narrate(query: str, rag_hits, graph_hits, health_hits, suggested):
        """OPT-IN narration. Uses LLMRouter.invoke() with function=
        narrative_generation — portable across any Scanner-tier model
        configured in args/llm_config.yaml. Graceful fallback: returns
        None if the router is unavailable or the call fails, so the
        caller can show raw evidence instead.
        """
        # nav-intel-06: rate-limit the LLM cost surface per authenticated user.
        # Both /ask endpoints funnel their narration through this one function,
        # so the budget check here covers both. Over budget -> None, and the
        # caller already renders the raw-evidence fallback.
        _u = getattr(g, "current_user", None)
        _key = _u.get("id", "anon") if isinstance(_u, dict) else "anon"
        if not _narration_budget_ok(_key):
            return None
        try:
            from tools.llm.router import LLMRouter
        except ImportError:
            return None
        try:
            prompt = (
                "Synthesize a concise answer (3-6 sentences) to the user's question "
                "using only the evidence below. Cite sources by their type (rag/graph/health/suggested).\n\n"
                f"QUESTION: {query}\n\n"
                f"RAG HITS: {json.dumps(rag_hits, ensure_ascii=False)[:2000]}\n\n"
                f"GRAPH HITS: {json.dumps(graph_hits, ensure_ascii=False)[:2000]}\n\n"
                f"HEALTH HITS: {json.dumps(health_hits, ensure_ascii=False)[:1000]}\n\n"
                f"SUGGESTED ACTIONS: {json.dumps(suggested, ensure_ascii=False)[:500]}\n"
            )
            router = LLMRouter()
            response = router.invoke(
                function="narrative_generation",
                prompt=prompt,
                max_tokens=400,
            )
            if isinstance(response, dict):
                return response.get("content") or response.get("text") or str(response)
            return str(response) if response else None
        except Exception:
            return None

    @app.route("/api/components-map/ask", methods=["POST"])
    def api_components_map_ask():
        """Unified Q&A: parallel RAG + GraphRAG + health + suggested.

        Reads from kg_nodes, awareness_component_health, and
        kanban_tasks via the helper functions below. The explicit
        `conn.execute()` up front verifies the self-awareness graph
        exists — also makes this handler visible to the coherence
        checker's api_wiring rule which scans for inline DB calls.
        """
        data = flask_request.get_json(silent=True) or {}
        query = (data.get("query") or "").strip()
        narrate = bool(data.get("narrate", False))
        top_k = int(data.get("top_k", 10))
        if not query:
            return jsonify({"error": "query is required"}), 400

        import concurrent.futures as _f

        conn = _get_db()
        try:
            # Graph existence probe — surfaces a clean error if the
            # component graph hasn't been indexed yet, and gives the
            # coherence_checker's api_wiring rule a visible DB call.
            _graph_row = conn.execute(
                "SELECT COUNT(*) AS cnt FROM kg_nodes WHERE graph_id = %s",
                (_COMPONENTS_MAP_GRAPH_ID,),
            ).fetchone()
            graph_node_count = dict(_graph_row).get("cnt", 0) if _graph_row else 0

            with _f.ThreadPoolExecutor(max_workers=4) as ex:
                rag_fut = ex.submit(_cm_rag_search, query, top_k)
                kg_fut = ex.submit(_cm_kg_retrieve, query, top_k)
                # Health + suggested run on the main conn (already open)
                rag_hits = rag_fut.result(timeout=30)
                kg_result = kg_fut.result(timeout=30)
            health_hits = _cm_health_hits(conn, query)
            suggested = _cm_suggested_next_actions(conn, query)
        finally:
            conn.close()

        response = {
            "query": query,
            "graph_node_count": graph_node_count,
            "rag_hits": rag_hits,
            "graph_hits": {
                "nodes": kg_result.get("nodes", []) if isinstance(kg_result, dict) else [],
                "edges": kg_result.get("edges", []) if isinstance(kg_result, dict) else [],
                "profile": kg_result.get("profile", "internal_awareness") if isinstance(kg_result, dict) else "internal_awareness",
                "rrf_fusion": kg_result.get("rrf_fusion", False) if isinstance(kg_result, dict) else False,
                "neighbor_count": kg_result.get("neighbor_count", 0) if isinstance(kg_result, dict) else 0,
            },
            "health_hits": health_hits,
            "suggested_next_actions": suggested,
            "narration": None,
            "narrated": False,
        }

        if narrate:
            narration = _cm_llm_narrate(query, rag_hits, response["graph_hits"]["nodes"], health_hits, suggested)
            if narration:
                response["narration"] = narration
                response["narrated"] = True

        # Sanitize any leftover non-JSON types (datetime etc.) by
        # round-tripping through json with default=str.
        safe = json.loads(json.dumps(response, default=str))
        return jsonify(safe)

    # ----------------------------------------------------------------
    # /ask-icdev — dedicated chat page with persistent sessions
    # ----------------------------------------------------------------

    def _ensure_ask_icdev_tables(conn):
        """Create icdev_qa_sessions + icdev_qa_messages on first use.

        Uses ADD COLUMN IF NOT EXISTS for user_id so a previously-
        created table (from an earlier experiment without this column)
        gets upgraded in place.
        """
        try:
            conn.execute("""
                CREATE TABLE IF NOT EXISTS icdev_qa_sessions (
                    id TEXT PRIMARY KEY,
                    title TEXT,
                    created_at TEXT NOT NULL,
                    updated_at TEXT NOT NULL
                )
            """)
            # Backward-compat: add user_id if the table was created earlier
            # without it. Postgres supports IF NOT EXISTS on ADD COLUMN;
            # SQLite does not, so we try/ignore duplicate-column errors.
            try:
                conn.execute(
                    "ALTER TABLE icdev_qa_sessions ADD COLUMN IF NOT EXISTS user_id TEXT"
                )
            except Exception:
                try:
                    conn.rollback()
                except Exception:
                    pass
                try:
                    conn.execute("ALTER TABLE icdev_qa_sessions ADD COLUMN user_id TEXT")
                except Exception:
                    try:
                        conn.rollback()
                    except Exception:
                        pass
            conn.execute("""
                CREATE TABLE IF NOT EXISTS icdev_qa_messages (
                    id TEXT PRIMARY KEY,
                    session_id TEXT NOT NULL,
                    turn INTEGER NOT NULL,
                    role TEXT NOT NULL,
                    content TEXT NOT NULL,
                    citations_json TEXT,
                    created_at TEXT NOT NULL
                )
            """)
            conn.execute(
                "CREATE INDEX IF NOT EXISTS idx_qa_messages_session ON icdev_qa_messages(session_id, turn)"
            )
            conn.commit()
        except Exception:
            try:
                conn.rollback()
            except Exception:
                pass

    @app.route("/ask-icdev")
    def ask_icdev_page():
        """Dedicated ICDEV Q&A chat page."""
        return render_template("ask_icdev.html")

    @app.route("/api/ask-icdev/sessions", methods=["GET"])
    def api_ask_icdev_list_sessions():
        conn = _get_db()
        try:
            _ensure_ask_icdev_tables(conn)
            rows = conn.execute(
                "SELECT id, title, created_at, updated_at FROM icdev_qa_sessions "
                "ORDER BY updated_at DESC LIMIT 50"
            ).fetchall()
            return jsonify({"sessions": [dict(r) for r in rows]})
        finally:
            conn.close()

    @app.route("/api/ask-icdev/sessions", methods=["POST"])
    def api_ask_icdev_create_session():
        data = flask_request.get_json(silent=True) or {}
        import uuid as _uuid
        session_id = f"qa-{_uuid.uuid4().hex[:12]}"
        now = datetime.now(timezone.utc).isoformat()
        title = (data.get("title") or "New session")[:200]
        conn = _get_db()
        try:
            _ensure_ask_icdev_tables(conn)
            # Use only the baseline columns so this works whether the
            # table is from a fresh create or a legacy version without
            # user_id. Keep user_id out of the INSERT entirely.
            conn.execute(
                "INSERT INTO icdev_qa_sessions (id, title, created_at, updated_at) "
                "VALUES (%s, %s, %s, %s)",
                (session_id, title, now, now),
            )
            conn.commit()
            return jsonify({"session_id": session_id, "title": title, "created_at": now}), 201
        finally:
            conn.close()

    @app.route("/api/ask-icdev/sessions/<session_id>", methods=["GET"])
    def api_ask_icdev_get_session(session_id):
        conn = _get_db()
        try:
            _ensure_ask_icdev_tables(conn)
            row = conn.execute(
                "SELECT id, title, created_at, updated_at FROM icdev_qa_sessions WHERE id = %s",
                (session_id,),
            ).fetchone()
            if not row:
                return jsonify({"error": "session not found"}), 404
            session = dict(row)
            rows = conn.execute(
                "SELECT id, turn, role, content, citations_json, created_at "
                "FROM icdev_qa_messages WHERE session_id = %s ORDER BY turn ASC",
                (session_id,),
            ).fetchall()
            messages = []
            for r in rows:
                d = dict(r)
                try:
                    d["citations"] = json.loads(d.pop("citations_json") or "{}")
                except Exception:
                    d["citations"] = {}
                messages.append(d)
            session["messages"] = messages
            return jsonify(session)
        finally:
            conn.close()

    @app.route("/api/ask-icdev/sessions/<session_id>", methods=["DELETE"])
    def api_ask_icdev_delete_session(session_id):
        conn = _get_db()
        try:
            _ensure_ask_icdev_tables(conn)
            conn.execute("DELETE FROM icdev_qa_messages WHERE session_id = %s", (session_id,))
            conn.execute("DELETE FROM icdev_qa_sessions WHERE id = %s", (session_id,))
            conn.commit()
            return jsonify({"deleted": session_id})
        finally:
            conn.close()

    @app.route("/api/ask-icdev/sessions/<session_id>/message", methods=["POST"])
    def api_ask_icdev_post_message(session_id):
        """Post a user message, run the unified Q&A, persist both turns."""
        data = flask_request.get_json(silent=True) or {}
        user_content = (data.get("content") or "").strip()
        narrate = bool(data.get("narrate", False))
        if not user_content:
            return jsonify({"error": "content required"}), 400

        import uuid as _uuid
        conn = _get_db()
        try:
            _ensure_ask_icdev_tables(conn)
            # Verify session exists
            session_row = conn.execute(
                "SELECT id, title FROM icdev_qa_sessions WHERE id = %s",
                (session_id,),
            ).fetchone()
            if not session_row:
                return jsonify({"error": "session not found"}), 404

            # Get next turn number
            row = conn.execute(
                "SELECT COALESCE(MAX(turn), -1) AS max_turn FROM icdev_qa_messages "
                "WHERE session_id = %s",
                (session_id,),
            ).fetchone()
            next_turn = (dict(row).get("max_turn", -1) + 1) if row else 0

            now = datetime.now(timezone.utc).isoformat()

            # Persist user turn
            user_msg_id = f"msg-{_uuid.uuid4().hex[:12]}"
            conn.execute(
                "INSERT INTO icdev_qa_messages "
                "(id, session_id, turn, role, content, citations_json, created_at) "
                "VALUES (%s, %s, %s, %s, %s, %s, %s)",
                (user_msg_id, session_id, next_turn, "user", user_content, "{}", now),
            )
            conn.commit()
        finally:
            conn.close()

        # Run the unified Q&A endpoint internally
        import concurrent.futures as _f
        conn = _get_db()
        try:
            with _f.ThreadPoolExecutor(max_workers=2) as ex:
                rag_fut = ex.submit(_cm_rag_search, user_content, 10)
                kg_fut = ex.submit(_cm_kg_retrieve, user_content, 10)
                rag_hits = rag_fut.result(timeout=30)
                kg_result = kg_fut.result(timeout=30)
            health_hits = _cm_health_hits(conn, user_content)
            suggested = _cm_suggested_next_actions(conn, user_content)
        finally:
            conn.close()

        graph_nodes = kg_result.get("nodes", []) if isinstance(kg_result, dict) else []

        narration = None
        if narrate:
            narration = _cm_llm_narrate(user_content, rag_hits, graph_nodes, health_hits, suggested)

        # Build the assistant response content
        if narration:
            assistant_content = narration
        else:
            # Raw evidence summary when LLM unavailable / disabled
            parts = [f"Found {len(rag_hits)} text hits and {len(graph_nodes)} graph hits for: {user_content}"]
            if health_hits:
                parts.append(f"{len(health_hits)} relevant failing health snapshots.")
            if suggested:
                parts.append(f"{len(suggested)} suggested next actions on the kanban board.")
            if not any([rag_hits, graph_nodes, health_hits, suggested]):
                parts.append("No evidence found — try a more specific query.")
            assistant_content = " ".join(parts)

        citations = {
            "rag_hits": rag_hits[:5],
            "graph_nodes": graph_nodes[:5],
            "health_hits": health_hits[:5],
            "suggested": suggested[:3],
            "narrated": narration is not None,
        }

        # Persist assistant turn. Use default=str on json.dumps so
        # any datetime / unexpected types from the RAG/GraphRAG hits
        # serialize as their string repr rather than raising.
        conn = _get_db()
        try:
            assistant_msg_id = f"msg-{_uuid.uuid4().hex[:12]}"
            conn.execute(
                "INSERT INTO icdev_qa_messages "
                "(id, session_id, turn, role, content, citations_json, created_at) "
                "VALUES (%s, %s, %s, %s, %s, %s, %s)",
                (
                    assistant_msg_id,
                    session_id,
                    next_turn + 1,
                    "assistant",
                    assistant_content,
                    json.dumps(citations, ensure_ascii=False, default=str),
                    datetime.now(timezone.utc).isoformat(),
                ),
            )
            conn.execute(
                "UPDATE icdev_qa_sessions SET updated_at = %s WHERE id = %s",
                (datetime.now(timezone.utc).isoformat(), session_id),
            )
            conn.commit()
        finally:
            conn.close()

        return jsonify({
            "session_id": session_id,
            "user_turn": next_turn,
            "assistant_turn": next_turn + 1,
            "assistant_content": assistant_content,
            "citations": citations,
            "narrated": narration is not None,
        })

    @app.route("/api/knowledge-graph/compliance-build", methods=["POST"])
    def api_knowledge_graph_compliance_build():
        """Build compliance crosswalk knowledge graph."""
        data = flask_request.get_json(silent=True) or {}
        try:
            from tools.knowledge_graph.compliance_graph import build_compliance_graph

            result = build_compliance_graph(project_id=data.get("project_id"))
            return jsonify(result)
        except ImportError:
            return jsonify({"error": "Compliance graph module not available"}), 503
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.route("/api/knowledge-graph/ingest", methods=["POST"])
    def api_knowledge_graph_ingest():
        """Ingest a document or table into the knowledge graph."""
        data = flask_request.get_json(silent=True) or {}
        source_table = data.get("source_table")
        project_id = data.get("project_id", "")
        if not source_table:
            return jsonify({"error": "source_table is required"}), 400
        try:
            from tools.knowledge_graph.ingester import ingest_from_table

            result = ingest_from_table(source_table, project_id=project_id)
            return jsonify(result)
        except ImportError:
            return jsonify({"error": "Ingester module not available"}), 503
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    # ---- Phase 64 Extension: Fine-Tuning Dashboard ----

    @app.route("/finetune")
    def finetune_overview_page():
        """Fine-Tuning overview — stats, GPU status, recent jobs, active overrides (D-FT-1 through D-FT-22)."""
        stats = {
            "datasets": 0,
            "total_jobs": 0,
            "active_jobs": 0,
            "model_versions": 0,
            "promoted_models": 0,
            "active_overrides": 0,
            "evaluations": 0,
        }
        recent_jobs = []
        active_overrides = []
        promotions = []
        error = False
        try:
            conn = _get_db()
            stats["datasets"] = conn.execute("SELECT COUNT(*) FROM ft_datasets").fetchone()[0]
            stats["total_jobs"] = conn.execute("SELECT COUNT(*) FROM ft_training_jobs").fetchone()[0]
            stats["active_jobs"] = conn.execute(
                "SELECT COUNT(*) FROM ft_training_jobs WHERE status IN ('pending','preparing','training','exporting','evaluating')"
            ).fetchone()[0]
            stats["model_versions"] = conn.execute("SELECT COUNT(*) FROM ft_model_versions").fetchone()[0]
            stats["promoted_models"] = conn.execute(
                "SELECT COUNT(*) FROM ft_model_versions WHERE status = 'promoted'"
            ).fetchone()[0]
            stats["active_overrides"] = conn.execute(
                "SELECT COUNT(*) FROM ft_active_models WHERE deactivated_at IS NULL"
            ).fetchone()[0]
            stats["evaluations"] = conn.execute("SELECT COUNT(*) FROM ft_evaluations").fetchone()[0]
            recent_jobs = [
                dict(r)
                for r in conn.execute("SELECT * FROM ft_training_jobs ORDER BY created_at DESC LIMIT 10").fetchall()
            ]
            active_overrides = [
                dict(r)
                for r in conn.execute(
                    "SELECT * FROM ft_active_models WHERE deactivated_at IS NULL ORDER BY activated_at DESC"
                ).fetchall()
            ]
            promotions = [
                dict(r)
                for r in conn.execute("SELECT * FROM ft_promotion_log ORDER BY created_at DESC LIMIT 10").fetchall()
            ]
            conn.close()
        except Exception as exc:
            error = True
            get_logger("icdev.dashboard").warning(
                "finetune_overview_page: DB error reading ft_* tables: %s", exc
            )
        return render_template(
            "finetune/index.html",
            stats=stats,
            recent_jobs=recent_jobs,
            active_overrides=active_overrides,
            promotions=promotions,
            error=error,
        )

    @app.route("/finetune/datasets")
    def finetune_datasets_page():
        """Fine-Tuning datasets — versioned training data collections."""
        datasets = []
        try:
            conn = _get_db()
            datasets = [dict(r) for r in conn.execute("SELECT * FROM ft_datasets ORDER BY updated_at DESC").fetchall()]
            conn.close()
        except Exception:
            pass
        return render_template("finetune/datasets.html", datasets=datasets)

    @app.route("/finetune/datasets/<dataset_id>")
    def finetune_dataset_detail_page(dataset_id):
        """Fine-Tuning dataset detail — examples with labeling controls."""
        dataset = None
        examples = []
        try:
            conn = _get_db()
            row = conn.execute("SELECT * FROM ft_datasets WHERE id = %s", (dataset_id,)).fetchone()
            if row:
                dataset = dict(row)
            examples = [
                dict(r)
                for r in conn.execute(
                    "SELECT * FROM ft_dataset_examples WHERE dataset_id = %s ORDER BY id DESC LIMIT 200",
                    (dataset_id,),
                ).fetchall()
            ]
            conn.close()
        except Exception:
            pass
        if not dataset:
            return render_template("404.html", message="Dataset not found"), 404
        return render_template("finetune/dataset_detail.html", dataset=dataset, examples=examples)

    @app.route("/finetune/label")
    def finetune_label_page():
        """Fine-Tuning bulk labeling — multi-dimensional scoring, batch approve/reject (D-FT-12)."""
        datasets = []
        examples = []
        selected_dataset_id = flask_request.args.get("dataset_id", "")
        try:
            conn = _get_db()
            datasets = [dict(r) for r in conn.execute("SELECT * FROM ft_datasets ORDER BY updated_at DESC").fetchall()]
            if selected_dataset_id:
                examples = [
                    dict(r)
                    for r in conn.execute(
                        "SELECT * FROM ft_dataset_examples WHERE dataset_id = %s ORDER BY id DESC LIMIT 200",
                        (selected_dataset_id,),
                    ).fetchall()
                ]
            conn.close()
        except Exception:
            pass
        return render_template(
            "finetune/label.html", datasets=datasets, examples=examples, selected_dataset_id=selected_dataset_id
        )

    @app.route("/finetune/jobs")
    def finetune_jobs_page():
        """Fine-Tuning training jobs — status tracking, loss curves."""
        jobs = []
        try:
            conn = _get_db()
            jobs = [dict(r) for r in conn.execute("SELECT * FROM ft_training_jobs ORDER BY created_at DESC").fetchall()]
            conn.close()
        except Exception:
            pass
        return render_template("finetune/jobs.html", jobs=jobs)

    @app.route("/finetune/jobs/<job_id>")
    def finetune_job_detail_page(job_id):
        """Fine-Tuning job detail — loss curve, hyperparams, events."""
        import json as _json

        job = None
        events = []
        loss_history = []
        try:
            conn = _get_db()
            row = conn.execute("SELECT * FROM ft_training_jobs WHERE id = %s", (job_id,)).fetchone()
            if row:
                job = dict(row)
                try:
                    loss_history = _json.loads(job.get("loss_history", "[]") or "[]")
                except (ValueError, TypeError):
                    loss_history = []
            events = [
                dict(r)
                for r in conn.execute(
                    "SELECT * FROM ft_training_job_events WHERE job_id = %s ORDER BY created_at DESC",
                    (job_id,),
                ).fetchall()
            ]
            conn.close()
        except Exception:
            pass
        if not job:
            return render_template("404.html", message="Training job not found"), 404
        return render_template("finetune/job_detail.html", job=job, events=events, loss_history=loss_history)

    @app.route("/finetune/models")
    def finetune_models_page():
        """Fine-Tuning model versions — eval scores, promotion status."""
        models = []
        try:
            conn = _get_db()
            models = [
                dict(r) for r in conn.execute("SELECT * FROM ft_model_versions ORDER BY created_at DESC").fetchall()
            ]
            conn.close()
        except Exception:
            pass
        return render_template("finetune/models.html", models=models)

    @app.route("/finetune/models/<model_id>")
    def finetune_model_detail_page(model_id):
        """Fine-Tuning model detail — evaluation history, promotion log."""
        model = None
        evaluations = []
        promotions = []
        try:
            conn = _get_db()
            row = conn.execute("SELECT * FROM ft_model_versions WHERE id = %s", (model_id,)).fetchone()
            if row:
                model = dict(row)
            evaluations = [
                dict(r)
                for r in conn.execute(
                    "SELECT * FROM ft_evaluations WHERE model_version_id = %s ORDER BY evaluated_at DESC",
                    (model_id,),
                ).fetchall()
            ]
            promotions = [
                dict(r)
                for r in conn.execute(
                    "SELECT * FROM ft_promotion_log WHERE model_version_id = %s ORDER BY created_at DESC",
                    (model_id,),
                ).fetchall()
            ]
            conn.close()
        except Exception:
            pass
        if not model:
            return render_template("404.html", message="Model version not found"), 404
        return render_template(
            "finetune/model_detail.html", model=model, evaluations=evaluations, promotions=promotions
        )

    @app.route("/finetune/evaluate")
    def finetune_evaluate_page():
        """Fine-Tuning evaluations — BLEU, ROUGE-L, perplexity scoring (D-FT-14, D-FT-15)."""
        evaluations = []
        try:
            conn = _get_db()
            evaluations = [
                dict(r) for r in conn.execute("SELECT * FROM ft_evaluations ORDER BY evaluated_at DESC").fetchall()
            ]
            conn.close()
        except Exception:
            pass
        return render_template("finetune/evaluate.html", evaluations=evaluations)

    # ── ICDEV™ Pulse — Blog Engine ─────────────────────────────────────



    # ── Pulse API Endpoints ──────────────────────────────────────────




























    @app.route("/api/llm/dual-model", methods=["GET"])
    def api_llm_dual_model_status():
        """Get current dual-model mode status."""
        try:
            from tools.llm.router import LLMRouter

            active = LLMRouter.get_dual_model()
            return jsonify(
                {
                    "dual_model": active,
                    "mode": "speed" if active else "quality",
                    "description": "1.7B text-only + Gemma3 (both VRAM-resident)"
                    if active
                    else "9B multimodal (single model, higher quality)",
                }
            )
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.route("/api/llm/dual-model", methods=["POST"])
    def api_llm_dual_model_toggle():
        """Toggle dual-model mode. Body: {"enabled": true/false}."""
        try:
            from tools.llm.router import LLMRouter

            data = flask_request.get_json(silent=True) or {}
            enabled = data.get("enabled")
            if enabled is None:
                # Toggle current state
                enabled = not LLMRouter.get_dual_model()
            LLMRouter.set_dual_model(bool(enabled))
            return jsonify(
                {
                    "dual_model": LLMRouter.get_dual_model(),
                    "mode": "speed" if enabled else "quality",
                }
            )
        except Exception as e:
            return jsonify({"error": str(e)}), 500


    # ── Pulse SAM Bridge ──────────────────────────────────────────────














    # ---- File Sync (D-SYNC-1 through D-SYNC-12) ----

    @app.route("/filesync")
    def filesync_page():
        """File Sync — sync jobs, status, conflicts, activity log."""
        stats = {
            "total_jobs": 0,
            "active_jobs": 0,
            "watching_jobs": 0,
            "completed_syncs": 0,
            "failed_syncs": 0,
            "pending_conflicts": 0,
            "total_bytes": 0,
            "total_bytes_display": "0 B",
        }
        jobs = []
        log_entries = []
        # Open with a short timeout so lock contention fails fast rather than
        # blocking for 30+ seconds (sync_log has millions of rows and the DB
        # is often under write load from the Kanban scheduler).
        import sqlite3 as _sq
        try:
            conn = _sq.connect(str(DB_PATH), timeout=3)
            conn.row_factory = _sq.Row
            conn.execute("PRAGMA busy_timeout=3000")
        except Exception:
            conn = None
        if conn is None:
            return render_template("filesync.html", stats=stats, jobs=jobs, log_entries=log_entries)
        try:
            try:
                row = conn.execute("SELECT COUNT(*) as cnt FROM sync_jobs").fetchone()
                stats["total_jobs"] = row["cnt"]
            except Exception:
                pass
            try:
                row = conn.execute(
                    "SELECT COUNT(*) as cnt FROM sync_jobs WHERE status IN ('scanning', 'syncing', 'watching')"
                ).fetchone()
                stats["active_jobs"] = row["cnt"]
            except Exception:
                pass
            try:
                row = conn.execute("SELECT COUNT(*) as cnt FROM sync_jobs WHERE status = 'watching'").fetchone()
                stats["watching_jobs"] = row["cnt"]
            except Exception:
                pass
            try:
                row = conn.execute("SELECT COUNT(*) as cnt FROM sync_log WHERE action = 'sync_completed'").fetchone()
                stats["completed_syncs"] = row["cnt"]
            except Exception:
                pass
            try:
                row = conn.execute("SELECT COUNT(*) as cnt FROM sync_log WHERE action = 'error'").fetchone()
                stats["failed_syncs"] = row["cnt"]
            except Exception:
                pass
            try:
                row = conn.execute("SELECT COUNT(*) as cnt FROM sync_conflicts WHERE resolution = 'pending'").fetchone()
                stats["pending_conflicts"] = row["cnt"]
            except Exception:
                pass
            # Skip SUM(bytes_transferred) over full sync_log — too expensive on millions of rows.
            # Use SUM from sync_jobs.bytes_transferred (per-job aggregate) as a fast proxy.
            try:
                row = conn.execute("SELECT COALESCE(SUM(bytes_transferred), 0) as total FROM sync_jobs").fetchone()
                total_bytes = row["total"]
                stats["total_bytes"] = total_bytes
                if total_bytes >= 1073741824:
                    stats["total_bytes_display"] = f"{total_bytes / 1073741824:.1f} GB"
                elif total_bytes >= 1048576:
                    stats["total_bytes_display"] = f"{total_bytes / 1048576:.1f} MB"
                elif total_bytes >= 1024:
                    stats["total_bytes_display"] = f"{total_bytes / 1024:.1f} KB"
                else:
                    stats["total_bytes_display"] = f"{total_bytes} B"
            except Exception:
                pass
            try:
                rows = conn.execute("SELECT * FROM sync_jobs ORDER BY created_at DESC LIMIT 50").fetchall()
                jobs = [dict(r) for r in rows]
            except Exception:
                pass
            try:
                rows = conn.execute("SELECT * FROM sync_log ORDER BY created_at DESC LIMIT 30").fetchall()
                log_entries = [dict(r) for r in rows]
            except Exception:
                pass
        finally:
            conn.close()
        return render_template("filesync.html", stats=stats, jobs=jobs, log_entries=log_entries)

    try:
        from tools.security.exceptions import TierAccessDenied as _TierAccessDenied

        @app.errorhandler(_TierAccessDenied)
        def tier_gate_handler(e):
            from tools.billing.tier import get_active_tier as _get_active_tier
            _active = _get_active_tier()
            if flask_request.is_json or flask_request.path.startswith("/api/"):
                return jsonify({
                    "error": "TierAccessDenied",
                    "canvas": e.canvas_key,
                    "required_tier": e.required_tier,
                    "active_tier": _active,
                    "message": str(e),
                }), 403
            return render_template(
                "errors/tier_gate.html",
                canvas_key=e.canvas_key,
                required_tier=e.required_tier,
                active_tier=_active,
            ), 403
    except ImportError:
        pass

    @app.route("/api/license/tier", methods=["GET"])
    def api_license_tier():
        """Return active tier, available features, and locked features."""
        from tools.billing.tier import get_active_tier as _gat, tier_satisfies as _ts
        _active = _gat()
        _available: list[str] = []
        _locked: list[str] = []
        _degraded = False
        _detail = None
        try:
            for _comp in _REGISTRY.iter_canvases():
                if _ts(_active, _comp.min_tier):
                    _available.append(_comp.key)
                else:
                    _locked.append(_comp.key)
        except Exception as exc:
            # nav-misc-02: an empty feature list must not read as "nothing licensed".
            _degraded = True
            _detail = str(exc)
            get_logger("icdev.dashboard").warning("api_license_tier: registry enumeration failed: %s", exc)
        _resp = {
            "active_tier": _active,
            "available_features": sorted(_available),
            "locked_features": sorted(_locked),
        }
        if _degraded:
            _resp["error"] = True
            _resp["detail"] = _detail
        return jsonify(_resp)

    @app.errorhandler(401)
    def unauthorized(e):
        if flask_request.is_json or flask_request.path.startswith("/api/"):
            return jsonify({"error": "Unauthorized", "message": "Valid API key required"}), 401
        return redirect(url_for("login_page"))

    @app.errorhandler(403)
    def forbidden(e):
        if flask_request.is_json or flask_request.path.startswith("/api/"):
            return jsonify({"error": "Forbidden", "message": "Insufficient permissions"}), 403
        return render_template("404.html", message="You do not have permission to access this page."), 403

    @app.errorhandler(404)
    def not_found(e):
        return render_template("404.html", message="Page not found"), 404

    @app.errorhandler(413)
    def payload_too_large(e):
        # cnr-plat-02: graceful JSON 413 when a request body exceeds
        # MAX_CONTENT_LENGTH (or a per-route cap that aborts 413).
        _limit = app.config.get("MAX_CONTENT_LENGTH")
        _limit_mb = round(_limit / (1024 * 1024), 1) if _limit else None
        if flask_request.is_json or flask_request.path.startswith("/api/"):
            return jsonify({
                "error": "Payload too large",
                "code": "PAYLOAD_TOO_LARGE",
                "max_upload_mb": _limit_mb,
                "message": (
                    f"Upload exceeds the maximum allowed size ({_limit_mb} MB)."
                    if _limit_mb else "Upload exceeds the maximum allowed size."
                ),
            }), 413
        return render_template(
            "404.html",
            message=(
                f"Upload exceeds the maximum allowed size ({_limit_mb} MB)."
                if _limit_mb else "Upload exceeds the maximum allowed size."
            ),
        ), 413

    @app.errorhandler(500)
    def internal_error(e):
        if flask_request.is_json or flask_request.path.startswith("/api/"):
            return jsonify({"error": "Internal server error", "message": str(e)}), 500
        return render_template("500.html", message=str(e)), 500

    # -------------------------------------------------------------------
    # Health check endpoints (P2 — monitoring / load-balancer)
    # -------------------------------------------------------------------

    @app.route("/api/health")
    def api_health():
        try:
            from tools.db.storage import get_connection as _gc
            _gc().execute("SELECT 1").fetchone()
            db_ok = True
        except Exception:
            db_ok = False
        return jsonify({"status": "ok" if db_ok else "degraded", "db": db_ok})

    @app.route("/api/status")
    def api_status():
        return jsonify({
            "status": "running",
            "service": "ICDEV™ Dashboard",
            "version": "1.0",
        })

    @app.route("/robots.txt")
    def robots_txt():
        from flask import Response
        body = "User-agent: *\nDisallow: /\n"
        return Response(body, mimetype="text/plain")

    # -------------------------------------------------------------------
    # CLI Generator (cherry-picked from icdev main)
    # -------------------------------------------------------------------

    @app.route("/api/cli-generator/generate", methods=["POST"])
    def api_cli_generator_generate():
        try:
            from tools.harness.cli_generator import generate

            data = flask_request.get_json(force=True)
            spec_path = data.get("spec_path", "")
            if not spec_path:
                return jsonify({"status": "error", "error": "spec_path is required"}), 400
            result = generate(
                spec_path=spec_path,
                output_dir=data.get("output_dir"),
                name=data.get("name"),
                dry_run=data.get("dry_run", False),
            )
            return jsonify(result)
        except Exception as e:
            return jsonify({"status": "error", "error": str(e)}), 500

    # ═══════════════════════════════════════════════════════════════════════════
    # MCP Wrapper Generator (Phase 3)
    # ═══════════════════════════════════════════════════════════════════════════

    @app.route("/mcp-wrapper")
    def mcp_wrapper_page():
        return render_template("mcp_wrapper.html", app_name="SparkPilot")

    @app.route("/api/mcp-wrapper/scan")
    def api_mcp_wrapper_scan():
        try:
            from tools.harness.mcp_wrapper_generator import scan_tools

            return jsonify(scan_tools())
        except Exception as e:
            return jsonify({"status": "error", "error": str(e), "discovered": [], "total": 0, "with_json_flag": 0})

    @app.route("/api/mcp-wrapper/list")
    def api_mcp_wrapper_list():
        try:
            from tools.harness.mcp_wrapper_generator import list_wrapped

            return jsonify(list_wrapped())
        except Exception as e:
            return jsonify({"wrappers": [], "count": 0, "error": str(e)})

    @app.route("/api/mcp-wrapper/wrap", methods=["POST"])
    def api_mcp_wrapper_wrap():
        try:
            from tools.harness.mcp_wrapper_generator import wrap_tool

            data = flask_request.get_json(force=True)
            tool_path = data.get("tool_path", "")
            if not tool_path:
                return jsonify({"status": "error", "error": "tool_path is required"}), 400
            return jsonify(wrap_tool(tool_path, dry_run=data.get("dry_run", False)))
        except Exception as e:
            return jsonify({"status": "error", "error": str(e)}), 500

    @app.route("/api/mcp-wrapper/wrap-all", methods=["POST"])
    def api_mcp_wrapper_wrap_all():
        try:
            from tools.harness.mcp_wrapper_generator import wrap_all

            data = flask_request.get_json(force=True) if flask_request.data else {}
            return jsonify(
                wrap_all(
                    dry_run=data.get("dry_run", False),
                    limit=data.get("limit", 20),
                )
            )
        except Exception as e:
            return jsonify({"status": "error", "error": str(e)}), 500

    # ── MCP Monitor status ───────────────────────────────────────────
    @app.route("/api/mcp/status", methods=["GET"])
    def api_mcp_status():
        """GET /api/mcp/status — snapshot of the MCP tool registry and wrapper activity."""
        import time as _t
        from collections import Counter as _Counter

        try:
            from tools.mcp.tool_registry import TOOL_REGISTRY
        except Exception:
            TOOL_REGISTRY = {}

        # Category breakdown
        cats = _Counter(v.get("category", "unknown") for v in TOOL_REGISTRY.values())
        top_cats = [{"name": k, "count": v} for k, v in cats.most_common(8)]

        # MCP debug log activity
        log_path = BASE_DIR / ".tmp" / "mcp_debug.log"
        log_age_secs = None
        wrapper_starts = 0
        error_count = 0
        if log_path.exists():
            log_age_secs = int(_t.time() - log_path.stat().st_mtime)
            try:
                with open(log_path, encoding="utf-8", errors="ignore") as _f:
                    today = str(datetime.now(timezone.utc).date())
                    for line in _f:
                        if "MCP wrapper started" in line and today in line:
                            wrapper_starts += 1
                        if "ERROR" in line or "Exception" in line:
                            error_count += 1
            except Exception:
                pass

        return jsonify({
            "tool_count": len(TOOL_REGISTRY),
            "category_count": len(cats),
            "top_categories": top_cats,
            "log_age_secs": log_age_secs,
            "wrapper_starts_today": wrapper_starts,
            "error_count": error_count,
        })

    # ── Page Agent Copilot API ───────────────────────────────────────
    # Inspired by alibaba/page-agent: text-based DOM navigation + AI copilot

    _PAGE_AGENT_ROUTE_MAP = {
        "home": "/",
        "dashboard": "/",
        "missions": "/missions",
        "simulator": "/simulator",
        "fleet": "/devices",
        "devices": "/devices",
        "firmware": "/firmware",
        "edge ai": "/edge-ai",
        "self-heal": "/crashes",
        "agents": "/agents",
        "govcon": "/govcon",
        "writeguard": "/writeguard",
        "pulse": "/pulse",
        "databridge": "/databridge",
        "messaging": "/databridge/messaging",
        "cloudforge": "/cloudforge",
        "knowledge": "/knowledge-graph",
        "knowledge graph": "/knowledge-graph",
        "marketplace": "/marketplace",
        "research": "/research",
        "harness": "/harness",
        "codelens": "/container-lens",
        "forge studio": "/forge-studio",
        "dochub": "/dochub",
        "resilience": "/resilience",
        "architecture": "/architecture",
        "compliance": "/compliance-accel",
        "agent evolution": "/agent-evolution",
        "intelligence": "/intelligence",
        "maturity": "/maturity",
        "decisions": "/decisions",
        "security": "/security-scan",
    }

    @app.route("/api/page-agent/message", methods=["POST"])
    def api_page_agent_message():
        """Process a Page Agent copilot message — navigation, search, or contextual help.

        Route map loaded from DB (page_agent_routes) with fallback to
        _PAGE_AGENT_ROUTE_MAP hardcoded dict.
        """
        # Load custom routes from DB if available
        try:
            conn = _get_db()
            rows = conn.execute("SELECT keyword, route FROM page_agent_routes").fetchall()
            conn.close()
            if rows:
                for r in rows:
                    _PAGE_AGENT_ROUTE_MAP[r["keyword"]] = r["route"]
        except Exception:
            pass  # Table may not exist — use defaults
        try:
            data = flask_request.get_json(force=True) if flask_request.is_json else {}
            message = data.get("message", "").strip()
            page = data.get("page", "/")
            if not message:
                return jsonify({"error": "message required"}), 400

            lower = message.lower().strip()

            # Navigation intent
            for prefix in ("go to ", "navigate to ", "show me ", "open "):
                if lower.startswith(prefix):
                    target = lower[len(prefix) :].strip()
                    route = _PAGE_AGENT_ROUTE_MAP.get(target)
                    if route:
                        return jsonify(
                            {
                                "response": f"Navigating to **{target}**...",
                                "action": "navigate",
                                "route": route,
                            }
                        )
                    # Fuzzy match
                    best, best_score = None, 0
                    for key in _PAGE_AGENT_ROUTE_MAP:
                        score = _bigram_similarity(target, key)
                        if score > best_score and score > 0.4:
                            best_score = score
                            best = key
                    if best:
                        return jsonify(
                            {
                                "response": f"Did you mean **{best}**? Navigating...",
                                "action": "navigate",
                                "route": _PAGE_AGENT_ROUTE_MAP[best],
                            }
                        )
                    return jsonify(
                        {
                            "response": f"Page not found: `{target}`. Try asking `show pages`.",
                            "suggestions": ["show pages", "help"],
                        }
                    )

            # Help
            if lower in ("help", "what can you do", "commands"):
                return jsonify(
                    {
                        "response": (
                            "**Commands:** `go to <page>`, `search <text>`, "
                            "`show pages`, `where am i`, `describe this page`, "
                            "`scroll up/down`, `click <element>`, `fill <value> in <field>`"
                        ),
                        "suggestions": ["go to compliance", "show pages", "describe this page"],
                    }
                )

            # Page listing
            if "show pages" in lower or "list pages" in lower or "list routes" in lower:
                pages = sorted(_PAGE_AGENT_ROUTE_MAP.keys())
                lines = [f"- `{p}` → {_PAGE_AGENT_ROUTE_MAP[p]}" for p in pages]
                return jsonify(
                    {
                        "response": f"**Available pages ({len(pages)}):**\n" + "\n".join(lines),
                    }
                )

            # Context-aware suggestions based on current page
            suggestions = _page_suggestions(page)
            return jsonify(
                {
                    "response": (
                        f"I understand your request: *{message}*. "
                        "For best results, try specific commands like `go to agents` or `search <keyword>`."
                    ),
                    "suggestions": suggestions,
                }
            )
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    def _bigram_similarity(a, b):
        """Bigram (Dice) similarity for fuzzy page matching."""
        if a == b:
            return 1.0
        if len(a) < 2 or len(b) < 2:
            return 0.0
        a_bigrams = {}
        for i in range(len(a) - 1):
            bg = a[i : i + 2]
            a_bigrams[bg] = a_bigrams.get(bg, 0) + 1
        matches = 0
        for i in range(len(b) - 1):
            bg = b[i : i + 2]
            if a_bigrams.get(bg, 0) > 0:
                matches += 1
                a_bigrams[bg] -= 1
        return (2.0 * matches) / (len(a) + len(b) - 2)

    def _page_suggestions(current_page):
        """Return contextual suggestions based on current page."""
        suggestions_map = {
            "/": ["go to agents", "go to compliance", "go to missions"],
            "/agents": ["go to agent evolution", "go to harness", "go to intelligence"],
            "/compliance-accel": ["go to dochub", "go to resilience", "go to maturity"],
            "/devices": ["go to firmware", "go to edge ai", "go to simulator"],
            "/knowledge-graph": ["go to research", "go to intelligence", "search compliance"],
            "/cloudforge": ["go to databridge", "go to marketplace", "go to forge studio"],
            "/genesis": ["run research", "run audit", "show promoter stats"],
        }
        return suggestions_map.get(current_page, ["help", "show pages", "go to agents"])

    # ── Proposal Genesis — Autonomous Capture Pipeline Dashboard ─────────────

    @app.route("/proposal-genesis")
    @require_role("admin", "pm", "bd", "capture_mgr")
    def proposal_genesis():
        """Proposal Genesis — autonomous capture-to-delivery pipeline dashboard."""
        status = {}
        summary = {}
        try:
            import subprocess

            _utf8_env = {**os.environ, "PYTHONIOENCODING": "utf-8"}
            result = subprocess.run(
                [sys.executable, "tools/proposal_genesis/daemon.py", "--status", "--json"],
                capture_output=True,
                text=True,
                timeout=15,
                cwd=BASE_DIR,
                env=_utf8_env,
            )
            stdout = result.stdout.strip()
            json_start = stdout.find("{")
            if json_start >= 0:
                status = json.loads(stdout[json_start:])
        except Exception as exc:
            status = {"error": str(exc)}
        # Summary stats
        degraded = False
        try:
            conn = _get_db()
            try:
                summary["opportunities"] = conn.execute(
                    "SELECT COUNT(*) as cnt FROM proposal_opportunities WHERE status IN ('tracking', 'drafting')"
                ).fetchone()["cnt"]
            except Exception:
                summary["opportunities"] = 0
            try:
                summary["shall_statements"] = conn.execute(
                    "SELECT COUNT(*) as cnt FROM rfp_shall_statements"
                ).fetchone()["cnt"]
            except Exception:
                summary["shall_statements"] = 0
            try:
                summary["drafts"] = conn.execute(
                    "SELECT COUNT(*) as cnt FROM proposal_section_drafts WHERE status = 'draft'"
                ).fetchone()["cnt"]
            except Exception:
                summary["drafts"] = 0
            try:
                row = conn.execute(
                    "SELECT AVG(composite_score) as avg_score FROM pg_proposal_quality_scores"
                ).fetchone()
                summary["avg_quality"] = round(row["avg_score"] or 0, 3)
            except Exception:
                summary["avg_quality"] = 0
            try:
                summary["pulse_links"] = conn.execute("SELECT COUNT(*) as cnt FROM pg_pulse_proposal_links").fetchone()[
                    "cnt"
                ]
            except Exception:
                summary["pulse_links"] = 0
            conn.close()
        except Exception as exc:
            # nav-misc-02: a DB-connection failure here dropped the whole summary
            # to zeros silently. Log + flag so the empty pipeline reads as an outage.
            degraded = True
            get_logger("icdev.dashboard").warning("proposal_genesis: summary DB read failed: %s", exc)
        return render_template("proposal_genesis.html", status=status, summary=summary, degraded=degraded)

    # ── Genesis v2.0 — Autonomous Research Lab Dashboard ──────────────────────

    # Registry of all Genesis-enabled apps (app_key → config)
    GENESIS_APPS = {
        "icdev": {
            "name": "ICDEV™",
            "root": str(BASE_DIR),
            "daemon": "tools/genesis/daemon.py",
            "promoter": "tools/genesis/promoter.py",
            "env_var": "ICDEV_GENESIS_ENABLED",
            "db": str(BASE_DIR / "data" / "icdev.db"),
        },
        "govchain": {
            "name": "GovChain",
            "root": str(Path(BASE_DIR).parent / "govchain"),
            "daemon": "tools/genesis/daemon.py",
            "promoter": None,
            "env_var": "GOVCHAIN_GENESIS_ENABLED",
            "db": str(Path(BASE_DIR).parent / "govchain" / "data" / "govchain.db"),
        },
        "govproposal": {
            "name": "GovProposal",
            "root": str(Path(BASE_DIR).parent / "GovProposal"),
            "daemon": "tools/genesis/daemon.py",
            "promoter": None,
            "env_var": "GOVPROPOSAL_GENESIS_ENABLED",
            "db": str(Path(BASE_DIR).parent / "GovProposal" / "data" / "govproposal.db"),
        },
        "trading-engine": {
            "name": "Trading Engine",
            "root": str(Path(BASE_DIR).parent / "trading-engine"),
            "daemon": "tools/genesis/daemon.py",
            "promoter": None,
            "env_var": "TRADING_GENESIS_ENABLED",
            "db": str(Path(BASE_DIR).parent / "trading-engine" / "data" / "trading-engine.db"),
        },
        "trading-strategy": {
            "name": "Trading Strategy",
            "root": str(Path(BASE_DIR).parent / "Trading_Strategy"),
            "daemon": "tools/genesis/daemon.py",
            "promoter": "tools/genesis/promoter.py",
            "env_var": "TRADING_GENESIS_ENABLED",
            "db": str(Path(BASE_DIR).parent / "Trading_Strategy" / "data" / "trading_strategy.db"),
        },
        "ninjaflow": {
            "name": "NinjaFlow",
            "root": str(Path(BASE_DIR).parent / "ninjaflow-ai" / "ninjaflow-ai"),
            "daemon": "tools/genesis/daemon.py",
            "promoter": None,
            "env_var": "NINJAFLOW_GENESIS_ENABLED",
            "db": str(Path(BASE_DIR).parent / "ninjaflow-ai" / "ninjaflow-ai" / "data" / "ninjaflow-ai.db"),
        },
        "signalforge": {
            "name": "SignalForge",
            "root": str(Path(BASE_DIR).parent / "signalforge"),
            "daemon": "tools/genesis/daemon.py",
            "promoter": None,
            "env_var": "SIGNALFORGE_GENESIS_ENABLED",
            "db": str(Path(BASE_DIR).parent / "signalforge" / "data" / "signalforge.db"),
        },
    }

    def _genesis_app(app_key):
        """Get Genesis app config, default to icdev."""
        return GENESIS_APPS.get(app_key, GENESIS_APPS["icdev"])

    def _genesis_run(app_key, args, timeout=15):
        """Run a Genesis daemon command for a given app."""
        import subprocess as _sp

        cfg = _genesis_app(app_key)
        app_root = cfg["root"]
        daemon_path = cfg["daemon"]
        env = {
            **os.environ,
            "PYTHONIOENCODING": "utf-8",
            "PYTHONPATH": app_root,
            cfg["env_var"]: "true",
            "PYTHONUNBUFFERED": "1",
        }
        result = _sp.run(
            [sys.executable, daemon_path] + args,
            capture_output=True,
            text=True,
            timeout=timeout,
            cwd=app_root,
            env=env,
        )
        stdout = result.stdout.strip()
        json_start = stdout.find("{")
        if json_start >= 0:
            return json.loads(stdout[json_start:])
        return {"error": "parse_failed", "stderr": result.stderr[:500] if result.stderr else ""}

    def _genesis_db(app_key):
        """Get a DB connection for a Genesis app."""
        _genesis_app(app_key)  # Validate app_key exists
        conn = get_connection(db_path=str(DB_PATH))
        conn.execute("PRAGMA journal_mode=WAL")
        return conn

    # ── Contact Form Submissions ─────────────────────────────────────────────

    @app.route("/api/contact/submit", methods=["POST", "OPTIONS"])
    def api_contact_submit():
        """Public endpoint — receives contact form submissions from icdev.ai."""
        # CORS for cross-origin from icdev.ai
        if flask_request.method == "OPTIONS":
            resp = app.make_default_options_response()
            resp.headers["Access-Control-Allow-Origin"] = "https://icdev.ai"
            resp.headers["Access-Control-Allow-Methods"] = "POST, OPTIONS"
            resp.headers["Access-Control-Allow-Headers"] = "Content-Type"
            return resp

        try:
            # Accept both JSON and form data
            if flask_request.is_json:
                data = flask_request.get_json()
            else:
                data = flask_request.form.to_dict()

            name = (data.get("name") or "").strip()
            email = (data.get("email") or "").strip()
            if not name or not email:
                return jsonify({"error": "Name and email are required"}), 400

            sub_id = f"lead-{uuid.uuid4().hex[:12]}"
            now = __import__("datetime").datetime.now(__import__("datetime").timezone.utc).isoformat()

            conn = _get_db()
            try:
                conn.execute(
                    "INSERT INTO contact_submissions "
                    "(id, name, email, organization, role, interest, message, status, created_at) "
                    "VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)",
                    (
                        sub_id,
                        name,
                        email,
                        (data.get("organization") or "").strip(),
                        (data.get("role") or "").strip(),
                        (data.get("interest") or "").strip(),
                        (data.get("message") or "").strip(),
                        "new",
                        now,
                    ),
                )
                conn.commit()
            finally:
                conn.close()

            resp = jsonify({"ok": True, "id": sub_id})
            resp.headers["Access-Control-Allow-Origin"] = "https://icdev.ai"
            return resp
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.route("/leads")
    def leads_page():
        """Contact form submissions dashboard."""
        conn = _get_db()
        try:
            rows = conn.execute("SELECT * FROM contact_submissions ORDER BY created_at DESC LIMIT 100").fetchall()
            submissions = [dict(r) for r in rows]
            stats = {
                "total": len(submissions),
                "new": sum(1 for s in submissions if s.get("status") == "new"),
                "contacted": sum(1 for s in submissions if s.get("status") == "contacted"),
                "closed": sum(1 for s in submissions if s.get("status") == "closed"),
            }
        except Exception:
            submissions = []
            stats = {"total": 0, "new": 0, "contacted": 0, "closed": 0}
        finally:
            conn.close()
        return render_template("leads.html", submissions=submissions, stats=stats)

    @app.route("/api/leads/<lead_id>/status", methods=["POST"])
    def api_lead_update_status(lead_id):
        """Update a lead's status."""
        data = flask_request.get_json(silent=True) or {}
        new_status = data.get("status", "")
        notes = data.get("notes", "")
        if new_status not in ("new", "contacted", "qualified", "closed"):
            return jsonify({"error": "Invalid status"}), 400
        now = __import__("datetime").datetime.now(__import__("datetime").timezone.utc).isoformat()
        conn = _get_db()
        try:
            conn.execute(
                "UPDATE contact_submissions SET status = %s, notes = %s, updated_at = %s WHERE id = %s",
                (new_status, notes, now, lead_id),
            )
            conn.commit()
        finally:
            conn.close()
        return jsonify({"ok": True, "id": lead_id, "status": new_status})

    # ── Genesis v2.0 — Autonomous Research Lab Dashboard ──────────────────────

    @app.route("/notifications")
    def notifications_page():
        """Notification Gateway — adapter config, delivery history, routing rules."""
        try:
            from tools.notifications.gateway import NotificationGateway

            gw = NotificationGateway()
            health = gw.health()
        except Exception:
            health = {"enabled": False, "adapters": {}, "error": "Gateway unavailable"}
        # Recent delivery log
        history = []
        degraded = not health.get("enabled", True) and bool(health.get("error"))
        try:
            conn = _get_db()
            history = conn.execute("SELECT * FROM notification_log ORDER BY created_at DESC LIMIT 50").fetchall()
            conn.close()
        except Exception as exc:
            # nav-misc-02: a delivery-log read failure must not render an empty
            # history that implies "no notifications ever sent".
            degraded = True
            get_logger("icdev.dashboard").warning("notifications_page: delivery-log read failed: %s", exc)
        return render_template("notifications.html", health=health, history=history, degraded=degraded)

    @app.route("/genesis")
    def genesis():
        """Genesis v2.0 — Autonomous Research Lab dashboard (multi-app)."""
        app_key = flask_request.args.get("app", "icdev")
        # Gather status for all apps
        all_status = {}
        for key, cfg in GENESIS_APPS.items():
            if Path(cfg["root"]).exists():
                try:
                    all_status[key] = _genesis_run(key, ["--status", "--json"])
                    all_status[key]["_name"] = cfg["name"]
                    all_status[key]["_available"] = True
                except Exception as exc:
                    all_status[key] = {"_name": cfg["name"], "_available": False, "error": str(exc)}
            else:
                all_status[key] = {"_name": cfg["name"], "_available": False, "error": "Directory not found"}
        # Active app status
        status = all_status.get(app_key, all_status.get("icdev", {}))
        return render_template(
            "genesis.html", status=status, all_apps=all_status, active_app=app_key, genesis_apps=GENESIS_APPS
        )

    @app.route("/api/genesis/status", methods=["GET"])
    def api_genesis_status():
        app_key = flask_request.args.get("app", "icdev")
        try:
            return jsonify(_genesis_run(app_key, ["--status", "--json"]))
        except Exception as exc:
            # DB fallback: query genesis_runs for last known status
            try:
                conn = _get_db()
                row = conn.execute(
                    "SELECT * FROM genesis_runs WHERE app_key = %s ORDER BY started_at DESC LIMIT 1",
                    (app_key,),
                ).fetchone()
                conn.close()
                if row:
                    return jsonify(
                        {"status": "cached", "app": app_key, "last_run": dict(row), "daemon_error": str(exc)}
                    )
            except Exception:
                pass
            return jsonify({"error": str(exc), "app": app_key}), 500

    @app.route("/api/genesis/all-status", methods=["GET"])
    def api_genesis_all_status():
        """Get status for all Genesis apps."""
        results = {}
        for key, cfg in GENESIS_APPS.items():
            if Path(cfg["root"]).exists():
                try:
                    results[key] = _genesis_run(key, ["--status", "--json"])
                    results[key]["_name"] = cfg["name"]
                except Exception as exc:
                    results[key] = {"_name": cfg["name"], "error": str(exc)}
            else:
                results[key] = {"_name": cfg["name"], "error": "not_found"}
        return jsonify(results)

    @app.route("/api/genesis/reflex/<name>", methods=["POST"])
    def api_genesis_run_reflex(name):
        """Run a single Genesis reflex on-demand."""
        app_key = flask_request.args.get("app", "icdev")
        allowed = [
            "research",
            "scout",
            "audit",
            "report",
            "comply",
            "ingest",
            "market",
            "publish",
            "test",
            "learn",
            "heal",
            "evolve",
            "docs",
        ]
        if name not in allowed:
            return jsonify({"error": f"Unknown reflex: {name}"}), 400
        try:
            result = _genesis_run(app_key, ["--reflex", name, "--json"], timeout=300)
            # Log to DB for audit trail
            try:
                conn = _get_db()
                conn.execute(
                    "INSERT INTO audit_trail (event_type, action, details, created_at) "
                    "VALUES (%s, %s, %s, datetime('now'))",
                    ("config_changed", f"genesis_reflex:{name}", json.dumps({"app": app_key, "reflex": name})),
                )
                conn.commit()
                conn.close()
            except Exception as _exc:  # noqa: BLE001 - best-effort persistence; logged, never raised
                logger.warning(
                    "api_genesis_run_reflex: best-effort INSERT into audit_trail failed (non-blocking): %s",
                    _exc,
                )
            return jsonify(result)
        except Exception as exc:
            return jsonify({"error": str(exc)}), 500

    @app.route("/api/genesis/promoter/stats", methods=["GET"])
    def api_genesis_promoter_stats():
        app_key = flask_request.args.get("app", "icdev")
        cfg = _genesis_app(app_key)
        if not cfg.get("promoter"):
            # No promoter — query DB directly for GKP counts
            try:
                conn = _genesis_db(app_key)
                try:
                    total = conn.execute("SELECT COUNT(*) FROM genesis_gkp").fetchone()[0]
                    by_status = {}
                    for row in conn.execute(
                        "SELECT promotion_status, COUNT(*) as cnt FROM genesis_gkp GROUP BY promotion_status"
                    ).fetchall():
                        by_status[row[0]] = row[1]
                    return jsonify({"total_gkps": total, "by_status": by_status})
                finally:
                    conn.close()
            except Exception as exc:
                return jsonify({"total_gkps": 0, "by_status": {}, "note": str(exc)})
        try:
            import subprocess

            _utf8_env = {**os.environ, "PYTHONIOENCODING": "utf-8", "PYTHONPATH": cfg["root"]}
            result = subprocess.run(
                [sys.executable, cfg["promoter"], "--stats", "--json"],
                capture_output=True,
                text=True,
                timeout=15,
                cwd=cfg["root"],
                env=_utf8_env,
            )
            stdout = result.stdout.strip()
            json_start = stdout.find("{")
            if json_start >= 0:
                return jsonify(json.loads(stdout[json_start:]))
            return jsonify({"error": "parse_failed"}), 500
        except Exception as exc:
            return jsonify({"error": str(exc)}), 500

    def _gkp_hidden_sources():
        """Load source patterns marked hide_from_dashboard in genesis auto_promote rules."""
        try:
            cfg_path = BASE_DIR / "args" / "genesis_config.yaml"
            if cfg_path.exists():
                import yaml

                cfg = yaml.safe_load(cfg_path.read_text(encoding="utf-8")) or {}
                rules = cfg.get("promoter", {}).get("auto_promote", [])
                return [
                    r["source_contains"].lower()
                    for r in rules
                    if r.get("hide_from_dashboard") and r.get("source_contains")
                ]
        except Exception:
            pass
        return []

    @app.route("/api/genesis/gkps", methods=["GET"])
    def api_genesis_gkps():
        """List GKPs with optional status filter."""
        app_key = flask_request.args.get("app", "icdev")
        status_filter = flask_request.args.get("status", None)
        show_hidden = flask_request.args.get("show_hidden", "false") == "true"
        limit = int(flask_request.args.get("limit", "100"))
        try:
            conn = _genesis_db(app_key)
            try:
                if status_filter:
                    rows = conn.execute(
                        "SELECT * FROM genesis_gkp WHERE promotion_status = %s ORDER BY created_at DESC LIMIT %s",
                        (status_filter, limit),
                    ).fetchall()
                else:
                    rows = conn.execute(
                        "SELECT * FROM genesis_gkp ORDER BY created_at DESC LIMIT %s",
                        (limit,),
                    ).fetchall()
                gkps = [dict(r) for r in rows]

                # Filter out hidden sources unless explicitly requested
                if not show_hidden:
                    hidden = _gkp_hidden_sources()
                    if hidden:

                        def _is_hidden(g):
                            try:
                                p = json.loads(g["payload"]) if isinstance(g["payload"], str) else (g["payload"] or {})
                                src = (p.get("source", "") or "").lower()
                                return any(h in src for h in hidden)
                            except Exception:
                                return False

                        gkps = [g for g in gkps if not _is_hidden(g)]

                return jsonify({"gkps": gkps, "count": len(gkps)})
            finally:
                conn.close()
        except Exception as exc:
            return jsonify({"gkps": [], "count": 0, "note": str(exc)})

    @app.route("/api/genesis/gkps/<gkp_id>", methods=["GET"])
    def api_genesis_gkp_detail(gkp_id):
        """Get a single GKP by ID."""
        app_key = flask_request.args.get("app", "icdev")
        try:
            conn = _genesis_db(app_key)
            try:
                row = conn.execute("SELECT * FROM genesis_gkp WHERE id = %s", (gkp_id,)).fetchone()
                if not row:
                    return jsonify({"error": "GKP not found"}), 404
                return jsonify(dict(row))
            finally:
                conn.close()
        except Exception as exc:
            return jsonify({"error": str(exc)}), 500

    def _audit_gkp_mutation(event_type, action, gkp_id, app_key, extra=None):
        """Append a GKP promote/reject decision to the append-only audit_trail.

        Best-effort (mirrors api_genesis_run_reflex): a missing/locked audit
        table must never break the mutation itself. The acting user is the
        resolved session user (g.current_user), never a request-body field.
        """
        try:
            user = getattr(g, "current_user", None)
            actor = (user.get("id") if isinstance(user, dict) else None) or "unknown"
            details = {"app": app_key, "gkp_id": gkp_id, "actor": actor}
            if extra:
                details.update(extra)
            conn = _get_db()
            conn.execute(
                "INSERT INTO audit_trail (event_type, action, details, created_at) "
                "VALUES (%s, %s, %s, datetime('now'))",
                (event_type, action, json.dumps(details)),
            )
            conn.commit()
            conn.close()
        except Exception as exc:  # noqa: BLE001 - best-effort persistence; logged, never raised
            logger.warning("_audit_gkp_mutation: best-effort INSERT into audit_trail failed (non-blocking): %s", exc)

    @app.route("/api/genesis/gkps/<gkp_id>/promote", methods=["POST"])
    @require_role("admin", "pm")
    def api_genesis_promote_gkp(gkp_id):
        """Promote a GKP to v1.x. State-changing — admin/pm only (nav-plat-05)."""
        app_key = flask_request.args.get("app", "icdev")
        _audit_gkp_mutation("approval_granted", f"genesis_gkp_promote:{gkp_id}", gkp_id, app_key)
        cfg = _genesis_app(app_key)
        if not cfg.get("promoter"):
            # Manual DB update for apps without a promoter
            try:
                conn = _genesis_db(app_key)
                try:
                    conn.execute(
                        "UPDATE genesis_gkp SET promotion_status = 'promoted', promoted_at = datetime('now') WHERE id = %s",
                        (gkp_id,),
                    )
                    conn.commit()
                    return jsonify({"status": "promoted", "gkp_id": gkp_id})
                finally:
                    conn.close()
            except Exception as exc:
                return jsonify({"error": str(exc)}), 500
        try:
            import subprocess as _sp

            _utf8_env = {**os.environ, "PYTHONIOENCODING": "utf-8", "PYTHONPATH": cfg["root"]}
            result = _sp.run(
                [sys.executable, cfg["promoter"], "--promote", gkp_id, "--json"],
                capture_output=True,
                text=True,
                timeout=30,
                cwd=cfg["root"],
                env=_utf8_env,
            )
            stdout = result.stdout.strip()
            json_start = stdout.find("{")
            if json_start >= 0:
                return jsonify(json.loads(stdout[json_start:]))
            return jsonify({"error": result.stderr or "promote failed"}), 500
        except Exception as exc:
            return jsonify({"error": str(exc)}), 500

    @app.route("/api/genesis/gkps/<gkp_id>/reject", methods=["POST"])
    @require_role("admin", "pm")
    def api_genesis_reject_gkp(gkp_id):
        """Reject a GKP. State-changing — admin/pm only (nav-plat-05)."""
        app_key = flask_request.args.get("app", "icdev")
        cfg = _genesis_app(app_key)
        data = flask_request.get_json(silent=True) or {}
        reason = data.get("reason", "Rejected via dashboard")
        _audit_gkp_mutation(
            "approval_denied", f"genesis_gkp_reject:{gkp_id}", gkp_id, app_key, {"reason": reason}
        )
        if not cfg.get("promoter"):
            try:
                conn = _genesis_db(app_key)
                try:
                    conn.execute("UPDATE genesis_gkp SET promotion_status = 'rejected' WHERE id = %s", (gkp_id,))
                    conn.commit()
                    return jsonify({"status": "rejected", "gkp_id": gkp_id, "reason": reason})
                finally:
                    conn.close()
            except Exception as exc:
                return jsonify({"error": str(exc)}), 500
        try:
            import subprocess as _sp

            _utf8_env = {**os.environ, "PYTHONIOENCODING": "utf-8", "PYTHONPATH": cfg["root"]}
            result = _sp.run(
                [sys.executable, cfg["promoter"], "--reject", gkp_id, "--reason", reason, "--json"],
                capture_output=True,
                text=True,
                timeout=30,
                cwd=cfg["root"],
                env=_utf8_env,
            )
            stdout = result.stdout.strip()
            json_start = stdout.find("{")
            if json_start >= 0:
                return jsonify(json.loads(stdout[json_start:]))
            return jsonify({"error": result.stderr or "reject failed"}), 500
        except Exception as exc:
            return jsonify({"error": str(exc)}), 500

    @app.route("/api/genesis/gkps/auto-promote", methods=["POST"])
    @require_role("admin", "pm")
    def api_genesis_auto_promote():
        """Auto-promote all eligible GKPs. State-changing — admin/pm only (nav-plat-05)."""
        app_key = flask_request.args.get("app", "icdev")
        _audit_gkp_mutation("approval_granted", "genesis_gkp_auto_promote", "*", app_key)
        cfg = _genesis_app(app_key)
        if not cfg.get("promoter"):
            # DB fallback: check for pending GKPs directly
            try:
                conn = _get_db()
                pending = conn.execute(
                    "SELECT COUNT(*) FROM genesis_knowledge_packets WHERE status = 'pending'"
                ).fetchone()[0]
                conn.close()
                return jsonify({"error": "No promoter configured", "auto_promoted": 0, "pending_gkps": pending}), 400
            except Exception:
                return jsonify({"error": "No promoter configured for this app", "auto_promoted": 0}), 400
        try:
            import subprocess as _sp

            _utf8_env = {**os.environ, "PYTHONIOENCODING": "utf-8", "PYTHONPATH": cfg["root"]}
            result = _sp.run(
                [sys.executable, cfg["promoter"], "--auto-promote", "--json"],
                capture_output=True,
                text=True,
                timeout=30,
                cwd=cfg["root"],
                env=_utf8_env,
            )
            stdout = result.stdout.strip()
            json_start = stdout.find("{")
            if json_start >= 0:
                return jsonify(json.loads(stdout[json_start:]))
            return jsonify({"auto_promoted": 0, "results": []})
        except Exception as exc:
            return jsonify({"error": str(exc)}), 500

    @app.route("/api/genesis/feedback/priorities", methods=["GET"])
    def api_genesis_feedback_priorities():
        try:
            import subprocess

            result = subprocess.run(
                [sys.executable, "tools/genesis/feedback_collector.py", "--priorities", "--json"],
                capture_output=True,
                text=True,
                timeout=15,
                cwd=BASE_DIR,
            )
            stdout = result.stdout.strip()
            json_start = stdout.find("{")
            if json_start >= 0:
                return jsonify(json.loads(stdout[json_start:]))
            return jsonify({"error": "parse_failed"}), 500
        except Exception as exc:
            return jsonify({"error": str(exc)}), 500

    # ---- Phase 67: Engineering Review Board ----
    @app.route("/review-board")
    def review_board_page():
        """Engineering Review Board — multi-persona analysis dashboard."""
        conn = _get_db()
        health_score = None
        health_grade = "N/A"
        health_trend = "stable"
        health_trend_data = []
        correlation_groups = []
        remediation_stats = {}
        try:
            # Reflex states
            try:
                reflex_rows = conn.execute("SELECT * FROM review_board_reflex_state ORDER BY reflex_name").fetchall()
                reflexes = [dict(r) for r in reflex_rows]
            except Exception:
                reflexes = []

            # Recent findings
            try:
                finding_rows = conn.execute(
                    "SELECT * FROM review_board_findings ORDER BY created_at DESC LIMIT 100"
                ).fetchall()
                findings = [dict(r) for r in finding_rows]
            except Exception:
                findings = []

            # Severity summary
            try:
                severity_rows = conn.execute(
                    "SELECT severity, COUNT(*) as cnt FROM review_board_findings GROUP BY severity"
                ).fetchall()
                severity_summary = {r[0]: r[1] for r in severity_rows}
            except Exception:
                severity_summary = {}

            # Recent audit events
            try:
                audit_rows = conn.execute(
                    "SELECT * FROM review_board_audit ORDER BY created_at DESC LIMIT 20"
                ).fetchall()
                audit_events = [dict(r) for r in audit_rows]
            except Exception:
                audit_events = []

            total_findings = sum(severity_summary.values())

            # Health score + trend
            try:
                latest_health = conn.execute(
                    "SELECT score, grade, trend FROM review_board_health_history ORDER BY created_at DESC LIMIT 1"
                ).fetchone()
                if latest_health:
                    health_score = latest_health[0]
                    health_grade = latest_health[1]
                    health_trend = latest_health[2]
                trend_rows = conn.execute(
                    "SELECT score, created_at FROM review_board_health_history ORDER BY created_at DESC LIMIT 20"
                ).fetchall()
                health_trend_data = [{"score": r[0], "created_at": r[1]} for r in reversed(list(trend_rows))]
            except Exception:
                pass

            # Correlation groups
            try:
                from tools.review_board.correlator import correlate_findings

                corr = correlate_findings()
                correlation_groups = corr.get("groups", [])
            except Exception:
                pass

            # Remediation stats
            try:
                rem_row = conn.execute(
                    "SELECT COUNT(*) FROM review_board_remediation_log "
                    "WHERE tier = 'auto_fix' AND status IN ('fixed', 'verified') "
                    "AND created_at > datetime('now', '-1 hour')"
                ).fetchone()
                remediation_stats = {"auto_fixes_last_hour": rem_row[0] if rem_row else 0}
            except Exception:
                pass

            return render_template(
                "review_board.html",
                reflexes=reflexes,
                findings=findings,
                severity_summary=severity_summary,
                total_findings=total_findings,
                audit_events=audit_events,
                health_score=health_score,
                health_grade=health_grade,
                health_trend=health_trend,
                health_trend_data=health_trend_data,
                correlation_groups=correlation_groups,
                remediation_stats=remediation_stats,
            )
        except Exception as e:
            import traceback

            traceback.print_exc()
            return render_template(
                "review_board.html",
                reflexes=[],
                findings=[],
                severity_summary={},
                total_findings=0,
                audit_events=[],
                error=str(e),
            )
        finally:
            conn.close()

    @app.route("/api/review-board/status", methods=["GET"])
    def api_review_board_status():
        """Review Board JSON status — daemon CLI with DB fallback."""
        try:
            import subprocess as _sp

            _utf8_env = {**os.environ, "PYTHONIOENCODING": "utf-8"}
            result = _sp.run(
                [sys.executable, "tools/review_board/daemon.py", "--status", "--json"],
                capture_output=True,
                text=True,
                timeout=15,
                cwd=str(BASE_DIR),
                env=_utf8_env,
            )
            if result.returncode == 0 and result.stdout.strip():
                return jsonify(json.loads(result.stdout))
        except Exception:
            pass
        # DB fallback: query review_board_findings for summary
        conn = _get_db()
        try:
            total = conn.execute("SELECT COUNT(*) FROM review_board_findings").fetchone()[0]
            by_sev = {}
            for row in conn.execute(
                "SELECT severity, COUNT(*) as cnt FROM review_board_findings GROUP BY severity"
            ).fetchall():
                by_sev[row["severity"]] = row["cnt"]
            return jsonify({"status": "db_fallback", "total_findings": total, "by_severity": by_sev})
        except Exception as exc:
            return jsonify({"error": str(exc), "status": "unavailable"}), 500
        finally:
            conn.close()

    @app.route("/api/review-board/findings", methods=["GET"])
    def api_review_board_findings():
        """Get review board findings with optional severity filter."""
        severity = flask_request.args.get("severity")
        limit = int(flask_request.args.get("limit", "100"))
        conn = _get_db()
        try:
            if severity:
                rows = conn.execute(
                    "SELECT * FROM review_board_findings WHERE severity = %s ORDER BY created_at DESC LIMIT %s",
                    (severity, limit),
                ).fetchall()
            else:
                rows = conn.execute(
                    "SELECT * FROM review_board_findings ORDER BY created_at DESC LIMIT %s",
                    (limit,),
                ).fetchall()
            return jsonify({"findings": [dict(r) for r in rows], "count": len(rows)})
        except Exception as exc:
            return jsonify({"findings": [], "count": 0, "error": str(exc)})
        finally:
            conn.close()

    @app.route("/api/review-board/reflex/<name>", methods=["POST"])
    def api_review_board_run_reflex(name):
        """Run a single Review Board reflex on-demand — daemon CLI with audit trail."""
        allowed = ["sre", "qa", "security", "perf", "ux", "docs", "product"]
        if name not in allowed:
            return jsonify({"error": f"Unknown reflex: {name}"}), 400
        try:
            import subprocess as _sp

            _utf8_env = {**os.environ, "PYTHONIOENCODING": "utf-8"}
            result = _sp.run(
                [sys.executable, "tools/review_board/daemon.py", "--reflex", name, "--json"],
                capture_output=True,
                text=True,
                timeout=300,
                cwd=str(BASE_DIR),
                env=_utf8_env,
            )
            # Log to audit trail
            try:
                conn = _get_db()
                conn.execute(
                    "INSERT INTO audit_trail (event_type, action, details, created_at) "
                    "VALUES (%s, %s, %s, datetime('now'))",
                    ("config_changed", f"review_board_reflex:{name}", json.dumps({"returncode": result.returncode})),
                )
                conn.commit()
                conn.close()
            except Exception as _exc:  # noqa: BLE001 - best-effort persistence; logged, never raised
                logger.warning(
                    "api_review_board_run_reflex: best-effort INSERT into audit_trail failed (non-blocking): %s",
                    _exc,
                )
            if result.returncode == 0 and result.stdout.strip():
                return jsonify(json.loads(result.stdout))
            return jsonify({"status": "completed", "stdout": result.stdout[:500]}), 200
        except Exception as exc:
            return jsonify({"error": str(exc)}), 500

    # ── Bayesian Autoresearch Dashboard (Phase 67) ────────────────────────────

    _AUTORESEARCH_PLACEHOLDER_NOTE = (
        "Experiment metrics are measured against an identity baseline — the engine "
        "does not apply real code modifications between pre/post measurement, so "
        "deltas are placeholder/heuristic, not real experiment evaluations."
    )

    @app.route("/autoresearch")
    def autoresearch_page():
        """Bayesian Autoresearch — autonomous experiment dashboard."""
        return render_template("autoresearch.html")

    @app.route("/api/autoresearch/summary", methods=["GET"])
    def api_autoresearch_summary():
        """Get autoresearch summary stats."""
        try:
            conn = get_connection(db_path=str(DB_PATH))
            total = conn.execute("SELECT COUNT(*) as cnt FROM experiment_results").fetchone()
            kept = conn.execute("SELECT COUNT(*) as cnt FROM experiment_results WHERE decision = 'keep'").fetchone()
            domains = conn.execute("SELECT DISTINCT domain FROM experiment_results").fetchall()
            best = conn.execute(
                "SELECT MAX(improvement_pct) as best FROM experiment_results WHERE decision = 'keep'"
            ).fetchone()
            conn.close()

            total_count = total["cnt"] if total else 0
            kept_count = kept["cnt"] if kept else 0
            return jsonify(
                {
                    "total_experiments": total_count,
                    "acceptance_rate": round(kept_count / max(total_count, 1) * 100, 1),
                    "active_domains": len(domains) if domains else 0,
                    "best_improvement": round(best["best"] or 0, 2) if best else 0,
                    "placeholder_metrics": True,
                    "heuristic": True,
                    "placeholder_note": _AUTORESEARCH_PLACEHOLDER_NOTE,
                }
            )
        except Exception:
            return jsonify(
                {
                    "total_experiments": 0,
                    "acceptance_rate": 0,
                    "active_domains": 0,
                    "best_improvement": 0,
                    "placeholder_metrics": True,
                    "heuristic": True,
                    "placeholder_note": _AUTORESEARCH_PLACEHOLDER_NOTE,
                }
            )

    @app.route("/api/autoresearch/experiments", methods=["GET"])
    def api_autoresearch_experiments():
        """Get experiment results list."""
        try:
            conn = get_connection(db_path=str(DB_PATH))
            rows = conn.execute("SELECT * FROM experiment_results ORDER BY created_at DESC LIMIT 100").fetchall()
            conn.close()
            return jsonify({
                "experiments": [dict(r) for r in rows],
                "placeholder_metrics": True,
                "heuristic": True,
                "placeholder_note": _AUTORESEARCH_PLACEHOLDER_NOTE,
            })
        except Exception:
            return jsonify({
                "experiments": [],
                "placeholder_metrics": True,
                "heuristic": True,
                "placeholder_note": _AUTORESEARCH_PLACEHOLDER_NOTE,
            })

    # ================================================================
    # Chat: /analyze command — URL fetch + LLM analysis
    # ================================================================

    @app.route("/api/chat/analyze", methods=["POST"])
    def api_chat_analyze():
        """Fetch a URL and return a structured LLM analysis.

        Body: {url: str, canvas_type?: str}
        Returns: {reply, url, source_type, error}
        """
        data = flask_request.get_json(silent=True) or {}
        url = (data.get("url") or "").strip()
        if not url:
            return jsonify({"error": "url required"}), 400
        canvas_type = (data.get("canvas_type") or "intake").strip().lower()
        try:
            from tools.chat_router.url_analyzer import analyze
            result = analyze(url, canvas_type)
            return jsonify(result)
        except Exception as exc:
            app.logger.warning("url_analyzer error: %s", exc)
            return jsonify({"reply": f"[Analyze error: {exc}]", "error": str(exc)}), 500

    # ================================================================
    # Phase 69: Chat Personas API (D-CU-3)
    # ================================================================

    @app.route("/api/chat/route-intent", methods=["POST"])
    def api_chat_route_intent():
        """Classify a user message to a canvas mode for intent-based routing.

        Body: {message: str, context_id: str (optional)}
        Returns: {mode, canvas_type, confidence, reason}
        """
        data = flask_request.get_json(silent=True) or {}
        message = (data.get("message") or "").strip()
        if not message:
            return jsonify({"mode": "intake", "canvas_type": None, "confidence": 1.0, "reason": "empty message"})
        try:
            from tools.chat_router.intent_classifier import classify
            result = classify(message)
            return jsonify(result)
        except Exception as exc:
            app.logger.warning("intent_classifier error: %s", exc)
            return jsonify({"mode": "intake", "canvas_type": None, "confidence": 0.5, "reason": "classifier unavailable"})

    @app.route("/api/chat/personas")
    def api_chat_personas():
        """Return agent persona registry as JSON."""
        try:
            import yaml

            personas_path = BASE_DIR / "args" / "chat_personas.yaml"
            if personas_path.exists():
                with open(personas_path, encoding="utf-8") as f:
                    data = yaml.safe_load(f) or {}
                return jsonify(data)
            return jsonify({"personas": {}})
        except Exception as exc:
            return jsonify({"error": str(exc)}), 500

    @app.route("/api/chat/upload", methods=["POST"])
    def api_chat_upload():
        """Ingest an uploaded document into RAG vector store and Knowledge Graph.

        Multipart form fields:
            file        — the document (txt, md, pdf, docx)
            context_id  — chat context (used to tag chunks with project_id)
            project_id  — optional project scope
            tenant_id   — optional tenant scope
        """
        import hashlib
        import tempfile
        import uuid as _uuid

        file = flask_request.files.get("file")
        if not file or not file.filename:
            return jsonify({"error": "No file provided"}), 400

        context_id = flask_request.form.get("context_id", "")
        project_id = flask_request.form.get("project_id", "")
        tenant_id = flask_request.form.get("tenant_id", "")
        filename = file.filename

        # --- Extract text ---
        _IMAGE_EXTS = frozenset({".png", ".jpg", ".jpeg", ".gif", ".webp", ".svg", ".bmp", ".tiff"})
        try:
            ext = Path(filename).suffix.lower()
            raw_bytes = file.read()
            file_size_kb = round(len(raw_bytes) / 1024, 1)
            is_image = ext in _IMAGE_EXTS

            if is_image:
                # Images: synthesise a searchable text stub so they appear in RAG + KG
                stem = Path(filename).stem.replace("_", " ").replace("-", " ")
                text = (
                    f"Image document: {filename}\n"
                    f"Type: {ext.lstrip('.')} image\n"
                    f"Size: {file_size_kb} KB\n"
                    f"Description: {stem}\n"
                    f"This is a diagram, figure, or image file uploaded to the chat context."
                )
            elif ext in (".txt", ".md", ".rst", ".yaml", ".yml", ".json", ".xml", ".csv"):
                text = raw_bytes.decode("utf-8", errors="replace")
            elif ext == ".pdf":
                try:
                    from tools.rag.pdf_provider import extract_text
                    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                        tmp.write(raw_bytes)
                        tmp_path = tmp.name
                    text = extract_text(tmp_path)
                    Path(tmp_path).unlink(missing_ok=True)
                except Exception as exc:
                    return jsonify({"error": f"PDF extraction failed: {exc}"}), 422
            elif ext in (".docx", ".doc"):
                try:
                    import io
                    import docx as _docx
                    doc = _docx.Document(io.BytesIO(raw_bytes))
                    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                    # Also extract table content
                    for tbl in doc.tables:
                        for row in tbl.rows:
                            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                            if row_text:
                                paragraphs.append(row_text)
                    text = "\n".join(paragraphs)
                except Exception as exc:
                    return jsonify({"error": f"DOCX extraction failed: {exc}"}), 422
            else:
                # Generic fallback: try UTF-8, mark as unknown type
                text = raw_bytes.decode("utf-8", errors="replace")

            if not text.strip():
                return jsonify({"error": "Document contains no extractable text"}), 422
        except Exception as exc:
            return jsonify({"error": f"Text extraction error: {exc}"}), 500

        doc_id = hashlib.sha256(text[:512].encode()).hexdigest()[:16]
        rag_chunks_stored = 0
        kg_result = {}

        # --- RAG ingestion: chunk + embed + upsert ---
        try:
            from tools.rag.chunker import chunk_content
            from tools.rag.vector_store_factory import VectorStoreFactory

            chunks = chunk_content(
                content=text,
                source_type="chat_upload",
                source_id=doc_id,
                source_table="chat_uploads",
                metadata={"filename": filename, "context_id": context_id},
                tenant_id=tenant_id,
                project_id=project_id,
                classification="CUI",
            )

            try:
                from tools.llm import get_embedding_provider
                emb_provider = get_embedding_provider()
            except Exception:
                emb_provider = None

            store = VectorStoreFactory.create(tenant_id=tenant_id)
            for chunk in chunks:
                if not chunk.chunk_id:
                    chunk.chunk_id = str(_uuid.uuid4())
                chunk.compute_content_hash()
                if emb_provider:
                    try:
                        chunk.embedding = emb_provider.embed(chunk.content)
                    except Exception:
                        pass
            rag_chunks_stored = store.upsert(chunks)
        except Exception as exc:
            app.logger.warning("RAG ingestion failed for upload %s: %s", filename, exc)

        # --- KG ingestion: entity + relationship extraction ---
        try:
            from tools.knowledge_graph.ingester import ingest_file as kg_ingest_file
            with tempfile.NamedTemporaryFile(
                suffix=Path(filename).suffix or ".txt",
                mode="w",
                encoding="utf-8",
                delete=False,
            ) as tmp:
                tmp.write(text)
                tmp_path = tmp.name
            kg_result = kg_ingest_file(
                file_path=tmp_path,
                project_id=project_id or context_id or "chat",
                graph_name=Path(filename).stem,
            )
            Path(tmp_path).unlink(missing_ok=True)
        except Exception as exc:
            app.logger.warning("KG ingestion failed for upload %s: %s", filename, exc)
            kg_result = {"status": "unavailable", "error": str(exc)}

        return jsonify({
            "status": "ok",
            "filename": filename,
            "doc_id": doc_id,
            "file_type": "image" if is_image else ext.lstrip("."),
            "file_size_kb": file_size_kb,
            "rag_chunks": rag_chunks_stored,
            "kg": kg_result,
            "text_length": len(text),
        })

    @app.route("/api/chat/sources")
    def api_chat_sources():
        """List documents indexed into RAG from chat uploads."""
        tenant_id = flask_request.args.get("tenant_id", "")
        context_id = flask_request.args.get("context_id", "")
        try:
            with get_connection() as conn:
                sources = _aggregate_chat_sources(conn, tenant_id, context_id)
                return jsonify({"sources": sources, "total": len(sources)})
        except Exception as exc:
            app.logger.warning("api_chat_sources error: %s", exc)
            return jsonify({"sources": [], "total": 0})

    # ================================================================
    # Chat Use Cases catalog (FORGE-pattern — reads args/use_cases.yaml)
    # ================================================================

    # ---- Use Case catalog helpers (YAML defaults + DB overrides) ----

    def _uc_load_yaml_cases():
        import yaml as _yaml
        _uc_path = BASE_DIR / "args" / "use_cases.yaml"
        try:
            with open(_uc_path, "r", encoding="utf-8") as _fh:
                return (_yaml.safe_load(_fh) or {}).get("use_cases", [])
        except Exception:
            return []

    def _uc_init_table(_conn):
        _conn.execute("""CREATE TABLE IF NOT EXISTS use_case_overrides (
            id TEXT PRIMARY KEY,
            label TEXT, description TEXT, icon TEXT, badge TEXT,
            agent_model TEXT, ricoas INTEGER, boost_threshold INTEGER,
            system_prompt TEXT, seed_message TEXT,
            canvas_wiring TEXT, quick_actions TEXT,
            updated_at TEXT, updated_by TEXT,
            classification TEXT DEFAULT NULL
        )""")
        try:
            _conn.execute("ALTER TABLE use_case_overrides ADD COLUMN classification TEXT DEFAULT NULL")
        except Exception:
            pass
        _conn.commit()

    def _uc_apply_override(_base, _row):
        import json as _json
        if not _row:
            return _base
        _result = dict(_base)
        for _col in ("label", "description", "icon", "badge", "agent_model",
                     "system_prompt", "seed_message"):
            if _row[_col] is not None:
                _result[_col] = _row[_col]
        if _row["ricoas"] is not None:
            _result["ricoas"] = bool(_row["ricoas"])
        if _row["boost_threshold"] is not None:
            _result["boost_threshold"] = _row["boost_threshold"]
        for _jcol in ("canvas_wiring", "quick_actions"):
            if _row[_jcol] is not None:
                try:
                    _result[_jcol] = _json.loads(_row[_jcol])
                except Exception:
                    pass
        return _result

    @app.route("/api/chat/use-cases", methods=["GET"])
    def api_chat_use_cases():
        """Return use case catalog (YAML defaults merged with DB overrides)."""
        _cases = _uc_load_yaml_cases()
        _category = flask_request.args.get("category", "").strip().lower()
        _q = flask_request.args.get("q", "").strip().lower()

        _overrides = {}
        try:
            with get_connection() as _conn:
                _uc_init_table(_conn)
                for _r in _conn.execute("SELECT * FROM use_case_overrides").fetchall():
                    _overrides[_r["id"]] = _r
        except Exception as _exc:
            app.logger.debug("use_case_overrides load: %s", _exc)

        _merged = [_uc_apply_override(c, _overrides.get(c.get("id", ""))) for c in _cases]
        if _category:
            _merged = [c for c in _merged if c.get("category", "").lower() == _category]
        if _q:
            _merged = [c for c in _merged if _q in c.get("label", "").lower()
                       or _q in (c.get("description") or "").lower()]

        _summary = [
            {
                "id": c.get("id", ""),
                "label": c.get("label", ""),
                "category": c.get("category", ""),
                "icon": c.get("icon", ""),
                "description": (c.get("description") or "").strip(),
                "badge": c.get("badge", ""),
                "agent_model": c.get("agent_model", "sonnet"),
                "ricoas": c.get("ricoas", False),
                "boost_threshold": c.get("boost_threshold", 70),
                "canvas_wiring": c.get("canvas_wiring", []),
                "quick_actions": c.get("quick_actions", []),
            }
            for c in _merged
        ]
        return jsonify({"use_cases": _summary, "total": len(_summary)})

    @app.route("/api/chat/use-cases/<use_case_id>", methods=["GET"])
    def api_chat_use_case_detail(use_case_id):
        """Return full use case definition (YAML + DB override merged)."""
        _base = next((c for c in _uc_load_yaml_cases() if c.get("id") == use_case_id), None)
        if not _base:
            return jsonify({"error": "Use case not found"}), 404
        try:
            with get_connection() as _conn:
                _uc_init_table(_conn)
                _row = _conn.execute(
                    "SELECT * FROM use_case_overrides WHERE id=%s", (use_case_id,)
                ).fetchone()
        except Exception:
            _row = None
        return jsonify(_uc_apply_override(dict(_base), _row))

    @app.route("/api/chat/use-cases/<use_case_id>", methods=["PUT"])
    def api_chat_use_case_update(use_case_id):
        """Persist user overrides for a use case (YAML unchanged, overrides in DB)."""
        import json as _json
        from datetime import datetime, timezone
        _body = flask_request.get_json(silent=True) or {}
        _now = datetime.now(timezone.utc).isoformat()
        _cw = _body.get("canvas_wiring")
        _qa = _body.get("quick_actions")
        try:
            with get_connection() as _conn:
                _uc_init_table(_conn)
                _conn.execute("""
                    INSERT INTO use_case_overrides
                        (id,label,description,icon,badge,agent_model,ricoas,
                         boost_threshold,system_prompt,seed_message,
                         canvas_wiring,quick_actions,updated_at,updated_by)
                    VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
                    ON CONFLICT(id) DO UPDATE SET
                        label=excluded.label, description=excluded.description,
                        icon=excluded.icon, badge=excluded.badge,
                        agent_model=excluded.agent_model, ricoas=excluded.ricoas,
                        boost_threshold=excluded.boost_threshold,
                        system_prompt=excluded.system_prompt,
                        seed_message=excluded.seed_message,
                        canvas_wiring=excluded.canvas_wiring,
                        quick_actions=excluded.quick_actions,
                        updated_at=excluded.updated_at,
                        updated_by=excluded.updated_by
                """, (
                    use_case_id,
                    _body.get("label"), _body.get("description"), _body.get("icon"),
                    _body.get("badge"), _body.get("agent_model"),
                    1 if _body.get("ricoas") else 0,
                    _body.get("boost_threshold"),
                    _body.get("system_prompt"), _body.get("seed_message"),
                    _json.dumps(_cw) if _cw is not None else None,
                    _json.dumps(_qa) if _qa is not None else None,
                    _now, _body.get("updated_by", "dashboard-user"),
                ))
                _conn.commit()
        except Exception as _exc:
            app.logger.error("use_case_overrides save: %s", _exc)
            return jsonify({"error": str(_exc)}), 500
        return jsonify({"ok": True, "id": use_case_id, "updated_at": _now})

    @app.route("/api/chat/use-cases/<use_case_id>/override", methods=["DELETE"])
    def api_chat_use_case_reset(use_case_id):
        """Delete DB override — restores YAML factory defaults."""
        try:
            with get_connection() as _conn:
                _uc_init_table(_conn)
                _conn.execute("DELETE FROM use_case_overrides WHERE id=%s", (use_case_id,))
                _conn.commit()
        except Exception as _exc:
            return jsonify({"error": str(_exc)}), 500
        return jsonify({"ok": True, "id": use_case_id, "reset": True})

    # ================================================================
    # Phase 69: Codebase Assistant API (D-CA-5 to D-CA-8)
    # ================================================================

    @app.route("/api/assistant/query", methods=["POST"])
    def api_assistant_query():
        """Handle codebase assistant queries."""
        try:
            from tools.dashboard.assistant_manager import query

            data = flask_request.get_json(force=True, silent=True) or {}
            result = query(
                question=data.get("question", ""),
                scope=data.get("scope"),
                context_id=data.get("context_id"),
                page_path=data.get("page_path"),
            )
            return jsonify(result)
        except Exception as exc:
            app.logger.error("Assistant query error: %s", exc)
            return jsonify({"answer": f"Error: {exc}", "citations": [], "source": "error"}), 500

    @app.route("/api/assistant/status")
    def api_assistant_status():
        """Return codebase indexer status."""
        try:
            from tools.dashboard.assistant_manager import get_status

            return jsonify(get_status())
        except Exception as exc:
            return jsonify({"indexed_files": 0, "index_status": "unavailable", "error": str(exc)})

    @app.route("/api/assistant/scope", methods=["POST"])
    def api_assistant_scope():
        """Set assistant scope to a specific module."""
        try:
            from tools.dashboard.assistant_config import files_in_scope

            data = flask_request.get_json(force=True, silent=True) or {}
            scope = data.get("scope", "")
            files = files_in_scope(scope) if scope else []
            return jsonify({"ok": True, "files_in_scope": len(files)})
        except Exception as exc:
            return jsonify({"ok": False, "error": str(exc)}), 500

    @app.route("/api/assistant/suggestions")
    def api_assistant_suggestions():
        """Return contextual question suggestions for the widget."""
        try:
            from tools.dashboard.assistant_manager import get_suggestions

            page_path = flask_request.args.get("page_path", "")
            return jsonify({"suggestions": get_suggestions(page_path)})
        except Exception:
            return jsonify(
                {
                    "suggestions": [
                        "How is the ICDEV™ codebase structured?",
                        "What does the LLM router do?",
                        "How does the RAG retriever work?",
                    ]
                }
            )

    # ── ClawHub Skill Browser (Phase 69) ───────────────────────────────

    @app.route("/clawhub")
    def clawhub():
        """ClawHub — discover and import OpenClaw skills."""
        imports = []
        enabled = os.environ.get("ICDEV_OPENCLAW_ENABLED", "").lower() in ("true", "1", "yes")
        try:
            from tools.marketplace.openclaw_bridge import list_quarantine

            result = list_quarantine()
            if result.get("success"):
                imports = [
                    (
                        i.get("id", ""),
                        i.get("skill_name", ""),
                        i.get("author", i.get("openclaw_author", "")),
                        i.get("scan_status", ""),
                        i.get("status", ""),
                        i.get("trust_score", 0.3),
                        i.get("has_scripts", i.get("has_executable_content", False)),
                        i.get("review_required", False),
                        str(i.get("created_at", ""))[:19],
                        i.get("rejected_by", ""),
                        i.get("rejected_reason", ""),
                        i.get("failed_gates", []),
                    )
                    for i in result.get("imports", [])
                ]
        except Exception:
            imports = []
        return render_template("clawhub.html", imports=imports, enabled=enabled)













    # ── ICDEV™ Studio (Phase 72) ─────────────────────────────────────

    @app.route("/studio/workflows")
    def studio_workflows():
        """Studio — Visual Workflow Editor."""
        return render_template("studio/workflow_studio.html")

    @app.route("/studio/marketplace")
    def studio_marketplace():
        """Studio — Marketplace Storefront."""
        return render_template("studio/marketplace.html")

    @app.route("/studio/app-builder")
    def studio_app_builder():
        """Studio — NL App Builder."""
        return render_template("studio/app_builder.html")

    @app.route("/studio/dashboards")
    def studio_dashboards():
        """Studio — Dashboard Builder."""
        return render_template("studio/dashboards.html")

    @app.route("/studio/automations")
    def studio_automations():
        """Studio — Citizen Automation Studio."""
        return render_template("studio/automations.html")

    @app.route("/studio/forms")
    def studio_forms():
        """Studio — Form Builder."""
        return render_template("studio/forms.html")

    @app.route("/studio/cases")
    def studio_cases():
        """Studio — Case Management."""
        return render_template("studio/cases.html")

    # ---- Platform Health page + API (Phase 73) ----
    @app.route("/platform-health", methods=["GET"])
    def platform_health_page():
        """Platform Health detail page — all 10 domains with drill-down."""
        return render_template("platform_health.html")

    try:
        # Only get_platform_health is used here; get_domain_health is imported
        # locally inside the per-domain handler below (see _gdh alias at L7175).
        from tools.dashboard.platform_health import get_platform_health  # noqa: E402

        @app.route("/api/platform/health", methods=["GET"])
        def api_platform_health():
            """GET /api/platform/health — Composite platform health across 10 domains."""
            from tools.dashboard.platform_health import _invalidate_cache  # noqa: E402

            if flask_request.args.get("invalidate") == "1":  # noqa: F821
                _invalidate_cache()
            result = get_platform_health()
            # Shape domains for API response (omit all_findings for brevity)
            return jsonify(
                {
                    "composite_score": result["composite_score"],
                    "composite_status": result["composite_status"],
                    "cached_at": result.get("cached_at"),
                    "domains": result["domains"],
                }
            )

        @app.route("/api/platform/health/<domain>", methods=["GET"])
        def api_platform_health_domain(domain: str):
            """GET /api/platform/health/<domain> — Detailed findings for one domain.

            Each domain probe in tools/dashboard/platform_health.py opens a
            storage connection and runs deterministic SELECT COUNT(*) checks
            against backend tables (PG or SQLite). Results are not cached so
            every call hits the database.
            """
            from tools.dashboard.platform_health import get_domain_health as _gdh  # noqa: E402

            detail = _gdh(domain)
            return jsonify(
                {
                    "domain": domain,
                    **detail,
                }
            )

    except ImportError as _ph_err:
        import logging as _logging

        _logging.getLogger(__name__).warning("Platform health module unavailable: %s", _ph_err)

    # ---- Air-gap Next.js static export ----
    try:
        from tools.airgap.detector import is_airgap as _is_airgap

        if _is_airgap():
            _next_out = str(Path(__file__).resolve().parent.parent.parent / "frontend" / "out")

            @app.route("/next/<path:path>")
            def next_static(path):
                return send_from_directory(_next_out, path)

            app.logger.info("Air-gap Next.js static route registered at /next/")
    except Exception as _ag_err:
        import logging as _logging

        _logging.getLogger(__name__).warning("Air-gap Next.js route skipped: %s", _ag_err)

    # ── Projects-in-Flight registry ─────────────────────────────────────────
    # Config-driven: args/projects.yaml. Each project renders as a collapsible
    # card on Home below the Task Board via tools/dashboard/templates/
    # _projects_in_flight.html. A project auto-hides when every task in its
    # prefix is done (total > 0 and done == total) or when no tasks match
    # the prefix yet (total == 0). Adding a project is a YAML edit.
    #
    # Invariants enforced here:
    #   * task_prefix and every epic.key MUST be present — silently dropped
    #     entries would produce empty cards that confuse operators.
    #   * LIKE wildcards (% and _) in prefix / keys are escaped with \ —
    #     prevents cross-project row leakage and SQL-wildcard surprises.
    #   * Within a project, no epic.key may be a prefix of another when both
    #     use the '-' separator (e.g. 'idc' + 'idc-iac' would double-count).
    #   * Across projects, no task_prefix may be a prefix of another — that
    #     would mean project A's query silently captures project B's tasks.
    #
    # Violations are logged via the app logger and cause the offending entry
    # (or epic) to be dropped from the rendered output so the rest still
    # works.
    import logging as _proj_logging
    _proj_log = _proj_logging.getLogger(__name__ + ".projects")

    _LIKE_ESCAPE = "\\"

    def _escape_like(s: str) -> str:
        r"""Escape LIKE metacharacters (% _ \) for use with `LIKE ? ESCAPE '\\'`."""
        if not isinstance(s, str):
            return ""
        return (s.replace(_LIKE_ESCAPE, _LIKE_ESCAPE + _LIKE_ESCAPE)
                 .replace("%", _LIKE_ESCAPE + "%")
                 .replace("_", _LIKE_ESCAPE + "_"))

    def _validate_projects(raw: list) -> list:
        """Validate + normalize project entries. Drops invalid ones with a
        logged warning so the page keeps rendering the rest.

        Nested prefixes (`aadc-` alongside `aadc-enh-`) are a legitimate
        parent/child namespace, not an error: the parent keeps rendering and
        gets an `exclude_prefixes` list so its queries subtract every
        registered child. Only an EXACT duplicate prefix is unresolvable and
        skipped. Previously the later entry was dropped outright, which
        silently hid whichever card happened to appear last in the YAML.
        """
        from tools.project.prefix_scope import child_prefixes

        out: list = []
        seen_prefixes: list = []
        seen_keys: set = set()
        # Every valid prefix in the file, needed up-front so a parent entry can
        # claim its children's exclusions regardless of YAML ordering.
        all_prefixes = {
            (p.get("task_prefix") or "").strip()
            for p in raw
            if isinstance(p, dict) and (p.get("task_prefix") or "").strip()
        }
        for i, p in enumerate(raw):
            if not isinstance(p, dict):
                _proj_log.warning("projects.yaml entry #%d is not a dict — skipping", i)
                continue
            key = (p.get("key") or "").strip()
            prefix = (p.get("task_prefix") or "").strip()
            if not key:
                _proj_log.warning("projects.yaml entry #%d missing 'key' — skipping", i)
                continue
            if not prefix:
                _proj_log.warning("projects.yaml '%s' missing 'task_prefix' — skipping", key)
                continue
            if key in seen_keys:
                _proj_log.warning("projects.yaml duplicate key '%s' — skipping", key)
                continue
            # Exact duplicate prefix — genuinely ambiguous, no way to split rows.
            if prefix in seen_prefixes:
                _proj_log.warning(
                    "projects.yaml '%s' prefix %r duplicates an earlier entry "
                    "— tasks would double-count. Skipping.",
                    key, prefix,
                )
                continue
            # Nested child prefixes: this entry is their parent, so subtract them.
            exclude_prefixes = child_prefixes(prefix, all_prefixes)
            # Within-project epic key prefix-of collision
            raw_epics = p.get("epics", []) or []
            clean_epics: list = []
            ekeys_sorted = sorted(
                [(e.get("key") or "").strip() for e in raw_epics if isinstance(e, dict)],
                key=len,
            )
            bad_ekeys: set = set()
            for a in range(len(ekeys_sorted)):
                for b in range(a + 1, len(ekeys_sorted)):
                    short, long = ekeys_sorted[a], ekeys_sorted[b]
                    if short and long and long.startswith(short + "-"):
                        _proj_log.warning(
                            "projects.yaml '%s' epic keys '%s' and '%s' overlap "
                            "(one is a prefix of the other under '-' separator). "
                            "Keeping '%s' only.",
                            key, short, long, long,
                        )
                        bad_ekeys.add(short)
            for ep in raw_epics:
                if not isinstance(ep, dict):
                    continue
                ek = (ep.get("key") or "").strip()
                if not ek or ek in bad_ekeys:
                    continue
                if not ep.get("title"):
                    _proj_log.warning(
                        "projects.yaml '%s' epic '%s' missing title — skipping epic",
                        key, ek,
                    )
                    continue
                clean_epics.append(ep)
            p2 = dict(p)
            p2["task_prefix"] = prefix
            p2["exclude_prefixes"] = exclude_prefixes
            p2["epics"] = clean_epics
            out.append(p2)
            seen_keys.add(key)
            seen_prefixes.append(prefix)
        return out

    def _load_projects_yaml() -> list:
        try:
            import yaml as _yaml  # PyYAML — declared dep
        except Exception as exc:
            _proj_log.warning("PyYAML import failed (%s); projects panel disabled", exc)
            return []
        cfg_path = Path(__file__).resolve().parent.parent.parent / "args" / "projects.yaml"
        if not cfg_path.exists():
            return []
        try:
            data = _yaml.safe_load(cfg_path.read_text(encoding="utf-8")) or {}
            raw = list(data.get("projects", []))
        except Exception as exc:
            _proj_log.warning("projects.yaml parse failed: %s", exc)
            return []
        return _validate_projects(raw)

    def _compute_project_progress(project: dict, conn) -> dict:
        """Query kanban_tasks for one project's epics + in-flight + failures.

        All LIKE patterns derived from YAML are escaped with `ESCAPE '\\'`
        so a malformed prefix or epic key can't leak across projects or
        match wildcards unintentionally.

        `exclude_prefixes` (set by _validate_projects) subtracts every nested
        child project, so a parent prefix like `aadc-` does not absorb
        `aadc-enh-` / `aadc-sp-` rows.
        """
        prefix = project.get("task_prefix", "")
        prefix_esc = _escape_like(prefix)
        excludes = project.get("exclude_prefixes") or []
        # One `AND id NOT LIKE ... ESCAPE` clause per nested child prefix.
        excl_sql = "".join(" AND id NOT LIKE %s ESCAPE '\\'" for _ in excludes)
        excl_params = tuple(f"{_escape_like(x)}%" for x in excludes)
        epics_out: list = []
        total_all = 0
        done_all = 0
        for ep in project.get("epics", []):
            ek_esc = _escape_like(ep["key"])
            pattern = f"{prefix_esc}{ek_esc}-%"
            rows = conn.execute(
                "SELECT status, COUNT(*) AS n FROM kanban_tasks "
                "WHERE id LIKE %s ESCAPE '\\'" + excl_sql + " GROUP BY status",
                (pattern,) + excl_params,
            ).fetchall()
            counts = {dict(r)["status"]: int(dict(r)["n"]) for r in rows}
            total = sum(counts.values())
            done = counts.get("done", 0)
            pct = int(round(100 * done / total)) if total else 0
            total_all += total
            done_all += done
            epics_out.append({
                "key": ep["key"],
                "title": ep["title"],
                "priority": ep.get("priority", "medium"),
                "total": total,
                "done": done,
                "in_progress": counts.get("in_progress", 0),
                "scheduled": counts.get("scheduled", 0),
                "backlog": counts.get("backlog", 0),
                "failed": counts.get("failed", 0),
                "needs_decomp": counts.get("needs_decomposition", 0),
                "pct": pct,
            })
        in_flight_rows = conn.execute(
            "SELECT id, title, status, priority, updated_at "
            "FROM kanban_tasks WHERE id LIKE %s ESCAPE '\\' " + excl_sql +
            "  AND status IN ('in_progress','scheduled') "
            "ORDER BY updated_at DESC LIMIT 15",
            (f"{prefix_esc}%",) + excl_params,
        ).fetchall()
        fail_rows = conn.execute(
            "SELECT id, title, status, failure_count, "
            "       last_failure_reason, updated_at "
            "FROM kanban_tasks WHERE id LIKE %s ESCAPE '\\' " + excl_sql +
            "  AND last_failure_reason IS NOT NULL "
            "ORDER BY updated_at DESC LIMIT 10",
            (f"{prefix_esc}%",) + excl_params,
        ).fetchall()
        return {
            "key": project.get("key"),
            "name": project.get("name"),
            "description": project.get("description", "").strip(),
            "default_open": bool(project.get("default_open", True)),
            "briefs": project.get("briefs", []),
            "epics": epics_out,
            "in_flight": [dict(r) for r in in_flight_rows],
            "recent_failures": [dict(r) for r in fail_rows],
            "total_tasks": total_all,
            "done_tasks": done_all,
            "overall_pct": int(round(100 * done_all / total_all)) if total_all else 0,
            "visible": total_all > 0 and done_all < total_all,
        }

    def _compute_triage_summary() -> dict:
        """Global failure_triage audit summary — applies to all projects."""
        summary = {"total": 0, "applied": 0, "rejected": 0, "verification_failed": 0}
        recent: list = []
        audit_dir = (Path(__file__).resolve().parent.parent.parent
                     / ".tmp" / "kanban" / "autofix-audit")
        if not audit_dir.exists():
            return {"summary": summary, "recent": recent}
        try:
            files = sorted(audit_dir.glob("*.json"),
                           key=lambda p: p.stat().st_mtime, reverse=True)
            summary["total"] = len(files)
            for f in files:
                try:
                    data = json.loads(f.read_text(encoding="utf-8"))
                    outcome = data.get("outcome") or ""
                    if outcome.startswith("applied_"):
                        summary["applied"] += 1
                    elif outcome.startswith("rejected_"):
                        summary["rejected"] += 1
                    elif outcome == "verification_failed":
                        summary["verification_failed"] += 1
                    if len(recent) < 5:
                        recent.append({
                            "task_id": data.get("task_id"),
                            "outcome": outcome,
                            "branch": data.get("branch"),
                            "started_at": data.get("started_at"),
                        })
                except Exception:
                    continue
        except Exception:  # pragma: no cover
            pass
        return {"summary": summary, "recent": recent}

    @app.route("/api/projects/progress")
    def api_projects_progress():
        """GET /api/projects/progress — JSON snapshot of every in-flight
        project. Auto-filters out projects with 0 tasks or 100% done."""
        from tools.db.storage import get_connection as _gc

        projects_cfg = _load_projects_yaml()
        out: list = []
        try:
            with _gc() as conn:
                for p in projects_cfg:
                    snap = _compute_project_progress(p, conn)
                    if snap["visible"]:
                        out.append(snap)
        except Exception as exc:
            return jsonify({"error": str(exc), "projects": []}), 500
        triage = _compute_triage_summary()
        return jsonify({"projects": out, "triage": triage})

    # ── Autonomous Recovery panel ───────────────────────────────────────────
    # Surfaces failure_triage activity + autofix branches + recent failures
    # on Home below Projects in Flight. Auto-hides when no activity —
    # "idle" means no triage markers in the last 24h, no autofix branches,
    # no unresolved failures in the last hour.
    @app.route("/api/autonomy/status")
    def api_autonomy_status():
        """GET /api/autonomy/status — autonomous-flow snapshot:
          * recent triage markers (last 24h)
          * active autofix branches
          * unresolved failures (last 1h)
          * failure_triage audit summary (global)
        All three empty → partial's host page hides the section entirely.
        """
        from tools.db.storage import get_connection as _gc
        import subprocess as _sp

        base_dir = Path(__file__).resolve().parent.parent.parent

        # 1. Recent triage markers (last 24h) — sorted newest-first
        triage_recent: list = []
        triaged_dir = base_dir / ".tmp" / "kanban" / "triaged"
        if triaged_dir.exists():
            cutoff = time.time() - 86400
            for f in sorted(
                triaged_dir.glob("*.marker"),
                key=lambda p: p.stat().st_mtime,
                reverse=True,
            )[:10]:
                if f.stat().st_mtime < cutoff:
                    continue
                try:
                    data = json.loads(f.read_text(encoding="utf-8"))
                    outcome = data.get("outcome") or {}
                    diag = outcome.get("diagnosis") or {}
                    gate = outcome.get("autofix_gate") or {}
                    triage_recent.append({
                        "task_id": data.get("task_id") or outcome.get("task_id"),
                        "title": outcome.get("title"),
                        "signature": data.get("sig"),
                        "recommendation": diag.get("recommendation"),
                        "confidence": diag.get("confidence"),
                        "gate_reason": gate.get("reason"),
                        "outcome": outcome.get("outcome"),
                        "ts": data.get("ts"),
                    })
                except Exception:
                    continue

        # 2. Autofix branches — each represents a real applied patch
        autofix_branches: list = []
        try:
            r = _sp.run(
                ["git", "branch", "--list", "autofix/*"],
                cwd=str(base_dir), capture_output=True, text=True, timeout=5,
            )
            for line in r.stdout.splitlines():
                name = line.replace("*", "").strip()
                if not name:
                    continue
                # Get the commit message for context
                try:
                    msg = _sp.run(
                        ["git", "log", "-1", "--format=%s|%ci", name],
                        cwd=str(base_dir), capture_output=True, text=True, timeout=5,
                    ).stdout.strip()
                except Exception:
                    msg = ""
                subject, _, when = msg.partition("|")
                autofix_branches.append({
                    "branch": name,
                    "subject": subject[:80],
                    "ts": when,
                })
        except Exception:
            pass

        # 3. Unresolved failures in the last hour — kanban tasks with
        # last_failure_reason set that are back in backlog/scheduled/failed
        unresolved_failures: list = []
        try:
            cutoff_iso = datetime.now(timezone.utc).replace(microsecond=0)
            from datetime import timedelta as _td
            cutoff_iso = (cutoff_iso - _td(hours=1)).isoformat()
            with _gc() as conn:
                rows = conn.execute(
                    "SELECT id, title, status, failure_count, "
                    "       last_failure_reason, updated_at "
                    "FROM kanban_tasks "
                    "WHERE last_failure_reason IS NOT NULL "
                    "  AND updated_at > %s "
                    "  AND status IN ('backlog','failed','scheduled') "
                    "ORDER BY updated_at DESC LIMIT 10",
                    (cutoff_iso,),
                ).fetchall()
                for r in rows:
                    d = dict(r)
                    unresolved_failures.append({
                        "id": d.get("id"),
                        "title": d.get("title"),
                        "status": d.get("status"),
                        "failure_count": d.get("failure_count"),
                        "reason": (d.get("last_failure_reason") or "")[:200],
                        "updated_at": d.get("updated_at"),
                    })
        except Exception:
            pass

        # 4. Global triage summary
        triage = _compute_triage_summary()

        visible = bool(triage_recent) or bool(autofix_branches) or bool(unresolved_failures)
        return jsonify({
            "visible": visible,
            "triage_recent": triage_recent,
            "autofix_branches": autofix_branches,
            "unresolved_failures": unresolved_failures,
            "triage_summary": triage.get("summary", {}),
        })

    @app.route("/digital-twin")
    def digital_twin_roadmap_legacy():
        """Legacy route — project moved inline under /kanban. Redirect to
        the anchor so old links still work."""
        return redirect("/kanban#project-digital-twin", code=302)

    # ── FathomDesk News Intelligence ─────────────────────────────────────────
    @app.route("/news")
    def news_page():
        """FathomDesk News — category-tab layout with sentiment sparklines."""
        return render_template("news.html")

    # ── FathomDesk Options Chain ──────────────────────────────────────────────
    @app.route("/options")
    def options_page():
        """Options page — IV Rank/Percentile badge and chain viewer."""
        ticker = flask_request.args.get("ticker", "SPY").upper()
        return render_template("options.html", ticker=ticker)

    # ── FathomDesk Trading Engine ─────────────────────────────────────────────
    @app.route("/fathomdesk")
    def fathomdesk_page():
        """FathomDesk — trading chart with volume profile overlay."""
        ticker = flask_request.args.get("ticker", "SPY").upper()
        return render_template("fathomdesk.html", ticker=ticker)

    @app.route("/fathomdesk/trap-events")
    def fathomdesk_trap_events():
        """FathomDesk — full trap event history with filters."""
        return render_template("fathomdesk_trap_events.html")

    @app.route("/analysis")
    def analysis_page():
        """Market Analysis — Macro Intelligence, IV Skew & Term Structure."""
        return render_template("analysis.html")

    @app.route("/quality-scores")
    def quality_scores_page():
        """Quality Scores — PE/NAV mispricing dashboard."""
        return render_template("quality_scores.html")

    @app.route("/api/macro/intelligence")
    def api_macro_intelligence():
        """Return macro regime badges for the /analysis page panel."""
        try:
            from tools.trading.data.macro_data import fetch_macro_context
            ctx = fetch_macro_context()
            return jsonify({
                "qeqt_phase": ctx.get("qeqt_phase", "NEUTRAL"),
                "credit_stress": ctx.get("credit_impulse", {}).get("label", "NEUTRAL"),
                "rotation_signal": ctx.get("regime", "UNKNOWN"),
                "macro_score": ctx.get("macro_score"),
                "summary": ctx.get("summary", ""),
                "fetched_at": ctx.get("fetched_at"),
            })
        except Exception as exc:
            # nav-plat-04: a data outage must NOT masquerade as a benign NEUTRAL
            # regime. Log it and return an explicit error state (HTTP 503) with
            # null badges so no consumer renders NEUTRAL for missing data.
            app.logger.exception("macro/intelligence fetch failed: %s", exc)
            return jsonify({
                "status": "error",
                "detail": str(exc),
                "qeqt_phase": None,
                "credit_stress": None,
                "rotation_signal": None,
                "macro_score": None,
                "summary": "",
                "fetched_at": None,
                "error": str(exc),
            }), 503

    @app.route("/api/trading/market")
    def api_trading_market():
        """Return macro context snapshot including qeqt_phase for FathomDesk overlay."""
        try:
            from tools.trading.data.macro_data import fetch_macro_context, fetch_extended_macro
            ctx = fetch_macro_context()
            ext = fetch_extended_macro()
            return jsonify({
                "macro_score": ctx.get("macro_score"),
                "regime": ctx.get("regime"),
                "qeqt_phase": ctx.get("qeqt_phase"),
                "qeqt_magnitude": ext.get("qeqt_magnitude"),
                "fed_bs_4w_roc_b": ext.get("fed_bs_4w_roc_b"),
                "fed_bs_13w_roc_b": ext.get("fed_bs_13w_roc_b"),
                "data_source": ctx.get("data_source"),
                "fetched_at": ctx.get("fetched_at"),
                "summary": ctx.get("summary"),
            })
        except Exception as exc:
            return jsonify({"error": str(exc)}), 500

    @app.route("/api/trading/chart/<ticker>")
    def api_trading_chart(ticker: str):
        """Return OHLCV bars, volume profile, patterns, and S/R levels for *ticker*."""
        ticker = ticker.upper()
        timeframe = flask_request.args.get("tf", "1D")
        limit = min(int(flask_request.args.get("limit", 120)), 500)
        try:
            from tools.trading.data.market_data import fetch_bars
            from tools.trading.ta.volume_profile import volume_profile as compute_vp
            from tools.trading.ta.swings import find_swings
            from tools.trading.ta.patterns import detect_patterns
            from tools.trading.ta.support_resistance import compute_sr

            bars = fetch_bars(ticker, timeframe, limit)
            provenance = _derive_chart_provenance(bars)
            vp = compute_vp(bars, bucket_count=40)
            swings = find_swings(bars)
            raw_patterns = detect_patterns(bars)
            sr_levels = compute_sr(bars, swings=swings)
            patterns = _enrich_chart_patterns(raw_patterns)
            return jsonify({
                "ticker": ticker,
                "bars": bars,
                "volume_profile": vp,
                "patterns": patterns,
                "sr_levels": sr_levels,
                # Top-level data provenance so the UI can flag synthetic bars
                # (nav-plat-01) — never render simulated data as real market data.
                "data_source": provenance["data_source"],
                "simulated": provenance["simulated"],
                "as_of": provenance["as_of"],
            })
        except Exception as exc:
            return jsonify({"error": str(exc)}), 500

    @app.route("/fathomdesk/api/traps")
    def fathomdesk_api_traps():
        """Return last 20 trap events from ad_trap_events for the Trap History panel."""
        def _confidence_to_severity(conf):
            if conf is None:
                return "medium"
            if conf >= 0.8:
                return "critical"
            if conf >= 0.6:
                return "high"
            return "medium"

        try:
            from tools.db.storage import get_connection
            conn = get_connection()
            try:
                rows = conn.execute(
                    "SELECT id, ticker, pattern, broken_level, confidence, "
                    "volume_ratio, timeframe, evidence_json, created_at "
                    "FROM ad_trap_events "
                    "ORDER BY created_at DESC "
                    "LIMIT 20"
                ).fetchall()
            finally:
                conn.close()

            events = []
            for r in rows:
                row = dict(r) if hasattr(r, "keys") else {
                    "id": r[0], "ticker": r[1], "pattern": r[2],
                    "broken_level": r[3], "confidence": r[4],
                    "volume_ratio": r[5], "timeframe": r[6],
                    "evidence_json": r[7], "created_at": r[8],
                }
                row["severity"] = _confidence_to_severity(row.get("confidence"))
                events.append(row)

            return jsonify({"events": events})
        except Exception as exc:
            return jsonify({"events": [], "error": str(exc)})

    @app.route("/fathomdesk/api/reflex-observations")
    def fathomdesk_api_reflex_observations():
        """Return recent reflex execution records."""
        from flask import request as _req
        try:
            limit = min(int(_req.args.get("limit", 50)), 200)
        except (ValueError, TypeError):
            limit = 50
        try:
            from tools.db.storage import get_connection
            conn = get_connection()
            ph = "%s" if getattr(conn, "_dialect", "sqlite") == "postgresql" else "?"
            try:
                rows = conn.execute(
                    f"SELECT id, reflex_name, started_at, duration_ms, status "  # nosec B608
                    f"FROM reflex_observations "
                    f"ORDER BY started_at DESC LIMIT {ph}",
                    [limit],
                ).fetchall()
            finally:
                conn.close()
            observations = []
            for r in rows:
                row = dict(r) if hasattr(r, "keys") else {
                    "id": r[0], "reflex_name": r[1], "started_at": r[2],
                    "duration_ms": r[3], "status": r[4],
                }
                observations.append({
                    "id": row["id"],
                    "reflex_name": row["reflex_name"],
                    "started_at": row["started_at"],
                    "duration_ms": row["duration_ms"],
                    "success": row["status"] == "done",
                })
            return jsonify({"observations": observations})
        except Exception as exc:
            return jsonify({"observations": [], "error": str(exc)})

    if _track_request is not None:
        _track_request(app)

    # ---- ECR-OBS-01: Prometheus /metrics endpoint ----
    try:
        from tools.observability.metrics import (
            wire_flask_metrics as _wire_metrics,
            registry as _prom_registry,
            generate_latest as _gen_latest,
            CONTENT_TYPE_LATEST as _PROM_CT,
        )

        _wire_metrics(app)

        @app.route("/metrics", methods=["GET"])
        def prometheus_metrics():
            if _prom_registry is None or _gen_latest is None:
                return Response("# prometheus_client not installed\n",
                                status=503, mimetype="text/plain")
            data = _gen_latest(_prom_registry)
            return Response(data, status=200, mimetype=_PROM_CT)
    except Exception as _prom_exc:
        app.logger.debug("Prometheus metrics init skipped: %s", _prom_exc)

    return app


app = create_app()

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="ICDEV™ Dashboard")
    parser.add_argument("--port", type=int, default=PORT, help="Port to run on (default: 5050)")
    parser.add_argument("--debug", action="store_true", default=DEBUG, help="Enable debug mode")
    args = parser.parse_args()

    # `app` is already built by the unconditional `app = create_app()` above
    # (needed for WSGI imports like `from tools.dashboard.app import app`).
    # Re-calling create_app() here ran the entire app-setup a second time in
    # every `python tools/dashboard/app.py` invocation — including CI's E2E
    # server start — which double-registered every canvas blueprint. Flask
    # blueprint objects are module-level singletons, so the second
    # registration pass raised on `.before_request()` for every canvas with
    # an RBAC/IL-level guard (caught as a generic "registration failed"
    # warning), silently disabling canvas-access enforcement.
    print(f"[ICDEV™ Dashboard] Starting on http://127.0.0.1:{args.port}")
    print(f"[ICDEV™ Dashboard] Database: {DB_PATH}")
    print(f"[ICDEV™ Dashboard] CUI Marking: {CUI_BANNER_TOP or '(none)'}")

    # ── Auto-start Kanban Scheduler (LLM-agnostic) ───────────────────
    # Launches kanban_scheduler.py as a child process so backlog tasks
    # are promoted and dispatched regardless of which LLM/IDE is in use.
    # The subprocess dies automatically when the dashboard exits.
    #
    # Dedup: the dashboard debug-reloader (and operators running
    # `python tools/dashboard/app.py` a second time) would otherwise
    # accumulate scheduler processes. A heartbeat file (refreshed every
    # cycle by kanban_scheduler.py) is the single source of truth; if a
    # fresh heartbeat exists, another instance is already running.
    try:
        import subprocess as _ks_sp
        import time as _ks_time

        _ks_script = Path(__file__).resolve().parent.parent / "genesis" / "kanban_scheduler.py"
        _ks_hb = BASE_DIR / ".tmp" / "kanban_scheduler.heartbeat"
        _ks_hb_fresh = False
        if _ks_hb.exists():
            try:
                _ks_hb_age = _ks_time.time() - _ks_hb.stat().st_mtime
                # Scheduler default interval is 60s; allow 3× slack for slow hosts.
                _ks_hb_fresh = _ks_hb_age < 180
            except OSError:
                _ks_hb_fresh = False

        if _ks_hb_fresh:
            print(
                "[ICDEV™ Dashboard] Kanban scheduler already running "
                f"(fresh heartbeat at {_ks_hb}) — not spawning another"
            )
        elif _ks_script.exists():
            _ks_log_dir = BASE_DIR / ".tmp"
            _ks_log_dir.mkdir(parents=True, exist_ok=True)
            _ks_log = open(str(_ks_log_dir / "kanban_scheduler.log"), "a", encoding="utf-8")  # noqa: SIM115
            _ks_sp.Popen(
                [sys.executable, str(_ks_script), "--interval", "60"],
                stdout=_ks_log,
                stderr=_ks_sp.STDOUT,
                cwd=str(BASE_DIR),
            )
            print("[ICDEV™ Dashboard] Kanban scheduler started (60s interval)")
        else:
            print("[ICDEV™ Dashboard] Kanban scheduler not found — skipping")

        # ── Watchdog: auto-restart scheduler if it dies ────────────────
        # Polls the heartbeat file every 120s. If it's > 300s old the
        # scheduler process crashed; spawn a fresh one.
        if _ks_script.exists():
            import threading as _ks_threading

            def _ks_watchdog():
                import time as _wt
                _POLL = 120   # check every 2 min
                _STALE = 300  # restart if heartbeat > 5 min old
                while True:
                    _wt.sleep(_POLL)
                    try:
                        age = _wt.time() - _ks_hb.stat().st_mtime if _ks_hb.exists() else 9999
                        if age > _STALE:
                            _log_dir = BASE_DIR / ".tmp"
                            _log_dir.mkdir(parents=True, exist_ok=True)
                            _lf = open(str(_log_dir / "kanban_scheduler.log"), "a", encoding="utf-8")  # noqa: SIM115
                            _ks_sp.Popen(
                                [sys.executable, str(_ks_script), "--interval", "60"],
                                stdout=_lf,
                                stderr=_ks_sp.STDOUT,
                                cwd=str(BASE_DIR),
                            )
                            print(
                                f"[ICDEV™ Dashboard] Kanban scheduler watchdog: "
                                f"heartbeat was {int(age)}s old — restarted scheduler"
                            )
                    except Exception as _we:
                        print(f"[ICDEV™ Dashboard] Kanban scheduler watchdog error: {_we}")

            _ks_wd = _ks_threading.Thread(target=_ks_watchdog, name="kanban-scheduler-watchdog", daemon=True)
            _ks_wd.start()
            print("[ICDEV™ Dashboard] Kanban scheduler watchdog started (polls every 120s, restarts if >300s stale)")

    except Exception as _ks_err:
        print(f"[ICDEV™ Dashboard] Kanban scheduler failed to start: {_ks_err}")

    # ── DIC freshness scan daemon (docmod-ux-01) ───────────────────────────
    # Keeps freshness scores + modernization findings live even when the
    # Genesis daemon isn't running. ICDEV_DIC_FRESHNESS_SCAN_HOURS: default 24,
    # 0 disables. Exception-safe per tick; never starts under pytest.
    try:
        _fs_hours = float(os.environ.get("ICDEV_DIC_FRESHNESS_SCAN_HOURS", "24") or 0)
        if _fs_hours > 0 and not os.environ.get("PYTEST_CURRENT_TEST"):
            import threading as _fs_threading

            def _dic_freshness_daemon():
                import time as _ft
                while True:
                    _ft.sleep(_fs_hours * 3600)
                    try:
                        from tools.document_intelligence.freshness_engine import scan_collection as _fs_scan
                        from tools.db.storage import get_connection as _fs_conn
                        _c = _fs_conn()
                        try:
                            _rows = _c.execute(
                                "SELECT collection_id FROM dic_collections"
                            ).fetchall()
                        finally:
                            _c.close()
                        for _r in _rows:
                            try:
                                _fs_scan(dict(_r)["collection_id"])
                            except Exception:
                                pass  # one collection must not kill the tick
                    except Exception as _fe:
                        print(f"[ICDEV™ Dashboard] DIC freshness daemon tick error: {_fe}")
                    try:  # docmod sweep hook — graceful before engine install
                        from tools.doc_modernization.scanner import scan_collection as _dm_scan
                        _dm_scan(collection_id=None, trigger="daemon")
                    except Exception:
                        pass

            _fs_t = _fs_threading.Thread(
                target=_dic_freshness_daemon, name="dic-freshness-scan-daemon", daemon=True
            )
            _fs_t.start()
            print(f"[ICDEV™ Dashboard] DIC freshness scan daemon started (every {_fs_hours}h)")
    except Exception as _fs_err:
        print(f"[ICDEV™ Dashboard] DIC freshness daemon failed to start: {_fs_err}")

    # Optional inbound TLS / mTLS (IL5+/GovCloud). Env vars:
    #   ICDEV_DASHBOARD_TLS_CERT      server certificate (PEM)
    #   ICDEV_DASHBOARD_TLS_KEY       server private key (PEM)
    #   ICDEV_DASHBOARD_TLS_CA_BUNDLE CA bundle — enables client-cert
    #                                 verification (mTLS) when set
    # When both cert+key are set the dashboard listens on HTTPS. When a CA
    # bundle is also set, clients must present a valid certificate signed by
    # that CA (CERT_REQUIRED). For dev/non-TLS deployments leave all three
    # unset.
    _ssl_context = None
    _tls_cert = os.environ.get("ICDEV_DASHBOARD_TLS_CERT")
    _tls_key = os.environ.get("ICDEV_DASHBOARD_TLS_KEY")
    _tls_ca = os.environ.get("ICDEV_DASHBOARD_TLS_CA_BUNDLE")
    if _tls_cert and _tls_key:
        import ssl as _ssl

        _ctx = _ssl.create_default_context(purpose=_ssl.Purpose.CLIENT_AUTH)
        _ctx.load_cert_chain(certfile=_tls_cert, keyfile=_tls_key)
        if _tls_ca:
            _ctx.load_verify_locations(cafile=_tls_ca)
            _ctx.verify_mode = _ssl.CERT_REQUIRED
            print(f"[ICDEV™ Dashboard] mTLS enabled (CA: {_tls_ca})")
        else:
            print("[ICDEV™ Dashboard] TLS enabled (server-only; no client CA)")
        _ssl_context = _ctx

    # Use SocketIO runner if available (D170), otherwise plain Flask
    # use_reloader=False: prevents Werkzeug's stat-based reloader from spawning
    # a second create_app() call and causing repeated restart loops on Windows.
    socketio = get_socketio()
    if socketio:
        print("[ICDEV™ Dashboard] WebSocket enabled (Flask-SocketIO)")
        if _ssl_context is not None:
            socketio.run(app, host=HOST, port=args.port, debug=args.debug, use_reloader=False, ssl_context=_ssl_context, allow_unsafe_werkzeug=True)  # nosec B104
        else:
            socketio.run(app, host=HOST, port=args.port, debug=args.debug, use_reloader=False, allow_unsafe_werkzeug=True)  # nosec B104 -- intentional bind-all for containerized/dev deployment
    else:
        print("[ICDEV™ Dashboard] WebSocket not available — using HTTP polling")
        if _ssl_context is not None:
            app.run(host=HOST, port=args.port, debug=args.debug, use_reloader=False, ssl_context=_ssl_context, threaded=True)  # nosec B104
        else:
            app.run(host=HOST, port=args.port, debug=args.debug, use_reloader=False, threaded=True)  # nosec B104 -- intentional bind-all for containerized/dev deployment
