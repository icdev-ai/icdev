# [CUI // SP-CTI]
"""ICDEV™ NDC — Security advisory extraction from PDF, DOCX, and EML files.

Public API:
- ``extract_advisory(file_path, file_format) -> dict``

Supported formats: ``pdf``, ``docx``, ``eml``.

PDF strategy: text extraction via pdfminer; if a vision-capable LLM is
available the first-page raster image is also passed so scanned/image-only
PDFs are handled without falling back silently to an empty result.

Output schema mirrors ``context/nql/advisory_extraction_prompt.md``.
"""
from __future__ import annotations

import base64
import email
import email.policy
import json
import re
import tempfile
from pathlib import Path
from typing import Optional

from tools.logging.icdev_logger import get_logger

logger = get_logger(__name__)

# ---------------------------------------------------------------------------
# System prompt (loaded from context file; inline fallback keeps it self-contained)
# ---------------------------------------------------------------------------

_PROMPT_PATH = Path(__file__).resolve().parents[4] / "context" / "nql" / "advisory_extraction_prompt.md"

def _load_system_prompt() -> str:
    try:
        text = _PROMPT_PATH.read_text(encoding="utf-8")
        # Strip the trailing placeholder line so it doesn't confuse the LLM
        return text.split("## Advisory Text")[0].strip()
    except Exception:
        return (
            "You are a security advisory parser. "
            "Extract structured fields from the advisory text and return only valid JSON "
            "matching the schema with keys: cve_ids, vendor, affected_models, affected_versions, "
            "fixed_versions, severity, cvss_score, cvss_vector, cvss_version, vulnerability_type, "
            "attack_vector, authentication_required, user_interaction_required, patch_available, "
            "workarounds, ioc_indicators, published_date, updated_date, advisory_id, title, "
            "summary, references, exploited_in_wild, exploit_public."
        )


# ---------------------------------------------------------------------------
# Format extractors
# ---------------------------------------------------------------------------

def _extract_text_pdf(file_path: str) -> tuple[str, Optional[bytes]]:
    """Return (text, first_page_png_bytes_or_None)."""
    from pdfminer.high_level import extract_text as pm_extract

    text = pm_extract(file_path) or ""

    first_page_bytes: Optional[bytes] = None
    try:
        import pypdfium2 as pdfium  # type: ignore
        doc = pdfium.PdfDocument(file_path)
        if len(doc) > 0:
            page = doc[0]
            bitmap = page.render(scale=2)
            pil_img = bitmap.to_pil()
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                pil_img.save(tmp.name, "PNG")
                tmp_path = tmp.name
            first_page_bytes = Path(tmp_path).read_bytes()
            Path(tmp_path).unlink(missing_ok=True)
    except Exception as e:
        logger.debug("advisory_extractor: first-page raster skipped (%s)", e)

    return text, first_page_bytes


def _extract_text_docx(file_path: str) -> str:
    from docx import Document  # type: ignore

    doc = Document(file_path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _extract_text_eml(file_path: str) -> str:
    with open(file_path, "rb") as fh:
        msg = email.message_from_binary_file(fh, policy=email.policy.default)
    return _walk_email_parts(msg)


def _walk_email_parts(msg) -> str:
    """Recursively walk MIME parts; prefer text/plain, fall back to text/html."""
    if msg.is_multipart():
        plain_parts: list[str] = []
        html_parts: list[str] = []
        for part in msg.iter_parts():
            ct = part.get_content_type()
            if ct == "text/plain":
                plain_parts.append(_decode_part(part))
            elif ct == "text/html":
                html_parts.append(_strip_html(_decode_part(part)))
            elif part.is_multipart():
                nested = _walk_email_parts(part)
                if nested:
                    plain_parts.append(nested)
        return "\n".join(plain_parts) if plain_parts else "\n".join(html_parts)
    ct = msg.get_content_type()
    if ct == "text/plain":
        return _decode_part(msg)
    if ct == "text/html":
        return _strip_html(_decode_part(msg))
    return ""


def _decode_part(part) -> str:
    try:
        return part.get_content() or ""
    except Exception:
        payload = part.get_payload(decode=True)
        if isinstance(payload, bytes):
            charset = part.get_content_charset() or "utf-8"
            return payload.decode(charset, errors="replace")
        return str(payload or "")


def _strip_html(html: str) -> str:
    return re.sub(r"<[^>]+>", " ", html)


# ---------------------------------------------------------------------------
# LLM invocation
# ---------------------------------------------------------------------------

_EMPTY_RESULT: dict = {
    "cve_ids": [],
    "vendor": "other",
    "affected_models": [],
    "affected_versions": [],
    "fixed_versions": [],
    "severity": None,
    "cvss_score": None,
    "cvss_vector": None,
    "cvss_version": None,
    "vulnerability_type": [],
    "attack_vector": None,
    "authentication_required": False,
    "user_interaction_required": False,
    "patch_available": False,
    "workarounds": [],
    "ioc_indicators": [],
    "published_date": None,
    "updated_date": None,
    "advisory_id": None,
    "title": "",
    "summary": "",
    "references": [],
    "exploited_in_wild": False,
    "exploit_public": False,
}


def _invoke_llm(text: str, first_page_bytes: Optional[bytes] = None) -> dict:
    try:
        from tools.llm import get_router
        from tools.llm.provider import LLMRequest
    except ImportError:
        logger.warning("advisory_extractor: LLMRouter unavailable; returning empty result")
        return dict(_EMPTY_RESULT)

    system_prompt = _load_system_prompt()
    user_text = f"## Advisory Text\n\n{text}"

    if first_page_bytes:
        b64 = base64.b64encode(first_page_bytes).decode("utf-8")
        user_content = [
            {"type": "image", "source": {"type": "base64", "media_type": "image/png", "data": b64}},
            {"type": "text", "text": user_text},
        ]
    else:
        user_content = [{"type": "text", "text": user_text}]

    request = LLMRequest(
        messages=[{"role": "user", "content": user_content}],
        system_prompt=system_prompt,
        max_tokens=2048,
        temperature=0.1,
        skip_injection_scan=True,
    )

    try:
        router = get_router()
        response = router.invoke("advisory_extraction", request)
        content = response.content or ""
        json_match = re.search(r"\{.*\}", content, re.DOTALL)
        if not json_match:
            logger.warning("advisory_extractor: no JSON in LLM response")
            return dict(_EMPTY_RESULT)
        return json.loads(json_match.group())
    except Exception as e:
        logger.error("advisory_extractor: LLM invocation failed: %s", e)
        return dict(_EMPTY_RESULT)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def extract_advisory(file_path: str, file_format: str) -> dict:
    """Extract structured security advisory fields from a file.

    Args:
        file_path: Absolute or relative path to the advisory document.
        file_format: One of ``'pdf'``, ``'docx'``, or ``'eml'``.

    Returns:
        Dict matching the advisory schema defined in
        ``context/nql/advisory_extraction_prompt.md``. On any extraction
        failure the dict is returned with empty/null fields rather than
        raising, so callers can safely index without checking for None.
    """
    fmt = file_format.lower().strip(".")
    first_page_bytes: Optional[bytes] = None

    try:
        if fmt == "pdf":
            text, first_page_bytes = _extract_text_pdf(file_path)
        elif fmt in ("docx", "doc"):
            text = _extract_text_docx(file_path)
        elif fmt in ("eml", "msg"):
            text = _extract_text_eml(file_path)
        else:
            logger.error("advisory_extractor: unsupported format %r", file_format)
            return dict(_EMPTY_RESULT)
    except Exception as e:
        logger.error("advisory_extractor: text extraction failed for %r: %s", file_path, e)
        return dict(_EMPTY_RESULT)

    if not text.strip():
        logger.warning("advisory_extractor: no text extracted from %r", file_path)
        return dict(_EMPTY_RESULT)

    return _invoke_llm(text, first_page_bytes)
