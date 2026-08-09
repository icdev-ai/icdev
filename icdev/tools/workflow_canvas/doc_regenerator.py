"""
Process-Ify Document Regenerator

After a workflow is processified and teams complete the tasks, this module
regenerates the original source document with updated content reflecting the
AI-enabled workflow — preserving the original document's style, tone, structure,
and formality (WriteGuard-style fingerprinting).

Supports:
- Single workflow: regenerate one document
- Process chain: regenerate per-phase OR one unified end-to-end document
"""
from __future__ import annotations

import json
import pathlib
import threading
import uuid
from datetime import datetime, timezone

from tools.logging.icdev_logger import get_logger

log = get_logger(__name__)

_ARTIFACTS_DIR = pathlib.Path("data/docgen/artifacts")

# In-memory job tracker — job_id -> {status, title, files, error, started_at, completed_at}
_REGEN_JOBS: dict[str, dict] = {}


# ── Helpers ────────────────────────────────────────────────────────────────────

def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _llm_generate(prompt: str, max_tokens: int = 4000) -> str:
    from tools.llm.router import LLMRouter
    from tools.llm.provider import LLMRequest

    router = LLMRouter()
    req = LLMRequest(
        messages=[{"role": "user", "content": prompt}],
        system_prompt="You are an expert technical writer. Generate complete, professional documents.",
        max_tokens=max_tokens,
    )
    result = router.invoke("document_generation", req)
    if hasattr(result, "content"):
        return result.content or ""
    if isinstance(result, dict):
        return result.get("content") or result.get("text") or str(result)
    return str(result or "")


# ── Style fingerprinting ───────────────────────────────────────────────────────

def extract_style_fingerprint(text: str) -> dict:
    """Analyze document text and return a style/structure profile dict.

    The fingerprint drives the regeneration prompt so the new document matches
    the original's formality, section numbering, vocabulary, and tone.
    """
    from tools.llm.router import LLMRouter
    from tools.llm.provider import LLMRequest

    excerpt = text[:3000]
    prompt = (
        "Analyze the document excerpt below and return ONLY a valid JSON object with these keys:\n"
        "  doc_type: one of [SOP, Agreement, Technical_Spec, Design_Doc, Plan, Policy, Report, Runbook, Other]\n"
        "  formality: one of [formal, semi_formal, informal]\n"
        "  voice: one of [active, passive, mixed]\n"
        "  section_numbering: exact numbering pattern seen (e.g. '1.0', '1.1.2', 'A.', 'I.', 'none')\n"
        "  section_headers: JSON array of up to 8 heading texts found in the document\n"
        "  tone_descriptors: JSON array of exactly 3 adjectives describing the writing tone\n"
        "  vocabulary_level: one of [executive, technical, general, legal]\n"
        "  uses_tables: true or false\n\n"
        f"Document excerpt:\n{excerpt}\n\n"
        "Return ONLY valid JSON. No markdown fences, no explanation."
    )
    try:
        router = LLMRouter()
        req = LLMRequest(
            messages=[{"role": "user", "content": prompt}],
            system_prompt="You are a document style analyzer. Return only valid JSON.",
            max_tokens=400,
        )
        result = router.invoke("style_analysis", req)
        raw = (result.content if hasattr(result, "content") else str(result)).strip()
        # Strip markdown fences if present
        if raw.startswith("```"):
            parts = raw.split("```")
            raw = parts[1][4:] if parts[1].startswith("json") else parts[1]
        return json.loads(raw)
    except Exception as exc:
        log.warning("style fingerprint extraction failed: %s", exc)
        return {
            "doc_type": "Other",
            "formality": "formal",
            "voice": "mixed",
            "section_numbering": "none",
            "section_headers": [],
            "tone_descriptors": ["professional", "clear", "neutral"],
            "vocabulary_level": "technical",
            "uses_tables": False,
        }


def _cache_fingerprint(workflow_id: str, fingerprint: dict) -> None:
    try:
        from tools.db.storage import get_connection, sql_placeholder
        conn = get_connection()
        ph = sql_placeholder(conn)
        conn.execute(
            f"UPDATE studio_workflows SET style_fingerprint = {ph} WHERE workflow_id = {ph}",
            (json.dumps(fingerprint), workflow_id),
        )
        conn.commit()
        conn.close()
    except Exception as exc:
        log.warning("could not cache fingerprint for %s: %s", workflow_id, exc)


# ── Prompt builders ────────────────────────────────────────────────────────────

def _build_single_prompt(workflow_data: dict, fingerprint: dict, original_excerpt: str) -> str:
    steps = workflow_data.get("steps", [])
    wf_name = workflow_data.get("name") or workflow_data.get("workflow_name") or "Workflow"
    industry = workflow_data.get("industry", "General")
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    fp = fingerprint or {}

    step_lines = []
    for i, s in enumerate(steps):
        checklist = ", ".join(s.get("checklist", s.get("checklist_items", [])) or [])
        form_fields = ", ".join(
            (f["label"] if isinstance(f, dict) else str(f))
            for f in (s.get("form_fields") or [])
        )
        step_lines.append(
            f"  Step {i+1}: {s.get('title', '')}\n"
            f"    Phase: {s.get('phase', '')} | "
            f"Assignee: {s.get('assignee_role', '')} | "
            f"Reviewer: {s.get('reviewer_role', '')} | "
            f"Approver: {s.get('approver_role', '')}"
            + (f"\n    Checklist: {checklist}" if checklist else "")
            + (f"\n    Form fields: {form_fields}" if form_fields else "")
        )

    style_lines = [
        f"Formality: {fp.get('formality', 'formal')}",
        f"Voice: {fp.get('voice', 'mixed')}",
        f"Section numbering: {fp.get('section_numbering', '1.0')}",
        f"Tone: {', '.join(fp.get('tone_descriptors', ['professional', 'clear']))}",
        f"Vocabulary: {fp.get('vocabulary_level', 'technical')}",
    ]
    if fp.get("section_headers"):
        style_lines.append(f"Use these section headings from the original: {', '.join(fp['section_headers'])}")

    orig_block = (
        f"ORIGINAL DOCUMENT EXCERPT (for style reference — do not copy content):\n"
        f"---\n{original_excerpt[:1500]}\n---\n\n"
        if original_excerpt else ""
    )

    return (
        f"You are rewriting an official {fp.get('doc_type', 'document')} to reflect "
        f"an AI-enabled, digitized process workflow.\n\n"
        f"STYLE REQUIREMENTS (match exactly):\n"
        + "\n".join(style_lines) + "\n\n"
        + orig_block
        + f"PROCESSIFIED WORKFLOW DATA:\n"
        f"Name: {wf_name} | Industry: {industry} | Total Steps: {len(steps)}\n\n"
        + "\n".join(step_lines) + "\n\n"
        f"TASK: Generate the complete updated document. Requirements:\n"
        f"1. Maintain the EXACT same section structure, numbering style, and tone as the original\n"
        f"2. Update the process steps section with the structured workflow above\n"
        f"3. Embed assignee, reviewer, and approver roles in each step section\n"
        f"4. Reference the forms and checklists created for each step\n"
        f"5. Keep Purpose and Scope sections unchanged unless workflow reveals corrections\n"
        f"6. Add a Version History entry: 'AI-enabled process update via Process-Ify ({today})'\n\n"
        f"Generate the complete document now."
    )


def _build_chain_prompt(chain: dict, phases_data: list[dict], fingerprint: dict, original_excerpt: str) -> str:
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    fp = fingerprint or {}

    phases_block = ""
    for p in phases_data:
        phases_block += (
            f"\n### Phase {p['phase_number']}: {p['name']}"
            + (f" (Team: {p['team_name']})" if p.get("team_name") else "")
            + "\n"
        )
        for wf in p.get("workflows", []):
            for i, s in enumerate(wf.get("steps", [])):
                checklist = ", ".join(s.get("checklist", s.get("checklist_items", [])) or [])
                phases_block += (
                    f"  Step {i+1}: {s.get('title', '')}\n"
                    f"    Assignee: {s.get('assignee_role', '')} | "
                    f"Reviewer: {s.get('reviewer_role', '')} | "
                    f"Approver: {s.get('approver_role', '')}"
                    + (f"\n    Checklist: {checklist}" if checklist else "")
                    + "\n"
                )

    orig_block = (
        f"Style reference from phase-1 document:\n---\n{original_excerpt[:800]}\n---\n\n"
        if original_excerpt else ""
    )

    return (
        f"Generate a complete, unified process document for this multi-phase workflow chain.\n\n"
        f"Chain: {chain['name']} | Industry: {chain.get('industry', 'General')}\n"
        f"Formality: {fp.get('formality', 'formal')} | "
        f"Tone: {', '.join(fp.get('tone_descriptors', ['professional', 'authoritative']))}\n\n"
        + orig_block
        + "Document structure (include all sections):\n"
        "1. Cover page with chain name, date, and classification\n"
        "2. Executive Summary (2-3 paragraphs synthesizing the full chain)\n"
        "3. Table of Contents\n"
        "4. One major section per phase with all workflow steps, roles, and forms\n"
        "5. Handoff checklist between each phase transition\n"
        "6. Appendix A: Forms Index (list all forms created per step)\n"
        "7. Appendix B: Role Matrix (who does what across all phases)\n"
        f"8. Version History: 'AI-enabled via Process-Ify — {today}'\n\n"
        f"PHASES AND WORKFLOWS:\n{phases_block}\n\n"
        "Generate the complete end-to-end document now."
    )


# ── Export ─────────────────────────────────────────────────────────────────────

def _export_artifacts(doc_text: str, title: str, out_dir: pathlib.Path, output_format: str) -> list[dict]:
    artifacts: list[dict] = []

    md_path = out_dir / "document.md"
    md_path.write_text(doc_text, encoding="utf-8", newline="")
    artifacts.append({"format": "markdown", "path": str(md_path), "filename": "document.md"})

    html_path = out_dir / "document.html"
    if output_format in ("html", "all", "pdf"):
        try:
            html_path.write_text(_to_html(doc_text, title), encoding="utf-8", newline="")
            if output_format in ("html", "all"):
                artifacts.append({"format": "html", "path": str(html_path), "filename": "document.html"})
        except Exception as exc:
            log.warning("html export failed: %s", exc)

    if output_format in ("docx", "all"):
        try:
            docx_path = out_dir / "document.docx"
            _to_docx(doc_text, title, docx_path)
            artifacts.append({"format": "docx", "path": str(docx_path), "filename": "document.docx"})
        except Exception as exc:
            log.warning("docx export failed: %s", exc)

    if output_format in ("pdf", "all"):
        try:
            from tools.network.pdf_export import export_to_pdf
            if not html_path.exists():
                html_path.write_text(_to_html(doc_text, title), encoding="utf-8", newline="")
            pdf_path = str(out_dir / "document.pdf")
            export_to_pdf(str(html_path), pdf_path)
            artifacts.append({"format": "pdf", "path": pdf_path, "filename": "document.pdf"})
        except Exception as exc:
            log.warning("pdf export failed: %s", exc)

    return artifacts


def _to_html(doc_text: str, title: str) -> str:
    try:
        import markdown as _md
        body = _md.markdown(doc_text, extensions=["tables", "fenced_code"])
    except ImportError:
        lines = []
        for line in doc_text.split("\n"):
            if line.startswith("# "):
                lines.append(f"<h1>{line[2:]}</h1>")
            elif line.startswith("## "):
                lines.append(f"<h2>{line[3:]}</h2>")
            elif line.startswith("### "):
                lines.append(f"<h3>{line[4:]}</h3>")
            elif line.strip():
                lines.append(f"<p>{line}</p>")
        body = "\n".join(lines)

    return (
        "<!DOCTYPE html><html lang='en'><head><meta charset='utf-8'>"
        f"<title>{title}</title><style>"
        "body{font-family:Arial,sans-serif;max-width:900px;margin:40px auto;padding:0 20px;"
        "line-height:1.6;color:#1a1a2e}"
        "h1{border-bottom:2px solid #1d72ff;padding-bottom:8px}"
        "h2{color:#1d72ff;margin-top:2em}"
        "h3{color:#374151;margin-top:1.5em}"
        "table{border-collapse:collapse;width:100%}"
        "th,td{border:1px solid #ccc;padding:8px;text-align:left}"
        "th{background:#f0f4ff}"
        ".banner{background:#1d72ff;color:#fff;text-align:center;padding:8px;font-weight:bold;"
        "letter-spacing:1px;margin-bottom:24px;font-size:13px}"
        "@media print{.banner{-webkit-print-color-adjust:exact}}"
        "</style></head><body>"
        f"<div class='banner'>PROCESS-IFY REGENERATED DOCUMENT</div>"
        f"<h1>{title}</h1>"
        f"{body}"
        "</body></html>"
    )


def _to_docx(doc_text: str, title: str, out_path: pathlib.Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Process-Ify Regenerated Document").alignment = WD_ALIGN_PARAGRAPH.CENTER

    for line in doc_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith(("- ", "* ")):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped and stripped[0].isdigit() and len(stripped) > 2 and stripped[1] in (". ", ") "):
            doc.add_paragraph(stripped, style="List Number")
        else:
            doc.add_paragraph(stripped)

    doc.save(str(out_path))


# ── Public API ─────────────────────────────────────────────────────────────────

def start_regen_workflow(workflow_id: str, output_format: str = "all") -> str:
    """Start background doc regen for a single processified workflow. Returns job_id."""
    job_id = f"regen-{workflow_id[:8]}-{uuid.uuid4().hex[:6]}"
    _REGEN_JOBS[job_id] = {"status": "running", "started_at": _now()}
    t = threading.Thread(
        target=_run_regen_workflow, args=(job_id, workflow_id, output_format), daemon=True
    )
    t.start()
    return job_id


def start_regen_chain(chain_id: str, phase_id: str | None, output_format: str = "all") -> str:
    """Start background doc regen for a process chain or single phase. Returns job_id."""
    job_id = f"regen-chain-{chain_id[:8]}-{uuid.uuid4().hex[:6]}"
    _REGEN_JOBS[job_id] = {"status": "running", "started_at": _now()}
    t = threading.Thread(
        target=_run_regen_chain, args=(job_id, chain_id, phase_id, output_format), daemon=True
    )
    t.start()
    return job_id


def get_job_status(job_id: str) -> dict | None:
    return _REGEN_JOBS.get(job_id)


# ── Background workers ─────────────────────────────────────────────────────────

def _run_regen_workflow(job_id: str, workflow_id: str, output_format: str) -> None:
    try:
        import yaml
        from tools.db.storage import get_connection, sql_placeholder

        conn = get_connection()
        ph = sql_placeholder(conn)
        row = conn.execute(
            f"SELECT * FROM studio_workflows WHERE workflow_id = {ph}", (workflow_id,)
        ).fetchone()
        conn.close()

        if not row:
            _REGEN_JOBS[job_id] = {"status": "error", "error": f"Workflow {workflow_id} not found"}
            return

        row = dict(row)
        try:
            workflow_data = yaml.safe_load(row.get("template_yaml") or "{}") or {}
        except Exception:
            workflow_data = {}

        source_doc_text = row.get("source_doc_text") or ""
        fp_raw = row.get("style_fingerprint")
        fingerprint = json.loads(fp_raw) if fp_raw else {}

        if not fingerprint and source_doc_text:
            fingerprint = extract_style_fingerprint(source_doc_text)
            _cache_fingerprint(workflow_id, fingerprint)

        prompt = _build_single_prompt(workflow_data, fingerprint, source_doc_text)
        doc_text = _llm_generate(prompt, max_tokens=4000)

        out_dir = _ARTIFACTS_DIR / job_id
        out_dir.mkdir(parents=True, exist_ok=True)
        title = f"{row.get('name', 'Document')} — Updated Process Document"
        artifacts = _export_artifacts(doc_text, title, out_dir, output_format)

        # Persist artifact path
        conn = get_connection()
        ph = sql_placeholder(conn)
        conn.execute(
            f"UPDATE studio_workflows SET regen_artifact_path = {ph}, "
            f"updated_at = CURRENT_TIMESTAMP WHERE workflow_id = {ph}",
            (str(out_dir), workflow_id),
        )
        conn.commit()
        conn.close()

        _REGEN_JOBS[job_id] = {
            "status": "done",
            "title": title,
            "artifact_dir": str(out_dir),
            "files": artifacts,
            "completed_at": _now(),
        }
        log.info("regen complete: workflow=%s job=%s files=%d", workflow_id, job_id, len(artifacts))
    except Exception as exc:
        log.error("regen workflow %s failed: %s", workflow_id, exc, exc_info=True)
        _REGEN_JOBS[job_id] = {"status": "error", "error": str(exc)}


def _run_regen_chain(job_id: str, chain_id: str, phase_id: str | None, output_format: str) -> None:
    try:
        import yaml
        from tools.db.storage import get_connection, get_canvas_connection, sql_placeholder

        # Canvas connection for wfc_* tables (no classification/tenant_id columns)
        cc = get_canvas_connection("ICDEV_WFC_ENABLED")
        cc_ph = sql_placeholder(cc)

        chain_row = cc.execute(
            f"SELECT * FROM wfc_process_chains WHERE id = {cc_ph}", (chain_id,)
        ).fetchone()
        if not chain_row:
            cc.close()
            _REGEN_JOBS[job_id] = {"status": "error", "error": f"Chain {chain_id} not found"}
            return
        chain = dict(chain_row)

        phases_rows = cc.execute(
            f"SELECT * FROM wfc_chain_phases WHERE chain_id = {cc_ph} ORDER BY phase_number",
            (chain_id,),
        ).fetchall()
        phases = [dict(r) for r in phases_rows]
        cc.close()

        if phase_id:
            phases = [p for p in phases if p["id"] == phase_id]
            if not phases:
                _REGEN_JOBS[job_id] = {"status": "error", "error": f"Phase {phase_id} not found"}
                return

        # Enrich each phase with its workflow data (studio_workflows has classification → regular conn)
        phases_data: list[dict] = []
        primary_fingerprint: dict = {}
        primary_excerpt = ""

        conn = get_connection()
        ph = sql_placeholder(conn)
        for phase in phases:
            wf_ids = json.loads(phase.get("workflow_ids") or "[]")
            phase_workflows = []
            for wf_id in wf_ids:
                wf_row = conn.execute(
                    f"SELECT * FROM studio_workflows WHERE workflow_id = {ph}", (wf_id,)
                ).fetchone()
                if wf_row:
                    wf_dict = dict(wf_row)
                    try:
                        wf_data = yaml.safe_load(wf_dict.get("template_yaml") or "{}") or {}
                    except Exception:
                        wf_data = {}
                    phase_workflows.append(wf_data)
                    if not primary_fingerprint:
                        fp_raw = wf_dict.get("style_fingerprint") or phase.get("style_fingerprint")
                        if fp_raw:
                            primary_fingerprint = json.loads(fp_raw)
                        src = wf_dict.get("source_doc_text") or ""
                        if src and not primary_excerpt:
                            primary_excerpt = src
            phases_data.append({
                "phase_number": phase["phase_number"],
                "name": phase["name"],
                "team_name": phase.get("team_name", ""),
                "workflows": phase_workflows,
            })
        conn.close()

        # Choose prompt based on scope
        if phase_id and len(phases_data) == 1 and len(phases_data[0]["workflows"]) == 1:
            prompt = _build_single_prompt(phases_data[0]["workflows"][0], primary_fingerprint, primary_excerpt)
            suffix = f"Phase {phases[0]['phase_number']} — {phases[0]['name']}"
            max_tokens = 4000
        else:
            prompt = _build_chain_prompt(chain, phases_data, primary_fingerprint, primary_excerpt)
            suffix = "End-to-End Process Document"
            max_tokens = 6000

        doc_text = _llm_generate(prompt, max_tokens=max_tokens)

        out_dir = _ARTIFACTS_DIR / job_id
        out_dir.mkdir(parents=True, exist_ok=True)
        title = f"{chain['name']} — {suffix}"
        artifacts = _export_artifacts(doc_text, title, out_dir, output_format)

        _REGEN_JOBS[job_id] = {
            "status": "done",
            "title": title,
            "artifact_dir": str(out_dir),
            "files": artifacts,
            "completed_at": _now(),
        }
        log.info("regen complete: chain=%s phase=%s job=%s files=%d", chain_id, phase_id, job_id, len(artifacts))
    except Exception as exc:
        log.error("regen chain %s failed: %s", chain_id, exc, exc_info=True)
        _REGEN_JOBS[job_id] = {"status": "error", "error": str(exc)}
