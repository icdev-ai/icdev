"""
ICDEV(tm) Executive AI Lab Documentation Generator  v2 — developer-first narrative
Produces: tools/presentations/output/ICDEV_AI_Lab_Documentation.docx
Message: A System That Builds Systems. Canvases, Workflows, Digital Twin, Use Cases.
Compliance / ATO = silent byproduct, not the lead.
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


NAVY  = RGBColor(0x0A, 0x16, 0x28)
GOLD  = RGBColor(0xC8, 0xA9, 0x51)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LGRAY = RGBColor(0x44, 0x44, 0x44)
MGRAY = RGBColor(0x1E, 0x3A, 0x5F)
GREEN = RGBColor(0x1A, 0x73, 0x48)


# ── XML helpers ──────────────────────────────────────────────────────────────

def _shd(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _gold_rule(para, color="C8A951", sz="6"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), sz)
    b.set(qn("w:space"), "4")
    b.set(qn("w:color"), color)
    pBdr.append(b)
    pPr.append(pBdr)


# ── Formatting helpers ────────────────────────────────────────────────────────

def _setup(doc):
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(1.0)
    sec.top_margin = sec.bottom_margin = Inches(0.9)
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(11)
    n.font.color.rgb = LGRAY


def _h1(doc, text):
    p = doc.add_heading(text, level=1)
    r = p.runs[0]
    r.font.color.rgb = NAVY
    r.font.size = Pt(18)
    r.font.bold = True
    _gold_rule(p)
    return p


def _h2(doc, text):
    p = doc.add_heading(text, level=2)
    r = p.runs[0]
    r.font.color.rgb = GOLD
    r.font.size = Pt(13)
    r.font.bold = True
    return p


def _h3(doc, text):
    p = doc.add_heading(text, level=3)
    r = p.runs[0]
    r.font.color.rgb = MGRAY
    r.font.size = Pt(11)
    r.font.bold = True
    return p


def _body(doc, text, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(6)
    return p


def _bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(4)


def _callout(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    r1 = p.add_run(f"{label}  ")
    r1.font.bold = True
    r1.font.color.rgb = GOLD
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.color.rgb = LGRAY


def _spacer(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)


def _pb(doc):
    doc.add_page_break()


def _table(doc, headers, rows, col_widths=None, hdr_bg="0A1628"):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        _shd(cell, hdr_bg)
        r = cell.paragraphs[0].runs[0]
        r.font.bold = True
        r.font.color.rgb = GOLD
        r.font.size = Pt(11)
        if col_widths:
            cell.width = Inches(col_widths[i])
    stripes = ["F5F7FA", "FFFFFF"]
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            cell.text = str(val)
            _shd(cell, stripes[ri % 2])
            r = cell.paragraphs[0].runs[0]
            r.font.size = Pt(10)
            r.font.color.rgb = LGRAY
            if col_widths:
                cell.width = Inches(col_widths[ci])
    doc.add_paragraph()
    return t


# ── Cover ─────────────────────────────────────────────────────────────────────

def cover(doc):
    _spacer(doc, 4)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ICDEV™")
    r.font.name = "Calibri"; r.font.size = Pt(52); r.font.bold = True
    r.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("A System That Builds Systems")
    r2.font.name = "Calibri"; r2.font.size = Pt(26); r2.font.bold = True
    r2.font.color.rgb = GOLD

    _spacer(doc)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("AI Lab Executive Documentation  ·  Adopt vs. Build Briefing Package")
    r3.font.size = Pt(13); r3.font.italic = True; r3.font.color.rgb = LGRAY

    _spacer(doc, 2)
    meta = [
        ("Classification", "CUI // SP-CTI"),
        ("Date",           "May 2026"),
        ("Version",        "2.0"),
        ("Audience",       "Senior Leadership / Executive Decision-Makers"),
        ("Contents",       "Platform overview  |  CBA  |  Lab requirements  |  Academy + GameDay playbooks"),
    ]
    t = doc.add_table(rows=len(meta), cols=2)
    t.style = "Table Grid"
    for i, (k, v) in enumerate(meta):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v
        for ci, c in enumerate(t.rows[i].cells):
            _shd(c, "0A1628")
            r = c.paragraphs[0].runs[0]
            r.font.color.rgb = GOLD if ci == 0 else WHITE
            r.font.size = Pt(11); r.font.bold = (ci == 0)
            c.width = Inches(2.5 if ci == 0 else 4.7)
    _pb(doc)


# ── Section 1: Executive Summary ─────────────────────────────────────────────

def sec_exec_summary(doc):
    _h1(doc, "1.  Executive Summary")
    _body(doc,
        "ICDEV™ is not a compliance tool. It is a developer platform that generates complete, "
        "working software systems from plain English requirements. Twelve design canvases give developers "
        "a visual thinking environment across every domain. The FORGE workflow engine and ANVIL build cycle "
        "convert intent into tested, deployable code. The Digital Program Twin simulates outcomes before "
        "a single line is written. Three pre-built use cases launch expert workflows in minutes. "
        "Compliance artifacts — SSP, POAM, SBOM, STIG checklists — are generated "
        "automatically as a byproduct. Developers never need to think about them."
    )
    _callout(doc, "The core insight:",
             "When developers build with ICDEV™, they get ATO-ready artifacts the same way they get "
             "unit tests — automatically, as a consequence of building correctly, not as extra work.")
    _h2(doc, "Recommendation")
    _body(doc,
        "Adopt ICDEV™. Stand up the AI Lab. Begin FORGE Academy Cohort 1 within 90 days. "
        "The platform is built, tested, and running. The only variable is the decision to move.")
    _h2(doc, "Headline Numbers")
    _table(doc,
        ["Metric", "Value"],
        [
            ["Year-1 Total Investment",         "$570K – $880K"],
            ["Year-1 Projected Return",          "$4M – $6M (3-program baseline)"],
            ["3-Year ROI (5 programs)",          "9×"],
            ["Time to First Feature",            "Day 1 (not after months of setup)"],
            ["Developer Delivery Speed",         "5× faster via 15 coordinating agents"],
            ["Developer Capacity (Year 1)",      "100 developers trained, 4 cohorts of 25"],
            ["Platform License Cost",            "$0 (Apache 2.0 open source)"],
            ["Compliance Artifact Generation",   "Automatic — zero developer effort"],
        ],
        col_widths=[3.2, 4.0],
    )
    _pb(doc)


# ── Section 2: The Developer Problem ─────────────────────────────────────────

def sec_problem(doc):
    _h1(doc, "2.  The Problem: Developers Spend More Time Getting Ready Than Building")
    _body(doc,
        "The constraint is not talent. Government software teams are full of capable developers. "
        "The constraint is the platform — or rather, the lack of one. "
        "Before the first line of mission code is written, a typical program team spends:"
    )
    for item in [
        "Weeks provisioning infrastructure and wiring CI/CD pipelines",
        "Months mapping compliance requirements that could have been derived from the code",
        "Repeated re-invention: every program starts from zero because knowledge lives in people, not systems",
        "Context switching across multi-cloud environments, six languages, AI governance, "
        "supply chain integrity, zero-trust architecture — simultaneously",
    ]:
        _bullet(doc, item)
    _body(doc,
        "By the time a team is ready to build, they are exhausted, behind schedule, "
        "and one key departure away from losing everything they know.")
    _h2(doc, "What a Platform Changes")
    _body(doc,
        "A platform shifts the question from “how do we set this up?” to "
        "“what are we building today?”. ICDEV™ is that platform. "
        "It handles setup, tooling, compliance, and knowledge transfer "
        "— so developers handle mission delivery.")
    _pb(doc)


# ── Section 3: Platform Overview ─────────────────────────────────────────────

def sec_platform(doc):
    _h1(doc, "3.  ICDEV™ Platform Overview")
    _body(doc,
        "ICDEV™ is a meta-builder: a system that builds systems. "
        "It generates complete, working applications from plain English requirements. "
        "Generated applications inherit the full ICDEV™ platform — FORGE framework, "
        "agent architecture, canvas suite, CI/CD, and knowledge graph — "
        "and can build their own next features using the same engine that built them.")

    _h2(doc, "3.1  FORGE — The Six-Layer Framework")
    _body(doc,
        "FORGE separates what to build (goals, args, context) from how to build it "
        "(orchestration, tools, hard prompts). AI reasoning is confined to orchestration only. "
        "Every tool call is deterministic Python. At 90% accuracy per step, a 5-step "
        "workflow degrades to ~59% without structure. FORGE keeps reliability above 95%.")
    _table(doc,
        ["Layer", "Location", "What It Does"],
        [
            ["Goals",         "goals/",       "66 workflow definitions: what to achieve, which tools, expected outputs"],
            ["Orchestration", "(AI layer)",   "Reads goal, decides tool order, applies args, handles errors"],
            ["Tools",         "icdev/tools/", "530+ Python scripts, one job each, deterministic execution"],
            ["Args",          "args/",        "YAML/JSON config: change behavior without editing code"],
            ["Context",       "context/",     "Static reference: catalogs, standards, tone rules"],
            ["Hard Prompts",  "hardprompts/", "Reusable LLM instruction templates across all workflows"],
        ],
        col_widths=[1.5, 1.8, 3.9],
    )

    _h2(doc, "3.2  ANVIL — The Five-Phase Build Cycle")
    _body(doc,
        "ANVIL runs automatically on every feature request. No manual steps.")
    _table(doc,
        ["Phase", "What Happens", "Output"],
        [
            ["Architect",  "Design and acceptance criteria defined up front",
             "Architecture doc, interface contracts"],
            ["Navigate",   "Map to existing tools and patterns; reuse before building",
             "Tool selection, dependency map"],
            ["Verify",     "Failing tests written first — the spec IS the test (RED)",
             "Test suite that fails correctly"],
            ["Integrate",  "Implementation generated until every test passes (GREEN)",
             "Working, tested code"],
            ["Launch",     "Refactor, security scan, compliance map, merge",
             "SBOM, scan report, compliance artifacts, merged PR"],
        ],
        col_widths=[1.2, 3.2, 2.8],
    )
    _callout(doc, "Byproduct:",
             "SBOM, STIG checklist, SSP delta, and compliance evidence are generated in the Launch phase "
             "automatically. Developers ship features. The platform produces the paperwork.")

    _h2(doc, "3.3  The 12 Design Canvases")
    _body(doc,
        "Canvases are the developer’s thinking environment. "
        "Each is a visual designer for a specific domain. "
        "Every canvas supports natural-language queries (IQE), "
        "simulation-readiness (Digital Twin), and one-click export to IaC or CI/CD config.")
    _table(doc,
        ["Code", "Canvas", "What a Developer Does Here"],
        [
            ["NDC",  "Network Design",       "Drag-drop topology, query blast radius, export Terraform"],
            ["SDC",  "Security Design",      "Run STRIDE analysis, surface gaps, generate mitigations"],
            ["PDC",  "Pipeline Design",      "Build CI/CD visually, score SLSA, estimate costs"],
            ["BDC",  "Boundary Design",      "Define system boundary, manage interface agreements"],
            ["DDC",  "Data Design",          "Classify data zones, track PII/CUI, visualize lineage"],
            ["ODC",  "Observability Design", "Map detection coverage, write Sigma rules, wire alerting"],
            ["IDC",  "Infrastructure Design","Design cloud resources across 6 CSPs, generate IaC"],
            ["AADC", "Agentic AI Design",    "Compose multi-agent systems from 7 solution packs"],
            ["QDC",  "Quality Design",       "Set quality gates, view UQS scores, surface code smells"],
            ["MDC",  "Migration Design",     "Run 7R assessment, track strangler fig, plan cutover"],
            ["AIMC", "AI/ML Model Design",   "Browse foundation model catalog, assess deployment options"],
            ["MSN",  "Mission Canvas",       "Plan operational missions, generate COA visualizations"],
        ],
        col_widths=[0.65, 1.85, 4.7],
    )
    _body(doc,
        "Common canvas workflow: design visually → query in plain English → "
        "run Digital Twin simulation → export to IaC or ANVIL build cycle.",
        italic=True)

    _h2(doc, "3.4  Unified Chat and Common Use Cases")
    _body(doc,
        "The /chat interface is a 3-pane, context-aware agent communication system. "
        "It is not a generic chatbot. Every conversation is typed: it carries project metadata, "
        "a use case template, canvas wiring, and a 7-dimension readiness score. "
        "The AI governance sidebar shows transparency, fairness, and accountability posture live.")
    _body(doc, "Three built-in use cases launch expert workflows in minutes:")
    _table(doc,
        ["Use Case", "What the Developer Starts With", "Canvas Wiring", "What Comes Out"],
        [
            ["General Modernization",
             "Describe the legacy system",
             "Migration, Digital Twin, Supply Chain",
             "Phased migration plan, 7Rs classification, rollback runbooks, IaC"],
            ["Year-End Budget Sprint",
             "Describe what needs to be purchased",
             "CPMP, Digital Twin, Supply Chain, Kanban",
             "BOM, IGCE, SOW, contract vehicle IDs — CSV/XLSX ready"],
            ["Crowd-Sourced Doc Refresh",
             "Point at a stale document",
             "Knowledge Graph, RAG, Audit",
             "Version-controlled doc, AI-rebuilt sections, auto-indexed in <24 hrs"],
        ],
        col_widths=[1.8, 2.1, 2.0, 2.3],
    )
    _callout(doc, "Adding use cases:",
             "Edit args/use_cases.yaml. No Python required. "
             "New use cases inherit expert persona, canvas wiring, and readiness scoring automatically.")

    _h2(doc, "3.5  Digital Program Twin")
    _body(doc,
        "The Digital Program Twin lets developers and program managers simulate outcomes "
        "before committing requirements. It runs 10,000 Monte Carlo iterations across six dimensions "
        "and produces confidence-banded forecasts at P10, P50, P80, and P90.")
    _table(doc,
        ["Dimension", "What It Measures"],
        [
            ["Architecture",  "Component count, API surface, coupling score, data flow complexity"],
            ["Compliance",    "Control coverage delta, boundary tier impact, POAM projection"],
            ["Supply Chain",  "New dependencies, vendor risk, SBOM delta, CVE exposure"],
            ["Schedule",      "PERT estimates, critical path, PI count via Monte Carlo"],
            ["Cost",          "T-shirt size roll-up, infrastructure delta, vendor licensing"],
            ["Risk",          "Compound risk score, risk interactions, mitigation effectiveness"],
        ],
        col_widths=[1.5, 5.7],
    )
    _body(doc, "Three Courses of Action generated automatically:")
    _table(doc,
        ["COA", "Timeline", "Cost", "Risk", "When to Choose"],
        [
            ["Speed",            "1–2 PIs", "Lowest",  "HIGH", "P1 only; fastest capability delivery; accept technical debt"],
            ["Balanced (★)","2–3 PIs", "Medium",  "MOD",  "P1+P2; recommended for most programs; best tradeoff"],
            ["Comprehensive",    "3–5 PIs", "Highest", "LOW",  "Full scope; lowest residual risk; longest timeline"],
        ],
        col_widths=[1.4, 1.0, 1.0, 0.8, 4.0],
    )

    _h2(doc, "3.6  15-Agent Autonomous Ecosystem")
    _body(doc,
        "ICDEV™ deploys 15 AI agents across three tiers. "
        "Agents coordinate continuously via JSON-RPC 2.0 over mutual TLS, "
        "eliminating the expert bottlenecks that stall government programs.")
    _table(doc,
        ["Tier", "Agent", "What It Does for Developers"],
        [
            ["Core",    "Orchestrator",    "Routes tasks, manages DAG execution, never blocks the developer"],
            ["Core",    "Architect",       "Designs systems, writes acceptance criteria, selects patterns"],
            ["Domain",  "Builder",         "Generates tested code in 6 languages via TDD"],
            ["Domain",  "Compliance",      "Produces SSP, POAM, STIG, SBOM, OSCAL — silently"],
            ["Domain",  "Security",        "Runs SAST, dependency audit, secret detection in every build"],
            ["Domain",  "Infrastructure",  "Generates Terraform, Ansible, K8s across 6 cloud providers"],
            ["Domain",  "Modernization",   "Runs 7R assessment, generates migration plan, bridges ATO"],
            ["Domain",  "Requirements",    "Drives conversational intake, detects gaps, scores readiness"],
            ["Domain",  "Supply Chain",    "Maintains SBOM, triages CVEs, checks Section 889"],
            ["Domain",  "Simulation",      "Runs Digital Twin, generates COAs, produces Monte Carlo output"],
            ["Domain",  "DevSecOps + ZTA", "Hardens pipelines, generates Zero Trust policies"],
            ["Support", "Knowledge",       "Records patterns, drives self-healing, feeds Academy curriculum"],
            ["Support", "Monitor",         "Tracks SLAs, surfaces anomalies, keeps agents healthy"],
        ],
        col_widths=[0.9, 1.6, 4.7],
    )
    _pb(doc)


# ── Section 4: CBA ────────────────────────────────────────────────────────────

def sec_cba(doc):
    _h1(doc, "4.  Cost-Benefit Analysis")
    _body(doc,
        "The CBA is framed around developer productivity and program delivery speed, "
        "not around compliance cost avoidance. Compliance savings are real and substantial — "
        "but they are a byproduct. The primary return is measured in programs delivered, "
        "features shipped, and teams that improve every sprint.")

    _h2(doc, "4.1  Benefits")
    _table(doc,
        ["Benefit Driver", "Without ICDEV™", "With ICDEV™", "Estimated Annual Value"],
        [
            ["Time to first feature",
             "Weeks of setup scaffolding",
             "Day 1 — canvases and agents ready",
             "$500K+ per program (velocity)"],
            ["Developer throughput",
             "1 feature / 2 weeks / developer",
             "5× faster via 15 coordinating agents",
             "$1M+ per program per year"],
            ["Knowledge retention",
             "Exits with departing personnel",
             "Knowledge graph + Genesis daemon",
             "$200K–$1M per key departure avoided"],
            ["Developer training (100 devs)",
             "External courses: ~$5K / developer",
             "FORGE Academy: ~$800 / developer",
             "$420K saved in Year 1"],
            ["Legacy modernization",
             "Greenfield rebuild required",
             "7Rs + migration canvas + ATO bridge",
             "$2M–$5M per program avoided"],
            ["Compliance artifacts (byproduct)",
             "2 ISSOs × 6 months manual",
             "Auto-generated in every build",
             "$400K–$800K / year"],
            ["TOTAL YEAR-1 BENEFIT (3 programs)", "", "", "$4M – $6M+"],
        ],
        col_widths=[2.4, 1.8, 2.0, 2.0],
    )

    _h2(doc, "4.2  Costs (Year 1)")
    _table(doc,
        ["Cost Item", "Year 1", "Notes"],
        [
            ["Infrastructure (K8s + cloud LLM API)",  "$150K – $300K", "AWS Bedrock or Azure OpenAI Gov (IL4+)"],
            ["ICDEV™ platform license",           "$0",                 "Apache 2.0 open source"],
            ["Personnel (4 FTEs, blended rate)",       "$300K – $400K", "Director, Engineer, Training Lead, ISSO"],
            ["FORGE Academy (4 cohorts × 25)",    "$80K – $120K",  "$20K per cohort"],
            ["AI GameDay (4 events × 100 devs)",  "$40K – $60K",   "$10K per event"],
            ["TOTAL YEAR-1 COST",                     "$570K – $880K", ""],
        ],
        col_widths=[3.0, 1.8, 2.4],
    )

    _h2(doc, "4.3  Three-Year ROI by Program Volume")
    _table(doc,
        ["Programs / Year", "Year 1 Net", "Year 2 Net", "Year 3 Net", "3-Year ROI"],
        [
            ["3 programs",  "+$3.0M",  "+$8.0M",  "+$14.0M", "5×"],
            ["5 programs",  "+$6.0M",  "+$16.0M", "+$28.0M", "9×"],
            ["10 programs", "+$16.0M", "+$36.0M", "+$60.0M", "18×"],
        ],
        col_widths=[1.8, 1.6, 1.6, 1.6, 1.6],
    )
    _callout(doc, "Compounding effect:",
             "Every program powered by ICDEV™ contributes patterns back to the shared knowledge base. "
             "Program N+1 starts smarter than Program N. ROI grows non-linearly.")
    _pb(doc)


# ── Section 5: Lab Requirements ──────────────────────────────────────────────

def sec_lab(doc):
    _h1(doc, "5.  AI Lab Build-Out Requirements")
    _body(doc,
        "The AI Lab is the home of the ICDEV™ platform, FORGE Academy, and AI GameDay. "
        "It is a developer-first environment — optimized for fast iteration, "
        "deep AI capability, and 100+ developer throughput.")

    _h2(doc, "5.1  Infrastructure")
    _table(doc,
        ["Component", "Specification", "Why It Matters"],
        [
            ["Kubernetes Cluster",   "2-node, 64 vCPU, 256 GB RAM, 10 TB NVMe",
             "Runs 15 ICDEV agents + dashboard + developer sandboxes"],
            ["LLM API",              "AWS Bedrock (Claude) or Azure OpenAI Gov",
             "IL4+ approved; powers all 15 agents and canvas intelligence"],
            ["Network",              "Isolated lab VLAN, CAC auth, VPN breakout",
             "Developer access without crossing production boundary"],
            ["Object Storage",       "S3 GovCloud or equiv, 10 TB",
             "Artifacts, SBOM, knowledge graph, audit logs"],
            ["CI/CD Platform",       "GitLab EE or GitHub Enterprise",
             "ANVIL pipeline integration for all developer builds"],
            ["Developer Sandboxes",  "Docker Compose per developer (or shared cloud VDI)",
             "Each developer runs a local ICDEV instance for hands-on building"],
        ],
        col_widths=[2.0, 2.6, 2.6],
    )

    _h2(doc, "5.2  Personnel (4 FTEs)")
    _table(doc,
        ["Role", "Responsibilities", "Clearance"],
        [
            ["AI Lab Director",
             "Strategy, executive liaison, program integration, Academy oversight",
             "TS/SCI preferred"],
            ["Platform / MLOps Engineer",
             "ICDEV deployment, 15-agent health, LLM integration, CI/CD, developer sandbox support",
             "Secret min."],
            ["Training Lead",
             "Academy curriculum delivery, cohort scheduling, GameDay facilitation, cert tracking",
             "Secret min."],
            ["ISSO",
             "Lab authorization to operate, STIG compliance, audit trail review, POA&M",
             "Secret req."],
        ],
        col_widths=[1.9, 4.3, 1.0],
    )

    _h2(doc, "5.3  30 / 60 / 90-Day Stand-Up Timeline")
    _table(doc,
        ["Window", "What Gets Done", "Who"],
        [
            ["Day 1–30\nInfrastructure",
             "K8s cluster provisioned. Bedrock/OpenAI Gov access granted. "
             "ICDEV™ deployed and health-checked. GitLab wired. Developer sandbox template created.",
             "Platform Eng + Lab Director"],
            ["Day 31–60\nFirst Program",
             "First program team onboarded to 2 canvases (NDC + AADC). "
             "RICOAS intake session run. Digital Twin simulation executed. "
             "Compliance artifacts generated automatically — no one had to ask.",
             "Platform Eng + ISSO"],
            ["Day 61–90\nAcademy + GameDay",
             "Cohort 1 (25 developers) begins 8-week Academy track. "
             "AI GameDay #1 executed — 4 teams, Operation CIPHER FORGE. "
             "AAR auto-published. Genesis daemon promotes GameDay patterns into new missions.",
             "Training Lead + Lab Director"],
        ],
        col_widths=[1.7, 4.3, 1.2],
    )
    _pb(doc)


# ── Section 6: Academy ────────────────────────────────────────────────────────

def sec_academy(doc):
    _h1(doc, "6.  FORGE Academy — Operating Playbook (100+ Developers)")
    _body(doc,
        "FORGE Academy turns developers into AI-native engineers. "
        "Every mission teaches a pattern that applies directly to ICDEV™ canvases and workflows. "
        "Developers do not learn about AI in the abstract — they build with it, "
        "on real tools, starting Day 1.")

    _h2(doc, "6.1  Program Structure")
    _table(doc,
        ["Attribute", "Detail"],
        [
            ["Role tracks",        "12: DevOps, DataOps, SecOps, SWE/Architect, NetOps, SRE, ISSO, ISSM, CISO, PM, Analyst, Leadership"],
            ["Missions",           "75 missions, 165 steps, 3 tiers (Fundamentals → Role-Specific → Capstone)"],
            ["Delivery",           "Hybrid: 2 hrs instructor-led + 8 hrs self-paced per week (10 hrs/wk total)"],
            ["Lab mode",           "Coding Lab (hands-on terminal) + Guided Lab (no-code for non-technical roles)"],
            ["Cohort size",        "25 developers per cohort"],
            ["Duration",           "8 weeks per cohort"],
            ["Facilitator ratio",  "1 Sensei per 25 learners (Sensei = internal L5-certified developer)"],
            ["Auto-currency",      "Genesis daemon promotes proven patterns into new draft missions weekly"],
        ],
        col_widths=[2.2, 5.0],
    )

    _h2(doc, "6.2  Developer Journey (Week by Week)")
    _table(doc,
        ["Week", "What the Developer Builds", "Canvas Used"],
        [
            ["1",    "First LLM feature via ICDEV™ router",                    "Chat"],
            ["2",    "Network topology + IQE natural-language query",               "NDC"],
            ["3",    "STRIDE threat model + gap detection",                         "SDC"],
            ["4",    "Multi-agent agentic system design",                           "AADC"],
            ["5",    "Digital Twin simulation — 3 COAs generated",             "Simulation"],
            ["6",    "CI/CD pipeline design + SLSA assessment",                     "PDC"],
            ["7",    "Data classification + PII tracking",                          "DDC"],
            ["8",    "Capstone: full ANVIL cycle — intake to deployed feature","All canvases"],
        ],
        col_widths=[0.7, 4.0, 2.5],
    )

    _h2(doc, "6.3  Rank Progression")
    _table(doc,
        ["Rank", "XP", "What the Developer Can Do"],
        [
            ["Recruit",    "0",       "Uses AI tools safely; understands LLM capabilities and limits"],
            ["Operative",  "500",     "Builds LLM features, simple RAG pipelines, single-tool agents"],
            ["Specialist", "2,000",   "Designs multi-agent systems; monitors production AI; owns AI features"],
            ["Architect",  "5,000",   "Builds AI platforms; leads AI initiatives; authors ICDEV™ tools"],
            ["Sensei",     "10,000+", "Defines AI strategy; teaches + certifies others; trains next cohort"],
        ],
        col_widths=[1.2, 0.9, 5.1],
    )

    _h2(doc, "6.4  100-Developer Cohort Schedule")
    _table(doc,
        ["Cohort", "Developers", "Weeks",   "Certification Window"],
        [
            ["Cohort 1", "25", "Wk 1–08",  "Foundation eligible: Month 3"],
            ["Cohort 2", "25", "Wk 5–12",  "Foundation eligible: Month 4"],
            ["Cohort 3", "25", "Wk 9–16",  "Foundation eligible: Month 5"],
            ["Cohort 4", "25", "Wk 13–20", "Foundation eligible: Month 6"],
            ["Total",    "100","Rolling",         "All 100 Foundation-certified by Month 6"],
        ],
        col_widths=[1.2, 1.4, 1.5, 3.1],
    )

    _h2(doc, "6.5  Cost per Cohort")
    _table(doc,
        ["Item", "Per Cohort (25 devs)", "4 Cohorts (100 devs)"],
        [
            ["Lab infrastructure (shared cloud sandbox)", "$5,000",  "$20,000"],
            ["Facilitator time (10 hrs/wk × 8 wks)", "$8,000",  "$32,000"],
            ["Learning tooling",                           "$3,000",  "$12,000"],
            ["Certification assessment",                   "$4,000",  "$16,000"],
            ["TOTAL",                                      "$20,000", "$80,000"],
        ],
        col_widths=[3.2, 1.8, 2.2],
    )
    _pb(doc)


# ── Section 7: GameDay ────────────────────────────────────────────────────────

def sec_gameday(doc):
    _h1(doc, "7.  AI GameDay — Operating Playbook (100+ Developers)")
    _body(doc,
        "AI GameDay is where FORGE Academy skills are tested under real pressure. "
        "It is a live, AI-scored, competitive tabletop exercise. "
        "Four teams race through five injects using ICDEV™ tools, "
        "with every response judged by an LLM against a published rubric. "
        "It is not a drill. It is a competition with a leaderboard, "
        "achievement ribbons, and an After-Action Report that feeds back into Academy curriculum.")

    _h2(doc, "7.1  Team Structure")
    _table(doc,
        ["Team", "Mission", "ICDEV™ Tools Used"],
        [
            ["RED — Threat",      "Attack design and threat emulation",
             "AADC canvas, Simulation, Security Design Canvas"],
            ["BLUE — Defense",    "Incident response and defense posture",
             "Observability canvas, Security canvas, IR runbooks"],
            ["GOLD — Innovation", "Novel AI modules and ML training pair generation",
             "AIMC canvas, Finetune tools, AADC canvas"],
            ["GREEN — Compliance","NIST audit verdicts and ethics assessment",
             "Security canvas, Compliance agents, Quality canvas"],
        ],
        col_widths=[1.7, 2.4, 3.1],
    )

    _h2(doc, "7.2  Five Injects — Operation CIPHER FORGE")
    _table(doc,
        ["Inject", "Duration", "What Teams Do"],
        [
            ["Signal Cluster",      "15 min",     "Analyze SIGINT feed; recommend threat posture using simulation canvas"],
            ["COA Posture",         "40 min",     "Design force posture using Digital Twin; compare 3 COAs"],
            ["Ransomware Cascade",  "Sequential", "Contain, eradicate, recover — using IR runbooks from Observability canvas"],
            ["Fine-Tune Sprint",    "Sequential", "Generate ML training pairs (Gold team); Red/Blue evaluate quality"],
            ["War Council Brief",   "Final",      "Executive synthesis: present findings from all injects to facilitator panel"],
        ],
        col_widths=[1.8, 1.2, 4.2],
    )

    _h2(doc, "7.3  Scoring Model")
    _table(doc,
        ["Dimension", "Weight", "How It Is Scored"],
        [
            ["Adversarial Effectiveness", "40%", "LLM judge evaluates Red/Blue tactical decisions against published rubric"],
            ["Innovation Score",          "25%", "Gold team’s ML modules evaluated by Evaluator role + auto-quality check"],
            ["Compliance Score",          "20%", "Green team NIST verdicts applied to all team artifacts"],
            ["Training Pairs",            "15%", "Gold team pairs auto-scored; quality > 0.60 ingested for model fine-tuning"],
        ],
        col_widths=[2.4, 0.9, 3.9],
    )

    _h2(doc, "7.4  Logistics (100-Person Event)")
    for item in [
        "4 teams × 25 participants; 1 Lead Facilitator + 4 Team Leads",
        "4-hour live session; pre-event 30-minute team formation (AI skill-to-role matching)",
        "Each team has LLM-backed ICDEV™ tool access; receipts server-verified for scoring integrity",
        "Live leaderboard visible to all participants throughout the event",
        "AAR auto-generated within 15 minutes of session close",
        "Kanban debrief card created automatically; lessons learned reviewed next Academy session",
        "GameDay completion required for FORGE Academy Practitioner certification",
        "Frequency: quarterly tournaments + ad hoc challenge events",
    ]:
        _bullet(doc, item)

    _h2(doc, "7.5  Cost per Event")
    _table(doc,
        ["Item", "Per Event (100 participants)", "Annual (4 events)"],
        [
            ["GameDay server burst compute", "$2,000",  "$8,000"],
            ["LLM API credits (4 teams × 4 hrs)", "$1,500", "$6,000"],
            ["Facilitation (Lead + 4 Team Leads)", "$4,000", "$16,000"],
            ["Venue / logistics (in-person)", "$2,500", "$10,000"],
            ["TOTAL", "$10,000", "$40,000"],
        ],
        col_widths=[3.0, 2.0, 2.2],
    )
    _pb(doc)


# ── Section 8: Risk ───────────────────────────────────────────────────────────

def sec_risk(doc):
    _h1(doc, "8.  Risk Analysis")
    _table(doc,
        ["Risk", "Likelihood", "Impact", "Mitigation"],
        [
            ["LLM API outage during GameDay", "Low", "High",
             "Local Ollama fallback built into ICDEV (air-gap mode); zero internet required"],
            ["First cohort skill gap", "Medium", "Medium",
             "Guided lab mode for non-technical roles; adaptive path adjusts automatically"],
            ["Infrastructure procurement delay", "Medium", "High",
             "Start on Bedrock cloud day 1; on-prem K8s arrives in parallel"],
            ["Canvas adoption curve", "Medium", "Medium",
             "Day-1 onboarding; Sensei-led sessions; canvas IQE makes querying natural"],
            ["LLM model deprecation", "Low", "Medium",
             "Vendor-agnostic routing: swap model in .env, no code changes"],
            ["Key person departure (Training Lead)", "Medium", "High",
             "Academy Sensei certifications create internal expert bench by Month 6"],
            ["Scope creep in first program", "High", "Medium",
             "Digital Twin Speed COA enforces P1-only boundary; COA gate before build starts"],
        ],
        col_widths=[2.5, 1.0, 0.8, 2.9],
    )
    _pb(doc)


# ── Section 9: Recommendation ─────────────────────────────────────────────────

def sec_recommendation(doc):
    _h1(doc, "9.  Recommendation")
    _body(doc,
        "Adopt ICDEV™. Stand up the AI Lab. Begin FORGE Academy Cohort 1 within 90 days.",
        bold=True, size=14)
    _spacer(doc)
    _body(doc,
        "The platform is built, tested, and running. "
        "Twelve canvases give developers the thinking environment they have never had. "
        "The Digital Program Twin eliminates requirements guesswork. "
        "Three use cases launch expert workflows in minutes. "
        "Fifteen agents handle the coordination that used to require fifteen specialists. "
        "And at the end of every build, compliance artifacts appear automatically — "
        "because that is what a well-built system produces.")
    _spacer(doc)
    _body(doc, "Three actions to approve:", bold=True)
    for item in [
        "Infrastructure: K8s cluster + Bedrock/OpenAI Gov access. 30-day procurement. $150K–$300K.",
        "Personnel: 4 FTEs (Lab Director, Platform Engineer, Training Lead, ISSO). $300K–$400K/yr.",
        "Budget: Academy (4 cohorts) + GameDay (4 events) = $120K–$160K.",
    ]:
        _bullet(doc, item)
    _spacer(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("One platform.  Every program.  Starting Day 1.")
    r.font.size = Pt(18); r.font.bold = True; r.font.color.rgb = NAVY
    _pb(doc)


# ── Appendices ────────────────────────────────────────────────────────────────

def appendices(doc):
    _h1(doc, "Appendix A: Canvas Reference Card")
    _table(doc,
        ["Code", "Canvas", "Route", "Top Developer Action"],
        [
            ["NDC",  "Network Design",       "/network",     "Drag topology → query blast radius → export Terraform"],
            ["SDC",  "Security Design",      "/security",    "Run STRIDE → map ATT&CK → generate mitigations"],
            ["PDC",  "Pipeline Design",      "/devops",      "Build CI/CD visually → score SLSA → estimate cost"],
            ["BDC",  "Boundary Design",      "/boundary",    "Define system boundary → manage ISAs → map PPS"],
            ["DDC",  "Data Design",          "/data",        "Classify zones → track PII/CUI → visualize lineage"],
            ["ODC",  "Observability Design", "/observability","Map detection → write Sigma rules → wire alerting"],
            ["IDC",  "Infrastructure Design","/infra",       "Design resources → compare 6 CSPs → generate IaC"],
            ["AADC", "Agentic AI Design",    "/agentic-ai",  "Compose agents → assign tools → assess risk"],
            ["QDC",  "Quality Design",       "/qdc",         "Set gates → view UQS → surface smells"],
            ["MDC",  "Migration Design",     "/migration",   "7R assess → track strangler fig → plan cutover"],
            ["AIMC", "AI/ML Model Design",   "/ai-ml",       "Browse catalog → assess IL → plan deployment"],
            ["MSN",  "Mission Canvas",       "/mission",     "Plan mission → generate COA viz → link to digital thread"],
        ],
        col_widths=[0.65, 1.85, 1.3, 3.4],
    )
    _pb(doc)

    _h1(doc, "Appendix B: Digital Twin COA Comparison Matrix")
    _table(doc,
        ["Dimension",     "Speed COA", "Balanced COA (★)", "Comprehensive COA"],
        [
            ["Timeline",  "1–2 PIs",   "2–3 PIs",   "3–5 PIs"],
            ["Cost",      "Lowest",          "Medium",          "Highest"],
            ["Risk",      "HIGH",            "MODERATE",        "LOW"],
            ["Scope",     "P1 only",         "P1 + P2",         "P1 + P2 + P3"],
            ["Tech Debt", "High",            "Moderate",        "Low"],
            ["Boundary",  "Least impact",    "Average",         "May trigger re-authorization"],
            ["Best for",  "Fast feedback",   "Most programs",   "Full delivery assurance"],
        ],
        col_widths=[1.8, 1.8, 2.0, 1.8],
    )
    _pb(doc)

    _h1(doc, "Appendix C: Use Case Reference")
    for name, badge, pre, wiring, output in [
        ("General Modernization",    "AI Boost",  "12 requirements",
         "Migration Canvas, Digital Twin, Supply Chain, Compliance",
         "Phased migration plan, 7Rs classification, rollback runbooks, IaC"),
        ("Year-End Budget Sprint",   "RICOAS",    "14 requirements",
         "CPMP, Digital Twin, Supply Chain, Kanban",
         "BOM with lead times, IGCE, SOW, contract vehicles (CSV/XLSX)"),
        ("Crowd-Sourced Doc Refresh","RAG + KG",  "10 requirements",
         "Knowledge Graph, RAG, Compliance, Audit",
         "Version-controlled doc, AI-rebuilt sections, auto-indexed <24 hrs"),
    ]:
        _body(doc, name, bold=True, color=GOLD)
        _bullet(doc, f"Badge: {badge}  ·  Pre-loaded: {pre}")
        _bullet(doc, f"Canvas wiring: {wiring}")
        _bullet(doc, f"Output: {output}")
        _spacer(doc)
    _body(doc,
        "Adding use cases: edit args/use_cases.yaml. No Python. "
        "Expert persona, canvas wiring, and readiness scoring applied automatically.",
        italic=True)
    _pb(doc)

    _h1(doc, "Appendix D: FORGE Academy Mission Sample — M01 LLM Fundamentals")
    _table(doc,
        ["Step", "Activity", "Mode", "XP"],
        [
            ["1", "5-min video: transformer architecture",                   "Guided", "10"],
            ["2", "Read: ICDEV™ LLM router overview",                   "Guided", "10"],
            ["3", "Guided Lab: send first prompt via ICDEV router",          "Guided", "20"],
            ["4", "Challenge: identify which ICDEV tool uses which LLM",     "Guided", "30"],
            ["5", "Assessment: 5 questions, 80% pass rate required",         "Guided", "30"],
        ],
        col_widths=[0.6, 4.2, 1.0, 0.7],
    )
    _body(doc, "Completion unlocks Operative rank track. DoD KSA: KS0001 (AI/ML Fundamentals).",
          italic=True)
    _pb(doc)

    _h1(doc, "Appendix E: Glossary")
    _table(doc,
        ["Term", "Definition"],
        [
            ["ANVIL",        "5-phase TDD build cycle: Architect, Navigate, Verify, Integrate, Launch"],
            ["ATO",          "Authority to Operate — a byproduct of building with ICDEV™, not the goal"],
            ["COA",          "Course of Action: one of three Digital Twin strategies (Speed / Balanced / Comprehensive)"],
            ["Digital Twin", "6-dimension what-if simulation engine for program planning"],
            ["FORGE",        "6-layer framework: Goals, Orchestration, Tools, Args, Context, Hard Prompts"],
            ["IQE",          "ICDEV Query Engine — natural-language queries over live canvas data"],
            ["Kanban",       "ICDEV™ task board: features flow from intake through ANVIL to shipped"],
            ["RICOAS",       "Requirements Intake, COA and Approval System (Phases 1–4)"],
            ["SBOM",         "Software Bill of Materials — generated automatically in every ANVIL Launch phase"],
            ["Sensei",       "L5-certified developer; trains and certifies the next Academy cohort"],
            ["TTX",          "Tabletop Exercise — what AI GameDay is, except competitive and AI-scored"],
        ],
        col_widths=[1.5, 5.7],
    )


# ── IDR export entry point ────────────────────────────────────────────────────

def generate_docx(title: str, content: str, output_path: str) -> str:
    """Generate a styled DOCX from arbitrary text content.

    Called by the IDR publish pipeline (tools/docgen/workflow._try_export_docx).

    Parameters
    ----------
    title:       Document title — rendered as an H1 cover heading.
    content:     Plain text body.  Blank lines become paragraph breaks;
                 lines starting with "# " / "## " become H1 / H2 headings;
                 lines starting with "- " / "* " become bullets.
    output_path: Absolute or relative path for the output .docx file.

    Returns
    -------
    str  — the resolved output path.
    """
    doc = Document()
    _setup(doc)

    # Cover heading
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.name = "Calibri"
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = NAVY
    _gold_rule(p)
    _spacer(doc)

    # Body — rudimentary Markdown-ish parsing
    for line in content.splitlines():
        stripped = line.rstrip()
        if not stripped:
            _spacer(doc)
        elif stripped.startswith("## "):
            _h2(doc, stripped[3:])
        elif stripped.startswith("# "):
            _h1(doc, stripped[2:])
        elif stripped.startswith(("- ", "* ")):
            _bullet(doc, stripped[2:])
        else:
            _body(doc, stripped)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return str(out)


# ── Main ─────────────────────────────────────────────────────────────────────

def main():
    out_dir = Path(__file__).parent / "output"
    out_dir.mkdir(exist_ok=True)

    doc = Document()
    _setup(doc)

    cover(doc)
    sec_exec_summary(doc)
    sec_problem(doc)
    sec_platform(doc)
    sec_cba(doc)
    sec_lab(doc)
    sec_academy(doc)
    sec_gameday(doc)
    sec_risk(doc)
    sec_recommendation(doc)
    appendices(doc)

    out = out_dir / "ICDEV_AI_Lab_Documentation.docx"
    doc.save(str(out))
    print(f"[OK] {out}")


if __name__ == "__main__":
    main()
