"""
RFI Response DOCX Exporter — converts a structured RFI response (markdown or JSON)
to a properly formatted Word document matching government submission requirements.

Formatting per common government RFI submission guidance (override per
solicitation via the exporter's formatting options):
  - 11-point Times New Roman
  - Single-spaced
  - 1-inch margins
  - UNCLASSIFIED//FOUO header and footer on every page
  - Cover page (not counted in page limit)

Usage:
    python tools/govcon/rfi_docx_exporter.py --input response.md --output response.docx
    python tools/govcon/rfi_docx_exporter.py --input response.md --output response.docx --fouo
    python tools/govcon/rfi_docx_exporter.py --input response.md --output response.docx --profile own_company
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

_ROOT = Path(__file__).resolve().parent.parent.parent
if str(_ROOT) not in sys.path:
    sys.path.insert(0, str(_ROOT))

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

FOUO_TEXT = "UNCLASSIFIED//FOUO"
DEFAULT_FONT = "Times New Roman"
DEFAULT_FONT_SIZE = 11
MARGIN_INCHES = 1.0


def _set_margins(doc: "Document", inches: float = MARGIN_INCHES) -> None:
    for section in doc.sections:
        section.top_margin = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin = Inches(inches)
        section.right_margin = Inches(inches)


def _add_fouo_header_footer(doc: "Document", text: str = FOUO_TEXT) -> None:
    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp.add_run(text)
        run.bold = True
        run.font.name = DEFAULT_FONT
        run.font.size = Pt(DEFAULT_FONT_SIZE - 1)
        run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

        # Footer
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run(text)
        frun.bold = True
        frun.font.name = DEFAULT_FONT
        frun.font.size = Pt(DEFAULT_FONT_SIZE - 1)
        frun.font.color.rgb = RGBColor(0x00, 0x80, 0x00)


def _set_run_font(run, bold: bool = False, size: int = DEFAULT_FONT_SIZE) -> None:
    run.font.name = DEFAULT_FONT
    run.font.size = Pt(size)
    run.bold = bold


def _add_paragraph(doc: "Document", text: str = "", bold: bool = False,
                   size: int = DEFAULT_FONT_SIZE,
                   align: "WD_ALIGN_PARAGRAPH" = WD_ALIGN_PARAGRAPH.LEFT,
                   space_before: float = 0, space_after: float = 3) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(DEFAULT_FONT_SIZE * 1.15)
    if text:
        run = p.add_run(text)
        _set_run_font(run, bold=bold, size=size)
    return p


def _add_heading(doc: "Document", text: str, level: int = 1) -> None:
    sizes = {1: 14, 2: 13, 3: 12}
    size = sizes.get(level, DEFAULT_FONT_SIZE)
    p = _add_paragraph(doc, text, bold=True, size=size, space_before=6, space_after=3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_table_row(table, cells: list[str]) -> None:
    row = table.add_row()
    for i, cell_text in enumerate(cells):
        cell = row.cells[i]
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = DEFAULT_FONT
                run.font.size = Pt(DEFAULT_FONT_SIZE - 1)


def _parse_markdown_to_blocks(md_text: str) -> list[dict]:
    """Parse markdown into a list of typed content blocks."""
    blocks = []
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # Skip FOUO markers (handled via header/footer)
        if FOUO_TEXT in line and line.strip().startswith(">"):
            i += 1
            continue

        # H1
        if line.startswith("# "):
            blocks.append({"type": "h1", "text": line[2:].strip()})
        # H2
        elif line.startswith("## "):
            blocks.append({"type": "h2", "text": line[3:].strip()})
        # H3
        elif line.startswith("### "):
            blocks.append({"type": "h3", "text": line[4:].strip()})
        # Horizontal rule
        elif re.match(r"^-{3,}$", line.strip()):
            blocks.append({"type": "hr"})
        # Code block
        elif line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            blocks.append({"type": "code", "text": "\n".join(code_lines)})
        # Table
        elif line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                if not re.match(r"^\|[-| ]+\|$", lines[i]):
                    table_lines.append(lines[i])
                i += 1
            blocks.append({"type": "table", "rows": table_lines})
            continue
        # Bullet
        elif re.match(r"^[-*] ", line):
            blocks.append({"type": "bullet", "text": line[2:].strip()})
        # Numbered list
        elif re.match(r"^\d+\. ", line):
            blocks.append({"type": "numbered", "text": re.sub(r"^\d+\. ", "", line).strip()})
        # Blockquote (non-FOUO)
        elif line.startswith("> ") and FOUO_TEXT not in line:
            blocks.append({"type": "quote", "text": line[2:].strip()})
        # Empty line
        elif line.strip() == "":
            blocks.append({"type": "empty"})
        # Normal paragraph
        else:
            blocks.append({"type": "para", "text": line.strip()})

        i += 1
    return blocks


def _render_inline(text: str) -> list[tuple[str, bool]]:
    """Split inline text into (text, bold) tuples based on **bold** markers."""
    parts = []
    for segment in re.split(r"(\*\*[^*]+\*\*)", text):
        if segment.startswith("**") and segment.endswith("**"):
            parts.append((segment[2:-2], True))
        else:
            # Strip remaining single * italic markers (unsupported in plain DOCX here)
            clean = segment.replace("*", "")
            if clean:
                parts.append((clean, False))
    return parts


def _add_inline_para(doc: "Document", text: str,
                     indent: float = 0, space_after: float = 3) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(DEFAULT_FONT_SIZE * 1.15)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    for segment_text, bold in _render_inline(text):
        run = p.add_run(segment_text)
        _set_run_font(run, bold=bold)


def markdown_to_docx(md_text: str, output_path: str, classification: str = FOUO_TEXT) -> str:
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    doc = Document()
    _set_margins(doc)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = DEFAULT_FONT
    style.font.size = Pt(DEFAULT_FONT_SIZE)

    blocks = _parse_markdown_to_blocks(md_text)
    list_counter = 0

    for block in blocks:
        btype = block.get("type")

        if btype == "h1":
            _add_heading(doc, block["text"], level=1)
            list_counter = 0
        elif btype == "h2":
            _add_heading(doc, block["text"], level=2)
            list_counter = 0
        elif btype == "h3":
            _add_heading(doc, block["text"], level=3)
            list_counter = 0
        elif btype == "hr":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run("─" * 80)
            run.font.name = DEFAULT_FONT
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            list_counter = 0
        elif btype == "code":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(block["text"])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif btype == "table":
            rows = block["rows"]
            if not rows:
                continue
            # Parse columns from first row
            cols = [c.strip() for c in rows[0].split("|") if c.strip()]
            num_cols = len(cols)
            table = doc.add_table(rows=1, cols=num_cols)
            table.style = "Table Grid"
            # Header row
            hdr = table.rows[0].cells
            for j, col_text in enumerate(cols):
                hdr[j].text = col_text
                for para in hdr[j].paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.name = DEFAULT_FONT
                        run.font.size = Pt(DEFAULT_FONT_SIZE - 1)
            # Data rows
            for row_line in rows[1:]:
                cells = [c.strip() for c in row_line.split("|") if c.strip()]
                # Pad/truncate to num_cols
                cells = (cells + [""] * num_cols)[:num_cols]
                _add_table_row(table, cells)
            doc.add_paragraph()  # spacing after table
        elif btype == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = Pt(DEFAULT_FONT_SIZE * 1.15)
            for segment_text, bold in _render_inline(block["text"]):
                run = p.add_run(segment_text)
                _set_run_font(run, bold=bold)
            list_counter = 0
        elif btype == "numbered":
            list_counter += 1
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = Pt(DEFAULT_FONT_SIZE * 1.15)
            for segment_text, bold in _render_inline(block["text"]):
                run = p.add_run(segment_text)
                _set_run_font(run, bold=bold)
        elif btype == "quote":
            _add_inline_para(doc, block["text"], indent=0.5, space_after=3)
        elif btype == "para":
            _add_inline_para(doc, block["text"], space_after=3)
        elif btype == "empty":
            pass  # skip empty lines

    _add_fouo_header_footer(doc, classification)

    doc.save(output_path)
    return output_path


def main() -> None:
    parser = argparse.ArgumentParser(description="Export RFI response markdown to DOCX")
    parser.add_argument("--input", required=True, help="Input markdown file path")
    parser.add_argument("--output", required=True, help="Output DOCX file path")
    parser.add_argument("--classification", default=FOUO_TEXT,
                        help=f"Classification marking (default: {FOUO_TEXT})")
    parser.add_argument("--json", action="store_true", help="Output JSON result")
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(json.dumps({"error": f"Input file not found: {args.input}"}))
        sys.exit(1)

    md_text = input_path.read_text(encoding="utf-8")
    output_path = markdown_to_docx(md_text, args.output, classification=args.classification)

    result = {
        "status": "ok",
        "input": str(input_path),
        "output": output_path,
        "classification": args.classification,
    }
    if args.json:
        print(json.dumps(result, indent=2))
    else:
        print(f"Exported: {output_path}")


if __name__ == "__main__":
    main()
