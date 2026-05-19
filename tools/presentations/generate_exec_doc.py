"""
ICDEV(tm) Executive AI Lab Documentation Generator
Produces: tools/presentations/output/ICDEV_AI_Lab_Documentation.docx
Audience: Senior leadership / executive decision-makers
Contents: CBA, AI Lab build requirements, FORGE Academy playbook, AI GameDay playbook
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Palette ─────────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x0A, 0x16, 0x28)
GOLD  = RGBColor(0xC8, 0xA9, 0x51)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LGRAY = RGBColor(0x44, 0x44, 0x44)
MGRAY = RGBColor(0x1E, 0x3A, 0x5F)
GREEN = RGBColor(0x1A, 0x73, 0x48)
RED   = RGBColor(0x9B, 0x1C, 0x1C)


# ── XML Helpers ──────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, **edges):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, params in edges.items():
        border = OxmlElement(f"w:{edge}")
        for k, v in params.items():
            border.set(qn(f"w:{k}"), v)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _para_border_bottom(para, color="C8A951", sz="6"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Document Helpers ─────────────────────────────────────────────────────────

def _setup(doc: Document):
    sec = doc.sections[0]
    sec.left_margin   = Inches(1.0)
    sec.right_margin  = Inches(1.0)
    sec.top_margin    = Inches(0.9)
    sec.bottom_margin = Inches(0.9)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = LGRAY

    for h in ["Heading 1", "Heading 2", "Heading 3"]:
        if h in styles:
            s = styles[h]
            s.font.name = "Calibri"
            s.font.color.rgb = NAVY if h == "Heading 1" else GOLD


def _h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = NAVY
    p.runs[0].font.size = Pt(18)
    p.runs[0].font.bold = True
    _para_border_bottom(p)
    return p


def _h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = GOLD
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.bold = True
    return p


def _h3(doc, text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = MGRAY
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.bold = True
    return p


def _body(doc, text, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(6)
    return p


def _bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(4)
    return p


def _spacer(doc, lines=1):
    for _ in range(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)


def _page_break(doc):
    doc.add_page_break()


def _table(doc, headers, rows, col_widths=None, header_bg="0A1628", stripe=True):
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        _set_cell_bg(cell, header_bg)
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = GOLD
        run.font.size = Pt(11)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        if col_widths:
            cell.width = Inches(col_widths[i])

    # Data rows
    stripe_colors = ["F5F7FA", "FFFFFF"]
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            cell.text = str(val)
            if stripe:
                _set_cell_bg(cell, stripe_colors[r_idx % 2])
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(10)
            run.font.color.rgb = LGRAY
            if col_widths:
                cell.width = Inches(col_widths[c_idx])

    doc.add_paragraph()   # spacing after table
    return table


def _callout(doc, label, text, bg="0A1628"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(f"{label}  ")
    run.font.bold = True
    run.font.color.rgb = GOLD
    run.font.size = Pt(11)
    run2 = p.add_run(text)
    run2.font.size = Pt(11)
    run2.font.color.rgb = LGRAY


# ── Cover Page ───────────────────────────────────────────────────────────────

def cover(doc):
    _spacer(doc, 4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ICDEV(tm)")
    run.font.name = "Calibri"
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("AI Lab Executive Documentation")
    r2.font.name = "Calibri"
    r2.font.size = Pt(24)
    r2.font.bold = True
    r2.font.color.rgb = GOLD

    _spacer(doc)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(
        "Cost-Benefit Analysis  |  Build Requirements  |  "
        "FORGE Academy Playbook  |  AI GameDay Playbook"
    )
    r3.font.size = Pt(13)
    r3.font.color.rgb = LGRAY
    r3.font.italic = True

    _spacer(doc, 2)

    meta = [
        ("Classification", "CUI // SP-CTI"),
        ("Date",           "May 2026"),
        ("Version",        "1.0"),
        ("Prepared By",    "ICDEV(tm) AI Lab Initiative"),
        ("Audience",       "Senior Leadership / Executive Decision-Makers"),
    ]
    t = doc.add_table(rows=len(meta), cols=2)
    t.style = "Table Grid"
    for i, (k, v) in enumerate(meta):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v
        for c in t.rows[i].cells:
            _set_cell_bg(c, "0A1628")
            run = c.paragraphs[0].runs[0]
            run.font.color.rgb = GOLD if c == t.rows[i].cells[0] else WHITE
            run.font.size = Pt(11)
            run.font.bold = (c == t.rows[i].cells[0])
            c.width = Inches(3.2)

    _page_break(doc)


# ── Section 1: Executive Summary ─────────────────────────────────────────────

def sec_executive_summary(doc):
    _h1(doc, "1.  Executive Summary")

    _body(doc,
        "ICDEV(tm) (Intelligent Certified Development) is an AI-powered meta-builder that generates "
        "complete, Authority to Operate (ATO)-ready government software systems from plain English "
        "requirements. It compresses the typical 12-18 month ATO timeline to 1-3 months, automates "
        "compliance across 42 frameworks through a dual-hub crosswalk engine, and deploys 15 coordinating "
        "AI agents that eliminate the expert bottlenecks that stall most government programs. "
        "The platform is open source (Apache 2.0), IL4+ ready, and maps to every current federal AI "
        "mandate including OMB M-25-21, OMB M-26-04, and the DoD AI Strategy."
    )

    _h2(doc, "Recommendation")
    _body(doc,
        "Adopt ICDEV(tm). Stand up the AI Lab. Begin FORGE Academy Cohort 1 within 90 days. "
        "The platform is built, tested, and compliance-mapped. The only variable is the decision to move."
    )

    _h2(doc, "Headline Numbers")
    _table(doc,
        ["Metric", "Value"],
        [
            ["Year-1 Total Investment",  "$570K - $880K"],
            ["Year-1 Projected Return",  "$4M - $6M (3-program baseline)"],
            ["3-Year ROI (5 programs)",  "9x"],
            ["ATO Timeline Reduction",   "12-18 months  -->  1-3 months"],
            ["Compliance Automation",    "42 frameworks, zero duplicate assessments"],
            ["Developer Capacity",       "100 developers trained in Year 1 (4 cohorts of 25)"],
            ["Platform License Cost",    "$0 (Apache 2.0 open source)"],
        ],
        col_widths=[3.0, 4.0],
        header_bg="0A1628"
    )

    _body(doc,
        "This document provides the supporting analysis for the AI Lab stand-up decision: a Cost-Benefit "
        "Analysis, infrastructure and personnel requirements, and operating playbooks for FORGE Academy "
        "and AI GameDay at 100+ developer scale.",
        italic=True, color=LGRAY
    )
    _page_break(doc)


# ── Section 2: Background ─────────────────────────────────────────────────────

def sec_background(doc):
    _h1(doc, "2.  Background and Problem Statement")

    _h2(doc, "The Status Quo is Failing")
    _body(doc,
        "Government software programs face three structural failures that repeat across every "
        "agency and every program office:"
    )
    for item in [
        "12-18 months to ATO: Programs spend over a year and millions of dollars in compliance "
        "activity before delivering a single line of production capability. The cost is not just "
        "financial -- it is mission impact deferred.",
        "Compliance discovered late: Security and compliance requirements are identified after "
        "development, forcing costly redesigns, POAM entries, and remediation cycles that follow "
        "the system through its entire operational life.",
        "Expert bottlenecks: One ISSO, one security architect, one CI/CD engineer. When any "
        "single expert is unavailable, entire programs stall. Knowledge lives in people, not systems.",
    ]:
        _bullet(doc, item)

    _h2(doc, "The Cost of the Status Quo")
    _table(doc,
        ["Item", "Typical Cost", "Frequency"],
        [
            ["ATO pursuit labor (ISSO + SSP writer)", "$300K - $600K", "Per program"],
            ["Redesign due to late compliance discovery", "$500K - $2M",  "Per program (frequent)"],
            ["Manual framework assessment (per framework)", "$50K - $200K", "Per framework per cycle"],
            ["External training for 100 developers",  "$400K - $500K", "Annual"],
            ["Expert turnover / knowledge loss",      "$200K - $1M+",  "Per departure"],
        ],
        col_widths=[3.0, 2.2, 2.0],
    )

    _h2(doc, "Why ICDEV(tm) and the AI Lab Solves This Structurally")
    _body(doc,
        "ICDEV(tm) is not a point solution. It addresses all three failures simultaneously:"
    )
    for item in [
        "FORGE framework confines AI reasoning to orchestration; 530+ deterministic Python tools "
        "execute reliably -- reducing probabilistic failure from ~59% (0.9^5) to >95%.",
        "Compliance automation is built into every build phase (ANVIL): SSP, POAM, STIG, SBOM, "
        "and OSCAL are generated automatically -- not authored manually after the fact.",
        "15 AI agents replace individual expert dependency. Builder, Compliance, Security, "
        "Infrastructure, Requirements, and Supply Chain agents coordinate continuously.",
    ]:
        _bullet(doc, item)

    _callout(doc, "Key Insight:",
             "Every program powered by ICDEV(tm) compounds returns across the enterprise -- "
             "shared knowledge base, reusable compliance evidence, and trained internal experts.")
    _page_break(doc)


# ── Section 3: Platform Overview ─────────────────────────────────────────────

def sec_platform_overview(doc):
    _h1(doc, "3.  ICDEV(tm) Platform Overview")

    _h2(doc, "3.1  The FORGE Framework (6 Layers)")
    _table(doc,
        ["Layer", "Location", "Role"],
        [
            ["Goals",         "goals/",          "Process definitions -- what to achieve, which tools, expected outputs"],
            ["Orchestration", "(Claude Code / multi-agent)", "AI reasoning: read goal, decide tool order, handle errors"],
            ["Tools",         "icdev/tools/",    "530+ Python scripts, one job each, deterministic execution"],
            ["Args",          "args/",           "YAML/JSON behavior settings -- change behavior without editing code"],
            ["Context",       "context/",        "42 compliance catalogs, tone rules, writing samples"],
            ["Hard Prompts",  "hardprompts/",    "Reusable LLM instruction templates"],
        ],
        col_widths=[1.6, 2.2, 3.4],
    )

    _h2(doc, "3.2  The ANVIL Workflow (5 Phases)")
    for phase, desc in [
        ("Architect", "Design and acceptance criteria. System architecture defined up front."),
        ("Navigate",  "Map to existing tools and patterns. No new code without cause."),
        ("Verify",    "Write failing tests first (RED). The spec IS the test."),
        ("Integrate", "Generate implementation until tests pass (GREEN)."),
        ("Launch",    "Refactor, security scan, compliance map, merge to main."),
    ]:
        p = doc.add_paragraph()
        run = p.add_run(f"{phase}:  ")
        run.font.bold = True
        run.font.color.rgb = GOLD
        run2 = p.add_run(desc)
        run2.font.color.rgb = LGRAY
        p.paragraph_format.space_after = Pt(4)

    _h2(doc, "3.3  12 Design Canvases")
    _table(doc,
        ["Code", "Canvas Name", "Route", "Key Capability"],
        [
            ["NDC",  "Network Design",       "/network",     "Cloud topology, FedRAMP/IL/CMMC mapping, COTS cost modeling"],
            ["SDC",  "Security Design",      "/security",    "STRIDE threat modeling, MITRE ATT&CK, SSP/POAM generator"],
            ["PDC",  "Pipeline Design",      "/devops",      "Visual CI/CD builder, SLSA assessment, OWASP coverage"],
            ["BDC",  "Boundary Design",      "/boundary",    "ATO boundary, ISA lifecycle, PPS matrix auto-generation"],
            ["DDC",  "Data Design",          "/data",        "Data classification zones, PII/PHI/CUI tracking, lineage"],
            ["ODC",  "Observability Design", "/observability","Detection coverage, Sigma rules, MITRE ATT&CK detection"],
            ["IDC",  "Infrastructure Design","/infra",       "IaC resource design, 6 CSPs, Terraform generation"],
            ["AADC", "Agentic AI Design",    "/agentic-ai",  "7 solution packs, 40+ node types, AI risk register"],
            ["QDC",  "Quality Design",       "/qdc",         "Code quality gates, UQS scoring, NIST SA-11 mapping"],
            ["MDC",  "Migration Design",     "/migration",   "7R assessment, strangler fig mapping, ATO bridge"],
            ["AIMC", "AI/ML Model Design",   "/ai-ml",       "Foundation model catalog, DoD RAI compliance, IL assessment"],
            ["MSN",  "Mission Canvas",       "/mission",     "Strategic mission planning, operational design, COA support"],
        ],
        col_widths=[0.6, 1.8, 1.4, 3.4],
    )

    _h2(doc, "3.4  Unified Chat and Common Use Cases")
    _body(doc,
        "The /chat interface is a 3-pane, context-aware agent communication system that merges "
        "structured requirements intake (RICOAS) with multi-stream parallel conversation. It is NOT "
        "a generic chatbot -- every conversation carries project metadata, compliance scope, and "
        "AI governance posture."
    )
    _body(doc, "Three built-in use cases provide expert-persona, pre-loaded requirements out of the box:")
    _table(doc,
        ["Use Case", "Pre-Loaded Reqs", "Canvas Wiring", "Output"],
        [
            ["General Modernization",    "12 requirements", "Migration, Digital Twin, Supply Chain, Compliance",
             "Phased migration plan with rollback runbooks"],
            ["Year-End Budget Sprint",   "14 requirements", "CPMP, Digital Twin, Supply Chain, Kanban",
             "CSV/XLSX BOM with lead times, IGCE, contract vehicles"],
            ["Crowd-Sourced Doc Refresh","10 requirements", "Knowledge, RAG, Compliance, Audit",
             "Version-controlled docs, auto-indexed to Knowledge Graph"],
        ],
        col_widths=[2.0, 1.4, 2.4, 2.0],
    )

    _h2(doc, "3.5  Digital Program Twin")
    _body(doc,
        "The Digital Program Twin is a 6-dimension what-if simulation engine that predicts program "
        "impact before requirements are committed. It runs 10,000 Monte Carlo iterations to produce "
        "P10/P50/P80/P90 confidence bands across architecture, compliance, supply chain, schedule, "
        "cost, and risk dimensions."
    )
    _body(doc, "Three Courses of Action (COAs) are auto-generated:")
    _table(doc,
        ["COA",           "Timeline",  "Cost",    "Risk",   "Description"],
        [
            ["Speed",         "1-2 PIs",   "Lowest",  "HIGH",   "P1 requirements only. Fast delivery. Technical debt."],
            ["Balanced (*)",  "2-3 PIs",   "Medium",  "MOD",    "P1+P2. Recommended. Best scope-risk tradeoff."],
            ["Comprehensive", "3-5 PIs",   "Highest", "LOW",    "Full scope. Complete coverage. Longest timeline."],
        ],
        col_widths=[1.3, 1.1, 1.1, 0.8, 3.3],
    )
    _callout(doc, "Differentiator:",
             "No commercial vendor offers program-level Digital Twin semantics with NIST/FedRAMP verdicts. ICDEV does.")

    _h2(doc, "3.6  42-Framework Compliance Crosswalk")
    _body(doc,
        "ICDEV(tm) uses a dual-hub crosswalk model. NIST 800-53 Rev 5 serves as the US hub; "
        "ISO/IEC 27001:2022 serves as the international hub. Implementing a control once "
        "automatically satisfies 30+ downstream frameworks."
    )
    _table(doc,
        ["Framework Category", "Frameworks"],
        [
            ["Federal",        "NIST 800-53, NIST 800-171, FedRAMP Mod/High/20x, CMMC L2/L3, FIPS 199/200, CNSSI 1253"],
            ["DoD-Specific",   "DoDI 5000.87 DES, DoD CSSP, MOSA (10 U.S.C. 4401), MITRE ATLAS v5.4"],
            ["AI/ML Security", "NIST AI RMF 1.0, NIST AI 600-1, OWASP LLM Top 10, OWASP Agentic AI, OMB M-25-21, M-26-04"],
            ["Healthcare",     "HIPAA Security Rule, HITRUST CSF v11"],
            ["Financial",      "PCI DSS v4.0, SOC 2 Type II"],
            ["International",  "ISO/IEC 27001:2022, ISO/IEC 42001:2023, EU AI Act Annex III"],
            ["Architecture",   "NIST SP 800-207 Zero Trust, CISA Secure by Design, IEEE 1012 IV&V"],
        ],
        col_widths=[2.2, 5.0],
    )

    _h2(doc, "3.7  15-Agent Autonomous Ecosystem")
    _table(doc,
        ["Tier", "Agent", "Port", "Specialty"],
        [
            ["Core",    "Orchestrator",       "8443", "Task routing, DAG execution, dispatcher mode"],
            ["Core",    "Architect",          "8444", "ANVIL A/T phases, system design, M-ANVIL for models"],
            ["Domain",  "Builder",            "8445", "TDD code generation -- 6 languages"],
            ["Domain",  "Compliance",         "8446", "ATO artifacts: SSP, POAM, STIG, SBOM, OSCAL, eMASS"],
            ["Domain",  "Security",           "8447", "SAST, dependency audit, secret detection, MITRE ATLAS"],
            ["Domain",  "Infrastructure",     "8448", "Terraform, Ansible, K8s, CI/CD, 6 CSPs"],
            ["Domain",  "MBSE",               "8451", "SysML, DOORS NG, digital thread, DES compliance"],
            ["Domain",  "Modernization",      "8452", "7R assessment, migration planning, ATO bridge"],
            ["Domain",  "Requirements",       "8453", "Conversational intake, gap detection, SAFe decomposition"],
            ["Domain",  "Supply Chain",       "8454", "SBOM, CVE triage, SCRM, Section 889 compliance"],
            ["Domain",  "Simulation",         "8455", "Digital Twin, Monte Carlo, COA generation"],
            ["Domain",  "DevSecOps + ZTA",    "8457", "Pipeline security, NIST 800-207 ZTA, policy-as-code"],
            ["Domain",  "Gateway",            "8458", "Remote commands, 8-gate security chain"],
            ["Support", "Knowledge",          "8449", "Self-healing patterns, ML recommendations, failure diagnosis"],
            ["Support", "Monitor",            "8450", "Logs, metrics, alerts, health checks, SLA tracking"],
        ],
        col_widths=[1.0, 1.6, 0.8, 3.8],
    )
    _page_break(doc)


# ── Section 4: CBA ────────────────────────────────────────────────────────────

def sec_cba(doc):
    _h1(doc, "4.  Cost-Benefit Analysis (CBA)")

    _h2(doc, "4.1  Benefits (Quantified)")
    _table(doc,
        ["Benefit", "Without ICDEV(tm)", "With ICDEV(tm)", "Annual Savings"],
        [
            ["ATO timeline",
             "12-18 months x $500K/month",
             "1-3 months",
             "$2M - $4M per program"],
            ["Compliance labor",
             "2 ISSOs x 6 months per framework",
             "Automated 42-framework crosswalk",
             "$400K - $800K/yr"],
            ["Developer velocity",
             "1 feature per 2 weeks",
             "5x faster via 15 agents",
             "$1M+ productivity gain"],
            ["Developer training (100 devs)",
             "External courses ~$5K/developer",
             "Academy ~$800/developer",
             "$420K savings"],
            ["Legacy modernization",
             "Greenfield rebuild required",
             "7Rs + ATO bridge pipeline",
             "$2M - $5M avoided"],
            ["TOTAL YEAR-1 BENEFIT", "", "", "$4M - $6M+"],
        ],
        col_widths=[2.4, 2.0, 2.0, 1.8],
        header_bg="0A1628",
    )

    _h2(doc, "4.2  Costs (Year 1 Estimate)")
    _table(doc,
        ["Cost Item", "Year 1 Estimate", "Notes"],
        [
            ["Infrastructure (K8s + cloud LLM API)", "$150K - $300K", "AWS Bedrock or Azure OpenAI Gov"],
            ["ICDEV(tm) platform",                   "$0",            "Apache 2.0 open source"],
            ["Personnel (4 FTEs, blended rate)",      "$300K - $400K", "Director, Engineer, Training Lead, ISSO"],
            ["FORGE Academy (4 cohorts)",             "$80K - $120K",  "25 developers per cohort, 8 weeks"],
            ["AI GameDay (4 events)",                 "$40K - $60K",   "$10K per event, 100 participants"],
            ["TOTAL YEAR-1 COST",                    "$570K - $880K", ""],
        ],
        col_widths=[3.0, 1.8, 2.4],
    )

    _h2(doc, "4.3  Return on Investment (3-Year Projection)")
    _body(doc,
        "ROI is calculated based on ATO acceleration savings per program, compliance automation "
        "labor avoidance, and developer productivity gains. The base assumption is $2.5M benefit "
        "per program per year (conservative: ATO savings + compliance + velocity)."
    )
    _table(doc,
        ["Programs/Year", "Year 1 Net", "Year 2 Net", "Year 3 Net", "3-Year ROI"],
        [
            ["3 programs",  "+$3.0M",  "+$8.0M",  "+$14.0M", "5x"],
            ["5 programs",  "+$6.0M",  "+$16.0M", "+$28.0M", "9x"],
            ["10 programs", "+$16.0M", "+$36.0M", "+$60.0M", "18x"],
        ],
        col_widths=[1.8, 1.6, 1.6, 1.6, 1.6],
    )
    _callout(doc, "Note:",
             "Year 1 net is lower due to platform stand-up costs. Year 2+ reflects steady-state "
             "operations with full cohort pipeline producing certified developers.")
    _page_break(doc)


# ── Section 5: Lab Requirements ──────────────────────────────────────────────

def sec_lab_requirements(doc):
    _h1(doc, "5.  AI Lab Build-Out Requirements")

    _h2(doc, "5.1  Infrastructure")
    _table(doc,
        ["Component", "Specification", "Purpose"],
        [
            ["Kubernetes Cluster",    "2-node, 64 vCPU, 256 GB RAM, 10 TB NVMe", "Host 15 ICDEV agents + dashboard"],
            ["LLM API Access",        "AWS Bedrock (Claude) or Azure OpenAI Gov", "IL4+ approved inference"],
            ["Network",               "Isolated lab VLAN, CAC auth, VPN breakout", "Security isolation"],
            ["Object Storage",        "S3 GovCloud or equivalent, 10 TB",          "Artifacts, SBOM, audit logs"],
            ["CI/CD Platform",        "GitLab EE or GitHub Enterprise",            "Pipeline automation"],
            ["Artifact Management",   "JFrog Artifactory or Nexus",               "Dependency management"],
            ["SIEM",                  "Splunk or OpenSearch",                      "Audit trail forwarding"],
            ["Developer Workstations","16 vCPU, 32 GB RAM (or shared cloud VDI)", "Local sandbox environments"],
        ],
        col_widths=[2.2, 2.6, 2.4],
    )

    _h2(doc, "5.2  Personnel (4 FTEs)")
    _table(doc,
        ["Role", "Key Responsibilities", "Clearance"],
        [
            ["AI Lab Director",
             "Strategy, stakeholder liaison, program integration, executive reporting",
             "TS/SCI preferred"],
            ["Platform / MLOps Engineer",
             "ICDEV deployment, 15-agent health, Bedrock integration, GitLab CI/CD",
             "Secret minimum"],
            ["Training Lead",
             "FORGE Academy curriculum, cohort scheduling, GameDay facilitation, cert tracking",
             "Secret minimum"],
            ["ISSO",
             "Lab ATO, STIG compliance, audit trail review, eMASS entries, POA&M management",
             "Secret required"],
        ],
        col_widths=[2.0, 4.2, 1.0],
    )

    _h2(doc, "5.3  30 / 60 / 90-Day Stand-Up Timeline")
    _table(doc,
        ["Milestone", "Actions", "Owner"],
        [
            ["Day 1-30: Infrastructure",
             "K8s cluster procured and provisioned. AWS Bedrock access granted. "
             "ICDEV deployed and health-checked. GitLab EE configured. ISSO assigned.",
             "Lab Director + Platform Engineer"],
            ["Day 31-60: First Program",
             "First program team onboarded. 2 canvases (NDC + SDC) active. "
             "Use Cases seeded. RICOAS intake session completed. "
             "Digital Twin run. ATO package auto-generated.",
             "Platform Engineer + ISSO"],
            ["Day 61-90: Training + GameDay",
             "FORGE Academy Cohort 1 begins (25 developers, 8-week track). "
             "AI GameDay #1 executed (4 teams, Operation CIPHER FORGE). "
             "AAR published. Lessons learned fed back into Academy curriculum.",
             "Training Lead + Lab Director"],
        ],
        col_widths=[2.0, 4.0, 1.8],
    )
    _page_break(doc)


# ── Section 6: Academy Playbook ───────────────────────────────────────────────

def sec_academy(doc):
    _h1(doc, "6.  FORGE Academy Operating Playbook (100+ Developers)")

    _h2(doc, "6.1  Program Overview")
    _body(doc,
        "FORGE Academy is a gamified, role-based AI training platform that teaches government and "
        "defense teams how to build, deploy, and govern AI systems. Every mission teaches a pattern "
        "that directly applies to the ICDEV(tm) platform. The curriculum is auto-updated by the "
        "Genesis daemon -- when new patterns are proven in production, draft missions are generated "
        "automatically for human review and activation."
    )

    _table(doc,
        ["Attribute", "Value"],
        [
            ["Role Tracks",     "12 (DevOps, DataOps, SecOps, SWE/Architect, NetOps, SRE, ISSO, ISSM, CISO, PM, Analyst, Leadership)"],
            ["Total Missions",  "75 missions, 165 steps across 3 tiers"],
            ["Delivery Model",  "Hybrid: 2 hrs instructor-led + 8 hrs self-paced per week"],
            ["Lab Mode",        "Coding Lab (hands-on terminal) + Guided Lab (no-code for non-technical roles)"],
            ["Cohort Size",     "25 developers per cohort"],
            ["Cohort Duration", "8 weeks"],
            ["Facilitator Ratio","1 Sensei (L5 certified) per 25 learners"],
        ],
        col_widths=[2.4, 4.8],
    )

    _h2(doc, "6.2  Rank Progression")
    _table(doc,
        ["Rank", "XP Required", "Competency Level", "Real-World Capability"],
        [
            ["Recruit",   "0 XP",      "L1 Aware",      "Understands LLMs; uses AI tools safely in daily workflow"],
            ["Operative", "500 XP",    "L2 Foundation",  "Builds basic AI features (LLM API, simple RAG, single-tool agents)"],
            ["Specialist","2,000 XP",  "L3 Practitioner","Designs multi-agent systems; monitors production AI; owns features"],
            ["Architect", "5,000 XP",  "L4 Expert",      "Builds AI platforms; leads initiatives; mentors; authors ICDEV tools"],
            ["Sensei",    "10,000+ XP","L5 Sensei",      "AI strategy at org level; teaches and certifies others; CAIO-level governance"],
        ],
        col_widths=[1.2, 1.2, 1.6, 3.2],
    )

    _h2(doc, "6.3  Certification Path")
    _table(doc,
        ["Certification", "Duration", "Requirements", "Expiry"],
        [
            ["Foundation",   "8 weeks",  "Tier 1 + role Tier 2 + 75% assessment score",              "2 years"],
            ["Practitioner", "+4 weeks", "Foundation + AADC canvas score >= 80 + 1 GameDay",          "2 years"],
            ["Expert",       "Capstone", "Practitioner + Tier 3 + GameDay top-50% finish",            "3 years"],
        ],
        col_widths=[1.5, 1.2, 3.8, 0.8],
    )

    _h2(doc, "6.4  100-Developer Cohort Schedule")
    _table(doc,
        ["Cohort", "Developers", "Weeks", "Start Window", "Certification Eligible"],
        [
            ["Cohort 1", "25", "Week 1-8",   "Month 3",  "Foundation -- Month 5"],
            ["Cohort 2", "25", "Week 5-12",  "Month 4",  "Foundation -- Month 6"],
            ["Cohort 3", "25", "Week 9-16",  "Month 5",  "Foundation -- Month 7"],
            ["Cohort 4", "25", "Week 13-20", "Month 6",  "Foundation -- Month 8"],
            ["Total",    "100", "--",         "--",       "All 100 Foundation-certified by Month 8"],
        ],
        col_widths=[1.2, 1.4, 1.6, 1.6, 2.5],
    )

    _h2(doc, "6.5  DoD and Federal Alignment")
    for item in [
        "All 5 competency levels are cross-walked to DoD AI Workforce Framework KSAs (KS0001, KS0010, etc.)",
        "OMB M-25-21 readiness: L1=Aware, L2=Capable, L3=Compliant, L4=Audit-Ready, L5=Leadership",
        "Organizational readiness score (0-100) ties cohort completion to AI-first migration phases",
        "Threshold guide: <40 = Phase 0 (foundation), 40-65 = Phase 1 (quick wins), 65-80 = Phase 2 (augmentation), >80 = Phase 3+ (AI-first)",
    ]:
        _bullet(doc, item)

    _h2(doc, "6.6  Cohort Cost Breakdown")
    _table(doc,
        ["Item", "Per Cohort (25 devs)", "4 Cohorts (100 devs)"],
        [
            ["Lab infrastructure (shared cloud sandbox)", "$5,000",  "$20,000"],
            ["Facilitator time (10 hrs/wk x 8 wks)",      "$8,000",  "$32,000"],
            ["Learning management / tooling",              "$3,000",  "$12,000"],
            ["Certification assessment",                   "$4,000",  "$16,000"],
            ["TOTAL",                                      "$20,000", "$80,000"],
        ],
        col_widths=[3.2, 1.8, 2.2],
    )
    _page_break(doc)


# ── Section 7: GameDay Playbook ───────────────────────────────────────────────

def sec_gameday(doc):
    _h1(doc, "7.  AI GameDay Operating Playbook (100+ Developers)")

    _h2(doc, "7.1  Program Overview")
    _body(doc,
        "AI GameDay transforms passive tabletop exercises (TTX) into live, AI-tooled, competitive "
        "wargames. Teams compete through scripted injects using ICDEV(tm) tools, and every response "
        "is scored by an LLM judge against a published rubric. Server-side receipt verification "
        "ensures scoring integrity. A live leaderboard tracks team performance in real time, and an "
        "auto-generated After-Action Report (AAR) provides formal lessons learned documentation."
    )

    _h2(doc, "7.2  Team Structure (4 Teams x 25 Participants)")
    _table(doc,
        ["Team", "Focus", "Roles", "Scoring Weight"],
        [
            ["RED   (Threat)",      "Attack design, threat emulation",
             "Scout, Threat Analyst, Exploit Engineer, Red Orchestrator",
             "Adversarial Effectiveness: 40%"],
            ["BLUE  (Defense)",     "Defense posture, incident response",
             "SOC Analyst, Security Architect, IR Responder, Blue Orchestrator",
             "Shared scoring -- IR response quality"],
            ["GOLD  (Innovation)",  "Novel AI/ML modules, training pair generation",
             "Researcher, Builder, Evaluator, Gold Orchestrator",
             "Innovation Score: 25%"],
            ["GREEN (Compliance)",  "Governance, ethics, regulatory verdicts",
             "NIST Auditor, Risk Assessor, Policy Advisor, Green Orchestrator",
             "Compliance Score: 20%"],
        ],
        col_widths=[1.4, 1.8, 2.4, 2.2],
    )

    _h2(doc, "7.3  Operation CIPHER FORGE -- 5 Injects")
    _table(doc,
        ["Inject", "Duration", "Challenge", "Scoring Dimension"],
        [
            ["Signal Cluster",      "15 min",     "SIGINT threat assessment in simulated environment",
             "Adversarial Effectiveness"],
            ["COA Posture",         "40 min",     "Military strategy: recommend force posture vs. threat",
             "Adversarial + Compliance"],
            ["Ransomware Cascade",  "Sequential", "IR: contain, eradicate, recover from ransomware event",
             "All dimensions"],
            ["Fine-Tune Sprint",    "Sequential", "ML: generate training pairs and evaluate quality",
             "Innovation (25%) + Training pairs (15%)"],
            ["War Council Brief",   "Final",      "Executive synthesis: strategic brief from all learnings",
             "All dimensions -- final standings"],
        ],
        col_widths=[1.8, 1.2, 2.8, 1.8],
    )

    _h2(doc, "7.4  Scoring and Judging")
    _table(doc,
        ["Dimension", "Weight", "Method"],
        [
            ["Adversarial Effectiveness", "40%", "LLM judge scores Red/Blue tactical decisions vs. rubric"],
            ["Innovation Score",          "25%", "Evaluator scores Gold team ML modules and training pair quality"],
            ["Compliance Score",          "20%", "Green team NIST auditor verdicts on all team artifacts"],
            ["Training Pairs",            "15%", "Auto-scored: pairs > 0.60 quality go to ft_datasets"],
            ["Time Bonus",               "Added", "Full bonus for < 8 min turnaround; zero for > 15 min"],
        ],
        col_widths=[2.4, 1.0, 3.8],
    )

    _h2(doc, "7.5  Event Logistics (100 Participants)")
    for item in [
        "Format: 4 teams x 25 participants, live timed session (4 hours)",
        "Facilitation: 1 Lead Facilitator + 4 Team Leads (1 per team)",
        "Infrastructure: ICDEV(tm) GameDay server with LLM API access per team",
        "Pre-event: 30-minute team registration + role assignment (skill-to-role AI matching)",
        "Post-event: AAR auto-generated within 15 minutes. Kanban debrief card created.",
        "Certification: GameDay completion required for FORGE Academy Practitioner certification",
        "Frequency: Quarterly tournaments + ad hoc challenge events",
        "Audit trail: NIST AC-2, AU-2, SI-10 compliant -- append-only ttx_api_log",
    ]:
        _bullet(doc, item)

    _h2(doc, "7.6  GameDay Cost Breakdown")
    _table(doc,
        ["Item", "Per Event (100 participants)", "Annual (4 events)"],
        [
            ["ICDEV(tm) GameDay server (burst compute)",  "$2,000",  "$8,000"],
            ["LLM API credits (4 teams x 4 hours)",       "$1,500",  "$6,000"],
            ["Facilitation (Lead + 4 Team Leads)",        "$4,000",  "$16,000"],
            ["Venue / logistics (if in-person)",          "$2,500",  "$10,000"],
            ["TOTAL",                                     "$10,000", "$40,000"],
        ],
        col_widths=[3.2, 1.8, 2.2],
    )
    _page_break(doc)


# ── Section 8: Risk Analysis ──────────────────────────────────────────────────

def sec_risk(doc):
    _h1(doc, "8.  Risk Analysis")

    _table(doc,
        ["Risk", "Likelihood", "Impact", "Mitigation"],
        [
            ["LLM API outage during GameDay",
             "Low", "High",
             "Local Ollama fallback built into ICDEV (air-gap mode) -- no internet required"],
            ["Skill gap in first cohort",
             "Medium", "Medium",
             "Adaptive learning paths + guided lab mode for non-technical roles (no-code path)"],
            ["Infrastructure procurement delay",
             "Medium", "High",
             "Start on AWS Bedrock cloud while on-prem K8s is procured (parallel tracks)"],
            ["ATO for the lab itself",
             "Low", "High",
             "ICDEV auto-generates its own SSP, POAM, SBOM on first build -- ~1 week to package"],
            ["LLM model deprecation",
             "Low", "Medium",
             "Vendor-agnostic routing: swap model in .env, no code changes required"],
            ["Key person dependency (Training Lead)",
             "Medium", "High",
             "Academy Sensei certifications create internal expert bench by Month 8"],
            ["Scope creep in first program",
             "High", "Medium",
             "Digital Twin COA gates scope: Speed COA enforces P1-only boundary up front"],
            ["Budget overrun",
             "Low", "Medium",
             "Apache 2.0 base: $0 license cost; infrastructure is the primary variable; Bedrock "
             "on-demand pricing avoids reserved capacity commitment year 1"],
        ],
        col_widths=[2.6, 1.0, 0.8, 2.8],
    )
    _page_break(doc)


# ── Section 9: Recommendation ────────────────────────────────────────────────

def sec_recommendation(doc):
    _h1(doc, "9.  Recommendation")

    _body(doc,
        "Adopt ICDEV(tm). Stand up the AI Lab. Begin FORGE Academy Cohort 1 within 90 days.",
        bold=True, size=14
    )
    _spacer(doc)
    _body(doc,
        "The platform is already built, tested, and compliance-mapped to 42 frameworks. "
        "The only variable is the decision to move. Three specific actions are required:"
    )
    for item in [
        "Authorize infrastructure procurement: 2-node K8s cluster + AWS Bedrock access. "
        "Timeline: 30 days. Budget: $150K-300K.",
        "Designate 4 FTEs: AI Lab Director, Platform Engineer, Training Lead, ISSO. "
        "Timeline: 30 days. Budget: $300K-400K/year.",
        "Approve Academy and GameDay budget: $80K-120K for 4 cohorts (100 developers) "
        "and $40K-60K for 4 quarterly GameDay events.",
    ]:
        _bullet(doc, item)

    _spacer(doc)
    _body(doc,
        "Year-1 total investment: $570K - $880K.\n"
        "Year-1 projected return: $4M - $6M (3-program baseline).\n"
        "Every additional program adds $2M-4M in returns. The ROI compounds.",
        bold=False, size=12
    )

    _spacer(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("One platform.  One decision.  Infinite programs.")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = NAVY
    _page_break(doc)


# ── Appendices ────────────────────────────────────────────────────────────────

def appendices(doc):
    _h1(doc, "Appendix A: ICDEV(tm) Compliance Framework Index (42 Frameworks)")
    frameworks = [
        ("Federal",       ["NIST SP 800-53 Rev 5 (all 18 control families)", "NIST SP 800-171 Rev 2",
                           "FedRAMP Moderate", "FedRAMP High", "FedRAMP 20x", "CMMC Level 2",
                           "CMMC Level 3", "FIPS 199/200", "CNSSI 1253 (IL6 SECRET overlay)"]),
        ("DoD",           ["DoDI 5000.87 DES (Digital Engineering Strategy)", "DoD CSSP (DI 8530.01)",
                           "DoD MOSA (10 U.S.C. 4401)", "MITRE ATLAS v5.4.0 (AI adversarial)"]),
        ("AI/ML/Gov",     ["NIST AI RMF 1.0", "NIST AI 600-1 (GenAI profile)", "OWASP LLM Top 10",
                           "OWASP Agentic AI Security", "OMB M-25-21 (High-Impact AI transparency)",
                           "OMB M-26-04 (Unbiased AI)", "GAO-21-519SP (AI Accountability)",
                           "ISO/IEC 42001:2023 (AI Management Systems)", "EU AI Act Annex III"]),
        ("Healthcare",    ["HIPAA Security Rule", "HITRUST CSF v11", "HL7/FHIR compliance"]),
        ("Financial",     ["PCI DSS v4.0", "SOC 2 Type II"]),
        ("Law Enforce.",  ["CJIS Security Policy"]),
        ("International", ["ISO/IEC 27001:2022 (US-International bridge)"]),
        ("Architecture",  ["NIST SP 800-207 Zero Trust Architecture", "CISA Secure by Design (14 domains)",
                           "IEEE 1012 IV&V (Verification and Validation)"]),
        ("Governance",    ["IV&V (DoD-specific)", "Section 889 (supply chain)",
                           "FAR/DFARS (procurement compliance)", "NDAA compliance"]),
    ]
    for category, fws in frameworks:
        _body(doc, category + ":", bold=True, color=GOLD)
        for fw in fws:
            _bullet(doc, fw, level=0)
    _page_break(doc)

    _h1(doc, "Appendix B: Digital Program Twin -- COA Comparison Matrix")
    _table(doc,
        ["Dimension",            "Speed COA",     "Balanced COA (*)", "Comprehensive COA"],
        [
            ["Timeline",         "1-2 PIs",       "2-3 PIs",          "3-5 PIs"],
            ["Cost",             "Lowest",        "Medium",           "Highest"],
            ["Risk",             "HIGH",          "MODERATE",         "LOW"],
            ["Scope",            "P1 only",       "P1 + P2",          "P1 + P2 + P3"],
            ["ATO Boundary",     "Best (lowest)", "Average",          "May trigger re-ATO"],
            ["Technical Debt",   "High",          "Moderate",         "Low"],
            ["Recommended?",     "Fast feedback", "YES (default)",    "Full coverage"],
            ["Architecture",     "Minimal viable","Balanced",         "Complete"],
            ["Supply Chain",     "Fewest vendors","Moderate vendors", "All vendors"],
            ["Compliance",       "P1 controls",   "P1+P2 controls",   "All controls"],
        ],
        col_widths=[1.8, 1.8, 2.0, 1.8],
    )
    _page_break(doc)

    _h1(doc, "Appendix C: ICDEV Chat -- Common Use Cases Reference")
    _body(doc, "Built-in use cases (seeded in args/use_cases.yaml):")
    for name, badge, reqs, wiring, output in [
        ("General Modernization",    "AI Boost",  "12",
         "Migration Canvas, Digital Twin, Supply Chain, Compliance",
         "Phased migration plan with 7Rs classification and rollback runbooks"),
        ("Year-End Budget Sprint",   "RICOAS",    "14",
         "CPMP, Digital Twin, Supply Chain, Kanban",
         "CSV/XLSX BOM with lead times, vendor quotes, IGCE, contract vehicles"),
        ("Crowd-Sourced Doc Refresh","RAG + KG",  "10",
         "Knowledge Graph, RAG, Compliance, Audit",
         "Version-controlled docs, auto-indexed to Knowledge Graph within 24 hours"),
    ]:
        _body(doc, name, bold=True, color=GOLD)
        _bullet(doc, f"Badge: {badge}  |  Pre-Loaded Requirements: {reqs}")
        _bullet(doc, f"Canvas Wiring: {wiring}")
        _bullet(doc, f"Output: {output}")
    _body(doc, "Adding new use cases: Edit args/use_cases.yaml only. No Python required.", italic=True)
    _page_break(doc)

    _h1(doc, "Appendix D: Sample FORGE Academy Mission -- M01 LLM Fundamentals")
    _body(doc, "Mission: M01 -- What is an LLM? (Tier 1, all roles)")
    _table(doc,
        ["Step", "Activity", "Mode", "XP Awarded"],
        [
            ["Step 1", "Watch: 5-minute video on transformer architecture",       "Guided", "10 XP"],
            ["Step 2", "Read: ICDEV LLM provider abstraction overview",           "Guided", "10 XP"],
            ["Step 3", "Guided Lab: Send first prompt to Claude via ICDEV router","Guided", "20 XP"],
            ["Step 4", "Challenge: Identify which ICDEV tool uses which LLM",     "Guided", "30 XP"],
            ["Step 5", "Assessment: 5-question quiz, 80% pass rate required",    "Guided", "30 XP"],
        ],
        col_widths=[0.8, 3.8, 1.0, 1.2],
    )
    _body(doc, "Completion unlocks: Operative rank track. DoD KSA: KS0001 (AI/ML Fundamentals).", italic=True)
    _page_break(doc)

    _h1(doc, "Appendix E: Sample AI GameDay After-Action Report (AAR Template)")
    _body(doc, "AUTO-GENERATED by ICDEV GameDay Engine post-session. Sections:")
    for section in [
        "Executive Summary: Session ID, scenario, date, duration, overall winner",
        "Team Performance: Points breakdown per team (receipt_pts, judge_pts, time_bonus_pts, total_pts)",
        "Inject Breakdown: Per-inject scoring, fastest team, best response excerpts",
        "Lessons Learned: Top 3 defensive gaps identified, top 3 innovative approaches",
        "Certification Actions: List of players who completed GameDay requirement for Practitioner cert",
        "Fine-Tune Dataset: Summary of Gold team training pairs approved for model improvement",
        "Recommended Follow-Up: Suggested missions in FORGE Academy based on gaps surfaced",
        "NIST Audit Evidence: Append-only ttx_api_log hash summary for AC-2, AU-2, SI-10 evidence",
    ]:
        _bullet(doc, section)
    _page_break(doc)

    _h1(doc, "Appendix F: Glossary")
    glossary = [
        ("ANVIL",    "5-phase TDD build workflow: Architect, Navigate, Verify, Integrate, Launch"),
        ("ATO",      "Authority to Operate -- federal approval to run a system in production"),
        ("COA",      "Course of Action -- one of three implementation strategies (Speed/Balanced/Comprehensive)"),
        ("CMMC",     "Cybersecurity Maturity Model Certification"),
        ("CUI",      "Controlled Unclassified Information"),
        ("Digital Twin", "6-dimension what-if simulation engine for program planning"),
        ("FORGE",    "6-layer framework: Goals, Orchestration, Tools, Args, Context, Hard Prompts"),
        ("IQE",      "ICDEV Query Engine -- SQL-injection-safe natural language query interface"),
        ("ISSO",     "Information System Security Officer"),
        ("NICE",     "National Initiative for Cybersecurity Education framework"),
        ("POAM",     "Plan of Action and Milestones -- compliance remediation tracker"),
        ("RAG",      "Retrieval-Augmented Generation -- knowledge-grounded LLM responses"),
        ("RICOAS",   "Requirements Intake, COA and Approval System (Phases 1-4)"),
        ("SBOM",     "Software Bill of Materials"),
        ("SCRM",     "Supply Chain Risk Management"),
        ("SSP",      "System Security Plan"),
        ("STIG",     "Security Technical Implementation Guide"),
        ("TTX",      "Tabletop Exercise -- structured team response simulation"),
        ("ZTA",      "Zero Trust Architecture -- NIST SP 800-207 framework"),
    ]
    _table(doc,
        ["Term", "Definition"],
        [[k, v] for k, v in glossary],
        col_widths=[1.5, 5.7],
    )


# ── Main ─────────────────────────────────────────────────────────────────────

def main():
    out_dir = Path(__file__).parent / "output"
    out_dir.mkdir(exist_ok=True)

    doc = Document()
    _setup(doc)

    cover(doc)
    sec_executive_summary(doc)
    sec_background(doc)
    sec_platform_overview(doc)
    sec_cba(doc)
    sec_lab_requirements(doc)
    sec_academy(doc)
    sec_gameday(doc)
    sec_risk(doc)
    sec_recommendation(doc)
    appendices(doc)

    out_path = out_dir / "ICDEV_AI_Lab_Documentation.docx"
    doc.save(str(out_path))
    print(f"[OK] Saved: {out_path}")


if __name__ == "__main__":
    main()
