#!/usr/bin/env python3
from __future__ import annotations

from tools.logging.icdev_logger import get_logger
# CUI // SP-CTI
"""Dashboard API: WriteGuard — Content Quality Analysis, Export & History.

Exposes POST /api/writeguard/analyze for text or file upload analysis.
Exposes POST /api/writeguard/export for exporting analysis in txt/md/json/docx/pdf.
Exposes GET  /api/writeguard/history for analysis history.
Exposes GET  /api/writeguard/history/<id> for single result with findings.
Runs 5 deterministic quality dimensions via tools/pulse/writeguard.py.
Persists results to wg_analysis_results / wg_analysis_findings tables.
"""

import hashlib
import io
import json as json_mod
import sys
import time
import uuid
from datetime import datetime, timezone
from pathlib import Path

from flask import Blueprint, jsonify, request, send_file

BASE_DIR = Path(__file__).resolve().parent.parent.parent.parent
if str(BASE_DIR) not in sys.path:
    sys.path.insert(0, str(BASE_DIR))

from tools.db.storage import get_connection, sql_placeholder  # noqa: E402
from tools.dashboard.auth import require_role  # noqa: E402

logger = get_logger(__name__)

writeguard_api = Blueprint("writeguard_api", __name__, url_prefix="/api/writeguard")

# nav-intel-06: WriteGuard content-management mutations (glossary / taxonomy /
# style-profile create-update-delete) write shared, cross-user configuration
# that steers every subsequent analysis, so they are restricted to an
# admin/pm content-manager role (401 anon / 403 wrong role). Reads (GET) and the
# interactive per-user analysis tools (analyze / rewrite / export / classify /
# style-match / diff / standards / stride / parallel / precedents / word-limits)
# stay open to any authenticated user — they take input text per request and do
# not mutate shared configuration. The /analyze best-effort history write is
# per-request activity logging, not shared config, so it does not gate the tool.
_WG_CRUD_ROLES = ("admin", "pm")

# ---------------------------------------------------------------------------
# Text extraction helpers
# ---------------------------------------------------------------------------


def _extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract plain text from a .docx file using python-docx."""
    try:
        import docx  # python-docx

        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs]
        return "\n".join(paragraphs)
    except ImportError:
        logger.warning("python-docx not installed — cannot extract .docx content")
        return ""
    except Exception as exc:
        logger.error("Failed to extract .docx: %s", exc)
        return ""


def _get_text_from_request() -> tuple[str, str | None]:
    """Extract text from request: JSON body or file upload.

    Returns:
        (text, error_message) — error_message is None on success.
    """
    # ── File upload path ────────────────────────────────────────────────────
    if "file" in request.files:
        f = request.files["file"]
        filename = (f.filename or "").lower()
        raw = f.read()

        if filename.endswith(".docx"):
            text = _extract_text_from_docx(raw)
            if not text.strip():
                return "", "Could not extract text from .docx file"
        elif filename.endswith((".txt", ".md")):
            try:
                text = raw.decode("utf-8", errors="replace")
            except Exception:
                return "", "Could not decode file as UTF-8"
        else:
            return "", "Unsupported file type. Accepted: .txt, .md, .docx"

        return text, None

    # ── JSON body path ───────────────────────────────────────────────────────
    body = request.get_json(silent=True) or {}
    text = body.get("text", "").strip()
    if not text:
        return "", "No text provided. Send JSON {\"text\": \"...\"} or upload a file."

    return text, None


# ---------------------------------------------------------------------------
# Persistence helpers
# ---------------------------------------------------------------------------

# Map WriteGuard dimension keys to DB category values accepted by the CHECK
# constraint on wg_analysis_findings.category.
_DIMENSION_TO_DB_CATEGORY = {
    "grammar": "grammar",
    "readability": "readability",
    "tone": "tone",
    "passive_voice": "passive_voice",
    "overused_words": "style",
    "cliches": "style",
    "consistency": "coherence",
    "plagiarism": "plagiarism",
    "ai_detection": "ai_content",
    "portion_marking": "cui_marking",
}


def _persist_result(
    text: str,
    result: dict,
    dimensions: dict,
    all_findings: list,
    elapsed_ms: int,
) -> str | None:
    """Save analysis result + findings to wg_analysis_results / wg_analysis_findings.

    Returns the result id on success, None on failure.
    """
    result_id = str(uuid.uuid4())
    input_hash = hashlib.sha256(text.encode("utf-8")).hexdigest()
    details = result.get("details", {})

    # Extract dimension-level scores for DB columns
    readability = details.get("readability", {})
    grammar = details.get("grammar", {})
    passive = details.get("passive_voice", {})
    tone = details.get("tone", {})
    coherence = details.get("consistency", {})
    plagiarism = details.get("plagiarism", {})
    ai_det = details.get("ai_detection", {})

    try:
        conn = get_connection()
        conn.execute(
            """INSERT INTO wg_analysis_results (
                   id, input_text_hash, input_length, mode,
                   readability_flesch, readability_gunning_fog, readability_grade_level,
                   grammar_error_count, passive_voice_pct, avg_sentence_length,
                   tone_profile, coherence_score, plagiarism_max_similarity,
                   ai_content_score, overall_quality_score,
                   findings_count, analysis_duration_ms, classification
               ) VALUES (
                   %s, %s, %s, %s,
                   %s, %s, %s,
                   %s, %s, %s,
                   %s, %s, %s,
                   %s, %s,
                   %s, %s, %s
               )""",
            [
                result_id, input_hash, len(text), "inline",
                readability.get("score", 0.0), 0.0, 0.0,
                grammar.get("finding_count", len(grammar.get("findings", []))),
                passive.get("score", 0.0),
                0.0,
                json_mod.dumps(tone.get("score", {})) if isinstance(tone.get("score"), dict) else str(tone.get("score", 0)),
                coherence.get("score", 0.0),
                plagiarism.get("score", 0.0),
                ai_det.get("score", 0.0),
                result.get("overall_score", 0.0),
                len(all_findings), elapsed_ms, "CUI",
            ],
        )

        # Persist individual findings
        for finding in all_findings:
            dim_key = finding.get("dimension_key", "")
            db_category = _DIMENSION_TO_DB_CATEGORY.get(dim_key, "custom_rule")
            severity = finding.get("severity", "info")
            # Clamp severity to valid CHECK values
            if severity not in ("critical", "high", "medium", "low", "info"):
                severity = "info"
            message = finding.get("message", "")
            suggestion = finding.get("suggestion", "")
            context_text = finding.get("context", "")

            conn.execute(
                """INSERT INTO wg_analysis_findings (
                       result_id, category, severity, message,
                       suggestion, context
                   ) VALUES (%s, %s, %s, %s, %s, %s)""",
                [result_id, db_category, severity, message, suggestion, context_text],
            )

        conn.commit()
        return result_id
    except Exception as exc:
        logger.error("Failed to persist WriteGuard result: %s", exc)
        return None


# ---------------------------------------------------------------------------
# Analyze endpoint
# ---------------------------------------------------------------------------


@writeguard_api.route("/modes", methods=["GET"])
def get_modes():
    """GET /api/writeguard/modes — Return available content-type modes.

    Query param ``opp_id`` merges opportunity-specific style guides
    (scope='project' in wg_style_guides) into the mode list.
    """
    try:
        from tools.writing.content_modes import list_modes
        modes = list_modes()
    except Exception as exc:
        logger.error("Failed to load content modes: %s", exc)
        modes = [{"key": "default", "label": "Default", "description": "Standard WriteGuard rules."}]

    opp_id = request.args.get("opp_id", "").strip()
    if opp_id:
        try:
            conn = get_connection()
            rows = conn.execute(
                "SELECT guide_name, id FROM wg_style_guides WHERE scope='project' AND scope_id=%s AND is_active=1 ORDER BY guide_name",
                (opp_id,),
            ).fetchall()
            for r in rows:
                modes.append({
                    "key": r["id"],
                    "label": r["guide_name"],
                    "description": "Project-specific style guide",
                    "source": "project",
                    "opp_id": opp_id,
                })
            conn.close()
        except Exception as exc:
            logger.warning("Failed to load opportunity modes: %s", exc)

    return jsonify({"modes": modes})


# ---------------------------------------------------------------------------
# Glossary CRUD (D-WG-14)
# ---------------------------------------------------------------------------


@writeguard_api.route("/glossary", methods=["GET"])
def list_glossary():
    """GET /api/writeguard/glossary — List glossary entries.

    Query params: opp_id (scopes to project), term_type, domain, q (search).
    """
    conn = get_connection()
    ph = sql_placeholder(conn)
    try:
        opp_id = request.args.get("opp_id", "").strip()
        term_type = request.args.get("term_type", "").strip()
        domain = request.args.get("domain", "").strip()
        q = request.args.get("q", "").strip()

        clauses = ["is_active = 1"]
        params = []
        if opp_id:
            clauses.append(f"scope = 'project' AND scope_id = {ph}")
            params.append(opp_id)
        else:
            clauses.append("scope = 'platform'")
        if term_type:
            clauses.append(f"term_type = {ph}")
            params.append(term_type)
        if domain:
            clauses.append(f"domain = {ph}")
            params.append(domain)
        if q:
            clauses.append(f"(term LIKE {ph} OR definition LIKE {ph} OR replacement LIKE {ph})")
            params.extend([f"%{q}%", f"%{q}%", f"%{q}%"])

        query = "SELECT * FROM wg_glossary WHERE " + " AND ".join(clauses) + " ORDER BY term_type, term"
        rows = conn.execute(query, params).fetchall()
        return jsonify({"entries": [dict(r) for r in rows]})
    finally:
        conn.close()


@writeguard_api.route("/glossary", methods=["POST"])
@require_role(*_WG_CRUD_ROLES)
def create_glossary():
    """POST /api/writeguard/glossary — Create a glossary entry."""
    data = request.get_json(silent=True) or {}
    required = ["term", "term_type"]
    missing = [f for f in required if not data.get(f)]
    if missing:
        return jsonify({"error": f"Missing fields: {', '.join(missing)}"}), 400

    conn = get_connection()
    try:
        opp_id = data.get("opp_id", "").strip()
        entry_id = data.get("id") or str(uuid.uuid4())
        conn.execute(
            """INSERT INTO wg_glossary
               (id, term, term_type, expansion, replacement, definition, domain,
                scope, scope_id, case_sensitive, enforcement, source, approved_by,
                created_by, created_at, is_active, classification)
               VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)""",
            (
                entry_id,
                data["term"],
                data["term_type"],
                data.get("expansion", ""),
                data.get("replacement", ""),
                data.get("definition", ""),
                data.get("domain", "general"),
                "project" if opp_id else "platform",
                opp_id,
                1 if data.get("case_sensitive", True) else 0,
                data.get("enforcement", "suggest"),
                data.get("source", "admin"),
                data.get("approved_by", ""),
                data.get("created_by", "dashboard_api"),
                datetime.now(timezone.utc).isoformat(),
                1,
                data.get("classification", "CUI"),
            ),
        )
        conn.commit()
        return jsonify({"id": entry_id, "status": "created"})
    finally:
        conn.close()


@writeguard_api.route("/glossary/<entry_id>", methods=["GET"])
def get_glossary_entry(entry_id):
    conn = get_connection()
    try:
        row = conn.execute("SELECT * FROM wg_glossary WHERE id = %s", (entry_id,)).fetchone()
        if not row:
            return jsonify({"error": "Not found"}), 404
        return jsonify(dict(row))
    finally:
        conn.close()


@writeguard_api.route("/glossary/<entry_id>", methods=["PUT"])
@require_role(*_WG_CRUD_ROLES)
def update_glossary_entry(entry_id):
    data = request.get_json(silent=True) or {}
    conn = get_connection()
    try:
        row = conn.execute("SELECT * FROM wg_glossary WHERE id = %s", (entry_id,)).fetchone()
        if not row:
            return jsonify({"error": "Not found"}), 404
        updatable = {
            "term": data.get("term", row["term"]),
            "term_type": data.get("term_type", row["term_type"]),
            "expansion": data.get("expansion", row["expansion"]),
            "replacement": data.get("replacement", row["replacement"]),
            "definition": data.get("definition", row["definition"]),
            "domain": data.get("domain", row["domain"]),
            "case_sensitive": 1 if data.get("case_sensitive", row["case_sensitive"]) else 0,
            "enforcement": data.get("enforcement", row["enforcement"]),
            "is_active": 1 if data.get("is_active", row["is_active"]) else 0,
        }
        conn.execute(
            """UPDATE wg_glossary SET
               term=%s, term_type=%s, expansion=%s, replacement=%s, definition=%s,
               domain=%s, case_sensitive=%s, enforcement=%s, is_active=%s
               WHERE id=%s""",
            tuple(updatable[k] for k in [
                "term", "term_type", "expansion", "replacement", "definition",
                "domain", "case_sensitive", "enforcement", "is_active"
            ]) + (entry_id,),
        )
        conn.commit()
        return jsonify({"id": entry_id, "status": "updated"})
    finally:
        conn.close()


@writeguard_api.route("/glossary/<entry_id>", methods=["DELETE"])
@require_role(*_WG_CRUD_ROLES)
def delete_glossary_entry(entry_id):
    conn = get_connection()
    try:
        conn.execute("UPDATE wg_glossary SET is_active = 0 WHERE id = %s", (entry_id,))
        conn.commit()
        return jsonify({"id": entry_id, "status": "deleted"})
    finally:
        conn.close()


# ---------------------------------------------------------------------------
# Taxonomy CRUD (D-WG-14)
# ---------------------------------------------------------------------------


@writeguard_api.route("/taxonomy", methods=["GET"])
def list_taxonomy():
    """GET /api/writeguard/taxonomy — List proposal taxonomy nodes.

    Query param: opp_id (required).
    """
    opp_id = request.args.get("opp_id", "").strip()
    if not opp_id:
        return jsonify({"error": "opp_id is required"}), 400
    conn = get_connection()
    try:
        rows = conn.execute(
            "SELECT * FROM proposal_taxonomy WHERE opportunity_id = %s AND is_active = 1 ORDER BY label",
            (opp_id,),
        ).fetchall()
        return jsonify({"entries": [dict(r) for r in rows]})
    finally:
        conn.close()


@writeguard_api.route("/taxonomy", methods=["POST"])
@require_role(*_WG_CRUD_ROLES)
def create_taxonomy():
    data = request.get_json(silent=True) or {}
    opp_id = data.get("opportunity_id", "").strip()
    if not opp_id:
        return jsonify({"error": "opportunity_id is required"}), 400
    if not data.get("label"):
        return jsonify({"error": "label is required"}), 400

    conn = get_connection()
    try:
        entry_id = data.get("id") or str(uuid.uuid4())
        conn.execute(
            """INSERT INTO proposal_taxonomy
               (id, opportunity_id, parent_id, label, description, weight, is_active, created_at, classification)
               VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)""",
            (
                entry_id,
                opp_id,
                data.get("parent_id") or None,
                data["label"],
                data.get("description", ""),
                float(data.get("weight", 1.0)),
                1,
                datetime.now(timezone.utc).isoformat(),
                data.get("classification", "CUI"),
            ),
        )
        conn.commit()
        return jsonify({"id": entry_id, "status": "created"})
    finally:
        conn.close()


@writeguard_api.route("/taxonomy/<entry_id>", methods=["GET"])
def get_taxonomy_entry(entry_id):
    conn = get_connection()
    try:
        row = conn.execute("SELECT * FROM proposal_taxonomy WHERE id = %s", (entry_id,)).fetchone()
        if not row:
            return jsonify({"error": "Not found"}), 404
        return jsonify(dict(row))
    finally:
        conn.close()


@writeguard_api.route("/taxonomy/<entry_id>", methods=["PUT"])
@require_role(*_WG_CRUD_ROLES)
def update_taxonomy_entry(entry_id):
    data = request.get_json(silent=True) or {}
    conn = get_connection()
    try:
        row = conn.execute("SELECT * FROM proposal_taxonomy WHERE id = %s", (entry_id,)).fetchone()
        if not row:
            return jsonify({"error": "Not found"}), 404
        conn.execute(
            """UPDATE proposal_taxonomy SET
               label=%s, description=%s, weight=%s, parent_id=%s, is_active=%s
               WHERE id=%s""",
            (
                data.get("label", row["label"]),
                data.get("description", row["description"]),
                float(data.get("weight", row["weight"])),
                data.get("parent_id", row["parent_id"]),
                1 if data.get("is_active", row["is_active"]) else 0,
                entry_id,
            ),
        )
        conn.commit()
        return jsonify({"id": entry_id, "status": "updated"})
    finally:
        conn.close()


@writeguard_api.route("/taxonomy/<entry_id>", methods=["DELETE"])
@require_role(*_WG_CRUD_ROLES)
def delete_taxonomy_entry(entry_id):
    conn = get_connection()
    try:
        conn.execute("UPDATE proposal_taxonomy SET is_active = 0 WHERE id = %s", (entry_id,))
        conn.commit()
        return jsonify({"id": entry_id, "status": "deleted"})
    finally:
        conn.close()


@writeguard_api.route("/analyze", methods=["POST"])
def analyze():
    """POST /api/writeguard/analyze — Run all 5 WriteGuard quality checks.

    Accepts:
      - multipart/form-data with a 'file' field (.txt, .md, .docx)
      - application/json with {"text": "...", "mode": "..."}

    Optional ``mode`` param selects a content-type rule override:
      default, blameless_postmortem, policy_language, adr,
      bluf_exec, runbook, rca

    Returns JSON with overall_score, passed flag, per-dimension scores,
    and aggregated findings list.
    """
    text, err = _get_text_from_request()
    if err:
        return jsonify({"error": err}), 400

    if len(text.strip()) < 10:
        return jsonify({"error": "Text is too short to analyze (minimum 10 characters)"}), 400

    # Extract content mode (JSON body or multipart form field)
    if request.content_type and "multipart" in request.content_type:
        content_mode = request.form.get("mode", "default").strip() or "default"
    else:
        body = request.get_json(silent=True) or {}
        content_mode = body.get("mode", "default").strip() or "default"

    t0 = time.monotonic()
    try:
        from tools.pulse.writeguard import run_full_quality_check

        result = run_full_quality_check(text)
    except Exception as exc:
        logger.error("WriteGuard analysis failed: %s", exc)
        return jsonify({"error": f"Analysis error: {exc}"}), 500
    elapsed_ms = int((time.monotonic() - t0) * 1000)

    # Apply content-mode overrides (no-op for "default")
    if content_mode != "default":
        try:
            from tools.writing.content_modes import apply_mode_overrides
            result = apply_mode_overrides(result, content_mode, text)
        except Exception as exc:
            logger.warning("Content mode override failed (%s): %s", content_mode, exc)

    # Flatten all findings into a single list with dimension label
    all_findings = []
    details = result.get("details", {})
    dimension_labels = {
        "grammar": "Grammar",
        "readability": "Readability",
        "tone": "Tone",
        "passive_voice": "Passive Voice",
        "overused_words": "Word Choice",
        "cliches": "Cliches",
        "consistency": "Consistency",
        "glossary": "Glossary",
        "style_guide": "Style Guide",
        "pacing": "Pacing",
        "structure": "Structure",
        "evidence": "Evidence",
        "numeric_integrity": "Numeric Integrity",
        "bias": "Bias",
        "standards": "Standards",
        "alternatives": "Alternatives",
        "reproducibility": "Reproducibility",
        "neutrality": "Neutrality",
        "plagiarism": "Plagiarism",
        "ai_detection": "AI Detection",
        "portion_marking": "Portion Marking",
    }
    # Add content_mode dimension if present (injected by apply_mode_overrides)
    if "content_mode" in details:
        mode_label = result.get("content_mode_label", content_mode)
        dimension_labels["content_mode"] = mode_label

    for dim_key, label in dimension_labels.items():
        dim = details.get(dim_key, {})
        for finding in dim.get("findings", []):
            if isinstance(finding, dict):
                finding_copy = dict(finding)
            else:
                # String finding from analysis_engine
                finding_copy = {"message": str(finding)}
            finding_copy.setdefault("dimension", label)
            finding_copy.setdefault("severity", "medium")
            finding_copy["dimension_key"] = dim_key
            all_findings.append(finding_copy)

    # Build per-dimension score summary
    dimensions = {}
    for dim_key, label in dimension_labels.items():
        dim = details.get(dim_key, {})
        dimensions[dim_key] = {
            "label": label,
            "score": dim.get("score", 0),
            "status": dim.get("status", "unavailable"),
            "finding_count": len(dim.get("findings", [])),
            "suppressed": dim.get("suppressed", False),
        }

    # Portion marking detail — passed through verbatim for the UI
    pm = details.get("portion_marking", {})
    portion_marking_detail = {
        "has_marks": pm.get("has_marks", False),
        "all_marked": pm.get("all_marked", True),
        "highest_level": pm.get("highest_level"),
        "derived_banner": pm.get("derived_banner", "UNCLASSIFIED"),
        "designation_block": pm.get("designation_block", ""),
        "violations": pm.get("violations", []),
        "paragraphs": pm.get("paragraphs", []),
        "marked_count": pm.get("marked_count", 0),
        "unmarked_count": pm.get("unmarked_count", 0),
        "total_paragraphs": pm.get("total_paragraphs", 0),
    }

    # Persist to DB (best-effort — does not block response on failure)
    result_id = _persist_result(text, result, dimensions, all_findings, elapsed_ms)

    return jsonify(
        {
            "id": result_id,
            "overall_score": result.get("overall_score", 0),
            "passed": result.get("passed", False),
            "composite_status": "PASSED" if result.get("passed") else "NEEDS REVIEW",
            "content_mode": result.get("content_mode", "default"),
            "content_mode_label": result.get("content_mode_label", "Default"),
            "dimensions": dimensions,
            "findings": all_findings,
            "recommendations": result.get("recommendations", []),
            "portion_marking": portion_marking_detail,
            "word_count": len(text.split()),
            "char_count": len(text),
        }
    )


# ---------------------------------------------------------------------------
# History endpoints
# ---------------------------------------------------------------------------


@writeguard_api.route("/history", methods=["GET"])
def history():
    """GET /api/writeguard/history — Return recent analysis results.

    Query params:
        limit     — max rows (default 50, capped at 200)
        min_score — filter: overall_quality_score >= value
        max_score — filter: overall_quality_score <= value
    """
    try:
        limit = min(int(request.args.get("limit", 50)), 200)
    except (ValueError, TypeError):
        limit = 50

    min_score = request.args.get("min_score")
    max_score = request.args.get("max_score")

    clauses = []
    params: list = []

    if min_score is not None:
        try:
            clauses.append("overall_quality_score >= %s")
            params.append(float(min_score))
        except ValueError:
            pass

    if max_score is not None:
        try:
            clauses.append("overall_quality_score <= %s")
            params.append(float(max_score))
        except ValueError:
            pass

    where = ""
    if clauses:
        where = "WHERE " + " AND ".join(clauses)

    query = f"""
        SELECT id, input_text_hash, overall_quality_score,
               findings_count, analysis_duration_ms,
               readability_flesch, grammar_error_count,
               passive_voice_pct, coherence_score,
               plagiarism_max_similarity, ai_content_score,
               tone_profile, mode, created_at
        FROM wg_analysis_results
        {where}
        ORDER BY created_at DESC
        LIMIT %s
    """
    params.append(limit)

    try:
        conn = get_connection()
        rows = conn.execute(query, params).fetchall()
    except Exception as exc:
        logger.error("Failed to fetch WriteGuard history: %s", exc)
        return jsonify({"error": f"Database error: {exc}"}), 500

    items = []
    for row in rows:
        # Build dimension_scores dict from individual columns
        dimension_scores = {
            "readability": row["readability_flesch"] if row["readability_flesch"] else 0,
            "grammar_errors": row["grammar_error_count"] if row["grammar_error_count"] else 0,
            "passive_voice": row["passive_voice_pct"] if row["passive_voice_pct"] else 0,
            "coherence": row["coherence_score"] if row["coherence_score"] else 0,
            "plagiarism": row["plagiarism_max_similarity"] if row["plagiarism_max_similarity"] else 0,
            "ai_content": row["ai_content_score"] if row["ai_content_score"] else 0,
        }
        overall = row["overall_quality_score"] if row["overall_quality_score"] else 0
        items.append({
            "id": row["id"],
            "input_hash": row["input_text_hash"],
            "overall_score": overall,
            "passed": overall >= 70,
            "dimension_scores": dimension_scores,
            "findings_count": row["findings_count"] if row["findings_count"] else 0,
            "analysis_duration_ms": row["analysis_duration_ms"] if row["analysis_duration_ms"] else 0,
            "mode": row["mode"],
            "created_at": row["created_at"],
        })

    return jsonify({"results": items, "count": len(items)})


@writeguard_api.route("/history/<result_id>", methods=["GET"])
def history_detail(result_id: str):
    """GET /api/writeguard/history/<id> — Return full result with findings."""
    try:
        conn = get_connection()
        row = conn.execute(
            "SELECT * FROM wg_analysis_results WHERE id = %s", [result_id]
        ).fetchone()
    except Exception as exc:
        logger.error("Failed to fetch WriteGuard result %s: %s", result_id, exc)
        return jsonify({"error": f"Database error: {exc}"}), 500

    if not row:
        return jsonify({"error": "Result not found"}), 404

    # Fetch associated findings
    try:
        finding_rows = conn.execute(
            """SELECT id, category, severity, message, suggestion, context,
                      line_number, char_offset, char_length, rule_id,
                      confidence, auto_fixable, created_at
               FROM wg_analysis_findings
               WHERE result_id = %s
               ORDER BY id""",
            [result_id],
        ).fetchall()
    except Exception as exc:
        logger.error("Failed to fetch findings for %s: %s", result_id, exc)
        finding_rows = []

    findings = []
    for fr in finding_rows:
        findings.append({
            "id": fr["id"],
            "category": fr["category"],
            "severity": fr["severity"],
            "message": fr["message"],
            "suggestion": fr["suggestion"],
            "context": fr["context"],
            "line_number": fr["line_number"],
            "char_offset": fr["char_offset"],
            "char_length": fr["char_length"],
            "rule_id": fr["rule_id"],
            "confidence": fr["confidence"],
            "auto_fixable": bool(fr["auto_fixable"]),
            "created_at": fr["created_at"],
        })

    overall = row["overall_quality_score"] if row["overall_quality_score"] else 0
    dimension_scores = {
        "readability": row["readability_flesch"] if row["readability_flesch"] else 0,
        "grammar_errors": row["grammar_error_count"] if row["grammar_error_count"] else 0,
        "passive_voice": row["passive_voice_pct"] if row["passive_voice_pct"] else 0,
        "coherence": row["coherence_score"] if row["coherence_score"] else 0,
        "plagiarism": row["plagiarism_max_similarity"] if row["plagiarism_max_similarity"] else 0,
        "ai_content": row["ai_content_score"] if row["ai_content_score"] else 0,
    }

    return jsonify({
        "id": row["id"],
        "input_hash": row["input_text_hash"],
        "input_length": row["input_length"],
        "mode": row["mode"],
        "overall_score": overall,
        "passed": overall >= 70,
        "dimension_scores": dimension_scores,
        "tone_profile": row["tone_profile"],
        "style_guide_id": row["style_guide_id"],
        "section_id": row["section_id"],
        "opportunity_id": row["opportunity_id"],
        "document_name": row["document_name"],
        "findings_count": row["findings_count"] if row["findings_count"] else 0,
        "analysis_duration_ms": row["analysis_duration_ms"] if row["analysis_duration_ms"] else 0,
        "classification": row["classification"],
        "created_at": row["created_at"],
        "findings": findings,
    })


@writeguard_api.route("/rewrite", methods=["POST"])
def rewrite():
    """Apply deterministic fixes to text and return rewritten version.

    Accepts JSON: { "text": "..." }
    Returns: { rewritten_text, changes_made, change_count, transitions, original_score, rewritten_score }
    """
    data = request.get_json(silent=True) or {}
    text = data.get("text", "").strip()

    if not text:
        return jsonify({"error": "No text provided"}), 400

    try:
        from tools.writing.rewriter import apply_deterministic_fixes, analyze_transitions

        # Apply fixes
        fix_result = apply_deterministic_fixes(text)
        rewritten = fix_result.get("rewritten_text", text)

        # Analyze transitions on rewritten text
        transitions = analyze_transitions(rewritten)

        # Score before and after
        from tools.pulse.writeguard import run_full_quality_check
        original_check = run_full_quality_check(text)
        rewritten_check = run_full_quality_check(rewritten)

        return jsonify({
            "rewritten_text": rewritten,
            "changes_made": fix_result.get("changes_made", []),
            "change_count": fix_result.get("change_count", 0),
            "rewrite_instructions": fix_result.get("rewrite_instructions", []),
            "transitions": {
                "count": transitions.get("transition_count", 0),
                "per_paragraph": transitions.get("transitions_per_paragraph", []),
                "pacing_score": transitions.get("pacing_score", 0),
                "findings": transitions.get("findings", []),
            },
            "original_score": original_check.get("overall_score", 0),
            "rewritten_score": rewritten_check.get("overall_score", 0),
            "score_improvement": round(
                rewritten_check.get("overall_score", 0) - original_check.get("overall_score", 0), 1
            ),
        })
    except Exception as exc:
        logger.error("WriteGuard rewrite failed: %s", exc)
        return jsonify({"error": f"Rewrite error: {exc}"}), 500


# ---------------------------------------------------------------------------
# Export helpers
# ---------------------------------------------------------------------------

_CLASSIFICATION_COLORS = {
    "SECRET": "red",
    "CUI": "blue",
    "UNCLASSIFIED": "green",
    "U": "green",
}

_DIMENSION_LABELS_EXPORT = {
    "grammar": "Grammar",
    "readability": "Readability",
    "tone": "Tone",
    "passive_voice": "Passive Voice",
    "overused_words": "Word Choice",
    "cliches": "Cliches",
    "consistency": "Consistency",
    "glossary": "Glossary",
    "style_guide": "Style Guide",
    "pacing": "Pacing",
    "plagiarism": "Plagiarism",
    "ai_detection": "AI Detection",
    "portion_marking": "Portion Marking",
}


def _run_analysis_for_export(text: str) -> dict:
    """Run WriteGuard analysis and return the structured result dict."""
    from tools.pulse.writeguard import run_full_quality_check

    result = run_full_quality_check(text)

    all_findings = []
    details = result.get("details", {})
    for dim_key, label in _DIMENSION_LABELS_EXPORT.items():
        dim = details.get(dim_key, {})
        for finding in dim.get("findings", []):
            if isinstance(finding, dict):
                finding_copy = dict(finding)
            else:
                finding_copy = {"message": str(finding)}
            finding_copy.setdefault("dimension", label)
            finding_copy.setdefault("severity", "medium")
            all_findings.append(finding_copy)

    dimensions = {}
    for dim_key, label in _DIMENSION_LABELS_EXPORT.items():
        dim = details.get(dim_key, {})
        dimensions[dim_key] = {
            "label": label,
            "score": dim.get("score", 0),
            "status": dim.get("status", "unavailable"),
            "finding_count": len(dim.get("findings", [])),
        }

    pm = details.get("portion_marking", {})
    portion_marking_detail = {
        "has_marks": pm.get("has_marks", False),
        "all_marked": pm.get("all_marked", True),
        "highest_level": pm.get("highest_level"),
        "derived_banner": pm.get("derived_banner", "UNCLASSIFIED"),
        "designation_block": pm.get("designation_block", ""),
        "violations": pm.get("violations", []),
        "paragraphs": pm.get("paragraphs", []),
        "marked_count": pm.get("marked_count", 0),
        "unmarked_count": pm.get("unmarked_count", 0),
        "total_paragraphs": pm.get("total_paragraphs", 0),
    }

    return {
        "overall_score": result.get("overall_score", 0),
        "passed": result.get("passed", False),
        "composite_status": "PASSED" if result.get("passed") else "NEEDS REVIEW",
        "dimensions": dimensions,
        "findings": all_findings,
        "recommendations": result.get("recommendations", []),
        "portion_marking": portion_marking_detail,
        "word_count": len(text.split()),
        "char_count": len(text),
    }


def _get_banner(analysis: dict) -> str:
    """Return classification banner string from analysis."""
    pm = analysis.get("portion_marking", {})
    return pm.get("derived_banner", "UNCLASSIFIED")


def _export_txt(analysis: dict) -> io.BytesIO:
    """Export analysis as plain text with portion marks + banner header/footer."""
    banner = _get_banner(analysis)
    now_str = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
    lines = [
        f"// {banner}",
        "",
        "WriteGuard Analysis Report",
        f"Generated: {now_str}",
        f"Classification: {banner}",
        (
            f"Overall Score: {analysis['overall_score']}/100"
            f" -- {analysis['composite_status']}"
        ),
        "",
        "=" * 60,
        "DIMENSION SCORES",
        "=" * 60,
    ]
    for _dim_key, dim in analysis.get("dimensions", {}).items():
        lines.append(
            f"  {dim['label']:<20s} {dim['score']:>6.1f}  ({dim['status']})"
        )
    lines.extend(["", "=" * 60, "FINDINGS", "=" * 60])
    for f in analysis.get("findings", []):
        sev = f.get("severity", "medium").upper()
        dim = f.get("dimension", "General")
        msg = f.get("message", str(f))
        lines.append(f"  [{sev}] ({dim}) {msg}")
    if not analysis.get("findings"):
        lines.append("  No findings.")
    lines.extend(["", "=" * 60, "RECOMMENDATIONS", "=" * 60])
    for r in analysis.get("recommendations", []):
        lines.append(f"  - {r}")
    if not analysis.get("recommendations"):
        lines.append("  No recommendations.")
    lines.extend(["", f"// {banner}"])
    buf = io.BytesIO(("\n".join(lines)).encode("utf-8"))
    buf.seek(0)
    return buf


def _export_md(analysis: dict) -> io.BytesIO:
    """Export analysis as markdown with YAML frontmatter."""
    banner = _get_banner(analysis)
    now_str = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    lines = [
        "---",
        f"classification: {banner}",
        f"date: {now_str}",
        f"overall_score: {analysis['overall_score']}",
        f"status: {analysis['composite_status']}",
        "dimensions:",
    ]
    for dim_key, dim in analysis.get("dimensions", {}).items():
        lines.append(f"  {dim_key}: {dim['score']}")
    lines.extend([
        "---",
        "",
        "# WriteGuard Analysis Report",
        "",
        f"> **Classification:** {banner}  ",
        (
            f"> **Score:** {analysis['overall_score']}/100"
            f" -- {analysis['composite_status']}  "
        ),
        f"> **Date:** {now_str}",
        "",
        "## Dimension Scores",
        "",
        "| Dimension | Score | Status |",
        "|-----------|-------|--------|",
    ])
    for _dim_key, dim in analysis.get("dimensions", {}).items():
        lines.append(
            f"| {dim['label']} | {dim['score']:.1f} | {dim['status']} |"
        )
    lines.extend(["", "## Findings", ""])
    if analysis.get("findings"):
        for f in analysis["findings"]:
            sev = f.get("severity", "medium").upper()
            dim = f.get("dimension", "General")
            msg = f.get("message", str(f))
            lines.append(f"- **[{sev}]** ({dim}) {msg}")
    else:
        lines.append("No findings.")
    lines.extend(["", "## Recommendations", ""])
    if analysis.get("recommendations"):
        for r in analysis["recommendations"]:
            lines.append(f"- {r}")
    else:
        lines.append("No recommendations.")
    lines.append("")
    buf = io.BytesIO(("\n".join(lines)).encode("utf-8"))
    buf.seek(0)
    return buf


def _export_json(analysis: dict) -> io.BytesIO:
    """Export full machine-readable analysis result as JSON."""
    export = dict(analysis)
    export["generated_utc"] = datetime.now(timezone.utc).strftime(
        "%Y-%m-%dT%H:%M:%SZ"
    )
    buf = io.BytesIO(
        json_mod.dumps(export, indent=2, default=str).encode("utf-8")
    )
    buf.seek(0)
    return buf


def _export_docx(analysis: dict) -> io.BytesIO:
    """Export analysis as a formatted .docx with classification banners."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        raise ImportError(
            "python-docx is required for .docx export. "
            "Install with: pip install python-docx"
        )

    banner = _get_banner(analysis)
    banner_color_name = _CLASSIFICATION_COLORS.get(banner, "green")
    banner_rgb = {
        "red": RGBColor(0xFF, 0x00, 0x00),
        "blue": RGBColor(0x00, 0x00, 0xFF),
        "green": RGBColor(0x00, 0x80, 0x00),
    }.get(banner_color_name, RGBColor(0x00, 0x80, 0x00))

    doc = Document()

    # Classification banner header
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run(f"// {banner} //")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = banner_rgb

    # Title
    title = doc.add_heading("WriteGuard Analysis Report", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    now_str = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
    meta = doc.add_paragraph()
    meta.add_run(f"Generated: {now_str}\n")
    meta.add_run(f"Classification: {banner}\n")
    score_run = meta.add_run(
        f"Overall Score: {analysis['overall_score']}/100"
        f" -- {analysis['composite_status']}"
    )
    score_run.bold = True

    # WriteGuard score summary table
    doc.add_heading("Dimension Scores", level=2)
    dims = analysis.get("dimensions", {})
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Dimension"
    hdr[1].text = "Score"
    hdr[2].text = "Status"
    for _dim_key, dim in dims.items():
        row = table.add_row().cells
        row[0].text = dim["label"]
        row[1].text = f"{dim['score']:.1f}"
        row[2].text = dim["status"]

    # Findings with severity colors
    doc.add_heading("Findings", level=2)
    findings = analysis.get("findings", [])
    sev_rgb_map = {
        "critical": RGBColor(0xFF, 0x00, 0x00),
        "high": RGBColor(0xFF, 0x8C, 0x00),
        "medium": RGBColor(0xDA, 0xA5, 0x20),
        "low": RGBColor(0x00, 0x80, 0x00),
        "info": RGBColor(0x80, 0x80, 0x80),
    }
    if findings:
        for f in findings:
            sev = f.get("severity", "medium")
            dim = f.get("dimension", "General")
            msg = f.get("message", str(f))
            p = doc.add_paragraph()
            sev_run = p.add_run(f"[{sev.upper()}] ")
            sev_run.bold = True
            sev_run.font.color.rgb = sev_rgb_map.get(
                sev, RGBColor(0xDA, 0xA5, 0x20)
            )
            p.add_run(f"({dim}) {msg}")
    else:
        doc.add_paragraph("No findings.")

    # Recommendations
    doc.add_heading("Recommendations", level=2)
    recs = analysis.get("recommendations", [])
    if recs:
        for r in recs:
            doc.add_paragraph(r, style="List Bullet")
    else:
        doc.add_paragraph("No recommendations.")

    # Portion marking section (if document has marks)
    pm = analysis.get("portion_marking", {})
    if pm.get("has_marks"):
        doc.add_heading("Portion Marking", level=2)
        doc.add_paragraph(
            f"Marked: {pm.get('marked_count', 0)} / "
            f"{pm.get('total_paragraphs', 0)} paragraphs"
        )
        doc.add_paragraph(
            f"Highest level: {pm.get('highest_level', 'N/A')}"
        )
        for v in pm.get("violations", []):
            p = doc.add_paragraph()
            vrun = p.add_run("VIOLATION: ")
            vrun.bold = True
            vrun.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            p.add_run(str(v))

    # Classification banner footer
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run(f"// {banner} //")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = banner_rgb

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _export_pdf(analysis: dict) -> io.BytesIO:
    """Export analysis as a formatted PDF report (air-gap safe, no network)."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.colors import HexColor, white
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer,
        )
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER
    except ImportError:
        raise ImportError(
            "reportlab is required for PDF export. "
            "Install with: pip install reportlab"
        )

    banner = _get_banner(analysis)
    banner_hex = {
        "red": "#FF0000", "blue": "#0000FF", "green": "#008000",
    }.get(_CLASSIFICATION_COLORS.get(banner, "green"), "#008000")
    banner_color = HexColor(banner_hex)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        topMargin=0.75 * inch, bottomMargin=0.75 * inch,
        leftMargin=0.75 * inch, rightMargin=0.75 * inch,
    )
    styles = getSampleStyleSheet()
    normal = styles["Normal"]

    banner_style = ParagraphStyle(
        "BannerStyle", parent=normal,
        fontSize=14, alignment=TA_CENTER, textColor=white,
        backColor=banner_color, spaceAfter=6, spaceBefore=6, leading=20,
    )
    title_style = ParagraphStyle(
        "TitleCenter", parent=styles["Title"],
        alignment=TA_CENTER, spaceAfter=12,
    )
    heading_style = ParagraphStyle(
        "HeadingSub", parent=styles["Heading2"],
        spaceAfter=6, spaceBefore=12,
    )
    sev_hex = {
        "critical": "#FF0000", "high": "#FF8C00", "medium": "#DAA520",
        "low": "#008000", "info": "#808080",
    }

    elements = []

    # Banner header
    elements.append(Paragraph(f"<b>// {banner} //</b>", banner_style))
    elements.append(Spacer(1, 12))

    # Title + metadata
    elements.append(Paragraph("WriteGuard Analysis Report", title_style))
    now_str = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
    elements.append(Paragraph(
        f"Generated: {now_str}",
        ParagraphStyle("Meta", parent=normal, alignment=TA_CENTER),
    ))
    elements.append(Spacer(1, 6))

    # Score gauge (overall + color-coded)
    score = analysis["overall_score"]
    status = analysis["composite_status"]
    score_color = (
        "#008000" if score >= 80
        else ("#DAA520" if score >= 60 else "#FF0000")
    )
    elements.append(Paragraph(
        f'<font size="18" color="{score_color}"><b>{score:.1f}</b></font>'
        f'<font size="12"> / 100 &mdash; {status}</font>',
        ParagraphStyle(
            "ScoreGauge", parent=normal, alignment=TA_CENTER, spaceAfter=12,
        ),
    ))

    # Dimension breakdown table
    elements.append(Paragraph("Dimension Scores", heading_style))
    dims = analysis.get("dimensions", {})
    table_data = [["Dimension", "Score", "Status"]]
    for _dim_key, dim in dims.items():
        s = dim["score"]
        s_color = (
            "#008000" if s >= 80
            else ("#DAA520" if s >= 60 else "#FF0000")
        )
        table_data.append([
            dim["label"],
            Paragraph(
                f'<font color="{s_color}"><b>{s:.1f}</b></font>', normal,
            ),
            dim["status"],
        ])
    t = Table(table_data, colWidths=[2.5 * inch, 1.5 * inch, 2 * inch])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), HexColor("#4472C4")),
        ("TEXTCOLOR", (0, 0), (-1, 0), white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#CCCCCC")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [HexColor("#F2F2F2"), white]),
        ("ALIGN", (1, 0), (1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    elements.append(t)
    elements.append(Spacer(1, 12))

    # Findings with severity
    elements.append(Paragraph("Findings", heading_style))
    findings = analysis.get("findings", [])
    if findings:
        for f in findings:
            sev = f.get("severity", "medium")
            dim = f.get("dimension", "General")
            msg = f.get("message", str(f))
            sc = sev_hex.get(sev, "#DAA520")
            elements.append(Paragraph(
                f'<font color="{sc}"><b>[{sev.upper()}]</b></font>'
                f' ({dim}) {msg}',
                ParagraphStyle(
                    "Finding", parent=normal, spaceAfter=4, leftIndent=12,
                ),
            ))
    else:
        elements.append(Paragraph("No findings.", normal))
    elements.append(Spacer(1, 8))

    # Recommendations
    elements.append(Paragraph("Recommendations", heading_style))
    recs = analysis.get("recommendations", [])
    if recs:
        for r in recs:
            elements.append(Paragraph(
                f"&bull; {r}",
                ParagraphStyle(
                    "Rec", parent=normal, spaceAfter=3, leftIndent=12,
                ),
            ))
    else:
        elements.append(Paragraph("No recommendations.", normal))
    elements.append(Spacer(1, 12))

    # Banner footer
    elements.append(Paragraph(f"<b>// {banner} //</b>", banner_style))

    doc.build(elements)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Export endpoint
# ---------------------------------------------------------------------------

_VALID_FORMATS = {"txt", "md", "json", "docx", "pdf"}
_MIME_TYPES = {
    "txt": "text/plain",
    "md": "text/markdown",
    "json": "application/json",
    "docx": (
        "application/vnd.openxmlformats-officedocument"
        ".wordprocessingml.document"
    ),
    "pdf": "application/pdf",
}


@writeguard_api.route("/export", methods=["POST"])
def export():
    """POST /api/writeguard/export -- Analyze and export as file download.

    Query params:
        format: txt | md | json | docx | pdf (default: json)

    Accepts same body as /analyze:
        - multipart/form-data with a 'file' field (.txt, .md, .docx)
        - application/json with {"text": "..."}

    Returns: file download in the requested format.
    """
    fmt = request.args.get("format", "json").lower().strip()
    if fmt not in _VALID_FORMATS:
        return jsonify({
            "error": (
                f"Unsupported format '{fmt}'. "
                f"Valid: {', '.join(sorted(_VALID_FORMATS))}"
            ),
        }), 400

    text, err = _get_text_from_request()
    if err:
        return jsonify({"error": err}), 400

    if len(text.strip()) < 10:
        return jsonify({
            "error": "Text is too short to analyze (minimum 10 characters)",
        }), 400

    try:
        analysis = _run_analysis_for_export(text)
    except Exception as exc:
        logger.error("WriteGuard analysis failed: %s", exc)
        return jsonify({"error": f"Analysis error: {exc}"}), 500

    exporters = {
        "txt": _export_txt,
        "md": _export_md,
        "json": _export_json,
        "docx": _export_docx,
        "pdf": _export_pdf,
    }

    try:
        buf = exporters[fmt](analysis)
    except ImportError as exc:
        return jsonify({"error": str(exc)}), 501
    except Exception as exc:
        logger.error("Export to %s failed: %s", fmt, exc)
        return jsonify({"error": f"Export error: {exc}"}), 500

    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    filename = f"writeguard_report_{timestamp}.{fmt}"

    return send_file(
        buf,
        mimetype=_MIME_TYPES.get(fmt, "application/octet-stream"),
        as_attachment=True,
        download_name=filename,
    )


# ---------------------------------------------------------------------------
# Word-limit endpoint
# ---------------------------------------------------------------------------


@writeguard_api.route("/word-limits", methods=["POST"])
def word_limits():
    """Check per-section word limits against configured caps.

    Request JSON:
        {
          "text":        str,   content to analyse
          "limits_yaml": str    YAML string with section → word-limit mapping
        }

    Response JSON:
        {
          "sections":     [ {name, word_count, limit, pct, status}, ... ],
          "violations":   [ {name, word_count, limit, excess}, ... ],
          "total_words":  int,
          "total_limit":  int | null,
          "total_pct":    float,
          "total_status": str,
          "passed":       bool
        }
    """
    try:
        from tools.writing.word_limit import check_limits, parse_limits
    except ImportError as exc:
        logger.error("word_limit import failed: %s", exc)
        return jsonify({"error": "word_limit module unavailable"}), 500

    payload = request.get_json(silent=True) or {}
    text: str = payload.get("text", "")
    limits_yaml: str = payload.get("limits_yaml", "")

    if not isinstance(text, str):
        return jsonify({"error": "text must be a string"}), 400
    if not text.strip():
        return jsonify({"error": "text is required"}), 400

    try:
        limits = parse_limits(limits_yaml)
        result = check_limits(text, limits)
    except Exception as exc:  # noqa: BLE001
        logger.error("word_limits check failed: %s", exc)
        return jsonify({"error": f"word limit check failed: {exc}"}), 500

    return jsonify(result)


# ---------------------------------------------------------------------------
# WG 10b — Style Profile CRUD
# ---------------------------------------------------------------------------
@writeguard_api.route("/profile", methods=["POST"])
@require_role(*_WG_CRUD_ROLES)
def create_profile():
    """Extract and save a style profile from sample text.

    Accepts JSON: {"text": "...", "name": "My Style", "content_type": "proposal"}
    or multipart file upload.

    Returns: {id, profile, confidence, source_word_count}
    """
    from tools.writing.style_profiler import extract_profile
    import hashlib as _hl
    import json as _json

    name = ""
    content_type = ""
    text = ""

    if request.content_type and "multipart" in request.content_type:
        f = request.files.get("file")
        if f:
            raw = f.read()
            text = raw.decode("utf-8", errors="replace")
        name = request.form.get("name", "")
        content_type = request.form.get("content_type", "")
    else:
        data = request.get_json(silent=True) or {}
        text = data.get("text", "") or ""
        name = data.get("name", "")
        content_type = data.get("content_type", "")

    if not text.strip():
        return jsonify({"error": "text is required"}), 400
    if len(text.split()) < 100:
        return jsonify({"error": "Minimum 100 words required (1500+ recommended)"}), 400

    try:
        profile = extract_profile(text, name=name or "Style Profile")
    except Exception as exc:  # noqa: BLE001
        logger.error("profile extraction failed: %s", exc)
        return jsonify({"error": f"extraction failed: {exc}"}), 500

    profile_id = str(uuid.uuid4())
    source_hash = _hl.sha256(text.encode("utf-8")).hexdigest()[:16]

    try:
        conn = get_connection()
        try:
            conn.execute(
                "INSERT INTO wg_style_profiles "
                "(id, name, content_type, source_word_count, source_hash, "
                " confidence, profile_json, created_at, is_active, classification) "
                "VALUES (%s, %s, %s, %s, %s, %s, %s, NOW(), 1, %s)",
                (profile_id, profile.get("name", name), content_type,
                 profile.get("source_word_count", 0), source_hash,
                 profile.get("confidence", "low"),
                 _json.dumps(profile),
                 "CUI"),
            )
            conn.commit()
        finally:
            conn.close()
    except Exception as exc:  # noqa: BLE001
        logger.error("profile persistence failed: %s", exc)
        return jsonify({"error": f"save failed: {exc}"}), 500

    return jsonify({
        "id": profile_id,
        "name": profile.get("name"),
        "source_word_count": profile.get("source_word_count"),
        "confidence": profile.get("confidence"),
        "profile": profile,
    })


@writeguard_api.route("/profiles", methods=["GET"])
def list_profiles():
    """List all active style profiles."""
    try:
        conn = get_connection()
        try:
            rows = conn.execute(
                "SELECT id, name, content_type, source_word_count, confidence, "
                "created_at FROM wg_style_profiles WHERE is_active = 1 "
                "ORDER BY created_at DESC LIMIT 100"
            ).fetchall()
        finally:
            conn.close()
    except Exception as exc:  # noqa: BLE001
        logger.error("list_profiles failed: %s", exc)
        return jsonify({"error": str(exc)}), 500

    profiles = []
    for r in rows:
        profiles.append({
            "id": r["id"],
            "name": r["name"],
            "content_type": r["content_type"],
            "source_word_count": r["source_word_count"],
            "confidence": r["confidence"],
            "created_at": r["created_at"],
        })
    return jsonify({"profiles": profiles, "count": len(profiles)})


@writeguard_api.route("/profiles/<profile_id>", methods=["GET"])
def get_profile(profile_id: str):
    """Get full details of a style profile including all 27 features."""
    import json as _json
    try:
        conn = get_connection()
        try:
            row = conn.execute(
                "SELECT * FROM wg_style_profiles WHERE id = %s AND is_active = 1",
                (profile_id,),
            ).fetchone()
        finally:
            conn.close()
    except Exception as exc:  # noqa: BLE001
        return jsonify({"error": str(exc)}), 500

    if not row:
        return jsonify({"error": "profile not found"}), 404

    try:
        profile = _json.loads(row["profile_json"])
    except Exception:
        profile = {}

    return jsonify({
        "id": row["id"],
        "name": row["name"],
        "content_type": row["content_type"],
        "source_word_count": row["source_word_count"],
        "confidence": row["confidence"],
        "profile": profile,
        "created_at": row["created_at"],
    })


@writeguard_api.route("/profiles/<profile_id>", methods=["DELETE"])
@require_role(*_WG_CRUD_ROLES)
def delete_profile(profile_id: str):
    """Soft-delete a style profile."""
    try:
        conn = get_connection()
        try:
            r = conn.execute(
                "UPDATE wg_style_profiles SET is_active = 0, updated_at = NOW() "
                "WHERE id = %s AND is_active = 1",
                (profile_id,),
            )
            conn.commit()
            deleted = getattr(r, "rowcount", 0)
        finally:
            conn.close()
    except Exception as exc:  # noqa: BLE001
        return jsonify({"error": str(exc)}), 500

    if deleted == 0:
        return jsonify({"error": "profile not found"}), 404
    return jsonify({"id": profile_id, "deleted": True})


# ---------------------------------------------------------------------------
# WG 11d — Auto-Classification Scan
# ---------------------------------------------------------------------------
@writeguard_api.route("/classify", methods=["POST"])
def classify():
    """Run auto-classification on text. Returns per-paragraph marks + banner.

    Accepts JSON {"text": "..."}.
    Returns: result dict from auto_mark_document with paragraphs, banner, etc.
    """
    from tools.writing.auto_marker import auto_mark_document

    data = request.get_json(silent=True) or {}
    text = data.get("text", "").strip()
    if not text:
        return jsonify({"error": "text is required"}), 400

    try:
        result = auto_mark_document(text)
    except Exception as exc:  # noqa: BLE001
        logger.error("classify failed: %s", exc)
        return jsonify({"error": str(exc)}), 500

    return jsonify(result)


@writeguard_api.route("/style-match", methods=["POST"])
def style_match():
    """Score text against a saved profile and generate rewrite instructions.

    Accepts JSON {"text": "...", "profile_id": "..."}.
    Returns: {style_match_score, per_feature_deltas, instructions}
    """
    from tools.writing.style_matcher import score_against_profile, generate_style_instructions
    import json as _json

    data = request.get_json(silent=True) or {}
    text = data.get("text", "").strip()
    profile_id = data.get("profile_id", "")

    if not text:
        return jsonify({"error": "text is required"}), 400
    if not profile_id:
        return jsonify({"error": "profile_id is required"}), 400

    try:
        conn = get_connection()
        try:
            row = conn.execute(
                "SELECT profile_json FROM wg_style_profiles WHERE id = %s AND is_active = 1",
                (profile_id,),
            ).fetchone()
        finally:
            conn.close()
        if not row:
            return jsonify({"error": "profile not found"}), 404
        profile = _json.loads(row["profile_json"])
    except Exception as exc:  # noqa: BLE001
        return jsonify({"error": str(exc)}), 500

    score = score_against_profile(text, profile)
    instructions = generate_style_instructions(text, profile)

    return jsonify({
        **score,
        "instructions": instructions,
        "profile_name": profile.get("name", ""),
    })


# ---------------------------------------------------------------------------
# WG 20 — Diff scorer
# ---------------------------------------------------------------------------
@writeguard_api.route("/diff", methods=["POST"])
def diff_analysis():
    """Compare two texts or two stored analyses via diff_scorer.

    Accepts JSON: {"old_text": "...", "new_text": "..."}  OR  {"old_id": "...", "new_id": "..."}
    """
    from tools.writing.diff_scorer import diff_analyses, diff_by_id

    data = request.get_json(silent=True) or {}
    try:
        if data.get("old_id") and data.get("new_id"):
            result = diff_by_id(data["old_id"], data["new_id"])
        else:
            old_text = data.get("old_text", "").strip()
            new_text = data.get("new_text", "").strip()
            if not old_text or not new_text:
                return jsonify({"error": "Provide old_text/new_text or old_id/new_id"}), 400
            result = diff_analyses(old_text, new_text)
        return jsonify(result)
    except Exception as exc:  # noqa: BLE001
        logger.error("diff failed: %s", exc)
        return jsonify({"error": str(exc)}), 500


# ---------------------------------------------------------------------------
# WG 22 — Standards citation validation
# ---------------------------------------------------------------------------
@writeguard_api.route("/standards", methods=["POST"])
def standards_check():
    """Validate standards citations (NIST/TOGAF/OWASP/CIS/ISO/PCI/SABSA)."""
    from tools.writing.standards_validator import validate_citations

    data = request.get_json(silent=True) or {}
    text = data.get("text", "").strip()
    if not text:
        return jsonify({"error": "text is required"}), 400
    try:
        return jsonify(validate_citations(text))
    except Exception as exc:  # noqa: BLE001
        return jsonify({"error": str(exc)}), 500


# ---------------------------------------------------------------------------
# WG 25 — STRIDE threat model coverage
# ---------------------------------------------------------------------------
@writeguard_api.route("/stride", methods=["POST"])
def stride_check():
    """Run STRIDE threat model coverage check."""
    from tools.writing.stride_checker import stride_coverage

    data = request.get_json(silent=True) or {}
    text = data.get("text", "").strip()
    if not text:
        return jsonify({"error": "text is required"}), 400
    try:
        return jsonify(stride_coverage(text))
    except Exception as exc:  # noqa: BLE001
        return jsonify({"error": str(exc)}), 500


# ---------------------------------------------------------------------------
# WG 28 — Technical ↔ Executive parallel generator
# ---------------------------------------------------------------------------
@writeguard_api.route("/parallel", methods=["POST"])
def parallel_generate():
    """Generate side-by-side technical + executive BLUF versions."""
    from tools.writing.parallel_generator import generate_parallel

    data = request.get_json(silent=True) or {}
    text = data.get("text", "").strip()
    if not text:
        return jsonify({"error": "text is required"}), 400
    try:
        return jsonify(generate_parallel(text))
    except Exception as exc:  # noqa: BLE001
        return jsonify({"error": str(exc)}), 500


# ---------------------------------------------------------------------------
# WG 23 — Precedent checker
# ---------------------------------------------------------------------------
@writeguard_api.route("/precedents", methods=["POST"])
def check_precedents_api():
    """Check prior board decisions for similar topics / contradictions."""
    from tools.writing.precedent_checker import check_precedents

    data = request.get_json(silent=True) or {}
    text = data.get("text", "").strip()
    board_type = data.get("board_type")
    if not text:
        return jsonify({"error": "text is required"}), 400
    try:
        return jsonify(check_precedents(text, board_type=board_type))
    except Exception as exc:  # noqa: BLE001
        return jsonify({"error": str(exc)}), 500
