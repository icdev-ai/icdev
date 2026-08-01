"""
Tests for tools/govcon/rfi_docx_exporter.py

Exercises markdown → DOCX conversion. The exporter's public API:
  - _parse_markdown_to_blocks(md_text: str) -> list[dict]
      Block types: h1, h2, h3, hr, code, table, bullet, numbered, quote, para, empty
  - markdown_to_docx(md_text: str, output_path: str, classification: str) -> str
      Returns the output_path string (not a dict).
"""

import sys
from pathlib import Path

import pytest

_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(_ROOT))

try:
    from docx import Document as DocxDocument
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

pytestmark = pytest.mark.skipif(
    not _DOCX_AVAILABLE,
    reason="python-docx not installed",
)

from tools.govcon.rfi_docx_exporter import markdown_to_docx, _parse_markdown_to_blocks


# ── _parse_markdown_to_blocks ─────────────────────────────────────────────────

def test_parse_heading1():
    blocks = _parse_markdown_to_blocks("# Title")
    assert any(b["type"] == "h1" for b in blocks)


def test_parse_heading1_text():
    blocks = _parse_markdown_to_blocks("# Title")
    h1 = next(b for b in blocks if b["type"] == "h1")
    assert h1["text"] == "Title"


def test_parse_heading2():
    blocks = _parse_markdown_to_blocks("## Section")
    assert any(b["type"] == "h2" for b in blocks)


def test_parse_heading3():
    blocks = _parse_markdown_to_blocks("### Subsection")
    assert any(b["type"] == "h3" for b in blocks)


def test_parse_paragraph():
    blocks = _parse_markdown_to_blocks("This is a paragraph.")
    assert any(b["type"] == "para" for b in blocks)


def test_parse_bullet():
    blocks = _parse_markdown_to_blocks("- item one\n- item two")
    bullets = [b for b in blocks if b["type"] == "bullet"]
    assert len(bullets) == 2


def test_parse_numbered_list():
    blocks = _parse_markdown_to_blocks("1. first\n2. second\n3. third")
    numbered = [b for b in blocks if b["type"] == "numbered"]
    assert len(numbered) == 3


def test_parse_table():
    md = "| Col A | Col B |\n|-------|-------|\n| val1 | val2 |"
    blocks = _parse_markdown_to_blocks(md)
    assert any(b["type"] == "table" for b in blocks)


def test_parse_code_block():
    md = "```python\nprint('hello')\n```"
    blocks = _parse_markdown_to_blocks(md)
    assert any(b["type"] == "code" for b in blocks)


def test_parse_code_block_content():
    md = "```\nsome code here\n```"
    blocks = _parse_markdown_to_blocks(md)
    code_block = next(b for b in blocks if b["type"] == "code")
    assert "some code here" in code_block["text"]


def test_parse_empty_string():
    blocks = _parse_markdown_to_blocks("")
    assert isinstance(blocks, list)


def test_parse_mixed_content():
    md = "# Heading\n\nParagraph text.\n\n- bullet A\n- bullet B\n\n## Sub\n\nMore text."
    blocks = _parse_markdown_to_blocks(md)
    types = [b["type"] for b in blocks]
    assert "h1" in types
    assert "para" in types
    assert "bullet" in types


def test_parse_fouo_marker_skipped():
    """FOUO blockquote markers should be skipped (handled via header/footer)."""
    md = "> UNCLASSIFIED//FOUO\n\n# Content"
    blocks = _parse_markdown_to_blocks(md)
    # The FOUO marker line should not produce a quote block
    fouo_quotes = [b for b in blocks if b.get("text", "") == "UNCLASSIFIED//FOUO"]
    assert len(fouo_quotes) == 0


# ── markdown_to_docx ─────────────────────────────────────────────────────────

_SAMPLE_MD = """\
> UNCLASSIFIED//FOUO

# Part 1: Administrative

**Entity Name:** Example Defense Corp

## 1.1 Business Size

Large Business, NAICS 541512.

## 1.2 Clearances

- TS/SCI personnel available
- SCIF-capable facilities

## 1.3 FOCI Status

No FOCI (Veritas Capital, US-based PE ownership).

# Part 2: Technical Approach

| Objective | ICDEV Capability | Grade |
|-----------|-----------------|-------|
| A | Three-Tier Routing | L |
| B | Hot-Reload YAML | M |

```
FORGE Architecture
  Policy Brain → Rule Engine / CoD / CoT
  Cloud: AWS GovCloud / Azure Gov / Ollama
```

> UNCLASSIFIED//FOUO
"""


def test_markdown_to_docx_creates_file(tmp_path):
    out_file = tmp_path / "response.docx"
    markdown_to_docx(_SAMPLE_MD, str(out_file))
    assert out_file.exists()
    assert out_file.stat().st_size > 1000


def test_markdown_to_docx_returns_output_path(tmp_path):
    out_file = tmp_path / "r.docx"
    result = markdown_to_docx(_SAMPLE_MD, str(out_file))
    assert result == str(out_file)


def test_markdown_to_docx_has_headings(tmp_path):
    out_file = tmp_path / "h.docx"
    markdown_to_docx(_SAMPLE_MD, str(out_file))
    doc = DocxDocument(str(out_file))
    all_text = " ".join(p.text for p in doc.paragraphs)
    assert "Part 1" in all_text or "Administrative" in all_text


def test_markdown_to_docx_has_paragraph_text(tmp_path):
    out_file = tmp_path / "p.docx"
    markdown_to_docx(_SAMPLE_MD, str(out_file))
    doc = DocxDocument(str(out_file))
    all_text = " ".join(p.text for p in doc.paragraphs)
    assert "Large Business" in all_text


def test_markdown_to_docx_has_bullet_content(tmp_path):
    out_file = tmp_path / "b.docx"
    markdown_to_docx(_SAMPLE_MD, str(out_file))
    doc = DocxDocument(str(out_file))
    all_text = " ".join(p.text for p in doc.paragraphs)
    assert "TS/SCI" in all_text


def test_markdown_to_docx_empty_content(tmp_path):
    out_file = tmp_path / "empty.docx"
    result = markdown_to_docx("", str(out_file))
    assert result == str(out_file)
    assert out_file.exists()


def test_markdown_to_docx_custom_classification(tmp_path):
    out_file = tmp_path / "classified.docx"
    markdown_to_docx("# Hello\n\nWorld.", str(out_file), classification="SECRET//NOFORN")
    doc = DocxDocument(str(out_file))
    # Classification appears in header/footer — check it's in the document XML
    header_text = " ".join(p.text for p in doc.sections[0].header.paragraphs)
    assert "SECRET" in header_text


def test_markdown_to_docx_table_rendered(tmp_path):
    md = "| A | B |\n|---|---|\n| 1 | 2 |"
    out_file = tmp_path / "table.docx"
    markdown_to_docx(md, str(out_file))
    doc = DocxDocument(str(out_file))
    assert len(doc.tables) >= 1


def test_markdown_to_docx_code_block_rendered(tmp_path):
    md = "```\nsome code\n```"
    out_file = tmp_path / "code.docx"
    markdown_to_docx(md, str(out_file))
    doc = DocxDocument(str(out_file))
    all_text = " ".join(p.text for p in doc.paragraphs)
    assert "some code" in all_text
