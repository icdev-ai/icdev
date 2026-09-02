"""The RFI workbench actually calls a function the DOCX exporter defines.

THE DEFECT THIS GUARDS. `assemble_and_export` called
`export_to_docx(md, path, rfi_number=..., entity_name=...)`. The exporter module
defines `markdown_to_docx(md_text, output_path, classification=...)` and nothing
named `export_to_docx`. The ImportError was caught by a broad `except Exception`,
logged at WARNING, and recorded as a SUCCESSFUL `md` export -- so every DOCX
request since the exporter was written silently produced markdown. Measured
2026-09-01: python-docx 1.2.0 is installed, the exporter produces a valid 37KB
document, and `.tmp/rfi_exports/` contained only `.md` files.

A unit test of the exporter alone passes -- it works. A unit test of the
workbench alone passes -- it returns a path. Only the SEAM between them was
broken, and only a test of the seam can see it.

This is the same shape as the three broken MCP handlers in
`tools/mcp/compliance_server.py` (`stig_check` importing `check_project` from a
module exporting `run_stig_check`), so the structural check here is written to
generalise: it reads what the caller imports and asserts the callee defines it.
"""

from __future__ import annotations

import ast
import inspect
import tempfile
from pathlib import Path

from icdev.core.paths import repo_root

_WORKBENCH = repo_root(__file__) / "tools" / "govcon" / "rfi_workbench.py"
_EXPORTER_MODULE = "tools.govcon.rfi_docx_exporter"


def _imported_from_exporter() -> set[str]:
    """Names rfi_workbench imports from the DOCX exporter, per its own source."""
    tree = ast.parse(_WORKBENCH.read_text(encoding="utf-8"))
    names: set[str] = set()
    for node in ast.walk(tree):
        if isinstance(node, ast.ImportFrom) and node.module == _EXPORTER_MODULE:
            names.update(a.name for a in node.names)
    return names


class TestSeam:
    def test_workbench_imports_something_from_the_exporter(self):
        """A guard that finds no imports would pass vacuously forever."""
        assert _imported_from_exporter(), (
            "rfi_workbench.py no longer imports from "
            f"{_EXPORTER_MODULE} -- this test can no longer see the seam it guards"
        )

    def test_every_imported_name_actually_exists(self):
        """The defect: the caller named a function the exporter never defined."""
        mod = __import__(_EXPORTER_MODULE, fromlist=["*"])
        missing = sorted(n for n in _imported_from_exporter() if not hasattr(mod, n))
        assert not missing, (
            f"rfi_workbench.py imports {missing} from {_EXPORTER_MODULE}, "
            f"which does not define them. Defined: "
            f"{sorted(n for n in vars(mod) if not n.startswith('_'))}"
        )

    def test_export_call_matches_the_exporter_signature(self):
        """The second half of the defect: wrong kwargs, not just a wrong name.

        The old call passed `rfi_number=` and `entity_name=`, which
        `markdown_to_docx` has never accepted.
        """
        from tools.govcon.rfi_docx_exporter import markdown_to_docx

        params = inspect.signature(markdown_to_docx).parameters
        for required in ("md_text", "output_path"):
            assert required in params, f"markdown_to_docx lost its {required!r} parameter"
        for never_accepted in ("rfi_number", "entity_name"):
            assert never_accepted not in params, (
                f"markdown_to_docx gained {never_accepted!r}; if the exporter's contract "
                "changed, update the workbench call rather than relying on this test"
            )


class TestFailureIsNotRecordedAsSuccess:
    def test_source_never_records_md_inside_the_docx_failure_path(self):
        """A DOCX failure must not be recorded as a successful markdown export.

        Recording `md` there is what hid this bug for the exporter's whole
        lifetime: the DB said the export succeeded in the format it fell back
        to, so nothing anywhere reported a problem.
        """
        src = _WORKBENCH.read_text(encoding="utf-8")
        assert '_record_export(session_id, "docx_failed"' in src, (
            "the DOCX failure path no longer records `docx_failed`; a failure "
            "recorded as `md` is indistinguishable from a markdown export"
        )

    def test_library_absent_and_export_failed_are_distinct(self):
        """Two causes, two behaviours -- they justify opposite actions.

        python-docx genuinely absent is a legitimate degradation on an air-gapped
        install. A failure WITH the library present is a defect.
        """
        src = _WORKBENCH.read_text(encoding="utf-8")
        assert "DOCX_AVAILABLE" in src, (
            "rfi_workbench no longer distinguishes 'python-docx absent' from "
            "'export failed' -- both would degrade silently to markdown again"
        )


class TestExporterActuallyProducesADocx:
    def test_markdown_to_docx_writes_a_readable_document(self):
        """End-to-end on the real exporter: the seam is worth fixing only if
        the thing on the other side works."""
        # A plain import, never importorskip: requirements.txt declares
        # `python-docx>=1.1` as a HARD dependency ("Hard imports, no try/except
        # fallback"), so its absence is a broken environment and must fail here.
        # A skip would also satisfy the CI coverage claim while asserting
        # nothing -- which is what the skip census refuses.
        import docx

        from tools.govcon.rfi_docx_exporter import markdown_to_docx

        md = (
            "# RFI Response\n\n"
            "## Part 1: Administrative Data\n\n"
            "| Field | Value |\n|-------|-------|\n| Company | Example Corp |\n\n"
            "Prose with a **bold** run.\n\n- bullet one\n- bullet two\n"
        )
        with tempfile.TemporaryDirectory() as tmp:
            out = Path(tmp) / "probe.docx"
            markdown_to_docx(md, str(out))
            assert out.is_file() and out.stat().st_size > 0
            doc = docx.Document(str(out))
            assert doc.paragraphs, "produced a document with no paragraphs"
            assert doc.tables, "the markdown table did not survive conversion"
            assert "RFI Response" in doc.paragraphs[0].text
